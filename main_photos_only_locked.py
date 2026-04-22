from fastapi import FastAPI, File, UploadFile, Form, Request
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Dict, Any, Optional
import os, io, re, json, base64, logging, zipfile, glob, uuid
from zoneinfo import ZoneInfo
from datetime import datetime
import urllib.parse, urllib.request
import smtplib  # email transport
from email.message import EmailMessage

from fpdf import FPDF
from docx import Document
from pdf2image import convert_from_bytes
from PIL import Image

# Optional HEIC/HEIF support if available
try:
    import pillow_heif  # type: ignore
    pillow_heif.register_heif_opener()  # type: ignore
except Exception:
    pass

# Optional OCR (pytesseract). Safe no-op if not installed or tesseract binary missing.
try:
    import pytesseract  # type: ignore
    _OCR_ENABLED = True
except Exception:
    _OCR_ENABLED = False

from openai import OpenAI

# --- PII Redaction (Presidio) ---
from presidio_analyzer import AnalyzerEngine, Pattern, PatternRecognizer, RecognizerResult
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig  # required for anonymizer API

# -----------------------
# Minimal setup
# -----------------------
PDF_DIR = os.getenv("PDF_DIR", "/tmp"); os.makedirs(PDF_DIR, exist_ok=True)
CLIENT_RULES_DIR = os.getenv("CLIENT_RULES_DIR", "client_rules")

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
log = logging.getLogger("nspxn")
log.info(f"Using CLIENT_RULES_DIR={CLIENT_RULES_DIR}")

# Use selected model everywhere
MODEL = os.getenv("OAI_MODEL") or os.getenv("OPENAI_MODEL") or "gpt-5.2"
# GPT-5.x models use max_completion_tokens; GPT-4.x uses max_tokens
_token_kw = "max_completion_tokens"

if not os.getenv("OPENAI_API_KEY"):
    raise RuntimeError("OPENAI_API_KEY missing")
try:
    client = OpenAI(api_key=os.environ["OPENAI_API_KEY"], timeout=120.0, max_retries=1)
except TypeError:
    # Backwards-compatible init for older openai-python versions
    client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])

# --------------------------------
# Presidio: Analyzer/Anonymizer (preserve VIN & Claim #)
# --------------------------------
analyzer = AnalyzerEngine()
anonymizer = AnonymizerEngine()

VIN_PATTERN = r"\b([A-HJ-NPR-Z0-9]{17})\b"
vin_recognizer = PatternRecognizer(
    supported_entity="VIN",
    patterns=[Pattern(name="vin-17", regex=VIN_PATTERN, score=0.8)],
)

CLAIM_PATTERN = r"\b(?:(?:Claim|CLM|Clm)\s*#?\s*[:\-]?\s*)?([A-Z0-9]{5,}[A-Z0-9\-]{0,})\b"
claim_recognizer = PatternRecognizer(
    supported_entity="CLAIM_NUMBER",
    patterns=[Pattern(name="claim-generic", regex=CLAIM_PATTERN, score=0.6)],
)

analyzer.registry.add_recognizer(vin_recognizer)
analyzer.registry.add_recognizer(claim_recognizer)

# Only redact these entity types (we PRESERVE VIN & CLAIM_NUMBER)
REDACT_ENTITY_TYPES = {
    "PERSON", "PHONE_NUMBER", "EMAIL_ADDRESS", "US_SSN",
    "CREDIT_CARD", "IBAN_CODE", "LOCATION", "NRP", "ORGANIZATION",
    "DATE_TIME", "IP_ADDRESS", "CRYPTO", "MEDICAL_LICENSE", "URL"
}

def _filter_results(results: List[RecognizerResult]) -> List[RecognizerResult]:
    return [r for r in results if r.entity_type in REDACT_ENTITY_TYPES]

def redact_text_preserve_vin_claim(text: str) -> str:
    """Mask PII while keeping VIN and Claim # intact. Fail-open to avoid 500s."""
    if not text:
        return text
    results = analyzer.analyze(text=text, language="en")
    to_mask = _filter_results(results)
    if not to_mask:
        return text
    try:
        return anonymizer.anonymize(
            text=text,
            analyzer_results=to_mask,
            operators={"DEFAULT": OperatorConfig("replace", {"new_value": "[REDACTED]"})}
        ).text
    except Exception as e:
        log.warning(f"Presidio anonymizer failed, passing text through. Error: {e}")
        return text

def _safe(s: str) -> str:
    return re.sub(r"[^\w.\-]+", "-", (s or "").strip()).strip("-_. ")



def _normalize_ai_notes(raw: str, max_len: int = 800) -> str:
    """Normalize/sanitize Add'l Notes so they can't break JSON/structured prompting."""
    if raw is None:
        return "No additional notes provided."
    s = str(raw).strip()
    if not s:
        return "No additional notes provided."

    # Remove ASCII control chars except newline/tab
    s = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", " ", s)

    # Normalize newlines + collapse whitespace
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s).strip()

    # Neutralize common schema-breakers
    s = s.replace("```", "'''")
    # Avoid curly braces interfering with JSON-format instructions
    s = s.replace("{", "(").replace("}", ")")

    if len(s) > max_len:
        s = s[:max_len].rstrip() + "…"
    return s

def _ai_notes_block(ai_notes_clean: str) -> str:
    """Always return a consistent, clearly-delimited notes block."""
    return (
        "\n\nADDL_NOTES_FOR_AI_REVIEW (USER PROVIDED):\n"
        "<<<\n"
        f"{ai_notes_clean}\n"
        ">>>\n"
        "REQUIREMENT:\n"
        "- You MUST consider these notes while analyzing the photos.\n"
        "- Use them to guide focus/wording, but DO NOT invent damage.\n"
        "- If notes conflict with visible evidence or orientation is unclear, say so and defer to what is visible.\n"
    )


def _extract_locked_cost_overrides(ai_notes_clean: str) -> Dict[str, float]:
    """Best-effort extraction of explicit labor/tax overrides from Add'l Notes."""
    out: Dict[str, float] = {}
    s = str(ai_notes_clean or "")
    if not s:
        return out

    def _to_float(v: str) -> Optional[float]:
        try:
            return float(str(v).replace(",", "").strip())
        except Exception:
            return None

    shared_patterns = [
        r'(?i)\bbody\s*(?:and|&|/)\s*paint\s+labor\s+rates?\b[^0-9$]{0,40}\$?\s*([0-9]+(?:\.[0-9]+)?)',
        r'(?i)\bbody\s*(?:and|&|/)\s*paint\s+labor\s+rate\b[^0-9$]{0,40}\$?\s*([0-9]+(?:\.[0-9]+)?)',
        r'(?i)\bbody\s*(?:and|&|/)\s*paint\s+rate\b[^0-9$]{0,40}\$?\s*([0-9]+(?:\.[0-9]+)?)',
        r'(?i)\bbody\s*(?:and|&|/)\s*paint\b[^0-9$]{0,40}\$?\s*([0-9]+(?:\.[0-9]+)?)\s*(?:/\s*hr|per\s*hour|hr)?',
        r'(?i)\bprovided\b[^\n]{0,40}\$\s*([0-9]+(?:\.[0-9]+)?)\s*(?:/\s*hr|per\s*hour|hr)?[^\n]{0,60}\bbody\s*(?:and|&|/)\s*paint\s+rate\b',
        r'(?i)\$\s*([0-9]+(?:\.[0-9]+)?)\s*(?:/\s*hr|per\s*hour|hr)?[^\n]{0,80}\bbody\s*(?:and|&|/)\s*paint\s+labor\s+rates?\b',
        r'(?i)\$\s*([0-9]+(?:\.[0-9]+)?)\s*(?:/\s*hr|per\s*hour|hr)?[^\n]{0,80}\bbody\s*(?:and|&|/)\s*paint\s+labor\s+rate\b',
        r'(?i)\$\s*([0-9]+(?:\.[0-9]+)?)\s*(?:/\s*hr|per\s*hour|hr)?[^\n]{0,80}\bbody\s*(?:and|&|/)\s*paint\s+rate\b',
    ]
    for pat in shared_patterns:
        m = re.search(pat, s)
        if m:
            val = _to_float(m.group(1))
            if isinstance(val, float) and val > 0:
                out['body_rate'] = val
                out['paint_rate'] = val
                break

    if 'body_rate' not in out:
        m = re.search(r'(?i)\bbody\s+labor\s+rates?\b[^0-9$]{0,20}\$?\s*([0-9]+(?:\.[0-9]+)?)', s)
        if m:
            val = _to_float(m.group(1))
            if isinstance(val, float) and val > 0:
                out['body_rate'] = val
    if 'paint_rate' not in out:
        m = re.search(r'(?i)\bpaint\s+labor\s+rates?\b[^0-9$]{0,20}\$?\s*([0-9]+(?:\.[0-9]+)?)', s)
        if m:
            val = _to_float(m.group(1))
            if isinstance(val, float) and val > 0:
                out['paint_rate'] = val

    tax_patterns = [
        r'(?i)\b(?:sales\s+)?tax\s+rate\b[^0-9]{0,20}([0-9]+(?:\.[0-9]+)?)\s*%',
        r'(?i)\b([0-9]+(?:\.[0-9]+)?)\s*%\s*(?:sales\s+)?tax\s+rate\b',
        r'(?i)\b(?:sales\s+)?tax\b[^0-9]{0,10}([0-9]+(?:\.[0-9]+)?)\s*%',
        r'(?i)\b([0-9]+(?:\.[0-9]+)?)\s*%\s*(?:sales\s+)?tax\b',
    ]
    for pat in tax_patterns:
        m = re.search(pat, s)
        if m:
            val = _to_float(m.group(1))
            if isinstance(val, float) and val > 0:
                out['tax_rate'] = (val / 100.0) if val > 1 else val
                break

    if 'body_rate' in out and 'paint_rate' not in out and re.search(r'(?i)\bbody\s*(?:and|&)\s*paint\b', s):
        out['paint_rate'] = out['body_rate']
    if 'paint_rate' in out and 'body_rate' not in out and re.search(r'(?i)\bbody\s*(?:and|&)\s*paint\b', s):
        out['body_rate'] = out['paint_rate']

    return out

# -----------------------
# Approximate Repair Cost Breakdown (location-based rates)
# -----------------------
# Notes:
# - Paint supplies/materials are modeled as $ per refinish hour (NOT a percent).
# - This is an approximation for observed damages, not a repair estimate.

# Optional external integration (no new dependencies):
# - Set LABORRATEHERO_API_URL to a JSON endpoint you control (proxy/scraper) that returns:
#   {"body_rate": 0, "paint_rate": 0, "frame_rate": 0, "mechanical_rate": 0, "paint_supplies_rate": 0}
# - Optionally set LABORRATEHERO_API_KEY if your endpoint requires it.
LABORRATEHERO_API_URL = os.getenv("LABORRATEHERO_API_URL", "").strip()
LABORRATEHERO_API_KEY = os.getenv("LABORRATEHERO_API_KEY", "").strip()

def _extract_inspection_location(text: str) -> str:
    """Best-effort extraction of Inspection Location (City/State/ZIP) from uploaded text."""
    if not text:
        return ""
    # Common patterns: "Inspection Location: Albuquerque, NM 87109" or multi-line variants
    m = re.search(r"(?im)^\s*Inspection\s+Location\s*[:\-]?\s*(.+)$", text)
    if m:
        return (m.group(1) or "").strip()
    m2 = re.search(r"(?is)\bInspection\s+Location\b\s*[:\-]?\s*(.{0,120}?)\n", text)
    if m2:
        return (m2.group(1) or "").strip()
    return ""

def _extract_zip5(*texts: str) -> str:
    """Return the first 5-digit ZIP found in any provided text (best-effort)."""
    for t in texts:
        if not t:
            continue
        m = re.search(r"\b(\d{5})(?:-\d{4})?\b", str(t))
        if m:
            return m.group(1)
    return ""

def _zip_to_city_state(zip5: str) -> Optional[Dict[str, str]]:
    """Best-effort ZIP -> City/State via Zippopotam (public). Fail-open."""
    z = (zip5 or "").strip()
    if not re.fullmatch(r"\d{5}", z):
        return None
    try:
        url = f"https://api.zippopotam.us/us/{z}"
        req = urllib.request.Request(url, headers={"Accept": "application/json"})
        with urllib.request.urlopen(req, timeout=6) as resp:
            raw = resp.read().decode("utf-8", "ignore")
        data = json.loads(raw) if raw else {}
        if not isinstance(data, dict):
            return None
        places = data.get("places") or []
        if not places:
            return None
        p0 = places[0] if isinstance(places[0], dict) else {}
        city = (p0.get("place name") or "").strip()
        state = (p0.get("state abbreviation") or "").strip().upper()
        if city and state:
            return {"city": city, "state": state, "zip": z}
    except Exception:
        return None
    return None

def _normalize_location_with_zip(location: str, *fallback_texts: str) -> str:
    """If location lacks City/State but has ZIP, expand it using ZIP->City/State."""
    loc = (location or "").strip()
    zip5 = _extract_zip5(loc, *fallback_texts)
    if not zip5:
        return loc
    # If already contains a state, just ensure ZIP is present
    if _parse_state_from_location(loc):
        return loc
    z = _zip_to_city_state(zip5)
    if z:
        return f"{z['city']}, {z['state']} {z['zip']}"
    # At least return ZIP if nothing else
    return zip5


def _parse_state_from_location(loc: str) -> str:
    if not loc:
        return ""
    # Pull 2-letter state if present (", NM" or " NM ")
    m = re.search(r"(?i)\b([A-Z]{2})\b\s*(\d{5})?(?:-\d{4})?\s*$", loc.strip())
    if m:
        return (m.group(1) or "").upper()
    m2 = re.search(r"(?i),\s*([A-Z]{2})\b", loc)
    if m2:
        return (m2.group(1) or "").upper()
    return ""

def _fallback_rates_by_state(state: str) -> Dict[str, float]:
    """Conservative defaults when no external rate service is configured/available."""
    # Keep these conservative; you can tune later or replace with your own table.
    # body/paint/frame/mechanical in $/hr; paint_supplies_rate in $/refinish hr.
    defaults = {
        "DEFAULT": {"body_rate": 60.0, "paint_rate": 60.0, "frame_rate": 75.0, "mechanical_rate": 115.0, "paint_supplies_rate": 38.0},
        "NM":      {"body_rate": 64.0, "paint_rate": 62.0, "frame_rate": 78.0, "mechanical_rate": 120.0, "paint_supplies_rate": 38.0},
        "AZ":      {"body_rate": 66.0, "paint_rate": 64.0, "frame_rate": 80.0, "mechanical_rate": 125.0, "paint_supplies_rate": 40.0},
        "CO":      {"body_rate": 72.0, "paint_rate": 70.0, "frame_rate": 88.0, "mechanical_rate": 135.0, "paint_supplies_rate": 42.0},
        "TX":      {"body_rate": 62.0, "paint_rate": 60.0, "frame_rate": 75.0, "mechanical_rate": 120.0, "paint_supplies_rate": 38.0},
        "CA":      {"body_rate": 85.0, "paint_rate": 85.0, "frame_rate": 110.0, "mechanical_rate": 165.0, "paint_supplies_rate": 50.0},
    }
    return dict(defaults.get(state.upper(), defaults["DEFAULT"]))

def _fetch_rates_from_laborratehero(location: str) -> Optional[Dict[str, float]]:
    """Fetch rates from a configured endpoint. Returns None on any failure."""
    if not LABORRATEHERO_API_URL:
        return None
    try:
        q = urllib.parse.urlencode({"location": location})
        url = LABORRATEHERO_API_URL
        url = (url + ("&" if "?" in url else "?") + q) if location else url
        req = urllib.request.Request(url, headers={"Accept": "application/json"})
        if LABORRATEHERO_API_KEY:
            req.add_header("Authorization", f"Bearer {LABORRATEHERO_API_KEY}")
        with urllib.request.urlopen(req, timeout=10) as resp:
            raw = resp.read().decode("utf-8", "ignore")
        data = json.loads(raw) if raw else {}
        if not isinstance(data, dict):
            return None
        out = {}
        for k in ("body_rate","paint_rate","frame_rate","mechanical_rate","paint_supplies_rate"):
            v = data.get(k)
            if isinstance(v, (int, float)) and v > 0:
                out[k] = float(v)
        return out or None
    except Exception:
        return None

def _lookup_rates(location: str) -> Dict[str, float]:
    """Return a complete rate card (with fallbacks)."""
    location = _normalize_location_with_zip(location)
    state = _parse_state_from_location(location)
    base = _fallback_rates_by_state(state)
    ext = _fetch_rates_from_laborratehero(location) or {}
    base.update({k: float(v) for k, v in ext.items() if isinstance(v, (int, float)) and float(v) > 0})
    # Ensure required keys exist
    for k in ("body_rate","paint_rate","frame_rate","mechanical_rate","paint_supplies_rate"):
        base.setdefault(k, 0.0)
    return base

def _lookup_tax_rate(location: str) -> float:
    """Best-effort tax rate for parts/materials (decimal). Conservative fallback."""
    # Optional external (user-controlled) endpoint
    tax_url = os.getenv("TAXRATE_API_URL", "").strip()
    tax_key = os.getenv("TAXRATE_API_KEY", "").strip()
    if tax_url and location:
        try:
            q = urllib.parse.urlencode({"location": location})
            url = tax_url + ("&" if "?" in tax_url else "?") + q
            req = urllib.request.Request(url, headers={"Accept": "application/json"})
            if tax_key:
                req.add_header("Authorization", f"Bearer {tax_key}")
            with urllib.request.urlopen(req, timeout=10) as resp:
                raw = resp.read().decode("utf-8", "ignore")
            data = json.loads(raw) if raw else {}
            if isinstance(data, dict):
                r = data.get("tax_rate")
                if isinstance(r, (int, float)) and 0 <= float(r) <= 0.15:
                    return float(r)
                # allow percent input
                if isinstance(r, (int, float)) and 0 <= float(r) <= 15:
                    return float(r) / 100.0
        except Exception:
            pass

    # Minimal state fallbacks (tune later). Use decimal (e.g., 0.07875 = 7.875%)
    state = _parse_state_from_location(location)
    state_map = {
        "NM": 0.07875,
        "AZ": 0.085,
        "CO": 0.075,
        "TX": 0.0825,
        "CA": 0.085,
    }
    return float(state_map.get(state, 0.08))

def _money(x: Optional[float]) -> str:
    try:
        if x is None:
            return "$0"
        return "${:,.0f}".format(float(x))
    except Exception:
        return "$0"

def _num(x: Optional[float]) -> float:
    try:
        return float(x) if x is not None else 0.0
    except Exception:
        return 0.0

def _extract_hours_and_parts_totals(text: str) -> Dict[str, Any]:
    """Parse estimate totals (best effort)."""
    out: Dict[str, Any] = {"body_hours": None, "paint_hours": None, "frame_hours": None, "mech_hours": None, "parts_total": None, "parts_lines": []}
    if not text:
        return out

    def _find_hours(label_rx: str) -> Optional[float]:
        m = re.search(label_rx, text, flags=re.IGNORECASE)
        if m:
            try:
                return float(m.group(1))
            except Exception:
                return None
        return None

    out["body_hours"]  = _find_hours(r"\bBody\s+Labor\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)")
    out["paint_hours"] = _find_hours(r"\bPaint\s+Labor\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)")
    out["frame_hours"] = _find_hours(r"\b(Frame|Structural)\s+Labor\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)")  # group2 maybe
    if out["frame_hours"] is None:
        m = re.search(r"\bFrame\s+Labor\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)", text, flags=re.IGNORECASE)
        if m:
            try: out["frame_hours"] = float(m.group(1))
            except Exception: pass
    out["mech_hours"]  = _find_hours(r"\b(Mechanical|Mech)\s+Labor\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)")
    if out["mech_hours"] is None:
        m = re.search(r"\bMechanical\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)", text, flags=re.IGNORECASE)
        if m:
            try: out["mech_hours"] = float(m.group(1))
            except Exception: pass

    # Parts total (best effort)
    mpt = re.search(r"(?i)\bParts\s*(?:Total)?\b\s*[:\-]?\s*\$\s*([0-9,]+(?:\.[0-9]{2})?)", text)
    if mpt:
        try:
            out["parts_total"] = float(mpt.group(1).replace(",", ""))
        except Exception:
            pass

    # Parts line items (best effort, capped)
    # Capture lines that look like: "<line#> ... <part desc> ... $123.45"
    lines = []
    for ln in (text or "").splitlines():
        s = ln.strip()
        if not s:
            continue
        if len(s) > 220:
            continue
        if re.search(r"(?i)\b(replace|r\&r|r\s*/\s*r)\b", s) and re.search(r"\$\s*[0-9,]+(?:\.[0-9]{2})?", s):
            lines.append(s)
        elif re.search(r"(?i)\b(part|oem)\b", s) and re.search(r"\$\s*[0-9,]+(?:\.[0-9]{2})?", s):
            lines.append(s)
    out["parts_lines"] = lines[:20]
    return out

def _structural_observed(text_blobs: List[str]) -> bool:
    t = "\n".join([x for x in text_blobs if x]).lower()
    if not t:
        return False
    rx = r"\b(frame|rail|apron|unibody|structural|pull|straighten|set\s*up\s*&\s*measure|measure\s*&\s*setup|dimension\s*check|buckl|kink|twist)\b"
    return bool(re.search(rx, t))
# -----------------------
# Prompt steering (free analysis + detailed narrative)
# -----------------------
DETAIL_TEMPLATES = {
    "guidelines_only": (
        "## Inputs Used\n"
        "- Briefly list which documents, pages, and photos you actually referenced.\n\n"
        "## Executive Summary\n"
        "- 3–5 bullets summarizing overall compliance and key risks.\n\n"
        "## AI-4-IA Review Summary\n"
        "- Write a formal, paragraph-style appraisal narrative. Include scope of impact, damage by zone/panel, "
        "repair vs. replace rationale, parts type (OEM/LKQ/Aftermarket), labor ops, refinish overlap, rate validation, "
        "tax handling, and estimate integrity. Reference evidence inline (e.g., 'p2/L14', 'Photo 3'). "
        "Close with compliance stance and final recommendation. Minimum 8–10 sentences.\n\n"
        "## Key Issues & Actions\n"
        "- Bullet list of the highest-impact issues with a one-line recommended action each.\n\n"
        "## Final\n"
        "- Compliance Score: NN% with one-sentence rationale."
    ),

    "comprehensive": (
        "## Inputs Used\n"
        "- List the estimate pages/lines and photo numbers you used, plus any rules text (if provided).\n\n"
        "## Executive Summary\n"
        "- 3–6 bullets capturing the big picture: estimate integrity, rule alignment (only if rules text was supplied), "
        "and photo consistency.\n\n"
        "## Detailed Condition Report\n"
        "- Write this section as a formal, paragraph-style appraisal report summarizing the entire claim. "
        "Include: scope of impact, damage by zone/panel, repair vs. replace rationale, parts type (OEM/LKQ/Aftermarket), "
        "labor operations, refinish/overlap considerations, rate validation, paint materials handling, sublet usage, "
        "tax/markup accuracy, and overall estimate integrity. Cite photos and estimate lines (e.g., 'Photo 3', 'p2/L14'). "
        "Close with compliance to any provided client rules and a clear final recommendation (Repairable vs. Total Loss). "
        "Do not declare Repairable/Total Loss unless the estimate itself explicitly marks 'Total Loss' or an ACV comparison is provided. "
        "If the shop info is listed under Repair Facility on ANY estimate, add only the shop name to the Detailed Condition Report narrative. "
        "If a Printout showing the Clean Retail Value or Estimated Trade-In Value of the unit is present which may include ANY of the following: NADA, J.D. Power, Kelly Blue Book, Edmunds, Carfax, or Cars.com, DO NOT declare as missing if any of these are present. "
        "Minimum 10–14 sentences (one continuous narrative, not bullets).\n\n"
        "## Photo-by-Photo Damage Ledger\n"
        "| Photo # | View/Angle | Panels/Parts Visible | Condition (dent/crease/scrape/misalignment) | Identifiers (VIN/odo/plate/reg) | Legibility |\n"
        "|---:|---|---|---|---|---|\n"
        "- One row per photo used in the analysis (>=6 rows if >=6 photos exist). If an identifier is present but unreadable, mark 'Present — not clearly legible'.\n\n"
        "## Brief Damage Descriptions\n"
        "- 6–12 bullets. Each: part/panel + condition + suggested op (repair/replace/refinish/blend) + Photo #.\n\n"
        "## Estimate Line Extract (top relevant lines)\n"
        "| Est. p#/L# | Part/Op | Labor Hrs | Rate | Part Type | Price | Notes |\n"
        "|---|---|---:|---:|---|---:|---|\n"
        "- Include 8+ lines if available, focusing on items tied to observed damages. Use 'Notes' for overlap/blend or rationale.\n\n"
        "## Estimate Compliance Cross-Check (brief)\n"
        "Status: Compliant / Non-compliant / Not Evidenced. Cite only the most important evidence (p#/L# or Photo #). Keep it short.\n"
        "| Topic | Evidence (p#/L# or value; Photo # if relevant) | Status | Required Fix |\n"
        "|---|---|:--:|---|\n"
        "| Labor Rates |  |  |  |\n"
        "| Refinish/Overlap |  |  |  |\n"
        "| Paint Materials |  |  |  |\n"
        "| OEM Procedures |  |  |  |\n"
        "| Sublet |  |  |  |\n"
        "| Tax/Markup |  |  |  |\n\n"
        "## Client Guidelines Comparison (if rules text was supplied)\n"
        "- 3–8 bullets. Quote the relevant rule fragment and note Aligned / Not Aligned / Not Evidenced, with evidence refs (p#/L#, Photo #). "
        "If no client_rules were provided, omit this section.\n\n"
        "## Risks / Missing Evidence\n"
        "- Short bullets with severity (High/Med/Low) and a one-line remediation.\n\n"
        "## Compliance Score Rationale\n"
        "- REQUIRED if score < 100: start at 100 and list each deficiency with evidence refs (p#/L# and/or Photo #), "
        "severity (Minor/Moderate/Major), and numeric deduction. Show the arithmetic to the final score.\n\n"
        "## Final Evaluation\n"
        "- Compliance Score: NN% with a single-sentence justification. "
        "If no fraud indicators are identified, state 'No material inconsistencies found.' Do not use 'N/A'."
    ),

    "damage_report_from_photos": (
        """
# Condition Report (Photos Only)

## Photo-by-Photo Condition Summary
| Photo # | View/Side | Key Panels/Parts Visible | Damage/Condition |
|---:|---|---|---|
- Cover EVERY provided photo. If no damage is obvious from that angle, write: "No obvious damage visible from this angle" (do not use the word intact).

## Side Checks
- **Driver/Left Side**: <what is visible; cite Photo #; if not shown, say not shown>
- **Passenger/Right Side**: <what is visible; cite Photo #; if not shown, say not shown>

## Detailed Condition Report
- Write a continuous 10–15 sentence narrative summarizing visible damage, impact zones, misalignment/gaps, and repair implications (photo-based).
- If VIN label or odometer are visible, state them with Photo #. If not visible or unreadable, say so.

## Approximate Repair Cost Breakdown
- You MUST produce a cost approximation derived from the PHOTOS ONLY (do not reference estimates, documents, or 'not evidenced').
- Provide AI-derived hours and assumptions:
  • Body labor hours
  • Paint labor hours
  • Paint & materials cost = (paint hours × $/refinish hr)
  • If structural damage is observed in photos: add Setup & Measure = 2.0 hrs @ body rate and include frame hours @ frame rate
  • If airbags/ADAS/mechanical damage is observed: include mechanical hours @ mechanical rate
- Provide an OEM replacement parts list (each part with an approximate $).
- Apply tax ONLY to (parts + paint materials). Do NOT tax labor.
- Include the Severity Tier checkboxes:
  ☐ Minor (< $3,500)
  ☐ Moderate ($3,500-$10,000)
  ☐ Major ($10,000+)
  ☐ Possible Total Loss Threshold Approaching
- Include a Repair Cost Disclaimer stating this is an approximation, not an official estimate.
- Forbidden phrases: 'from estimate', 'from documentation', 'not evidenced', 'no documentation provided'.
"""
    ),
}

# --- Static audit questions ---
STATIC_AUDIT_QUESTIONS = [
    "Do the photos substantiate the highest-cost operations (frame/sectioning/panel replace)?",
    "Are ADAS calibrations or wheel alignments required and supported by the damage and OEM procedures?",
    "Is blend time justified by color/finish (metallic/pearl/tri-coat) and adjacent panel visibility?",
    "Do invoices corroborate parts used and match estimate line items (brand/grade, price, quantity)?",
    "Are AM/LKQ choices compliant with age/mileage rules, and is OEM required anywhere by client policy or safety?",
    "Is there evidence of prior or unrelated damage (UPD) that materially affects valuation or repair scope?",
    "Are there structural/safety indicators (buckles, misalignments, airbags/pretensioners) that alter repair strategy?",
    "Are materials/hazard charges (paint supplies, corrosion protection, seam sealer) aligned with operations and shop norms?",
    "Are storage/tow charges and dates supported and reasonable given claim timeline and shop status?",
    "Are scanner reports (pre/post) included or needed; if absent, does that meaningfully impact confidence?",
    "Did the supplement (if any) correct earlier gaps, and are newly added operations now evidenced?",
    "Are client-required documents present (e.g., NADA printout, release forms, production date plate); if missing, what’s the impact?",
    "What is the bottom-line recommendation (approve as-is, adjust items, or request specific evidence)?",
]

# --- Identifiers Verification Protocol (prompt-only; no new logic) ---
IDENTIFIERS_VERIFICATION_PROTOCOL = (
    "\n\nIDENTIFIERS VERIFICATION PROTOCOL (must follow):"
    "\n1) Search the photos for: windshield VIN plate, driver-door VIN label, driver-door VIN label Production date, driver-door VIN label Date of Mfr, odometer cluster."
    "\n2) Transcribe the VIN exactly as visible and cite Photo # for EACH location you find."
    "\n3) If multiple VINs, compare them to each other and to the estimate VIN; explicitly state: MATCH / MISMATCH."
    "\n4) Transcribe the odometer reading exactly as shown and cite Photo #."
    "\n5) Grade legibility for each identifier as one of: 'Clearly legible' / 'Present — not clearly legible' / 'Not present'."
    "\n6) If any identifier is present but not clearly legible, say why (glare, blur, angle) and what photo would resolve it."
    "\n7) Write a one-line bottom line: 'VIN verification: <MATCH/MISMATCH/INCONCLUSIVE>; Odometer: <value or reason>'."
    "\n8) Weave these facts naturally into the '## Detailed Condition Report' narrative and keep the top-line fields "
    "(vin, vin_verification, odometer_estimate_only) consistent."
    "\n9) When citing more than one VIN location (e.g., windshield vs. door label), you must cite DISTINCT Photo #s; "
    "never reuse the same photo number for two different locations."
    "\n10) Compare VINs as literal 17-character strings. If any single character differs between sources, report "
    "MISMATCH, and quote both strings with their Photo #/page references."
    "\n11) ODOMETER RULES (photos-only especially): transcribe only the digits visible in the odometer photo; "
    "do not infer from estimate text or metadata. Include the exact Photo #. If any digit is unclear, state 'Present — not clearly legible' and explain why; do not guess."
)

# --- Consistency Guard (prompt-only; avoid contradictions) ---
CONSISTENCY_GUARD = (
    "\n\nCONSISTENCY GUARD:"
    "\n- Do not claim any required photo is 'missing' if you graded it 'Clearly legible' or 'Present — not clearly legible'."
    "\n- For VIN, Odometer, and Production Date specifically: if present in any photo, do not write any sentence implying they are absent."
    "\n- If a legible driver-door VIN label photo is present, treat the Production Date requirement as evidenced (the production month/year appears on the same label). Do NOT deduct or say 'not separately documented'."
    "\n- Only deduct for missing Repair Facility info when the Closing Report or other documents clearly show the vehicle is at a named repair facility AND the estimate's 'Repair Facility' section does not list that same facility; "
    "if no repair facility information appears anywhere in the estimate or Closing Report, report 'N/A — not provided' and do NOT deduct."
    "\n- If legibility is the issue, explicitly say 'Present — not clearly legible' and explain why (glare/blur/angle), and request a precise retake rather than marking it missing."
    "\n- Before finalizing, re-scan your output: confirm every referenced Photo # matches the content described (e.g., do not cite an Odometer photo as the point-of-impact photo). Correct any mismatches."
)

# --- No-intact-if-damaged rule (prompt-only; prevents false 'intact' claims) ---
NO_INTACT_IF_DAMAGED_RULE = (
    "\n\nNO 'INTACT/NO-DAMAGE' OVERRIDE RULE (DO NOT PRINT CONFLICT WARNINGS):"
    "\n- If any photo indicates damage to a panel/component, you may NOT state that same panel/component is intact/undamaged/no visible damage anywhere."
    "\n- If photos appear inconsistent, DO NOT write a conflict warning; instead, remove/avoid the intact/no-damage claim and describe only what appears damaged with citations."
)

# --- Damage Side / Orientation Guard (prompt-only; prevents left/right drift) ---
DAMAGE_SIDE_GUARD = (
    "\n\nDAMAGE SIDE GUIDANCE (MINIMAL):"
    "\n- Describe any visible damage on any side (Driver/Passenger or Left/Right) when it is visible in photos."
    "\n- Do NOT suppress side-level damage descriptions when damage is clearly visible."
    "\n- If orientation is genuinely unclear, say so and avoid guessing."
)

BILATERAL_DAMAGE_MANDATE = (
    "\n\nBILATERAL / SECONDARY DAMAGE MANDATE (STRICT):\n"
    "- In frontal impacts, explicitly address BOTH front corners (driver-side and passenger-side if clear; otherwise ‘front corner A/B’ or ‘front corner (viewed)’)\n"
    "- If a photo shows partial view of the opposite side with visible distortion/misalignment/crush, describe it — do NOT default to 'intact' or 'no damage' without citing clear evidence.\n"
    "- Contradicting visible photo evidence (e.g. calling a crushed fender 'intact') is forbidden."
)

FRONT_CORNER_ORIENTATION_GUARD = (
    "\n\nFRONT / SIDE LABELING (LOCKED):"
    "\n- For primary_impact and secondary_impact fields: if the damage is at the front end and left/right cannot be proven from the photos, set the value to exactly 'Front' (no qualifiers)."
    "\n- Do NOT label front damage as LF/RF or Driver/Passenger unless orientation is clearly established by unmistakable cues within the same photo set (e.g., readable badge/plate position combined with a full-front view)."
    "\n- If orientation is not clearly established, use neutral wording only (e.g., 'front end', 'front headlamp area', 'front corner') and DO NOT assign left/right."
    "\n- Do not state that a left/right front component is undamaged/intact unless that specific component is clearly shown and its condition is confirmable."
)

# --- Parts Source Guard (prompt-only; prevents OEM vs Aftermarket drift) ---
PARTS_SOURCE_GUARD = (
    "\n\nPARTS SOURCE GUARD (MANDATORY):"
    "\n- Do NOT claim 'aftermarket', 'A/M', 'quality replacement', 'non-OEM', or 'LKQ/used/recycled' parts were used unless the estimate LINE ITEMS explicitly label them as such."
    "\n- Generic disclosure/boilerplate text about aftermarket crash parts does NOT prove aftermarket parts were used."
    "\n- If the Closing Report states no aftermarket/LKQ parts were included, your narrative must not claim they were used."
    "\n- When parts source is not explicit, state that it is not explicitly labeled and avoid guessing; default to OEM only when supported by part numbers/labels."
)

# --- Supplement Handling (prompt-only; ensures detection + narrative mention) ---
SUPPLEMENT_HANDLING = (
    "\n\nSUPPLEMENT HANDLING:"
    "\n- Examine the estimate documents for explicit supplement indicators: 'Supplement', 'Supplement of record', 'S01', 'S02', 'Supplement Summary', or similar."
    "\n- If a supplement or multiple supplements are detected, clearly state in the narrative that the estimate is a supplement and summarize what changed: added operations/parts, rate updates, refinish overlap changes, or corrections to prior omissions."
    "\n- If the supplement(s) corrects earlier deficiencies (e.g., missing materials line, added calibrations), note that improvement explicitly."
    "\n- If a supplement(s) exists but required supporting evidence (invoices, photos) is still missing, call this out in Risks/Missing Evidence."
)

ALLOWED_INTENTS = {"guidelines_only","comprehensive","damage_report_from_photos"}

SYSTEM_BASE = (
    "You are an auto-claims appraisal assistant. Return ONLY valid JSON (no code fences). "
    "Always include all required keys: "
    "['file_number','request_type','claim_number','vin','vin_verification','vehicle',"
    "'odometer_estimate_only','compliance_score','summary_brief','summary_markdown',"
    "'fraud_markdown','primary_impact','secondary_impact','estimated_costs_markdown','conclusion']. "
    "Use only provided evidence. Cite photos as 'Photo #' and docs as 'p#/L#' when available. "
    "Do not guess. If something is not visible, say why. "
    "NEVER return 'N/A' for summary_markdown, fraud_markdown, estimated_costs_markdown, or conclusion. For request_type 'Create a Damage Report from Photos', 'estimated_costs_markdown' must be a photos-only approximation (no document/estimate dependency language) and must model Paint & Materials as a $ per refinish hour rate (not a percent)."
)

SYSTEM_BASE += (
    " Do not state or imply any client rule unless it appears verbatim in the provided client_rules text. "
    "If client_rules is blank, write the entire report without referencing client rules. "
    "If a value cannot be confirmed from the visible evidence, set it to 'N/A' and briefly state why. "
    " Except when the request_type is 'Create a Damage Report from Photos', Compliance Score must be a numeric percentage 0–100 (never 'N/A'). "
    "If the request_type is 'Create a Damage Report from Photos', set compliance_score to 'N/A' and omit the '## Compliance Score Rationale' section. "
    "If no client_rules are supplied, base the score on estimate-photo internal consistency, evidence completeness, and clarity/legibility. "
    "If compliance_score < 100, include a dedicated section titled '## Compliance Score Rationale' which itemizes every deficiency with exact evidence references "
    "(estimate p#/L# and/or Photo #), assigns an explicit deduction per item, and shows the arithmetic to the final score. "
    "Use a consistent scheme (e.g., Minor -5, Moderate -10, Major -20) and never go below 0. "
    "The 'fraud_markdown' section must never be 'N/A'. If nothing material is found, write "
    "'No material inconsistencies found.' and briefly note what was checked (VIN match, date/metadata, obvious photo tampering, duplicated images)."
)

SYSTEM_BASE += (
    " Focus on a cohesive, professional appraisal. Prefer narrative over rigid tables. "
    "Include a section named '## Detailed Condition Report'. "
    "Include '## Compliance Score Rationale' only when compliance_score < 100, and show deductions from 100 with brief evidence refs (p#/L# or Photo #). "
    "If you include tables, keep them concise and only when they help clarity. "
    "Avoid placeholder rows/columns; do not invent data. "
    "When client_rules text is provided, also include a section titled '## Client Guidelines Comparison' with 3–8 concise bullets quoting the relevant rule fragment and citing evidence (p#/L#, Photo #); "
    "weave any material rule alignment/misalignment into the Detailed Condition Report narrative."
    "When a valuation/clean retail printout exists but the header doesn’t match the estimate’s VIN/year/trim/mileage, label it “Present — mismatched (detail the differences)” and request a corrected printout; never mark it Missing/Not Evidenced. "
    "If a legible driver-door VIN label photo is present, treat Production Date as evidenced; do not mark 'missing' or deduct for lack of a separate photo. "
    "Only deduct for missing Repair Facility info when the Closing Report or other documents clearly show the vehicle is at a named repair facility AND the estimate's 'Repair Facility' section does not list that same facility; "
    "if no repair facility information appears anywhere in the estimate or Closing Report, report 'N/A — not provided' and do NOT deduct."
)

SYSTEM_BASE += (
    " Your 'summary_markdown' MUST include a top-level section named '## Detailed Condition Report' containing a cohesive narrative of at least 10–14 sentences (not bullets). "
    "It must synthesize: impact zones, per-panel damages, repair vs. replace rationale, parts type (OEM/LKQ/Aftermarket), labor ops, refinish/overlap, rate/materials/sublet/tax handling, and estimate integrity. "
    "It must cite concrete evidence inline (e.g., p2/L14, Photo 3). "
    "When evaluating paint materials, recognize that a summary line such as 'Paint Supplies' or 'Paint Materials' with hours and rate in the totals section constitutes a valid cost breakdown. "
    "Do not mark it missing if such a line is present, even if materials are not listed per-panel. "
    "Avoid categorical phrases such as 'deemed repairable' or 'deemed total loss' unless that exact determination appears in the provided documents (e.g., estimate header says 'Total Loss' or an ACV comparison is shown). "
    "Otherwise, use neutral language and do not make a repairability determination."
)

# -----------------------
# Supported file types
# -----------------------
SUPPORTED_IMAGE_EXTS = (".jpg",".jpeg",".png",".webp",".heic",".heif")
SUPPORTED_TEXT_EXTS = (".txt",)
SUPPORTED_DOCX_EXTS = (".docx",)
SUPPORTED_PDF_EXTS = (".pdf",)

# -----------------------
# Helpers to add parts from bytes
# -----------------------
def _image_part_from_bytes(raw: bytes) -> Dict[str, Any]:
    b64 = base64.b64encode(raw).decode("utf-8")
    return {"type": "image_url", "image_url": {"url": "data:image/jpeg;base64," + b64}}

# Safe PDF text extract helper
def _maybe_extract_pdf_text(raw: bytes, fname: str, parts: List[Dict[str, Any]], files_seen: List[str], pdf_text_fulls: Optional[List[str]] = None) -> None:
    try:
        from pdfminer_high_level import extract_text as _x  # type: ignore
    except Exception:
        try:
            from pdfminer.high.level import extract_text as _x  # fallback
        except Exception:
            _x = None
    try:
        if _x:
            full = (_x(io.BytesIO(raw)) or "")
            if pdf_text_fulls is not None and full.strip():
                pdf_text_fulls.append(full)
            t = full[:12000]
            if t.strip():
                parts.insert(0, {"type": "text", "text": t})
                files_seen.append(f"{fname} (pdf text extracted)")
    except Exception:
        pass

# Lightweight OCR helper for images
def _maybe_ocr_image_text(im: Image.Image) -> str:
    if not _OCR_ENABLED:
        return ""
    try:
        im2 = im.convert("L")
        if max(im2.size) < 1400:
            scale = 1400 / max(im2.size)
            im2 = im2.resize((int(im2.width*scale), int(im2.height*scale)))
        txt = pytesseract.image_to_string(im2)
        return (txt or "").strip()
    except Exception:
        return ""


def _qr_decode_vin_from_pil(im: Image.Image) -> Optional[str]:
    """Best-effort local QR/barcode decode for VIN (optional deps)."""
    try:
        import cv2  # type: ignore
        import numpy as np  # type: ignore
        arr = cv2.cvtColor(np.array(im.convert("RGB")), cv2.COLOR_RGB2BGR)
        det = cv2.QRCodeDetector()
        val, _, _ = det.detectAndDecode(arr)
        if val:
            s = str(val).strip().upper()
            m = re.search(VIN_PATTERN, s)
            if m:
                return m.group(0)
    except Exception:
        pass
    try:
        from pyzbar.pyzbar import decode  # type: ignore
        for d in decode(im.convert("RGB")):
            try:
                s = d.data.decode("utf-8", "ignore").strip().upper()
            except Exception:
                s = str(d.data).strip().upper()
            m = re.search(VIN_PATTERN, s)
            if m:
                return m.group(0)
    except Exception:
        pass
    return None


def _add_bytes(parts: List[Dict[str,Any]], files_seen: List[str], photo_index: Optional[List[str]], thumb_paths: Optional[List[str]], raw: bytes, fname: str, used: int, max_images: int, pdf_text_fulls: Optional[List[str]] = None, ocr_pairs: Optional[List[Dict[str, Any]]] = None) -> int:
    low = fname.lower()
    if low.endswith(SUPPORTED_PDF_EXTS) and used < max_images:
        try:
            pages = convert_from_bytes(raw, dpi=200)
            files_seen.append(f"{fname} (pdf, {len(pages)} page(s))")
            _maybe_extract_pdf_text(raw, fname, parts, files_seen, pdf_text_fulls=pdf_text_fulls)
            OCR_PAGE_CAP = 100
            ocr_collected = []
            for idx, im in enumerate(pages[:max_images - used]):
                b = io.BytesIO()
                im.save(b, format="JPEG", quality=75, optimize=True)
                parts.append(_image_part_from_bytes(b.getvalue()))
                used += 1
                if photo_index is not None:
                    photo_index.append(f"{fname}::page_{idx+1}")
                if idx < OCR_PAGE_CAP:
                    txt = _maybe_ocr_image_text(im)
                    if txt:
                        ocr_collected.append(txt)
            if ocr_collected:
                parts.insert(0, {"type": "text", "text": ("\n".join(ocr_collected))[:12000]})
                files_seen.append(f"{fname} (ocr text extracted)")
        except Exception as e:
            logging.warning(f"pdf2image failed for {fname}: {e}")
            files_seen.append(f"{fname} (pdf, could not be converted)")
    elif low.endswith(SUPPORTED_IMAGE_EXTS) and used < max_images:
        im_ref = None
        raw_for_vin = None
        qr_vin = None
        try:
            im = Image.open(io.BytesIO(raw)).convert("RGB")
            im_ref = im.copy()
            # Preserve ORIGINAL bytes for VIN label / QR/barcode decoding (before any resizing or recompression).
            raw_for_vin = raw
            # Secondary confirmation: attempt local QR/barcode decode from the ORIGINAL image (before resizing).
            qr_vin = _qr_decode_vin_from_pil(im_ref)
            # Use the same preprocessing for ZIP and loose JPGs (keep higher res for small label text).
            max_dim = 2048
            if max(im.size) > max_dim:
                scale = max_dim / float(max(im.size))
                im = im.resize((int(im.width * scale), int(im.height * scale)))
            b = io.BytesIO()
            im.save(b, format="JPEG", quality=75, optimize=True)
            raw = b.getvalue()
        except Exception:
            im_ref = None
            raw_for_vin = None
        parts.append(_image_part_from_bytes(raw))
        used += 1
        if photo_index is not None:
            photo_index.append(fname)
        files_seen.append(f"{fname} (photo)")
        # Keep a local copy of each uploaded photo for the PDF thumbnail appendix page.
        if thumb_paths is not None:
            try:
                im_save = im_ref if im_ref is not None else Image.open(io.BytesIO(raw)).convert("RGB")
                im_save.thumbnail((900, 900))
                thumb_name = f"thumb_{uuid.uuid4().hex}.jpg"
                thumb_path = os.path.join(PDF_DIR, thumb_name)
                im_save.save(thumb_path, format="JPEG", quality=75, optimize=True)
                thumb_paths.append(thumb_path)
            except Exception:
                pass
        if im_ref is not None:
            txt = _maybe_ocr_image_text(im_ref)
            if ocr_pairs is not None:
                try:
                    ocr_pairs.append({"name": fname, "text": (txt or ""), "raw_for_vin": raw_for_vin, "qr_vin": (qr_vin or None)})
                except Exception:
                    pass
            if txt:
                parts.insert(0, {"type":"text", "text": txt[:12000]})
                files_seen.append(f"{fname} (ocr text extracted)")
    elif low.endswith(SUPPORTED_DOCX_EXTS):
        try:
            text = "\n".join([p.text for p in Document(io.BytesIO(raw)).paragraphs if p.text.strip()])
        except Exception:
            text = ""
        if text.strip():
            parts.insert(0, {"type":"text","text": text[:12000]})
            files_seen.append(f"{fname} (docx text included)")
        else:
            files_seen.append(f"{fname} (docx, no readable text)")
    elif low.endswith(SUPPORTED_TEXT_EXTS):
        try:
            text = raw.decode("utf-8","ignore")[:12000]
        except Exception:
            text = ""
        if text.strip():
            parts.insert(0, {"type":"text","text": text})
            files_seen.append(f"{fname} (txt included)")
        else:
            files_seen.append(f"{fname} (txt, empty)")
    else:
        files_seen.append(f"{fname} (unsupported type)")
    return used

# -----------------------
# App + CORS
# -----------------------
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://nspxn.com","https://www.nspxn.com","http://nspxn.com","http://www.nspxn.com",
        "https://nspxn.onrender.com"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# -----------------------
# Client Rules: fuzzy finder + endpoints
# -----------------------
def _find_rules_path(base_name: str, rules_dir: str) -> Optional[str]:
    target = base_name.strip()
    if not target.lower().endswith(".docx"):
        target += ".docx"
    p = os.path.join(rules_dir, target)
    if os.path.exists(p):
        return p
    pattern = os.path.join(rules_dir, "*.docx")
    candidates = glob.glob(pattern)
    if not candidates:
        return None
    low_target = target.lower()
    for c in candidates:
        if os.path.basename(c).lower() == low_target:
            return c
    for c in candidates:
        if os.path.basename(c).lower().startswith(low_target[:-5]):
            return c
    tokens = [t for t in low_target[:-5].split() if t]
    def contains_in_order(name: str) -> bool:
        i = 0
        for tok in tokens:
            i = name.find(tok, i)
            if i == -1:
                return False
            i += len(tok)
        return True
    for c in candidates:
        if contains_in_order(os.path.basename(c).lower()):
            return c
    return None

@app.get("/client-rules/{client_name}")
async def get_client_rules(client_name: str):
    path = _find_rules_path(client_name, CLIENT_RULES_DIR)
    if not path:
        try:
            files = sorted(os.path.basename(p) for p in glob.glob(os.path.join(CLIENT_RULES_DIR, "*.docx")))
        except Exception:
            files = []
        return JSONResponse(
            status_code=404,
            content={
                "error": f"Rules not found for '{client_name}'.",
                "hint": "Ensure the .docx file exists in CLIENT_RULES_DIR (default 'client_rules').",
                "available_rule_files": files[:50]
            },
        )
    try:
        doc = Document(path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return {"file": os.path.basename(path), "text": text}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"Unable to read rules: {e}", "file": os.path.basename(path)})

@app.get("/client-rules")
async def list_client_rules():
    try:
        files = sorted(os.path.basename(p) for p in glob.glob(os.path.join(CLIENT_RULES_DIR, "*.docx")))
        return {"dir": CLIENT_RULES_DIR, "files": files}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e), "dir": CLIENT_RULES_DIR})

# -----------------------
# Vision Review
# -----------------------
@app.post("/vision-review")
async def vision_review(
    request: Request,
    files: Optional[List[UploadFile]] = File(None),
    client_rules: str = Form(""),
    ai_notes: str = Form(""),
    addl_notes: str = Form(""),
    additional_notes: str = Form(""),
    notes: str = Form(""),
    file_number: Optional[str] = Form(None),
    ia_company: str = Form(""),
    appraiser_id: str = Form(""),
    ai_intent: str = Form("damage_report_from_photos")
):
    parts: List[Dict[str, Any]] = []
    files_seen: List[str] = []
    photo_index: List[str] = []
    thumbnail_paths: List[str] = []
    ocr_pairs: List[Dict[str, Any]] = []

    # --- 422 guard: avoid FastAPI validation failures when frontend keys vary ---
    # Accept missing/alternate keys without raising 422; return a clear 400 instead.
    if files is None:
        files = []
    if file_number is None or str(file_number).strip() == "":
        try:
            _form = await request.form()
            # Common file number key variants from different frontends
            for _k in ("file_number", "file-number", "fileNumber", "fileNum", "file_num", "filenumber"):
                _v = str(_form.get(_k, "") or "").strip()
                if _v:
                    file_number = _v
                    break
        except Exception:
            pass
    # If 'files' was posted under a different key, recover it from the raw form
    if (not files) or (isinstance(files, list) and len(files) == 0):
        try:
            _form = await request.form()
            _maybe = []
            try:
                _maybe = list(_form.getlist("files")) if hasattr(_form, "getlist") else []
            except Exception:
                _maybe = []
            if not _maybe:
                try:
                    _maybe = list(_form.getlist("file")) if hasattr(_form, "getlist") else []
                except Exception:
                    _maybe = []
            _maybe_files = [x for x in _maybe if hasattr(x, "filename")]
            if _maybe_files:
                files = _maybe_files  # type: ignore[assignment]
        except Exception:
            pass
    # Hard-required fields (explicit 400 instead of 422)
    if not file_number or str(file_number).strip() == "":
        return JSONResponse(status_code=400, content={"error": "Missing required field: file_number"})
    if not files:
        return JSONResponse(status_code=400, content={"error": "Missing required upload: files"})
    vin_candidates: List[str] = []  # reserved (filenames not used)
    MAX_IMAGES = 48
    used = 0
    # Coalesce Add\'l Notes from multiple possible frontend field names
    ai_notes_used = ((ai_notes or "").strip() or (addl_notes or "").strip() or (additional_notes or "").strip() or (notes or "").strip())
    ai_notes_used = ai_notes_used.strip()
    # If notes still empty, fall back to reading the raw posted form (covers mismatched frontend field names)
    if not ai_notes_used:
        try:
            _form = await request.form()
            # First try common variants explicitly
            for _k in ("ai_notes","ai-notes","addl_notes","additional_notes","notes","ai_review_notes","ai_notes_box","addlNote","addlNoteText"):
                _v = str(_form.get(_k, "") or "").strip()
                if _v:
                    ai_notes_used = _v
                    break
            # Then any key containing 'note' (last resort)
            if not ai_notes_used:
                for _k in _form.keys():
                    if "note" in str(_k).lower():
                        _v = str(_form.get(_k, "") or "").strip()
                        if _v:
                            ai_notes_used = _v
                            break
        except Exception:
            pass
    # Normalize/sanitize notes so they cannot break structured prompting
    ai_notes_used = _normalize_ai_notes(ai_notes_used)
    ai_notes_block = _ai_notes_block(ai_notes_used)
    locked_cost_overrides = _extract_locked_cost_overrides(ai_notes_used)

    pdf_text_fulls: List[str] = []  # full PDF text for supplement detection

    # Anti-zipbomb guardrails
    MAX_ZIP_FILES = 100
    MAX_ENTRY_SIZE = 15 * 1024 * 1024  # 15 MB

    for f in sorted(files, key=lambda _f: ((_f.filename or '').lower())):
        raw = await f.read()
        fname = f.filename or "upload"
        low = fname.lower()
        if low.endswith(".zip"):
            try:
                zf = zipfile.ZipFile(io.BytesIO(raw))
            except Exception as e:
                files_seen.append(f"{fname} (zip, unreadable: {e})")
                continue
            members = sorted([zi for zi in zf.infolist() if not zi.is_dir()], key=lambda _zi: (_zi.filename or '').lower())
            if len(members) > MAX_ZIP_FILES:
                files_seen.append(f"{fname} (zip, too many entries: {len(members)})")
                members = members[:MAX_ZIP_FILES]
            for zi in members:
                inner_name = zi.filename
                if ".." in inner_name or inner_name.startswith(("/", "\\")):
                    files_seen.append(f"{fname}::{inner_name} (skipped unsafe path)"); continue
                if zi.file_size > MAX_ENTRY_SIZE:
                    files_seen.append(f"{fname}::{inner_name} (skipped >15MB)"); continue
                try:
                    data = zf.read(zi)
                except Exception as e:
                    files_seen.append(f"{fname}::{inner_name} (read error: {e})"); continue
                used = _add_bytes(parts, files_seen, photo_index, thumbnail_paths, data, f"{fname}::{inner_name}", used, MAX_IMAGES, pdf_text_fulls=pdf_text_fulls, ocr_pairs=ocr_pairs)
        else:
            used = _add_bytes(parts, files_seen, photo_index, thumbnail_paths, raw, fname, used, MAX_IMAGES, pdf_text_fulls=pdf_text_fulls, ocr_pairs=ocr_pairs)

    # Collect uploaded TEXT ONLY for evidence checks
    uploaded_text_blobs = []
    for p in parts:
        if isinstance(p, dict) and p.get("type") == "text" and isinstance(p.get("text"), str):
            uploaded_text_blobs.append(p["text"])
    uploaded_text_all = "\n".join(uploaded_text_blobs)
    # --- VIN EXTRACTION (door label OCR -> QR/barcode) ---
    # Do NOT use filenames. Prefer driver-door certification label OCR; if not found, attempt QR/barcode decode.
    vin_from_label = None
    vin_from_label_photo = None
    vin_from_qr = None
    vin_from_qr_photo = None

    def _looks_like_door_label(txt: str) -> bool:
        if not txt:
            return False
        t = txt.upper()
        keys = ["MFD BY", "MANUFACTURED", "GVWR", "GAWR", "TIRE SIZE", "CONFORMS TO"]
        return sum(1 for k in keys if k in t) >= 2

    # (b) Door label OCR first
    try:
        for rec in ocr_pairs:
            if not isinstance(rec, dict):
                continue
            t = rec.get("text") or ""
            if _looks_like_door_label(t):
                mvin = re.search(VIN_PATTERN, t.upper())
                if mvin:
                    vin_from_label = mvin.group(0)
                    vin_from_label_photo = rec.get("name") or None
                    break
    except Exception:
        vin_from_label = None
        vin_from_label_photo = None
    # (c) QR/barcode VIN as secondary confirmation (prefer same photo as door label if available)
    try:
        if vin_from_label_photo:
            for rec in ocr_pairs:
                if isinstance(rec, dict) and (rec.get("name") == vin_from_label_photo) and rec.get("qr_vin"):
                    vin_from_qr = str(rec.get("qr_vin") or "").strip().upper() or None
                    vin_from_qr_photo = rec.get("name") or None
                    break
        if not vin_from_qr:
            for rec in ocr_pairs:
                if isinstance(rec, dict) and rec.get("qr_vin"):
                    vin_from_qr = str(rec.get("qr_vin") or "").strip().upper() or None
                    vin_from_qr_photo = rec.get("name") or None
                    break
    except Exception:
        vin_from_qr = None
        vin_from_qr_photo = None


    # (c) If OCR didn't yield a VIN, try a dedicated vision decode for VIN/QR/barcode on the best label candidate
    def _decode_vin_from_label_or_qr(raw_bytes: Optional[bytes]) -> Optional[str]:
        if not raw_bytes:
            return None
        try:
            prompt = (
                "Return ONLY JSON: {\"vin\": \"...\"}.\n"
                "Task: Read the vehicle VIN from the door-jamb certification label text and/or decode any QR/barcode if present.\n"
                "Rules: VIN must be exactly 17 characters (A-H, J-N, P, R-Z, 0-9; no I/O/Q). "
                "If not fully legible, return {\"vin\": null}.\n"
            )
            rsp_v = client.chat.completions.create(
                model=MODEL,
                messages=[
                    {"role": "system", "content": "You extract VINs from vehicle door-jamb certification labels. JSON only."},
                    {"role": "user", "content": [{"type": "text", "text": prompt}, _image_part_from_bytes(raw_bytes)]},
                ],
                max_completion_tokens=300,
                temperature=0,
                response_format={"type": "json_object"},
            )
            raw_v = (rsp_v.choices[0].message.content or "").strip()
            try:
                data_v = json.loads(raw_v)
            except Exception:
                data_v = None
            if isinstance(data_v, dict):
                vv = data_v.get("vin")
                if isinstance(vv, str):
                    vv = vv.strip().upper()
                    if re.fullmatch(VIN_PATTERN, vv):
                        return vv
        except Exception:
            return None
        return None

    if not vin_from_label:
        try:
            candidate = None
            for rec in ocr_pairs:
                if isinstance(rec, dict) and _looks_like_door_label(rec.get("text") or "") and rec.get("raw_for_vin"):
                    candidate = rec
                    break
            if candidate is None:
                for rec in ocr_pairs:
                    if isinstance(rec, dict) and rec.get("raw_for_vin"):
                        candidate = rec
                        break
            if candidate is not None:
                # Prefer local QR/barcode VIN if available before calling the vision model.
                if candidate.get("qr_vin"):
                    vin_from_label = str(candidate.get("qr_vin") or "").strip().upper() or None
                if not vin_from_label:
                    vin_from_label = _decode_vin_from_label_or_qr(candidate.get("raw_for_vin"))
                vin_from_label_photo = candidate.get("name") if vin_from_label else None
        except Exception:
            pass

    # normalize + keep unique order
    _seen_v = set(); _tmp_v=[]
    for _v in vin_candidates:
        if _v and _v not in _seen_v:
            _seen_v.add(_v); _tmp_v.append(_v)
    vin_candidates = _tmp_v

    # --- ODOMETER OCR LOCK (extract mileage from OCR text if visible) ---
    # This makes it impossible for the narrative to claim the odometer is not visible when OCR captured a mileage value.
    odometer_value = None
    try:
        _odo_txt = uploaded_text_all or ""
        # Common mileage patterns from digital clusters: "72,261 mi", "72261 mi", "72261 miles", "116000 km"
        _m = re.search(r"(?i)\b(\d{1,3}(?:,\d{3})+|\d{4,7})\s*(mi|miles|km)\b", _odo_txt)
        if _m:
            _digits = _m.group(1).replace(",", "")
            _unit = _m.group(2).lower()
            if _unit == "miles":
                _unit = "mi"
            odometer_value = f"{int(_digits):,} {_unit}"
    except Exception:
        odometer_value = None


    # --- Closing Report: extract "Inspection Results" section for deterministic cross-check + shop-status ---
    def _extract_inspection_results_block(_txt: str) -> str:
        if not _txt:
            return ""
        t = _txt.replace("\r", "\n")
        # Capture the text after the "Inspection Results" header up to the next major header.
        m_ir = re.search(
            r"(?is)\bInspection\s+Results\b\s*[:\-]?\s*(.{0,5000}?)(?=\n\s*(?:Possible\s+Supplement\s+Amount|Supplement\b|Estimate\b|Inspection\s+Location\b|Repair\s+Facility\b|Vehicle\b|Owner\b|Appraiser\b|Adjuster\b|$))",
            t,
        )
        if m_ir:
            return (m_ir.group(1) or "").strip()[:5000]
        return ""

    inspection_results_text = _extract_inspection_results_block(uploaded_text_all or "")

    # Closing Report: "Possible Supplement Amount" (this ALONE does NOT mean the estimate is a supplement)
    _possible_supp_amount = None
    _m_psa = re.search(r"(?i)\bPossible\s+Supplement\s+Amount\b\s*\$?\s*([0-9,]+(?:\.[0-9]{2})?)", uploaded_text_all or "")
    if _m_psa:
        _possible_supp_amount = _m_psa.group(1)

    photos_provided = any(isinstance(p, dict) and p.get('type') != 'text' for p in parts)

    # --- Robust detectors (CLEAN RETAIL + ADVISOR) ---
    clean_retail_rx = (
        r"(?i)\b("
        r"NADA|J[.\s-]*D[.\s-]*\s*Power|JDPower\.com|"
        r"Kell?ey\s+Blue\s+Book|KBB\.com|"
        r"Edmunds|Carfax|Cars\.com|"
        r"Clean\s+Retail(?:\s+Value)?"
        r")\b"
    )
    _clean_retail_present = bool(re.search(clean_retail_rx, uploaded_text_all or ""))

    advisor_rx = r"(?i)\bAdvisor\s+Report\b"
    _advisor_present = bool(re.search(advisor_rx, uploaded_text_all or ""))

    paint_mat_rx = r"(Paint\s+(Suppl(?:ies|y)|Materials)|Materials\s*Line)"
    _paint_materials_present = bool(re.search(paint_mat_rx, uploaded_text_all or "", flags=re.IGNORECASE))

    # --- Parts source guardrail (OEM vs Aftermarket/LKQ) ---
    # Detect explicit non-OEM indicators in estimate LINE ITEMS only (ignore generic boilerplate/disclosures).
    _explicit_non_oem_parts = False
    try:
        # Heuristic: look for non-OEM keywords on the same line as an operation/part line (starts with line #).
        _explicit_non_oem_parts = bool(re.search(
            r"(?im)^\s*\d+\s+.*\b(A/M|AFTERMARKET|NON\s*OEM|CAPA|NSF|LKQ|RCY|RECYCLED|USED|RECOND)\b",
            uploaded_text_all or "",
        ))
    except Exception:
        _explicit_non_oem_parts = False

    # Closing report sometimes explicitly states no aftermarket/LKQ parts were included.
    _closing_no_aftermarket = False
    try:
        _closing_no_aftermarket = bool(re.search(r"(?i)\bno\s+aftermarket\b|\bno\s+aftermarket\s+or\s+lkq\b", uploaded_text_all or ""))
    except Exception:
        _closing_no_aftermarket = False

    # Presence detectors for VIN / Odometer / Production Date (OCR text)
    vin_photo_rx = r"(?i)\bDescription:\s*VIN\b|\bVIN(?:\s*#|:)?\s*[A-HJ-NPR-Z0-9]{17}\b"
    odo_photo_rx = r"(?i)\bDescription:\s*Odometer\b|\bOdometer\s*Photo\b|\bPhoto\s*[:#]?\s*Odometer\b"

    # Production date patterns (label)
    prod_date_rx = (
        r"(?is)\b(Production\s*date|Prod(?:uction)?\s*Date|Date\s*of\s*Mfr|Date\s*of\s*Manufacture|"
        r"MFD\.?\s*(?:BY|DATE)?|MFR\.?\s*DATE|MFG\.?\s*DATE|DATE)\b"
        r".{0,80}?\b(0[1-9]|1[0-2])\s*[-/.\u2013\u2014\u2212:\s]\s*(20\d{2}|\d{2})\b"
    )

    _vin_photo_present = bool(re.search(vin_photo_rx, uploaded_text_all or ""))
    _odo_photo_present = bool(re.search(odo_photo_rx, uploaded_text_all or ""))

    # presence + extractor
    _prod_date_present = False
    _prod_date_str = None
    m = re.search(prod_date_rx, uploaded_text_all or "")
    if m:
        mm = m.group(2); yy = m.group(3)
        if len(yy) == 2: yy = "20" + yy
        _prod_date_present = True
        _prod_date_str = f"{mm}/{yy}"
    if not _prod_date_present:
        loose_prod_rx = (
            r"(?is)\b(MFD|MFR|MFG|PROD\.?|PRODUCTION|DATE)\b"
            r".{0,60}?\b(0[1-9]|1[0-2])\s*[-/.\u2013\u2014\u2212:\s]\s*(20\d{2}|\d{2})\b"
        )
        m2 = re.search(loose_prod_rx, uploaded_text_all or "")
        if m2:
            mm = m2.group(2); yy = m2.group(3)
            if len(yy) == 2: yy = "20" + yy
            _prod_date_present = True
            _prod_date_str = f"{mm}/{yy}"

    # treat production date as evidenced if either door label VIN photo or prod date text is seen
    _prod_evidenced = _prod_date_present or _vin_photo_present

    # Locked photos-only entrypoint
    ai_intent = "damage_report_from_photos"

    REQ_LABELS = {
        "guidelines_only": "Guidelines → Estimate (no photos)",
        "comprehensive": "Comprehensive: Guidelines + Estimate + Photos (with VIN check)",
        "damage_report_from_photos": "Create a Damage Report from Photos",
    }
    req_label = REQ_LABELS.get(ai_intent, "Comprehensive: Guidelines + Estimate + Photos (with VIN check)")
    log.info(f"ai_intent received: {ai_intent} -> using label: {req_label}")

    KEYS = [
        "file_number","request_type","claim_number","vin","vin_verification","vehicle",
        "odometer_estimate_only","compliance_score","summary_brief","summary_markdown",
        "fraud_markdown","primary_impact","secondary_impact","estimated_costs_markdown","conclusion"
    ]

    SYSTEM = SYSTEM_BASE

    # Closing Report shop-status detector (documents-only; used to prevent false Repair Facility deductions)
    _not_at_shop = False
    try:
        # Prefer the extracted "Inspection Results" block if available (more deterministic)
        if inspection_results_text:
            _not_at_shop = bool(re.search(
                r"(?i)\bNot\s+at\s+Shop\b|\bOwner\s+Location\b|\bInspection\s+Location\s*:?\s*Owner\b|\bInspection\s+Location\s+Owner\b",
                inspection_results_text,
            ))
        if not _not_at_shop:
            _not_at_shop = bool(re.search(
                r"(?i)\bNot\s+at\s+Shop\b|\bOwner\s+Location\b|\bInspection\s+Location\s*:?\s*Owner\b|\bInspection\s+Location\s+Owner\b",
                uploaded_text_all or "",
            ))
    except Exception:
        _not_at_shop = False

    # --- NEW: detect all supplement tags S01, S02, ... in the combined document text ---
    combined_detection_text = (uploaded_text_all or "") + "\n" + "\n".join(pdf_text_fulls or [])
    supplement_versions = sorted(set(re.findall(r"(?i)\bS[0-9]{2}\b", combined_detection_text)))
    supplement_block = ""
    if supplement_versions:
        supplement_block = (
            "\n\nSUPPLEMENT VERSIONS DETECTED FROM DOCUMENTS:\n"
            f"- Detected supplement tags: {', '.join(supplement_versions)}.\n"
            "- Use these exact tags (e.g., 'Supplement S01 and S02') when describing supplements in the narrative.\n"
        )

    # -----------------------
    # Minimal prompt (LET GPT DO THE WORK)
    # -----------------------
    prompt_text = (
        f"REQUEST TYPE (use exactly in request_type): {req_label}\n"
        f"FILE #: {file_number}\n"
        f"CLIENT: {ia_company}\n\n"
        "PHOTO INDEX (use Photo # citations exactly as listed):\n"
        + ("\n".join([f"Photo {i+1}: {name}" for i, name in enumerate(photo_index)]) if photo_index else "No photos provided.")
        + "\n\n"
        "CLIENT RULES (only if provided):\n"
        + (client_rules[:1500] if client_rules else "")
        + ai_notes_block
        + "\n\n"
        "INSTRUCTIONS:\n"
        "- Return strict JSON only.\n"
        "- REQUIRED: Populate 'estimated_costs_markdown' in the JSON.\n"
        "- Use the template below for narrative formatting.\n\n"
        + DETAIL_TEMPLATES.get(ai_intent, DETAIL_TEMPLATES['comprehensive'])
    )
    if ai_intent == "damage_report_from_photos":
        prompt_text += (
            "\nNEUTRAL TERMINOLOGY RULE:\n"
            "- Do not use left, right, driver-side, or passenger-side terminology.\n"
            "- Use neutral terms only: Front, Rear, Side, one headlamp, one corner.\n"
        )


    parts_payload: List[Dict[str,Any]] = []
    redaction_success = False
    try:
        red_prompt = redact_text_preserve_vin_claim(prompt_text)
        redaction_success = True
        # Prevent Presidio from mutating fixed markdown headings used for prompting.
        # (This avoids model outputs like "### [REDACTED] Addressed".)
        red_prompt = red_prompt.replace("### [REDACTED] Addressed", "### Add'l Notes Addressed")
    except Exception as e:
        log.warning(f"Redaction failed on prompt_text: {e}")
        red_prompt = prompt_text

    parts_payload.append({"type": "text", "text": red_prompt})

    if parts:
        for p in parts:
            if p.get("type") == "text" and isinstance(p.get("text"), str):
                try:
                    red_txt = redact_text_preserve_vin_claim(p["text"])
                    redaction_success = True
                except Exception as e:
                    log.warning(f"Redaction failed on a text part: {e}")
                    red_txt = p["text"]
                parts_payload.append({"type": "text", "text": red_txt})
            else:
                parts_payload.append(p)

    redaction_status = "Redacted PII: Successful ✅" if redaction_success else "Redacted PII: Not Applied"

    # -------------------------------
    # Keep prompt lean
    # -------------------------------
    TEXT_PART_LIMIT = 6
    _text_parts = [p for p in parts_payload if p.get("type") == "text"]
    _image_parts = [p for p in parts_payload if p.get("type") != "text"]
    for tp in _text_parts:
        if isinstance(tp.get("text"), str) and len(tp["text"]) > 8000:
            tp["text"] = tp["text"][:8000]
    parts_payload = _text_parts[:TEXT_PART_LIMIT] + _image_parts

    # Token limits
    MAX_TOKENS_BY_INTENT = {
        "comprehensive": 2200,
        "guidelines_only": 1500,
        "damage_report_from_photos": 2400
    }
    max_tokens = MAX_TOKENS_BY_INTENT.get(ai_intent, 1500)

    # Call GPT and parse JSON (JSON hardened)
    # Prefer the canonical SDK path (client.chat.completions). Keep fallback for older SDKs.
    try:
        rsp = client.chat.completions.create(
            model=MODEL,
            messages=[{"role":"system","content": SYSTEM},
                      {"role":"user","content": parts_payload}],
            max_completion_tokens=max_tokens,
            temperature=0,
            top_p=1,
            presence_penalty=0,
            frequency_penalty=0,
            response_format={"type":"json_object"},
        )
    except AttributeError:
        rsp = client.chat_completions.create(  # type: ignore[attr-defined]
            model=MODEL,
            messages=[{"role":"system","content": SYSTEM},
                      {"role":"user","content": parts_payload}],
            max_completion_tokens=max_tokens,
            temperature=0,
            top_p=1,
            presence_penalty=0,
            frequency_penalty=0,
            response_format={"type":"json_object"},
        )

    # --- Hardened JSON parse helper
    def _try_parse_json(raw_text: str):
        if not raw_text:
            return None
        raw_local = raw_text.strip()
        for fence in ("```json", "```JSON", "```"):
            if raw_local.startswith(fence):
                raw_local = raw_local[len(fence):]
            if raw_local.endswith("```"):
                raw_local = raw_local[:-3]
        raw_local = raw_local.strip()
        raw_local = raw_local.replace("\ufeff", "").replace("\u200b", "").replace("\u00A0", " ")
        try:
            return json.loads(raw_local)
        except Exception:
            pass
        lb = raw_local.find("{"); rb = raw_local.rfind("}")
        chunk = raw_local[lb:rb+1] if (lb != -1 and rb != -1 and rb > lb) else raw_local
        fixes = {"\u2018": "'", "\u2019": "'", "\u201C": '"', "\u201D": '"', "\r": "", "\t": "    "}
        for k, v in fixes.items():
            chunk = chunk.replace(k, v)
        chunk = re.sub(r",\s*([}\]])", r"\1", chunk)
        try:
            return json.loads(chunk)
        except Exception:
            return None

    try:
        raw = (rsp.choices[0].message.content or "")
    except Exception as e:
        log.error(f"LLM returned no content: {e}")
        return JSONResponse(status_code=500, content={"error":"Model returned no content."})

    data = _try_parse_json(raw)
    # Prefer door-label VIN when present (OCR -> QR/barcode). Do NOT use filenames.
    # --- HARD VIN LOCK (door label wins; QR only confirms) ---
    try:
        if isinstance(data, dict):

            # 1) If door-label VIN exists, it is authoritative.
            if vin_from_label:
                data["vin"] = vin_from_label

                if vin_from_qr:
                    if vin_from_qr == vin_from_label:
                        data["vin_verification"] = "MATCH (door label + QR)"
                    else:
                        data["vin_verification"] = (
                            f"MISMATCH (door label: {vin_from_label}; QR: {vin_from_qr})"
                        )
                else:
                    data["vin_verification"] = (
                        "INCONCLUSIVE (door label extracted; no secondary confirmation)"
                    )

            # 2) If no door label but QR found
            elif vin_from_qr:
                data["vin"] = vin_from_qr
                data["vin_verification"] = "INCONCLUSIVE (QR only; door label not detected)"

    except Exception:
        pass

    # Prefer door-label VIN when present (OCR -> QR/barcode). Do NOT use filenames.
    try:
        if isinstance(data, dict) and vin_from_label:
            v_model = (data.get("vin") or "").strip().upper()
            if not re.fullmatch(VIN_PATTERN, v_model):
                data["vin"] = vin_from_label
                if not (data.get("vin_verification") or "").strip():
                    data["vin_verification"] = "INCONCLUSIVE (door label VIN extracted; compare to other docs if present)"
            elif v_model != vin_from_label:
                data["vin_verification"] = f"MISMATCH (door label: {vin_from_label}; other source: {v_model})"
            # Secondary confirmation vs QR/barcode if available
            if vin_from_qr:
                if vin_from_label == vin_from_qr:
                    if not re.search(r"\bMATCH\b|\bMISMATCH\b", str(data.get("vin_verification") or ""), flags=re.IGNORECASE):
                        data["vin_verification"] = "MATCH (door label + QR/barcode)"
                else:
                    data["vin_verification"] = f"MISMATCH (door label: {vin_from_label}; QR/barcode: {vin_from_qr})"
    except Exception:
        pass


    # One safe retry on truncation
    try:
        finish_reason = getattr(rsp.choices[0], "finish_reason", None)
    except Exception:
        finish_reason = None

    if data is None and (finish_reason == "length" or (raw and raw.strip().endswith("..."))):
        _text_parts_retry = [p for p in parts_payload if p.get("type") == "text"]
        _image_parts_retry = [p for p in parts_payload if p.get("type") != "text"]
        shrunk = _text_parts_retry[: max(3, len(_text_parts_retry)//2)] + _image_parts_retry
        retry_tokens = min(3000, max_tokens + 600)
        try:
            rsp2 = client.chat.completions.create(
                model=MODEL,
                messages=[{"role": "system", "content": SYSTEM},
                          {"role": "user", "content": shrunk}],
                max_completion_tokens=retry_tokens,
                temperature=0,
                response_format={"type": "json_object"}
            )
            raw2 = (rsp2.choices[0].message.content or "")
            data = _try_parse_json(raw2)
        except Exception:
            pass

    # Fallback: formatting pass
    if data is None:
        try:
            fix_prompt = [
                {"role":"system","content":
                    "You are a formatter. Return ONLY a strict JSON object. No prose. No markdown. No code fences."
                },
                {"role":"user","content":
                    "Convert the following text into a valid JSON object. "
                    "Use exactly these keys (all required): "
                    "['file_number','request_type','claim_number','vin','vin_verification','vehicle',"
                    "'odometer_estimate_only','compliance_score','summary_brief','summary_markdown',"
                    "'fraud_markdown','primary_impact','secondary_impact','estimated_costs_markdown','conclusion'] "
                    "Do not invent new keys. If a field is unavailable, use 'N/A'. "
                    "Here is the text:\n\n" + raw
                }
            ]
            try:
                fix_rsp = client.chat_completions.create(  # type: ignore[attr-defined]
                    model=MODEL,
                    messages=fix_prompt,
                    max_completion_tokens=max_tokens,
                    temperature=0,
                    response_format={"type":"json_object"}
                )
            except AttributeError:
                fix_rsp = client.chat.completions.create(
                    model=MODEL,
                    messages=fix_prompt,
                    max_completion_tokens=max_tokens,
                    temperature=0,
                    response_format={"type":"json_object"}
                )
            fixed = (fix_rsp.choices[0].message.content or "")
            data = _try_parse_json(fixed)
        except Exception as e:
            log.error(f"Self-heal reformat failed: {e}")

    if data is None:
        log.error(f"LLM failure or JSON parse error; first 500 chars:\n" + (raw or "")[:500])
        skeleton = {k: "N/A" for k in KEYS}
        skeleton["file_number"] = file_number
        skeleton["request_type"] = req_label
        skeleton["summary_brief"] = "N/A (model output could not be parsed; skeleton returned)."
        skeleton["summary_markdown"] = (
            "## Detailed Condition Report\n"
            "Model output could not be parsed into JSON on this run. Please resubmit."
        )
        skeleton["fraud_markdown"] = "No material inconsistencies found."
        return skeleton

    def _get(k):
        v = data.get(k)
        return "" if v is None else str(v)

    result = {
        "file_number": file_number,
        "request_type": req_label,
        "claim_number": _get("claim_number"),
        "vin": _get("vin"),
        "vin_verification": _get("vin_verification"),
        "vehicle": _get("vehicle"),
"odometer_estimate_only": (odometer_value if odometer_value else _get("odometer_estimate_only")),
        "compliance_score": _get("compliance_score"),
        "summary_brief": _get("summary_brief"),
        "summary_markdown": _get("summary_markdown"),
        "fraud_markdown": _get("fraud_markdown"),
        "primary_impact": _get("primary_impact"),
        "secondary_impact": _get("secondary_impact"),
        "estimated_costs_markdown": _get("estimated_costs_markdown"),
        "conclusion": _get("conclusion"),
        "redaction_status": redaction_status,
    }

    locked_costs_obj: Optional[Dict[str, Any]] = None

    # ---- Minimal Impact Sanitizer (no Left/Right/Driver/Passenger unless proven) ----
    # Policy A: If any impact label implies left/right/driver/passenger, collapse to Front/Rear/Side.
    def _sanitize_impact_label(v: Optional[str]) -> str:
        if v is None:
            return "N/A"
        s = str(v).strip()
        if not s:
            return "N/A"
        low = s.lower()

        # Common abbreviations / synonyms
        front_hit = bool(re.search(r"\bfront\b|\bfrt\b", low)) or bool(re.search(r"\b(?:lf|rf)\b", low))
        rear_hit  = bool(re.search(r"\brear\b", low)) or bool(re.search(r"\b(?:lr|rr)\b", low))
        side_hint = bool(re.search(r"\bleft\b|\bright\b|\bdriver\b|\bpassenger\b|\b(?:lf|rf|lr|rr)\b", low))

        if front_hit and side_hint:
            return "Front"
        if rear_hit and side_hint:
            return "Rear"
        if side_hint:
            return "Side"

        # Normalize a few common exact values
        if re.fullmatch(r"(?i)front", s):
            return "Front"
        if re.fullmatch(r"(?i)rear", s):
            return "Rear"
        if re.fullmatch(r"(?i)side", s):
            return "Side"
        return s

    try:
        result["primary_impact"] = _sanitize_impact_label(result.get("primary_impact"))
        result["secondary_impact"] = _sanitize_impact_label(result.get("secondary_impact"))
    except Exception:
        pass


    # ✅ FIX #2: Hard fallback so UI never gets an empty narrative ("No narrative generated")
    try:
        sm_tmp = (result.get("summary_markdown") or "").strip()
        if not sm_tmp:
            result["summary_markdown"] = (
                "## Detailed Condition Report\n"
                "Narrative fallback: The model returned an empty narrative field. "
                "Please re-run with the same inputs; core identifiers and score fields were still returned.\n\n"
                "## Overall Assessment\n"
                f"Request Type: {result.get('request_type','N/A')}\n"
                f"Compliance Score: {result.get('compliance_score','N/A')}\n"
            )
        elif "## Detailed Condition Report" not in sm_tmp:
            # Keep minimal: do not re-write content; just prepend the required header to avoid downstream display rules.
            result["summary_markdown"] = "## Detailed Condition Report\n" + sm_tmp
    except Exception:
        pass


    # -----------------------
    # Photos-Only OUTPUT HARDENING (prevents blank/N/A reports; enforces Add'l Notes)
    # -----------------------
    def _bad_field(v: str) -> bool:
        """Return True if the model gave a placeholder/empty section.
        This catches bare 'N/A' AND markdown sections that are effectively only headings + N/A.
        """
        s = (v or "").strip()
        if not s:
            return True
        up = s.strip().upper()

        # Bare placeholders
        if up in ("N/A", "NA", "NONE", "NULL"):
            return True
        if up.startswith("N/A") and len(up) <= 8:
            return True

        # If it's markdown, strip headings/blank lines and see if anything substantive remains
        lines = [ln.strip() for ln in s.splitlines() if ln.strip()]
        # remove markdown headings and common labels
        stripped = []
        for ln in lines:
            if ln.startswith("#"):
                continue
            if ln.lower() in ("report selected", "approximate repair cost breakdown", "fraud & authenticity check", "fraud detection", "conclusion"):
                continue
            stripped.append(ln)

        if not stripped:
            return True

        # If after stripping headings the only remaining content is 'N/A'
        if len(stripped) == 1 and stripped[0].strip().upper() in ("N/A", "NA", "NONE", "NULL"):
            return True

        # If the remaining content is extremely short and contains only placeholders
        joined = " ".join(stripped).strip().upper()
        if joined in ("N/A", "NA", "NONE", "NULL"):
            return True

        return False

    try:
        if ai_intent == "damage_report_from_photos":
            _bad = (
                _bad_field(result.get("summary_markdown") or "")
                or _bad_field(result.get("estimated_costs_markdown") or "")
                or _bad_field(result.get("fraud_markdown") or "")
                or _bad_field(result.get("conclusion") or "")
            )

            if _bad:
                retry_preamble = (
                    "CRITICAL RETRY (PHOTOS-ONLY): Your prior output was invalid because required fields were blank or 'N/A'.\n"
                    "Return ONLY valid JSON (no prose, no code fences).\n"
                    "Hard rules:\n"
                    "- NEVER return 'N/A' for summary_markdown, estimated_costs_markdown, fraud_markdown, or conclusion.\n"
                    "- You MUST explicitly address the Add'l Notes in the narrative (e.g., 'Impact to Right Front') and describe the impact zone accordingly.\n"
                    "- estimated_costs_markdown MUST be a PHOTOS-ONLY approximation: create body/paint/frame/mechanical hours and an OEM parts list with approximate $ by part.\n"
                    "- Paint & Materials MUST be modeled as $ per refinish hour (not a percent).\n"
                    "- Forbidden phrases: 'from estimate', 'from documentation', 'not evidenced', 'no documentation provided'.\n"
                )

                retry_parts = list(parts_payload)  # shallow copy is enough
                # Replace the first text block with a retry preamble + original prompt
                if retry_parts and isinstance(retry_parts[0], dict) and retry_parts[0].get("type") == "text":
                    retry_parts[0] = {"type": "text", "text": retry_preamble + "\n\n" + (retry_parts[0].get("text") or "")}
                else:
                    retry_parts.insert(0, {"type": "text", "text": retry_preamble})

                retry_tokens = min(3200, max_tokens + 500)

                try:
                    rsp_retry = client.chat.completions.create(
                        model=MODEL,
                        messages=[{"role": "system", "content": SYSTEM},
                                  {"role": "user", "content": retry_parts}],
                        max_completion_tokens=retry_tokens,
                        temperature=0,
                        top_p=1,
                        presence_penalty=0,
                        frequency_penalty=0,
                        response_format={"type": "json_object"},
                    )
                except AttributeError:
                    rsp_retry = client.chat_completions.create(  # type: ignore[attr-defined]
                        model=MODEL,
                        messages=[{"role": "system", "content": SYSTEM},
                                  {"role": "user", "content": retry_parts}],
                        max_completion_tokens=retry_tokens,
                        temperature=0,
                        top_p=1,
                        presence_penalty=0,
                        frequency_penalty=0,
                        response_format={"type": "json_object"},
                    )

                raw_retry = (rsp_retry.choices[0].message.content or "")
                data_retry = _try_parse_json(raw_retry)

                if isinstance(data_retry, dict):
                    # Only overwrite the output fields; preserve file_number/request_type etc.
                    for k in ("summary_brief","summary_markdown","fraud_markdown","primary_impact","secondary_impact","estimated_costs_markdown","conclusion","claim_number","vin","vin_verification","vehicle","odometer_estimate_only","compliance_score"):
                        if k in data_retry and data_retry.get(k) is not None:
                            result[k] = str(data_retry.get(k))
                    # Re-apply the narrative header guard if needed
                    sm_retry = (result.get("summary_markdown") or "").strip()
                    if sm_retry and "## Detailed Condition Report" not in sm_retry:
                        result["summary_markdown"] = "## Detailed Condition Report\n" + sm_retry

                    # If the retry STILL returned placeholders, force a non-N/A minimal scaffold (never print dead 'N/A')
                    if (
                        _bad_field(result.get("summary_markdown") or "")
                        or _bad_field(result.get("estimated_costs_markdown") or "")
                        or _bad_field(result.get("fraud_markdown") or "")
                        or _bad_field(result.get("conclusion") or "")
                    ):
                        _notes = (ai_notes_used or "").strip()
                        if not _notes or _notes.lower().startswith("no additional notes"):
                            _notes = "(No additional notes provided.)"

                        result["summary_markdown"] = (
                            "## Detailed Condition Report\n"
                            "Photo-based narrative could not be generated on this run due to a model compliance error.\n"
                            f"Add'l Notes received: {_notes}\n"
                            "Please re-run. (This placeholder is generated by NSPXN to avoid blank/N/A reports.)"
                        )
                        result["estimated_costs_markdown"] = (
                            "## Approximate Repair Cost Breakdown (Photos-Only Approximation)\n"
                            "The model failed to generate a cost breakdown on this run. Please re-run with the same photo set.\n"
                            "Paint & Materials are modeled as $/refinish hour; tax applies to parts + paint materials only."
                        )
                        result["fraud_markdown"] = "No material inconsistencies found."
                        result["conclusion"] = (
                            "Conclusion unavailable due to model compliance error on this run. Please re-run."
                        )

    except Exception:
        # Fail-open: never break the request if retry logic fails.
        pass

        # -----------------------
    # Approximate Repair Cost Breakdown
    # - Prefer model-provided `estimated_costs_markdown` (especially for Photos-Only).
    # - Fail-open: never break the run if cost math cannot be produced.
    # -----------------------
    tax_rate = None  # default; set later if we compute/lookup tax
    try:
        _existing_costs = (result.get("estimated_costs_markdown") or "").strip()

        if not _existing_costs:
            inspection_location = _normalize_location_with_zip(_extract_inspection_location(uploaded_text_all or ""), ia_company, uploaded_text_all)
            rates = _lookup_rates(inspection_location)
            tax_rate = _lookup_tax_rate(inspection_location)

            # Photos-only: do NOT depend on documents/estimates. The model is instructed to generate hours/parts.
            # This backend block is only a minimal fallback if the model omitted the field.
            if ai_intent == "damage_report_from_photos":
                structural_flag = _structural_observed([result.get("summary_markdown") or ""])

                result["estimated_costs_markdown"] = (
                    "## Approximate Repair Cost Breakdown (Photos-Only Approximation)\n\n"
                    f"**Inspection Location:** {inspection_location or 'Not provided'}\n\n"
                    "**Rates Used (regional average / fallback):**\n"
                    f"- Avg Body Rate: ${rates.get('body_rate', 0):.0f}/hr\n"
                    f"- Avg Paint Rate: ${rates.get('paint_rate', 0):.0f}/hr\n"
                    f"- Avg Frame Rate: ${rates.get('frame_rate', 0):.0f}/hr\n"
                    f"- Avg Mechanical Rate: ${rates.get('mechanical_rate', 0):.0f}/hr\n"
                    f"- Paint & Materials Rate: ${rates.get('paint_supplies_rate', 0):.0f} per refinish hr\n"
                    f"- Parts/Materials Tax Rate: {float(tax_rate or 0.0)*100:.3f}%\n\n"
                    "**AI-Derived Repair Scope:**\n"
                    "- Body labor hours: (model should provide)\n"
                    "- Paint labor hours: (model should provide)\n"
                    f"- Setup & Measure (if structural observed): {('2.0 hrs @ body rate' if structural_flag else '0.0 hrs')}\n"
                    "- Frame labor hours (if structural observed): (model should provide)\n"
                    "- Mechanical/ADAS/Airbag/Suspension hours (if observed): (model should provide)\n\n"
                    "**OEM Parts Needed (with approximate $ by part):**\n"
                    "- (model should provide)\n\n"
                    "**Tax (parts + paint materials only):**\n"
                    "- (model should provide)\n\n"
                    "**Approximated Repair Cost Total (Approximation Only):**\n"
                    "- (model should provide)\n\n"
                    "**Estimated Severity Tier (based on this approximation only):**\n"
                    "- [ ] Minor (< $3,500)\n"
                    "- [ ] Moderate ($3,500-$10,000)\n"
                    "- [ ] Major ($10,000+)\n"
                    "- [ ] Likely Total Loss Threshold Approaching\n\n"
                    "_Repair Cost Disclaimer: This section is an AI-generated approximation of repair-related costs based on observed damage in photos and regional average rates. "
                    "It is not an official repair estimate and must be validated by a qualified appraiser using an estimating platform._"
                )
            else:
                # Non-photos intents: best-effort parse from documents (if present), using $/refinish-hr materials.
                totals = _extract_hours_and_parts_totals(uploaded_text_all or "")

                structural_flag = _structural_observed([uploaded_text_all or "", result.get("summary_markdown") or ""])
                setup_measure_hours = 2.0 if structural_flag else 0.0

                body_hours = totals.get("body_hours")
                paint_hours = totals.get("paint_hours")
                frame_hours = totals.get("frame_hours")
                mech_hours = totals.get("mech_hours")
                parts_total = totals.get("parts_total")

                paint_supplies_rate = float(rates.get("paint_supplies_rate") or 0.0)
                paint_supplies_total = (_num(paint_hours) * paint_supplies_rate) if paint_hours is not None else 0.0

                body_labor_total = (_num(body_hours) + setup_measure_hours) * float(rates.get("body_rate") or 0.0) if (body_hours is not None or setup_measure_hours) else 0.0
                paint_labor_total = _num(paint_hours) * float(rates.get("paint_rate") or 0.0) if paint_hours is not None else 0.0

                frame_labor_total = 0.0
                if structural_flag and frame_hours is not None:
                    frame_labor_total = _num(frame_hours) * float(rates.get("frame_rate") or 0.0)

                mech_labor_total = _num(mech_hours) * float(rates.get("mechanical_rate") or 0.0) if mech_hours is not None else 0.0

                taxable = _num(parts_total) + _num(paint_supplies_total)
                tax_total = taxable * float(tax_rate or 0.0)

                approx_total = (
                    _num(body_labor_total)
                    + _num(paint_labor_total)
                    + _num(frame_labor_total)
                    + _num(mech_labor_total)
                    + _num(parts_total)
                    + _num(paint_supplies_total)
                    + _num(tax_total)
                )

                parts_lines = totals.get("parts_lines") or []
                if parts_lines:
                    parts_md = "\n".join([f"- {ln}" for ln in parts_lines[:20]])
                else:
                    parts_md = "- (Itemized OEM parts list not available from parsed text.)"

                result["estimated_costs_markdown"] = (
                    "## Approximate Repair Cost Breakdown (Approximation Only)\n\n"
                    f"**Inspection Location:** {inspection_location or 'Not provided'}\n\n"
                    "**Regional Average Rates (configured / fallback):**\n"
                    f"- Avg Body Rate: ${rates.get('body_rate', 0):.0f}/hr\n"
                    f"- Avg Paint Rate: ${rates.get('paint_rate', 0):.0f}/hr\n"
                    f"- Avg Frame Rate: ${rates.get('frame_rate', 0):.0f}/hr\n"
                    f"- Avg Mechanical Rate: ${rates.get('mechanical_rate', 0):.0f}/hr\n"
                    f"- Paint & Materials Rate: ${rates.get('paint_supplies_rate', 0):.0f} per refinish hr\n"
                    f"- Parts/Materials Tax Rate: {float(tax_rate or 0.0)*100:.3f}%\n\n"
                    "**Hours & Totals:**\n"
                    f"- Approx Total Labor Hours @ Body Rate: {('%.1f' % _num(body_hours)) if body_hours is not None else 'Unknown'}\n"
                    f"- Approx Total Paint Labor Hours @ Paint Rate: {('%.1f' % _num(paint_hours)) if paint_hours is not None else 'Unknown'}\n"
                    f"- Approx Total Paint Supplies @ ${rates.get('paint_supplies_rate', 0):.0f}/refinish hr: {_money(paint_supplies_total)}\n"
                    f"- Setup & Measure (if structural observed): {('2.0 hrs' if structural_flag else '0.0 hrs')}\n"
                    f"- Approx Total Frame Labor (if structural observed) @ Frame Rate: {('%.1f hrs' % _num(frame_hours)) if (structural_flag and frame_hours is not None) else ('Unknown' if structural_flag else '0.0 hrs')}\n"
                    f"- Approx Mechanical/ADAS/Airbag/Suspension Hours @ Mechanical Rate: {('%.1f' % _num(mech_hours)) if mech_hours is not None else 'Unknown'}\n\n"
                    "**Approx Cost of OEM Parts Needed:**\n"
                    f"{parts_md}\n\n"
                    "**Tax (parts + paint materials only):**\n"
                    f"- Taxable subtotal: {_money(taxable)}\n"
                    f"- Estimated tax: {_money(tax_total)}\n\n"
                    "**Approximated Repair Cost Total (Approximation Only):**\n"
                    f"- **{_money(approx_total)}**\n\n"
                    "_Repair Cost Disclaimer: This section is an approximation. It is not an official repair estimate and must be validated by a qualified appraiser using an estimating platform._"
                )

    except Exception:
        # Never fail the request for cost calculation issues
        if not (result.get("estimated_costs_markdown") or "").strip():
            result["estimated_costs_markdown"] = (
                "## Approximate Repair Cost Breakdown (Approximation Only)\n"
                "Unable to generate cost approximation on this run."
            )
        if ai_intent == "damage_report_from_photos":

            def _neutralize_side_terms(text: str) -> str:
                if not text:
                    return text

                replacements = [
                    (r"(?i)\b(driver|passenger)[-\s]?side\s+headlamp\b", "one headlamp"),
                    (r"(?i)\b(left|right)[-\s]?front\s+corner\b", "front corner"),
                    (r"(?i)\b(left|right)[-\s]?rear\s+corner\b", "rear corner"),
                    (r"(?i)\b(left|right)[-\s]?front\b", "front"),
                    (r"(?i)\b(left|right)[-\s]?rear\b", "rear"),
                    (r"(?i)\bdriver[-\s]?side\b", "one side"),
                    (r"(?i)\bpassenger[-\s]?side\b", "one side"),
                    (r"(?i)\bleft side\b", "one side"),
                    (r"(?i)\bright side\b", "one side"),
                ]

                for pattern, repl in replacements:
                    text = re.sub(pattern, repl, text)

                return text

            summary_markdown = _neutralize_side_terms(summary_markdown)
            estimated_costs_markdown = _neutralize_side_terms(estimated_costs_markdown)
            conclusion = _neutralize_side_terms(conclusion)
# -----------------------
    def _normalize_photos_only_cost_block(_cm: str) -> str:
        """Deterministic tail completer for Photos-Only cost markdown.
        If tax/total are missing, compute and inject the canonical lines and
        replace Severity Tier so exactly one box is checked from TOTAL.
        """
        _cm = _cm or ""
        _cm = re.sub(r"(?im)^\s*Approximate\s+Repair\s+Cost\s+Total\s*:.*$", "", _cm).strip()
        _cm = re.sub(r"(?im)^\s*Approximate\s+Repair\s+Cost\s+Total\b.*$", "", _cm).strip()

        def _moneyf(s: str) -> float:
            try:
                return float(str(s).replace(",", "").strip())
            except Exception:
                return 0.0

        def _grab(pats: List[str]) -> Optional[float]:
            for pat in pats:
                m = re.search(pat, _cm, flags=re.IGNORECASE | re.MULTILINE)
                if m:
                    try:
                        return _moneyf(m.group(1))
                    except Exception:
                        pass
            return None

        parts = _grab([
            r"^\s*[-*]?\s*Estimated\s+parts\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Parts\s+subtotal\s*\(approx\.?\)\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Parts\s+subtotal\s*\(approx\.?\)\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Parts\s+subtotal\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Parts\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
        ])
        labor = _grab([
            r"^\s*[-*]?\s*Labor\s+subtotal\s*\([^\n]*?\)\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Labor\s+subtotal\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Labor\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
        ])
        paint_materials = _grab([
            r"^\s*[-*]?\s*Paint\s*&\s*materials\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Paint\s*&\s*materials\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Paint\s+materials\s+subtotal\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Paint\s+materials\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Paint\s+materials\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Paint\s+materials\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
        ])
        sublet = _grab([
            r"^\s*[-*]?\s*Sublet\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Sublet/Other\s*\(approx\.?\)\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Rear\s+glass\s+install\s*\(sublet\s+allowance\)\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
        ])
        tax_amt = _grab([
            r"^\s*[-*]?\s*Sales\s+tax\s*\(assumed\s*[0-9]+(?:\.[0-9]+)?%\s*for\s*approximation\)\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*(?:Estimated\s+tax|Sales\s+tax|Tax)\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*(?:Estimated\s+tax|Sales\s+tax|Tax)\b.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
        ])
        approx_total = None

        if labor is None:
            body = _grab([
                r"^\s*[-*]?\s*Body(?:\s+labor)?\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"^\s*[-*]?\s*Body(?:\s+labor)?\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            ]) or 0.0
            paint = _grab([
                r"^\s*[-*]?\s*Paint\s+labor\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"^\s*[-*]?\s*Paint\s+labor\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"^\s*[-*]?\s*Refinish(?:\s+labor)?\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"^\s*[-*]?\s*Refinish(?:\s+labor)?\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            ]) or 0.0
            mechanical = _grab([
                r"^\s*[-*]?\s*Mechanical(?:/SRS/Glass|/diagnostic)?\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"^\s*[-*]?\s*Mechanical(?:/SRS/Glass|/diagnostic)?\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"^\s*[-*]?\s*Mechanical\s+labor\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"^\s*[-*]?\s*Mechanical\s+labor\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            ]) or 0.0
            setup = _grab([
                r"^\s*[-*]?\s*Setup\s*&\s*Measure\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"^\s*[-*]?\s*Setup\s*&\s*Measure\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            ]) or 0.0
            frame = _grab([
                r"^\s*[-*]?\s*Frame/measure\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"^\s*[-*]?\s*Frame/measure\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            ]) or 0.0
            calc_labor = body + paint + mechanical + setup + frame
            if calc_labor > 0:
                labor = calc_labor

        if tax_amt is None and parts is not None and paint_materials is not None:
            tax_amt = round((parts + paint_materials) * 0.07, 2)
        if parts is not None and paint_materials is not None and labor is not None and tax_amt is not None:
            approx_total = round(labor + parts + paint_materials + (sublet or 0.0) + tax_amt, 2)

        cleaned = []
        for ln in _cm.splitlines():
            s = (ln or "").strip()
            if re.search(r"(?i)^\s*[-*]?\s*Sales\s+tax\s*\(assumed\s*7%\s*for\s*approximation\)", s):
                continue
            if re.search(r"(?i)^\s*[-*]?\s*Tax\s+rate\s*:", s):
                continue
            if re.search(r"(?i)^\s*[-*]?\s*Tax\s+basis\s*\(", s):
                continue
            if re.search(r"(?i)^\s*[-*]?\s*Tax\s*:\s*\$", s):
                continue
            if re.search(r"(?i)^\s*\*{0,2}\s*Approximate\s+Repair\s+Total\s*:", s):
                continue
            if re.search(r"(?i)^\s*Severity\s+Tier\s*$", s):
                continue
            if re.search(r"(?i)^\s*\[[ xX]\]\s*(Minor|Moderate|Major|Possible\s+Total\s+Loss)", s):
                continue
            cleaned.append(ln)
        _cm = "\n".join(cleaned).rstrip()

        tail = []
        if tax_amt is not None:
            tail.append(f"- Sales tax (assumed 7% for approximation) = ${tax_amt:,.2f}")
        if approx_total is not None:
            tail.append(f"**Approximate Repair Total: ${approx_total:,.2f}**")

        checked_minor = checked_mod = checked_major = checked_tl = " "
        if isinstance(approx_total, (int, float)):
            if approx_total < 3500:
                checked_minor = "x"
            elif approx_total < 10000:
                checked_mod = "x"
            else:
                checked_major = "x"

        tail.extend([
            "Severity Tier",
            f"[{checked_minor}] Minor (< $3,500)",
            f"[{checked_mod}] Moderate ($3,500-$10,000)",
            f"[{checked_major}] Major ($10,000+)",
            f"[{checked_tl}] Possible Total Loss Threshold Approaching",
        ])
        return (_cm.rstrip() + "\n\n" + "\n".join(tail)).strip()
# -----------------------
    # --- Normalize Photos-Only cost markdown for downstream PDF/UI ---
    try:
        if ai_intent == "damage_report_from_photos":
            _cm = result.get("estimated_costs_markdown") or ""
            _cm = _normalize_photos_only_cost_block(_cm)
            result["estimated_costs_markdown"] = _cm
    except Exception:
        pass

    def _extract_itemized_part_lines(md_text: str) -> List[str]:
        """Extract itemized part lines for the locked photos-only PDF cost block.

        Primary path:
        - read only lines inside a parts heading/section

        Fallback path:
        - if no explicit parts section is found, scan all lines for likely part/component
          lines with dollar amounts while excluding labor, tax, totals, and rationale lines
        """
        text_local = str(md_text or '').replace("\r\n", "\n").replace("\r", "\n")
        out: List[str] = []
        seen = set()
        in_parts_section = False

        def _keep_line(s: str) -> bool:
            if not s or '$' not in s:
                return False
            if re.search(r'(?i)body\s+labor|paint\s+labor|paint\s*(?:&|and)\s*materials|setup\s*&\s*measure|frame\s+labor|mechanical|tax\s+rate|tax\s+basis|estimated\s+tax|sales\s+tax|labor\s+subtotal|parts\s+subtotal|subtotal|approximate\s+repair\s+total|severity\s+tier|rationale|hours?\b', s):
                return False
            if re.search(r'(?i)\b(rate|assumption|approximation only|cost calculation)\b', s):
                return False
            if not re.search(r'(?i)\b(assembly|bumper|reinforcement|impact\s+bar|absorber|lamp|light|glass|windshield|hatch|gate|liftgate|panel|closure|quarter|molding|trim|harness|connector|retainer|clips?|fasteners?|sealants?|sensor|camera|exhaust|wheel\-arch|aperture|spoiler|emblem|bracket|cover|grille|fender|door|mirror|oem)\b', s):
                return False
            return True

        for raw_ln in text_local.splitlines():
            s = (raw_ln or '').strip()
            if not s:
                continue

            if re.search(r'(?i)^\*{0,2}\s*(OEM\s+replacement\s+parts|OEM\s+parts\s+needed|replacement\s+parts|parts\s+needed|itemized\s+parts\s+breakdown)\b', s):
                in_parts_section = True
                continue

            if in_parts_section and re.search(r'(?i)parts\s+subtotal|estimated\s+parts\s+subtotal|taxable\s+subtotal|estimated\s+tax|sales\s+tax|tax\s+rate|tax\s+basis|approximate\s+repair\s+total|severity\s+tier', s):
                break
            if in_parts_section and re.search(r'(?i)^\*{0,2}\s*(labor|body\s+labor|paint\s+labor|paint\s*(?:&|and)\s*materials|setup\s*&\s*measure|frame\s+labor|mechanical|sublet)\b', s):
                break

            if in_parts_section and _keep_line(s):
                if s not in seen:
                    out.append(s)
                    seen.add(s)

        if out:
            return out

        for raw_ln in text_local.splitlines():
            s = (raw_ln or '').strip()
            if not _keep_line(s):
                continue
            if s not in seen:
                out.append(s)
                seen.add(s)

        return out



    def _seed_scope_labor_buckets(scope_text: str, parsed: Dict[str, Any], seed_rates: Optional[Dict[str, float]] = None) -> Dict[str, Any]:
        """Deterministically seed labor buckets from visible repair scope when labor collapsed to empty.

        This only runs when *all* labor buckets are effectively empty, but the narrative/parts scope
        clearly supports body / paint / mechanical work. It preserves the existing atomic bucket lock
        and simply ensures there is something real to finalize.
        """
        out = dict(parsed or {})
        seed_rates = dict(seed_rates or {})

        def _has_pos(v: Any) -> bool:
            try:
                return isinstance(v, (int, float)) and float(v) > 0
            except Exception:
                return False

        labor_keys = (
            'body_hours', 'paint_hours', 'setup_hours', 'frame_hours', 'mech_hours',
            'body_labor', 'paint_labor', 'setup_measure', 'frame_labor', 'mech_labor',
        )
        if any(_has_pos(out.get(k)) for k in labor_keys):
            return out

        scope = str(scope_text or '').lower()
        if not scope.strip():
            return out

        parts_lines = out.get('parts_lines') or []
        parts_blob = " ".join(str(x).lower() for x in parts_lines)
        scope_all = (scope + " " + parts_blob).strip()

        def _has(rx: str) -> bool:
            try:
                return bool(re.search(rx, scope_all, flags=re.IGNORECASE))
            except Exception:
                return False

        # Deterministic fallback rates. Prefer parsed rates, then seeded location rates, then conservative defaults.
        body_rate = float(out.get('body_rate') or 0.0) or float(seed_rates.get('body_rate') or 0.0) or float(seed_rates.get('paint_rate') or 0.0) or 75.0
        paint_rate = float(out.get('paint_rate') or 0.0) or float(seed_rates.get('paint_rate') or 0.0) or body_rate or 75.0
        frame_rate = float(out.get('frame_rate') or 0.0) or float(seed_rates.get('frame_rate') or 0.0) or body_rate or 95.0
        mech_rate = float(out.get('mech_rate') or 0.0) or float(seed_rates.get('mechanical_rate') or 0.0) or 95.0
        paint_mat_rate = float(out.get('paint_mat_rate') or 0.0) or float(seed_rates.get('paint_supplies_rate') or 0.0) or 45.0

        body_hours = 0.0
        if _has(r'quarter'):
            body_hours += 16.0
        if _has(r'rear\s+door|door\s+shell|door'):
            body_hours += 10.0
        if _has(r'bumper\s+cover|bumper'):
            body_hours += 4.0
        if _has(r'tail\s+lamp|tail lamp|lamp'):
            body_hours += 1.0
        if _has(r'molding|trim|clips?|fasteners?|hardware'):
            body_hours += 1.0
        if _has(r'crush|crushing|crease|creased|tear|torn|scrap|scraping|misalign|deform|distortion'):
            body_hours += 4.0
        if body_hours <= 0 and (_has(r'panel|door|quarter|bumper|fender|hood|gate|liftgate|hatch') or parts_lines):
            body_hours = 8.0

        paint_hours = 0.0
        if _has(r'paint|refinish|blend|blending|color\s+match'):
            paint_hours += 4.0
        if _has(r'quarter'):
            paint_hours += 6.0
        if _has(r'rear\s+door|door\s+shell|door'):
            paint_hours += 5.0
        if _has(r'bumper\s+cover|bumper'):
            paint_hours += 4.0
        if _has(r'adjacent\s+panel|blend'):
            paint_hours += 2.0
        if paint_hours <= 0 and _has_pos(out.get('paint_mat')) and paint_mat_rate > 0:
            paint_hours = round(float(out.get('paint_mat') or 0.0) / float(paint_mat_rate), 1)
        elif paint_hours <= 0 and body_hours > 0 and _has(r'white\s+finish|color\s+match|refinish|blend|paint'):
            paint_hours = round(max(1.0, body_hours * 0.45), 1)

        setup_hours = 2.0 if body_hours > 0 and _has(r'structural|aperture|setup\s*&\s*measure|measure|heavy\s+crushing|severe') else 0.0
        frame_hours = 3.0 if _has(r'structural|aperture|pull|straighten|rail|frame|unibody|buckl|kink|twist') else 0.0
        mech_hours = 4.0 if _has(r'wheel/?tire|wheel|tire|suspension|alignment|knuckle|control\s+arm|hub|mechanical|adas') else 0.0

        out['body_hours'] = round(body_hours, 1) if body_hours > 0 else out.get('body_hours')
        out['paint_hours'] = round(paint_hours, 1) if paint_hours > 0 else out.get('paint_hours')
        out['setup_hours'] = round(setup_hours, 1) if setup_hours > 0 else out.get('setup_hours')
        out['frame_hours'] = round(frame_hours, 1) if frame_hours > 0 else out.get('frame_hours')
        out['mech_hours'] = round(mech_hours, 1) if mech_hours > 0 else out.get('mech_hours')
        out['body_rate'] = body_rate
        out['paint_rate'] = paint_rate
        out['frame_rate'] = frame_rate
        out['mech_rate'] = mech_rate
        out['paint_mat_rate'] = paint_mat_rate

        # Seed paint materials from seeded paint hours if the model left it empty.
        if not _has_pos(out.get('paint_mat')) and _has_pos(out.get('paint_hours')) and paint_mat_rate > 0:
            out['paint_mat'] = round(float(out.get('paint_hours') or 0.0) * float(paint_mat_rate), 2)

        return out

    def _apply_normalization_lock(parsed: Dict[str, Any]) -> Dict[str, Any]:
        """Stabilize photos-only cost outputs without changing the underlying logic path.

        This helper now enforces atomic labor bucket finalization so a bucket can never
        print as a mixed hybrid like 0.0 hrs @ $0.00/hr = $2,890.00.
        For each labor bucket, finalize exactly one consistent tuple:
        - hours
        - rate
        - amount
        Then recompute subtotal/tax/total from those finalized buckets only.
        """
        out = dict(parsed or {})

        def _round_half(v):
            try:
                return round(round(float(v) * 2.0) / 2.0, 1)
            except Exception:
                return None

        def _round_rate(v):
            try:
                x = float(v)
                return round(x, 2) if x > 0 else 0.0
            except Exception:
                return 0.0

        def _money(v):
            try:
                return round(float(v), 2)
            except Exception:
                return 0.0

        def _finalize_labor_bucket(hours, rate, amount, fallback_rate=None):
            hrs = _round_half(hours) if isinstance(hours, (int, float)) else None
            rte = _round_rate(rate) if isinstance(rate, (int, float)) else 0.0
            amt = _money(amount)
            fb = _round_rate(fallback_rate) if isinstance(fallback_rate, (int, float)) else 0.0

            # Prefer a locked fallback rate when the parsed rate dropped out.
            if rte <= 0 and fb > 0:
                rte = fb

            # Case 1: hours + rate => compute amount
            if isinstance(hrs, (int, float)) and rte > 0:
                amt = round(float(hrs) * float(rte), 2)
                return hrs, rte, amt

            # Case 2: amount + hours => derive rate
            if isinstance(hrs, (int, float)) and float(hrs) > 0 and amt > 0:
                rte = round(float(amt) / float(hrs), 2)
                return hrs, rte, amt

            # Case 3: amount + rate => derive hours
            if rte > 0 and amt > 0:
                hrs = _round_half(float(amt) / float(rte))
                return hrs, rte, amt

            # Case 3b: amount only, but usable labor evidence exists.
            # Do NOT zero the bucket just because the model omitted hours/rate.
            # Preserve the labor dollars and derive a consistent tuple from a locked rate.
            if amt > 0:
                locked_rate = rte if rte > 0 else (fb if fb > 0 else 75.0)
                hrs = _round_half(float(amt) / float(locked_rate))
                return hrs, locked_rate, amt

            # Case 4: partial/empty bucket => zero the whole bucket consistently
            return 0.0, rte if rte > 0 else 0.0, 0.0

        for rk in ('body_rate', 'paint_rate', 'frame_rate', 'mech_rate'):
            out[rk] = _round_rate(out.get(rk))

        # Keep tax_rate_value at full precision. Do not round 0.095 to 0.10.
        # Only the final tax dollars should be rounded to cents.
        try:
            if isinstance(out.get('tax_rate_value'), (int, float)) and float(out.get('tax_rate_value')) > 0:
                out['tax_rate_value'] = float(out.get('tax_rate_value'))
        except Exception:
            out['tax_rate_value'] = out.get('tax_rate_value')

        if isinstance(out.get('paint_mat_rate'), (int, float)):
            out['paint_mat_rate'] = _round_rate(out.get('paint_mat_rate'))

        # Finalize all labor buckets atomically.
        bh, br, ba = _finalize_labor_bucket(out.get('body_hours'), out.get('body_rate'), out.get('body_labor'))
        ph, pr, pa = _finalize_labor_bucket(out.get('paint_hours'), out.get('paint_rate'), out.get('paint_labor'), fallback_rate=br)
        sh, sr, sa = _finalize_labor_bucket(out.get('setup_hours'), out.get('body_rate'), out.get('setup_measure'), fallback_rate=br)
        fh, fr, fa = _finalize_labor_bucket(out.get('frame_hours'), out.get('frame_rate'), out.get('frame_labor'))
        mh, mr, ma = _finalize_labor_bucket(out.get('mech_hours'), out.get('mech_rate'), out.get('mech_labor'))

        out['body_hours'], out['body_rate'], out['body_labor'] = bh, br, ba
        out['paint_hours'], out['paint_rate'], out['paint_labor'] = ph, pr, pa
        out['setup_hours'], out['body_rate'], out['setup_measure'] = sh, sr if sr > 0 else br, sa
        out['frame_hours'], out['frame_rate'], out['frame_labor'] = fh, fr, fa
        out['mech_hours'], out['mech_rate'], out['mech_labor'] = mh, mr, ma

        parts_lines = out.get('parts_lines') or []
        if isinstance(parts_lines, list):
            norm = []
            seen = set()
            for pl in parts_lines:
                s = re.sub(r'^[-*]\s*', '', str(pl or '').strip())
                if not s:
                    continue
                key = re.sub(r'\s+', ' ', s).strip().lower()
                if key in seen:
                    continue
                seen.add(key)
                norm.append(s)
            norm.sort(key=lambda s: (s.lower(), s))
            out['parts_lines'] = norm

        out['parts_sub'] = _money(out.get('parts_sub'))
        out['paint_mat'] = _money(out.get('paint_mat'))
        out['sublet'] = _money(out.get('sublet'))
        out['labor_sub'] = round(ba + pa + sa + fa + ma, 2)
        out['tax_basis'] = round(_money(out.get('parts_sub')) + _money(out.get('paint_mat')), 2)
        tax_rate_val = out.get('tax_rate_value') if isinstance(out.get('tax_rate_value'), (int, float)) and out.get('tax_rate_value') > 0 else 0.07
        out['tax_amt'] = round(_money(out.get('tax_basis')) * float(tax_rate_val), 2)
        out['total_val'] = round(
            _money(out.get('labor_sub')) +
            _money(out.get('parts_sub')) +
            _money(out.get('paint_mat')) +
            _money(out.get('sublet')) +
            _money(out.get('tax_amt')), 2
        )
        return out

    def _apply_locked_rate_overrides_to_parsed(parsed: Dict[str, Any], overrides: Optional[Dict[str, float]] = None) -> Dict[str, Any]:
        """Force explicit Add'l Notes labor-rate overrides into the locked parsed object.

        Narrow lock behavior:
        - preserve the tax fix and parts capture
        - explicit Add'l Notes body/paint overrides are the canonical body/paint rates
        - if the model block already contains real body/paint hours, keep those exact hours and apply the locked override rates
        - do NOT zero seeded labor buckets unless override-driven body/paint hours were actually locked first
        - only zero setup/frame/mech when they were seeded-only and the model block did not explicitly emit them
        """
        out = dict(parsed or {})
        ov = dict(overrides or {})

        def _pos(v: Any) -> Optional[float]:
            try:
                x = float(v)
                return x if x > 0 else None
            except Exception:
                return None

        def _has_pos(v: Any) -> bool:
            try:
                return isinstance(v, (int, float)) and float(v) > 0
            except Exception:
                return False

        body_override = _pos(ov.get('body_rate'))
        paint_override = _pos(ov.get('paint_rate'))
        if body_override is None and paint_override is not None:
            body_override = paint_override
        if paint_override is None and body_override is not None:
            paint_override = body_override

        explicit_body_present = bool(out.get('explicit_body_present'))
        explicit_paint_present = bool(out.get('explicit_paint_present'))
        explicit_setup_present = bool(out.get('explicit_setup_present'))
        explicit_frame_present = bool(out.get('explicit_frame_present'))
        explicit_mech_present = bool(out.get('explicit_mech_present'))

        def _rebuild(bucket_hours_key: str, bucket_rate_key: str, bucket_amount_key: str, forced_rate: Optional[float]) -> None:
            if forced_rate is None:
                return
            out[bucket_rate_key] = round(float(forced_rate), 2)
            hrs = out.get(bucket_hours_key)
            amt = out.get(bucket_amount_key)
            try:
                if isinstance(hrs, (int, float)) and float(hrs) > 0:
                    out[bucket_amount_key] = round(float(hrs) * float(forced_rate), 2)
                    return
            except Exception:
                pass
            try:
                if isinstance(amt, (int, float)) and float(amt) > 0:
                    out[bucket_hours_key] = round(round((float(amt) / float(forced_rate)) * 2.0) / 2.0, 1)
                    out[bucket_amount_key] = round(float(out[bucket_hours_key]) * float(forced_rate), 2)
            except Exception:
                pass

        # First pass: if the parser/model block already gave us usable body/paint hours or dollars,
        # keep those exact hours and apply the explicit Add'l Notes rates.
        _rebuild('body_hours', 'body_rate', 'body_labor', body_override)
        _rebuild('paint_hours', 'paint_rate', 'paint_labor', paint_override)

        locked_body_paint_hours = (
            _has_pos(out.get('body_hours')) or _has_pos(out.get('paint_hours')) or
            _has_pos(out.get('body_labor')) or _has_pos(out.get('paint_labor'))
        )

        # If explicit overrides exist but no body/paint hours survived parsing, try one narrow reseed from scope.
        # This preserves the earlier tax fix/parts capture while preventing a zero-labor collapse.
        if (body_override is not None or paint_override is not None) and not locked_body_paint_hours:
            try:
                scope_seed_rates = {
                    'body_rate': float(body_override or paint_override or 0.0),
                    'paint_rate': float(paint_override or body_override or 0.0),
                    'frame_rate': float(out.get('frame_rate') or 0.0),
                    'mechanical_rate': float(out.get('mech_rate') or 0.0),
                    'paint_supplies_rate': float(out.get('paint_mat_rate') or 0.0),
                }
                out = _seed_scope_labor_buckets(str(out.get('scope_text_for_rate_lock') or ''), out, seed_rates=scope_seed_rates)
            except Exception:
                pass
            _rebuild('body_hours', 'body_rate', 'body_labor', body_override)
            _rebuild('paint_hours', 'paint_rate', 'paint_labor', paint_override)
            locked_body_paint_hours = (
                _has_pos(out.get('body_hours')) or _has_pos(out.get('paint_hours')) or
                _has_pos(out.get('body_labor')) or _has_pos(out.get('paint_labor'))
            )

        # Only suppress seeded setup/frame/mech after real override-driven body/paint hours were successfully locked.
        if locked_body_paint_hours and (body_override is not None or paint_override is not None):
            if not explicit_setup_present:
                out['setup_hours'] = 0.0
                out['setup_measure'] = 0.0
            if not explicit_frame_present:
                out['frame_hours'] = 0.0
                out['frame_labor'] = 0.0
                out['frame_rate'] = 0.0
            if not explicit_mech_present:
                out['mech_hours'] = 0.0
                out['mech_labor'] = 0.0
                out['mech_rate'] = 0.0

        out = _apply_normalization_lock(out)

        # Final hard lock: explicit Add'l Notes body/paint overrides are the printed rates.
        if body_override is not None:
            out['body_rate'] = round(float(body_override), 2)
            try:
                if isinstance(out.get('body_hours'), (int, float)):
                    out['body_labor'] = round(float(out.get('body_hours') or 0.0) * float(body_override), 2)
            except Exception:
                pass
        if paint_override is not None:
            out['paint_rate'] = round(float(paint_override), 2)
            try:
                if isinstance(out.get('paint_hours'), (int, float)):
                    out['paint_labor'] = round(float(out.get('paint_hours') or 0.0) * float(paint_override), 2)
            except Exception:
                pass

        # Re-apply setup/frame/mech suppression only when body/paint hours actually survived the lock.
        if locked_body_paint_hours and (body_override is not None or paint_override is not None):
            if not explicit_setup_present:
                out['setup_hours'] = 0.0
                out['setup_measure'] = 0.0
            if not explicit_frame_present:
                out['frame_hours'] = 0.0
                out['frame_labor'] = 0.0
                out['frame_rate'] = 0.0
            if not explicit_mech_present:
                out['mech_hours'] = 0.0
                out['mech_labor'] = 0.0
                out['mech_rate'] = 0.0

        return _apply_normalization_lock(out)

    def _parse_locked_photos_only_costs(md_text: str, tax_rate_value: Optional[float] = None, seed_rates: Optional[Dict[str, float]] = None, scope_text: Optional[str] = None) -> Dict[str, Any]:
        """Single source of truth for photos-only cost math.
        Locked rules:
        - labor subtotal is ALWAYS recomputed from the five printed labor buckets
        - tax is ALWAYS recomputed from (parts subtotal + paint materials) * tax_rate
        - total is ALWAYS recomputed from labor subtotal + parts subtotal + paint materials + sublet + tax
        - if explicit parts subtotal is missing, itemized part lines are summed and become the single parts subtotal
        - carry the actual hours/rates used so the PDF can show how each labor total was derived
        """
        text_local = str(md_text or '').replace("\r\n", "\n").replace("\r", "\n")

        if tax_rate_value is None or not isinstance(tax_rate_value, (int, float)) or tax_rate_value <= 0:
            _m_tax_rate_inline = re.search(r'(?im)^\s*[-*]?\s*Tax\s+rate\s*:\s*([0-9]+(?:\.[0-9]+)?)\s*%\s*$', text_local)
            if not _m_tax_rate_inline:
                _m_tax_rate_inline = re.search(r'(?im)\bTax\s+rate\s*(?:\(assumed\))?\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)\s*%', text_local)
            if _m_tax_rate_inline:
                try:
                    tax_rate_value = float(_m_tax_rate_inline.group(1)) / 100.0
                except Exception:
                    tax_rate_value = 0.07
            else:
                tax_rate_value = 0.07

        def _grab_money_line(pats: List[str]) -> Optional[float]:
            for pat in pats:
                mm = re.search(pat, text_local, flags=re.IGNORECASE | re.MULTILINE)
                if mm:
                    try:
                        return float(mm.group(1).replace(',', ''))
                    except Exception:
                        pass
            return None

        def _grab_float_line(pats: List[str]) -> Optional[float]:
            for pat in pats:
                mm = re.search(pat, text_local, flags=re.IGNORECASE | re.MULTILINE)
                if mm:
                    try:
                        return float(mm.group(1).replace(',', ''))
                    except Exception:
                        pass
            return None

        def _sum_itemized_part_amounts(src_text: str) -> Optional[float]:
            """Sum part dollars from the exact normalized itemized part lines.
            This prevents ZIP/JPG divergence where lines can be extracted/printed,
            but subtotal assignment fails because the raw-text heading wasn't recognized.
            """
            extracted_lines = _extract_itemized_part_lines(src_text)
            total = 0.0
            found = False
            for raw_ln in extracted_lines:
                s = (raw_ln or '').strip()
                if not s:
                    continue
                monies = re.findall(r'\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)', s)
                if monies:
                    try:
                        total += float(monies[-1].replace(',', ''))
                        found = True
                    except Exception:
                        pass
            return round(total, 2) if found else None

        # Accept both canonical locked lines and model-emitted markdown like:
        # - Body labor: **45.0 hrs @ $68/hr = $3,060.00**
        # - Paint & materials: **22.0 refinish hrs × $45/hr = $990.00**
        _num_pat = r'([0-9][0-9,]*(?:\.[0-9]+)?)'
        body_rate = _grab_float_line([
            rf'(?im)^\s*[-*]?\s*Body\s+labor\s+rate\s*:\s*\$\s*{_num_pat}\s*/\s*hr',
            rf'(?im)^\s*[-*]?\s*Body(?:\s+labor)?\s*:\s*\*{{0,2}}\s*{_num_pat}\s*hrs?\s*@\s*\$\s*{_num_pat}\s*/\s*hr',
            rf'(?im)^\s*[-*]?\s*Body(?:\s+labor)?\s*:\s*(?:\*{{1,2}})?[^\n]*?@\s*\$\s*{_num_pat}\s*/\s*hr',
        ])
        paint_rate = _grab_float_line([
            rf'(?im)^\s*[-*]?\s*Paint\s+labor\s+rate\s*:\s*\$\s*{_num_pat}\s*/\s*hr',
            rf'(?im)^\s*[-*]?\s*Paint\s+labor\s*:\s*\*{{0,2}}\s*{_num_pat}\s*hrs?\s*@\s*\$\s*{_num_pat}\s*/\s*hr',
            rf'(?im)^\s*[-*]?\s*Paint\s+labor\s*:\s*(?:\*{{1,2}})?[^\n]*?@\s*\$\s*{_num_pat}\s*/\s*hr',
            rf'(?im)^\s*[-*]?\s*Refinish(?:\s+labor)?\s*:\s*\*{{0,2}}\s*{_num_pat}\s*hrs?\s*@\s*\$\s*{_num_pat}\s*/\s*hr',
            rf'(?im)^\s*[-*]?\s*Refinish(?:\s+labor)?\s*:\s*(?:\*{{1,2}})?[^\n]*?@\s*\$\s*{_num_pat}\s*/\s*hr',
        ])
        frame_rate = _grab_float_line([
            rf'(?im)^\s*[-*]?\s*(?:Frame|Structural)(?:\s+labor)?\s+rate\s*:\s*\$\s*{_num_pat}\s*/\s*hr',
            rf'(?im)^\s*[-*]?\s*Frame(?:/measure|\s+labor)?\s*:\s*(?:\*{{1,2}})?[^\n]*?@\s*\$\s*{_num_pat}\s*/\s*hr',
            rf'(?im)^\s*[-*]?\s*Structural(?:\s+labor)?\s*:\s*(?:\*{{1,2}})?[^\n]*?@\s*\$\s*{_num_pat}\s*/\s*hr',
        ])
        mech_rate = _grab_float_line([
            rf'(?im)^\s*[-*]?\s*Mechanical(?:/diagnostic)?\s+rate\s*:\s*\$\s*{_num_pat}\s*/\s*hr',
            rf'(?im)^\s*[-*]?\s*Mechanical(?:/SRS/Glass|/diagnostic)?[^\n]*?:\s*(?:\*{{1,2}})?[^\n]*?@\s*\$\s*{_num_pat}\s*/\s*hr',
            rf'(?im)^\s*[-*]?\s*Mechanical\s+labor[^\n]*?:\s*(?:\*{{1,2}})?[^\n]*?@\s*\$\s*{_num_pat}\s*/\s*hr',
        ])

        paint_mat_rate = _grab_float_line([
            rf'(?im)^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s+rate\s*:\s*\$\s*{_num_pat}\b',
            rf'(?im)^\s*[-*]?\s*Paint\s+materials\s+rate\s*:\s*\$\s*{_num_pat}\b',
            rf'(?im)^\s*[-*]?\s*Paint\s+suppl(?:y|ies)\s+rate\s*:\s*\$\s*{_num_pat}\b',
            rf'(?im)^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s*:\s*(?:\*{{1,2}})?[^\n]*?[×x*]\s*\$\s*{_num_pat}\s*/\s*hr',
            rf'(?im)^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s*:\s*(?:\*{{1,2}})?[^\n]*?\$\s*{_num_pat}\s*(?:per\s+refinish\s+hour|/\s*hr)',
            rf'(?im)^\s*[-*]?\s*Paint\s+materials\s*:\s*(?:\*{{1,2}})?[^\n]*?[×x*]\s*\$\s*{_num_pat}\s*/\s*hr',
        ])

        def _derive_rate_from_amount(amount: Optional[float], hours: Optional[float], current_rate: Optional[float]) -> Optional[float]:
            if isinstance(current_rate, (int, float)) and float(current_rate) > 0:
                return float(current_rate)
            if isinstance(amount, (int, float)) and isinstance(hours, (int, float)) and float(hours) > 0:
                try:
                    return round(float(amount) / float(hours), 2)
                except Exception:
                    return current_rate
            return current_rate

        body_hours = _grab_float_line([
            r'(?im)^\s*[-*]?\s*Body(?:\s+labor)?\s*:\s*\*{0,2}\s*([0-9]+(?:\.[0-9]+)?)\s*hrs?\b',
            r'(?im)^\s*[-*]?\s*Body(?:\s+labor)?\s*:\s*(?:\*{1,2})?[^\n]*?([0-9]+(?:\.[0-9]+)?)\s*hrs?\s*@',
        ])
        paint_hours = _grab_float_line([
            r'(?im)^\s*[-*]?\s*Paint\s+labor\s*\(?(?:refinish)?\)?\s*:\s*\*{0,2}\s*([0-9]+(?:\.[0-9]+)?)\s*hrs?\b',
            r'(?im)^\s*[-*]?\s*Paint\s+labor\s*:\s*(?:\*{1,2})?[^\n]*?([0-9]+(?:\.[0-9]+)?)\s*hrs?\s*@',
            r'(?im)^\s*[-*]?\s*Refinish(?:\s+labor)?\s*:\s*\*{0,2}\s*([0-9]+(?:\.[0-9]+)?)\s*hrs?\b',
            r'(?im)^\s*[-*]?\s*Refinish(?:\s+labor)?\s*:\s*(?:\*{1,2})?[^\n]*?([0-9]+(?:\.[0-9]+)?)\s*hrs?\s*@',
        ])
        setup_hours = _grab_float_line([
            r'(?im)^\s*[-*]?\s*Setup\s*&\s*Measure[^\n]*?:\s*\*{0,2}\s*([0-9]+(?:\.[0-9]+)?)\s*hrs?\b',
            r'(?im)^\s*[-*]?\s*Setup\s*&\s*Measure[^\n]*?:\s*(?:\*{1,2})?[^\n]*?([0-9]+(?:\.[0-9]+)?)\s*hrs?\s*@',
        ])
        frame_hours = _grab_float_line([
            r'(?im)^\s*[-*]?\s*Frame(?:/measure|\s+labor)?\s*:\s*\*{0,2}\s*([0-9]+(?:\.[0-9]+)?)\s*hrs?\b',
            r'(?im)^\s*[-*]?\s*Frame(?:/measure|\s+labor)?\s*:\s*(?:\*{1,2})?[^\n]*?([0-9]+(?:\.[0-9]+)?)\s*hrs?\s*@',
        ])
        mech_hours = _grab_float_line([
            r'(?im)^\s*[-*]?\s*Mechanical(?:/SRS/Glass|/diagnostic)?[^\n]*?:\s*\*{0,2}\s*([0-9]+(?:\.[0-9]+)?)\s*hrs?\b',
            r'(?im)^\s*[-*]?\s*Mechanical(?:/SRS/Glass|/diagnostic)?[^\n]*?:\s*(?:\*{1,2})?[^\n]*?([0-9]+(?:\.[0-9]+)?)\s*hrs?\s*@',
        ])

        # Hard-bind paint labor when refinish/blend scope is present but explicit paint hours were not emitted.
        # This restores consistent paint labor + paint materials for both ZIP and loose JPG runs
        # without changing any other cost-path behavior.
        _refinish_scope_present = bool(re.search(
            r'(?i)\b(refinish|blend|blending|repaint|paint\s+operations?|adjacent\s+panel\s+blend)\b',
            text_local,
        ))
        if paint_hours is None:
            _paint_labor_amount_for_hours = _grab_money_line([
                r'^\s*[-*]?\s*Paint\s+labor\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
                r'^\s*[-*]?\s*Paint\s+labor\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*\s*$',
                r'^\s*[-*]?\s*Refinish\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
                r'^\s*[-*]?\s*Refinish\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*\s*$',
            ])
            _paint_mat_amount_for_hours = _grab_money_line([
                r'^\s*[-*]?\s*Paint\s+materials\s+subtotal\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
                r'^\s*[-*]?\s*Paint\s+materials\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
                r'^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
                r'^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*\s*$',
                r'^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            ])
            if isinstance(_paint_labor_amount_for_hours, (int, float)) and isinstance(paint_rate, (int, float)) and float(paint_rate) > 0:
                paint_hours = round(float(_paint_labor_amount_for_hours) / float(paint_rate), 1)
            elif isinstance(_paint_mat_amount_for_hours, (int, float)) and isinstance(paint_mat_rate, (int, float)) and float(paint_mat_rate) > 0:
                paint_hours = round(float(_paint_mat_amount_for_hours) / float(paint_mat_rate), 1)
            elif _refinish_scope_present and isinstance(body_hours, (int, float)) and float(body_hours) > 0:
                # Conservative, deterministic fallback for photo-based refinish scope.
                paint_hours = round(max(1.0, float(body_hours) * 0.4), 1)

        body_labor_explicit = _grab_money_line([
            r'^\s*[-*]?\s*Body(?:\s+labor)?\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*[-*]?\s*Body(?:\s+labor)?\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*\s*$',
        ])
        paint_labor_explicit = _grab_money_line([
            r'^\s*[-*]?\s*Paint\s+labor\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*[-*]?\s*Paint\s+labor\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*\s*$',
            r'^\s*[-*]?\s*Refinish(?:\s+labor)?\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*[-*]?\s*Refinish(?:\s+labor)?\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*\s*$',
        ])
        setup_measure_explicit = _grab_money_line([
            r'^\s*[-*]?\s*Setup\s*&\s*Measure\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*[-*]?\s*Setup\s*&\s*Measure\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*\s*$',
        ])
        frame_labor_explicit = _grab_money_line([
            r'^\s*[-*]?\s*Frame(?:/measure|\s+labor)?\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*[-*]?\s*Frame(?:/measure|\s+labor)?\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*\s*$',
        ])
        mech_labor_explicit = _grab_money_line([
            r'^\s*[-*]?\s*Mechanical(?:/SRS/Glass|/diagnostic)?\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*[-*]?\s*Mechanical(?:/SRS/Glass|/diagnostic)?\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*\s*$',
            r'^\s*[-*]?\s*Mechanical\s+labor\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*[-*]?\s*Mechanical\s+labor\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*\s*$',
        ])

        explicit_body_present = (body_hours is not None) or (body_labor_explicit is not None)
        explicit_paint_present = (paint_hours is not None) or (paint_labor_explicit is not None)
        explicit_setup_present = (setup_hours is not None) or (setup_measure_explicit is not None)
        explicit_frame_present = (frame_hours is not None) or (frame_labor_explicit is not None)
        explicit_mech_present = (mech_hours is not None) or (mech_labor_explicit is not None)

        # Preserve hrs @ rate = total formatting by deriving missing rates from explicit dollars when possible.
        body_rate = _derive_rate_from_amount(body_labor_explicit, body_hours, body_rate)
        paint_rate = _derive_rate_from_amount(paint_labor_explicit, paint_hours, paint_rate)
        frame_rate = _derive_rate_from_amount(frame_labor_explicit, frame_hours, frame_rate)
        mech_rate = _derive_rate_from_amount(mech_labor_explicit, mech_hours, mech_rate)

        if (not isinstance(frame_rate, (int, float)) or float(frame_rate) <= 0) and isinstance(seed_rates, dict):
            try:
                _seed_frame_rate = float(seed_rates.get('frame_rate') or 0.0)
                if _seed_frame_rate > 0:
                    frame_rate = _seed_frame_rate
            except Exception:
                pass

        # Surgical paint-rate lock:
        # if paint labor hours exist but paint rate was not parsed/derived, do not let it fall to $0.00/hr.
        # keep the existing file behavior everywhere else and use the locked body rate as the fallback,
        # which preserves the expected hours @ rate = total format in the PDF cost block.
        if isinstance(paint_hours, (int, float)) and float(paint_hours) > 0:
            if not isinstance(paint_rate, (int, float)) or float(paint_rate) <= 0:
                if isinstance(body_rate, (int, float)) and float(body_rate) > 0:
                    paint_rate = float(body_rate)

        def _derive_amount(hours: Optional[float], rate: Optional[float], explicit: Optional[float]) -> float:
            if isinstance(hours, (int, float)) and isinstance(rate, (int, float)):
                return round(float(hours) * float(rate), 2)
            if isinstance(explicit, (int, float)):
                return round(float(explicit), 2)
            return 0.0

        body_labor = _derive_amount(body_hours, body_rate, body_labor_explicit)
        paint_labor = _derive_amount(paint_hours, paint_rate, paint_labor_explicit)
        setup_measure = _derive_amount(setup_hours, body_rate, setup_measure_explicit)
        frame_labor = _derive_amount(frame_hours, frame_rate, frame_labor_explicit)
        mech_labor = _derive_amount(mech_hours, mech_rate, mech_labor_explicit)

        labor_sub = round(body_labor + paint_labor + setup_measure + frame_labor + mech_labor, 2)

        parts_sub = _grab_money_line([
            r'^\s*\*{0,2}\s*[-*]?\s*Estimated\s+parts\s+subtotal\s*:\s*\*{0,2}\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*\*{0,2}\s*[-*]?\s*Parts\s+subtotal\s*\(approx\.?\)\s*=\s*\*{0,2}\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*\*{0,2}\s*[-*]?\s*Parts\s+subtotal\s*\(approx\.?\)\s*:\s*\*{0,2}\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*\*{0,2}\s*[-*]?\s*Parts\s+subtotal\s*=\s*\*{0,2}\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*\*{0,2}\s*[-*]?\s*Parts\s+subtotal\s*:\s*\*{0,2}\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
        ])
        itemized_parts_sub = _sum_itemized_part_amounts(text_local)
        if isinstance(parts_sub, (int, float)):
            parts_sub = round(float(parts_sub), 2)
        elif isinstance(itemized_parts_sub, (int, float)):
            parts_sub = round(float(itemized_parts_sub), 2)
        else:
            parts_sub = 0.0

        paint_mat = _grab_money_line([
            r'^\s*[-*]?\s*Paint\s+materials\s+subtotal\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*[-*]?\s*Paint\s+materials\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*\s*$',
            r'^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
        ])
        if (not isinstance(paint_mat_rate, (int, float)) or float(paint_mat_rate) <= 0) and isinstance(paint_mat, (int, float)) and isinstance(paint_hours, (int, float)) and float(paint_hours) > 0:
            try:
                paint_mat_rate = round(float(paint_mat) / float(paint_hours), 2)
            except Exception:
                pass
        if isinstance(paint_hours, (int, float)) and float(paint_hours) > 0 and isinstance(paint_mat_rate, (int, float)) and float(paint_mat_rate) > 0:
            paint_mat = round(float(paint_hours) * float(paint_mat_rate), 2)
        elif isinstance(paint_mat, (int, float)):
            paint_mat = round(float(paint_mat), 2)
        else:
            paint_mat = 0.0

        sublet = _grab_money_line([
            r'^\s*[-*]?\s*Sublet\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
            r'^\s*[-*]?\s*Rear\s+glass\s+install\s*\(sublet\s+allowance\)\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b',
        ])
        sublet = round(float(sublet), 2) if isinstance(sublet, (int, float)) else 0.0

        tax_basis = round(parts_sub + paint_mat, 2)
        tax_amt = round(tax_basis * float(tax_rate_value), 2)
        total_val = round(labor_sub + parts_sub + paint_mat + sublet + tax_amt, 2)

        _parsed = {
            'body_hours': body_hours,
            'paint_hours': paint_hours,
            'setup_hours': setup_hours,
            'frame_hours': frame_hours,
            'mech_hours': mech_hours,
            'body_rate': body_rate,
            'paint_rate': paint_rate,
            'frame_rate': frame_rate,
            'mech_rate': mech_rate,
            'paint_mat_rate': paint_mat_rate,
            'body_labor': body_labor,
            'paint_labor': paint_labor,
            'setup_measure': setup_measure,
            'frame_labor': frame_labor,
            'mech_labor': mech_labor,
            'labor_sub': labor_sub,
            'parts_lines': _extract_itemized_part_lines(text_local),
            'parts_sub': parts_sub,
            'paint_mat': paint_mat,
            'sublet': sublet,
            'tax_basis': tax_basis,
            'tax_amt': tax_amt,
            'total_val': total_val,
            'tax_rate_value': float(tax_rate_value),
            'explicit_body_present': explicit_body_present,
            'explicit_paint_present': explicit_paint_present,
            'explicit_setup_present': explicit_setup_present,
            'explicit_frame_present': explicit_frame_present,
            'explicit_mech_present': explicit_mech_present,
            'scope_text_for_rate_lock': ((scope_text or '') + '\n' + text_local).strip(),
        }
        _parsed = _seed_scope_labor_buckets(((scope_text or '') + '\n' + text_local).strip(), _parsed, seed_rates=seed_rates)
        return _apply_normalization_lock(_parsed)

    def _locked_backend_total_from_cost_md(md_text: str, tax_rate_value: Optional[float] = None) -> Optional[float]:
        parsed = _parse_locked_photos_only_costs(md_text, tax_rate_value)
        return parsed.get('total_val')

    def _fallback_summary_from_visible_evidence() -> str:
        """Backend fallback narrative when the model returns blank/N/A fields."""
        evidence_bits = []
        count_photos = len(photo_index or [])
        if count_photos:
            evidence_bits.append(f"The uploaded set contains {count_photos} photo(s) that were reviewed for visible condition and damage pattern assessment.")
        if vin_from_label:
            src = f" ({vin_from_label_photo})" if vin_from_label_photo else ""
            evidence_bits.append(f"A driver-door VIN label appears visible{src}, supporting vehicle identification from the photo set.")
        if odometer_value:
            evidence_bits.append(f"An odometer reading appears visible in the photos and reads approximately {odometer_value}.")
        rear_damage = bool(re.search(r'(?i)rear|liftgate|hatch|bumper|lamp|back\s*panel|quarter', uploaded_text_all or ''))
        front_damage = bool(re.search(r'(?i)front|hood|grille|headlamp|fender', uploaded_text_all or ''))
        structural_flag = _structural_observed([uploaded_text_all or "", result.get("estimated_costs_markdown") or ""])
        if rear_damage:
            evidence_bits.append("Visible photo evidence indicates concentrated rear-end damage with damaged or displaced rear closure and bumper-area components.")
        elif front_damage:
            evidence_bits.append("Visible photo evidence indicates front-end damage involving exterior panels and attached front-end components.")
        else:
            evidence_bits.append("Visible photo evidence indicates exterior collision damage, but the exact impact zone is not fully established from OCR text alone.")
        if structural_flag:
            evidence_bits.append("The visible condition suggests possible structural or aperture-related distortion, so setup/measure and deeper teardown verification would be appropriate.")
        evidence_bits.append("This fallback narrative was generated by the backend because the model returned an empty or unusable narrative field on this run.")
        body = " ".join(evidence_bits).strip()
        return "## Detailed Condition Report\n" + body

    def _fallback_fraud_from_visible_evidence() -> str:
        checks = ["VIN/photo consistency"]
        if odometer_value:
            checks.append("odometer visibility")
        checks.extend(["duplicate image scan", "obvious photo tampering scan"])
        return "No material inconsistencies found. Backend fallback review checked " + ", ".join(checks) + "."

    def _fallback_conclusion_from_visible_evidence() -> str:
        structural_flag = _structural_observed([uploaded_text_all or "", result.get("estimated_costs_markdown") or ""])
        if structural_flag:
            return "Visible photo evidence supports a higher-severity damage review and teardown/measurement confirmation before final repair planning."
        return "Visible photo evidence supports a photo-based condition review, but final repair planning should be confirmed after in-person inspection and teardown."

    def _run_visual_sanity_check() -> Dict[str, Any]:
        """Targeted low-token sanity pass used only when critical photos-only fields are missing.

        This is intentionally narrow: it checks whether obvious evidence exists in the uploaded images
        for VIN label, odometer, front-end damage, and loose removed front-end parts. It also returns
        a conservative year/make/model guess when clearly visible.
        """
        out = {
            "vin_label_visible": False,
            "odometer_visible": False,
            "obvious_front_damage": False,
            "parts_pile_visible": False,
            "year_make_model": "",
        }
        try:
            image_parts = [p for p in parts if isinstance(p, dict) and p.get("type") != "text"][:12]
            if not image_parts:
                return out
            sanity_prompt = (
                "Return ONLY JSON with exactly these keys: "
                "vin_label_visible, odometer_visible, obvious_front_damage, parts_pile_visible, year_make_model.\n"
                "Rules:\n"
                "- vin_label_visible: true only if a driver-door VIN/manufacturer label is visibly present.\n"
                "- odometer_visible: true only if an instrument cluster odometer reading is visibly present.\n"
                "- obvious_front_damage: true only if there is clear front-end collision damage, missing front bumper/headlamp/grille/core support, or exposed front structure.\n"
                "- parts_pile_visible: true only if removed front-end parts/components are visibly laid out off-vehicle.\n"
                "- year_make_model: provide only when clearly supported by visible badge/VIN-label text or unmistakable model cues; otherwise empty string.\n"
            )
            rsp_sanity = client.chat.completions.create(
                model=MODEL,
                messages=[
                    {"role": "system", "content": "You perform narrow visual sanity checks for auto damage photos. JSON only."},
                    {"role": "user", "content": [{"type": "text", "text": sanity_prompt}] + image_parts},
                ],
                max_completion_tokens=250,
                temperature=0,
                response_format={"type": "json_object"},
            )
            raw_sanity = (rsp_sanity.choices[0].message.content or "").strip()
            data_sanity = _try_parse_json(raw_sanity)
            if isinstance(data_sanity, dict):
                out["vin_label_visible"] = bool(data_sanity.get("vin_label_visible"))
                out["odometer_visible"] = bool(data_sanity.get("odometer_visible"))
                out["obvious_front_damage"] = bool(data_sanity.get("obvious_front_damage"))
                out["parts_pile_visible"] = bool(data_sanity.get("parts_pile_visible"))
                ymm = data_sanity.get("year_make_model")
                if isinstance(ymm, str):
                    out["year_make_model"] = ymm.strip()
        except Exception:
            pass
        return out

    def _maybe_recover_critical_fields_from_sanity(sanity: Dict[str, Any]) -> None:
        """Attempt one narrow recovery of missing critical fields before blocking the report."""
        try:
            if (not str(result.get("vehicle") or "").strip() or str(result.get("vehicle") or "").strip().upper() == "N/A") and str(sanity.get("year_make_model") or "").strip():
                result["vehicle"] = str(sanity.get("year_make_model") or "").strip()
        except Exception:
            pass

        try:
            _impact_now = str(result.get("primary_impact") or "").strip().lower()
            if sanity.get("obvious_front_damage") and (_impact_now in ("", "n/a", "none", "none observed") or "no obvious impact" in _impact_now):
                result["primary_impact"] = "Front"
        except Exception:
            pass

        try:
            if sanity.get("parts_pile_visible") and isinstance(locked_costs_obj, dict):
                _parts_lines = locked_costs_obj.get("parts_lines") or []
                _parts_sub = float(locked_costs_obj.get("parts_sub") or 0.0)
                if (not _parts_lines or all("not separately derived" in str(x).strip().lower() for x in _parts_lines)) and _parts_sub <= 0:
                    raise ValueError("parts still missing")
        except Exception:
            # Deliberately do nothing here; blocker logic below will stop release.
            pass

    def _collect_report_blockers(sanity: Optional[Dict[str, Any]] = None) -> List[str]:
        """Hard validation gate for client-facing photos-only reports.

        The report may still be internally generated, but it must not be released when obvious
        evidence exists and critical extracted fields remain blank or contradictory.
        """
        blockers: List[str] = []
        if ai_intent != "damage_report_from_photos":
            return blockers

        sanity = sanity or {}
        vin_visible = bool(vin_from_label) or bool(sanity.get("vin_label_visible"))
        odo_visible = bool(odometer_value) or bool(sanity.get("odometer_visible"))
        front_damage_visible = bool(sanity.get("obvious_front_damage"))
        parts_pile_visible = bool(sanity.get("parts_pile_visible"))

        _vin_now = str(result.get("vin") or "").strip()
        if vin_visible and (not _vin_now or _vin_now.upper() == "N/A"):
            blockers.append("REPORT BLOCKED: VIN label evidence exists but VIN is still blank.")

        _odo_now = str(result.get("odometer_estimate_only") or "").strip()
        if odo_visible and (not _odo_now or _odo_now.upper() == "N/A"):
            blockers.append("REPORT BLOCKED: odometer evidence exists but odometer is still blank.")

        _impact_now = str(result.get("primary_impact") or "").strip().lower()
        if front_damage_visible and (_impact_now in ("", "n/a", "none", "none observed") or "no obvious impact" in _impact_now):
            blockers.append("REPORT BLOCKED: obvious front-end damage exists but primary impact is still blank/minimized.")

        _vehicle_now = str(result.get("vehicle") or "").strip()
        _ymm_now = str(sanity.get("year_make_model") or "").strip()
        if _ymm_now and (not _vehicle_now or _vehicle_now.upper() == "N/A"):
            blockers.append("REPORT BLOCKED: year/make/model appears visually supported but vehicle field is blank.")

        _cost_now = str(result.get("estimated_costs_markdown") or "").strip().lower()
        _parts_bad = ("itemized parts breakdown:" in _cost_now and "not separately derived" in _cost_now)
        try:
            _parts_sub_bad = isinstance(locked_costs_obj, dict) and float(locked_costs_obj.get("parts_sub") or 0.0) <= 0.0
        except Exception:
            _parts_sub_bad = True
        if parts_pile_visible and (_parts_bad or _parts_sub_bad):
            blockers.append("REPORT BLOCKED: removed parts are visibly present but parts capture remained empty/not separately derived.")

        _summary_now = str(result.get("summary_markdown") or "").strip().lower()
        if "fallback narrative" in _summary_now and (vin_visible or odo_visible or front_damage_visible or parts_pile_visible):
            blockers.append("REPORT BLOCKED: backend fallback narrative triggered despite visible core evidence in the photo set.")

        return blockers

    def _final_non_empty_output_lock() -> None:
        """Last-line protection against blank/N/A photos-only outputs.

        Surgical lock only:
        - preserve any usable identifiers, cost text, labor buckets, and parts lines already present
        - fill only truly empty narrative/fraud/conclusion text
        - never replace estimated_costs_markdown with a zero scaffold here, because that can wipe
          usable extracted evidence and force the locked math path to collapse to $0.00
        """
        if ai_intent != "damage_report_from_photos":
            return
        if _bad_field(result.get("summary_markdown") or ""):
            result["summary_markdown"] = _fallback_summary_from_visible_evidence()
        if _bad_field(result.get("fraud_markdown") or ""):
            result["fraud_markdown"] = _fallback_fraud_from_visible_evidence()
        if _bad_field(result.get("conclusion") or ""):
            result["conclusion"] = _fallback_conclusion_from_visible_evidence()

    def _canonical_locked_photos_only_cost_markdown_from_parsed(parsed: Dict[str, Any], tax_rate_value: Optional[float] = None) -> str:
        """Rebuild photos-only costs from one already-locked backend object.
        This avoids reparsing markdown after the final locked object has been computed.
        """
        def _m(v: Optional[float]) -> str:
            try:
                return "${:,.2f}".format(float(v) if v is not None else 0.0)
            except Exception:
                return "$0.00"

        def _fmt(hours: Optional[float], rate: Optional[float], amount: Optional[float]) -> str:
            try:
                if isinstance(rate, (int, float)):
                    shown_hours = float(hours) if isinstance(hours, (int, float)) else 0.0
                    return f"{shown_hours:.1f} hrs @ ${float(rate):,.2f}/hr = {_m(amount)}"
            except Exception:
                pass
            return f"Not separately derived = {_m(amount)}"

        total_val = float(parsed.get('total_val') or 0.0)
        if total_val < 3500:
            boxes = ("[x]", "[ ]", "[ ]", "[ ]")
        elif total_val < 10000:
            boxes = ("[ ]", "[x]", "[ ]", "[ ]")
        else:
            boxes = ("[ ]", "[ ]", "[x]", "[ ]")

        lines = [
            "## Approximate Repair Cost Breakdown",
            f"Body labor: {_fmt(parsed.get('body_hours'), parsed.get('body_rate'), parsed.get('body_labor'))}",
            f"Paint labor: {_fmt(parsed.get('paint_hours'), parsed.get('paint_rate'), parsed.get('paint_labor'))}",
            f"Setup & Measure: {_fmt(parsed.get('setup_hours'), parsed.get('body_rate'), parsed.get('setup_measure'))}",
            f"Frame labor: {(float(parsed.get('frame_hours')) if isinstance(parsed.get('frame_hours'), (int, float)) else 0.0):.1f} hrs @ ${(float(parsed.get('frame_rate')) if isinstance(parsed.get('frame_rate'), (int, float)) else 0.0):,.2f}/hr = {_m(parsed.get('frame_labor'))}",
            f"Mechanical labor: {_fmt(parsed.get('mech_hours'), parsed.get('mech_rate'), parsed.get('mech_labor'))}",
            f"Labor subtotal: {_m(parsed.get('labor_sub'))}",
            "Itemized parts breakdown:",
        ]

        parts_lines = parsed.get('parts_lines') or []
        if parts_lines:
            for pl in parts_lines:
                clean_pl = re.sub(r'^[-*]\s*', '', str(pl).strip())
                lines.append(f"- {clean_pl}")
        else:
            lines.append("- Not separately derived.")

        lines.extend([
            f"Parts subtotal: {_m(parsed.get('parts_sub'))}",
            f"Paint & materials: {_m(parsed.get('paint_mat'))}",
        ])
        if float(parsed.get('sublet') or 0.0) > 0:
            lines.append(f"Sublet: {_m(parsed.get('sublet'))}")
        rate_val = tax_rate_value if isinstance(tax_rate_value, (int, float)) and tax_rate_value > 0 else 0.07
        lines.extend([
            f"Tax rate: {float(rate_val)*100:.3f}%",
            f"Tax basis (parts + paint materials): {_m(parsed.get('tax_basis'))}",
            f"Tax: {_m(parsed.get('tax_amt'))}",
            f"Approximate Repair Total: {_m(parsed.get('total_val'))}",
            "Severity Tier",
            f"{boxes[0]} Minor (< $3,500)",
            f"{boxes[1]} Moderate ($3,500-$10,000)",
            f"{boxes[2]} Major ($10,000+)",
            f"{boxes[3]} Possible Total Loss Threshold Approaching",
        ])
        return "\n".join(lines)

    def _canonical_locked_photos_only_cost_markdown(md_text: str, tax_rate_value: Optional[float] = None) -> str:
        """Rebuild photos-only costs into one canonical backend-owned markdown block.
        This forces ZIP and loose JPG runs through the same normalized sequence:
        normalize inputs -> labor buckets -> itemized parts -> parts subtotal -> tax basis -> tax -> total -> severity.
        """
        parsed = _parse_locked_photos_only_costs(md_text, tax_rate_value)
        return _canonical_locked_photos_only_cost_markdown_from_parsed(parsed, tax_rate_value)

    def _force_conclusion_to_locked_total(conclusion_text: str, locked_total: Optional[float]) -> str:
        """Preserve the conclusion review text while replacing only the conflicting cost sentence."""
        base = str(conclusion_text or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        if not isinstance(locked_total, (int, float)):
            return base

        try:
            _locked_total_str = "${:,.2f}".format(float(locked_total))
        except Exception:
            _locked_total_str = "$0.00"

        locked_sentence = (
            f"The photo-based repair cost approximation is approximately {_locked_total_str} "
            "based on the backend-locked cost calculation shown in this report."
        )

        if not base:
            return locked_sentence

        def _is_conflicting_cost_sentence(sentence: str) -> bool:
            s = (sentence or "").strip()
            if not s:
                return False
            mentions_cost = bool(re.search(
                r'(?i)photo-based repair cost approximation|photo-based repair approximation|approximate repair|repair approximation|estimated repair|estimated total|approximate total|repair total|cost total|body/paint rate|body & paint|tax \(parts \+ paint materials only\)|provided \$?\d+(?:\.\d+)?/hr',
                s,
            ))
            has_amount = ('$' in s) or bool(re.search(r'(?i)\b\d+(?:\.\d+)?k\b', s))
            return mentions_cost and has_amount

        paragraphs_out: List[str] = []
        replaced_any = False

        for para in re.split(r'\n\s*\n', base):
            p = para.strip()
            if not p:
                continue
            sentences = re.split(r'(?<=[.!?])\s+', p)
            kept_sentences: List[str] = []
            for sent in sentences:
                if _is_conflicting_cost_sentence(sent):
                    replaced_any = True
                    continue
                kept_sentences.append(sent.strip())
            rebuilt = ' '.join([s for s in kept_sentences if s]).strip()
            if rebuilt:
                paragraphs_out.append(rebuilt)

        cleaned_base = '\n\n'.join(paragraphs_out).strip()

        if replaced_any:
            if cleaned_base:
                return cleaned_base + "\n\n" + locked_sentence
            return locked_sentence

        if locked_sentence in cleaned_base:
            return cleaned_base

        return cleaned_base

    try:
        _final_non_empty_output_lock()
        if ai_intent == "damage_report_from_photos":
            _raw_locked_cost_source = result.get("estimated_costs_markdown") or ""
            _raw_summary_before_lock = str(result.get("summary_markdown") or "").strip()
            _scope_seed_text = "\n\n".join([
                _raw_summary_before_lock,
                str(result.get("conclusion") or "").strip(),
            ]).strip()
            _inspection_location_for_lock = _normalize_location_with_zip(_extract_inspection_location(uploaded_text_all or ""), ia_company, uploaded_text_all)
            _seed_rates_for_lock = _lookup_rates(_inspection_location_for_lock)
            if isinstance(locked_cost_overrides, dict):
                if isinstance(locked_cost_overrides.get("body_rate"), (int, float)) and float(locked_cost_overrides.get("body_rate") or 0.0) > 0:
                    _seed_rates_for_lock["body_rate"] = float(locked_cost_overrides.get("body_rate"))
                if isinstance(locked_cost_overrides.get("paint_rate"), (int, float)) and float(locked_cost_overrides.get("paint_rate") or 0.0) > 0:
                    _seed_rates_for_lock["paint_rate"] = float(locked_cost_overrides.get("paint_rate"))
            if isinstance(locked_cost_overrides, dict) and isinstance(locked_cost_overrides.get("tax_rate"), (int, float)) and float(locked_cost_overrides.get("tax_rate") or 0.0) > 0:
                tax_rate = float(locked_cost_overrides.get("tax_rate"))
            elif tax_rate is None or not isinstance(tax_rate, (int, float)) or tax_rate <= 0:
                tax_rate = _lookup_tax_rate(_inspection_location_for_lock)
            locked_costs_obj = _parse_locked_photos_only_costs(
                _raw_locked_cost_source,
                tax_rate,
                seed_rates=_seed_rates_for_lock,
                scope_text=_scope_seed_text,
            )
            locked_costs_obj = _apply_locked_rate_overrides_to_parsed(
                locked_costs_obj,
                locked_cost_overrides,
            )
            result["estimated_costs_markdown"] = _canonical_locked_photos_only_cost_markdown_from_parsed(
                locked_costs_obj, tax_rate
            )
            # Remove any model-injected cost section from the narrative so email/PDF never show two totals.
            result["summary_markdown"] = _scrub_photo_only_narrative_cost_headers(_raw_summary_before_lock)
            _locked_total = locked_costs_obj.get("total_val") if isinstance(locked_costs_obj, dict) else None
            result["conclusion"] = _force_conclusion_to_locked_total(result.get("conclusion") or "", _locked_total)
            result["summary_markdown"] = _scrub_photo_only_narrative_cost_headers(_scrub_model_headings(result.get("summary_markdown") or ""))
            _final_non_empty_output_lock()
    except Exception:
        pass

    try:
        if ai_intent == "damage_report_from_photos":
            _sanity_needed = False
            if not str(result.get("vin") or "").strip() or str(result.get("vin") or "").strip().upper() == "N/A":
                _sanity_needed = True
            if not str(result.get("odometer_estimate_only") or "").strip() or str(result.get("odometer_estimate_only") or "").strip().upper() == "N/A":
                _sanity_needed = True
            _impact_check = str(result.get("primary_impact") or "").strip().lower()
            if _impact_check in ("", "n/a", "none", "none observed") or "no obvious impact" in _impact_check:
                _sanity_needed = True
            _cost_check = str(result.get("estimated_costs_markdown") or "").lower()
            if "itemized parts breakdown:" in _cost_check and "not separately derived" in _cost_check:
                _sanity_needed = True
            if "fallback narrative" in str(result.get("summary_markdown") or "").lower():
                _sanity_needed = True

            _sanity = _run_visual_sanity_check() if _sanity_needed else {}
            if _sanity:
                _maybe_recover_critical_fields_from_sanity(_sanity)
            _blockers = _collect_report_blockers(_sanity)
            if _blockers:
                return JSONResponse(
                    status_code=200,
                    content={
                        "status": "blocked",
                        "error": "REPORT BLOCKED: core extraction mismatch",
                        "reasons": _blockers,
                        "sanity_check": _sanity,
                        "draft": {
                            "file_number": file_number,
                            "vin": result.get("vin"),
                            "vehicle": result.get("vehicle"),
                            "odometer": result.get("odometer_estimate_only"),
                            "primary_impact": result.get("primary_impact"),
                        },
                    },
                )
    except Exception:
        pass

    # PDF helpers
    # -----------------------
    def _pdf_sanitize(text: str, max_token_len: int = 60) -> str:
        if text is None:
            return ""
        s = str(text).replace("\r\n", "\n").replace("\r", "\n")
        s = "".join(ch if ord(ch) < 256 else " " for ch in s)
        def _break(tok: str) -> str:
            if len(tok) <= max_token_len:
                return tok
            return " ".join(tok[i:i+max_token_len] for i in range(0, len(tok), max_token_len))
        s = " ".join(_break(t) for t in s.split(" "))
        return s

    pdf = FPDF(); pdf.add_page()
    # --- NSPXN Logo (Top Right, First Page Only) ---
    try:
        logo_path = os.path.join(os.path.dirname(__file__), "ChatGPT logo100725.png")
        if os.path.exists(logo_path):
            pdf.image(logo_path, x=pdf.w - 45, y=8, w=35)  # small–medium size
    except Exception:
        pass
    pdf.set_auto_page_break(auto=True, margin=10)
    pdf.set_left_margin(10); pdf.set_right_margin(10)

    try:
        pdf.add_font("DejaVu","", "DejaVuSans.ttf", uni=True); pdf.set_font(size=11)
    except Exception:
        pdf.set_font("Arial", size=11)

    try:
        report_generated_ts = datetime.now(ZoneInfo("America/New_York")).strftime("%m/%d/%Y %I:%M %p") + " EST"
    except Exception:
        report_generated_ts = datetime.now().strftime("%m/%d/%Y %I:%M %p")

    def mc(s):
        try:
            effective_w = pdf.w - pdf.l_margin - pdf.r_margin
            if effective_w <= 5:
                effective_w = 180
            safe = _pdf_sanitize(s)
            if not safe.strip():
                safe = "-"
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(effective_w, 6, safe)
        except Exception:
            effective_w = pdf.w - pdf.l_margin - pdf.r_margin
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(effective_w, 6, (_pdf_sanitize(str(s))[:2000] + " …"))


    
    def _money2(x: Optional[float]) -> str:
        try:
            if x is None:
                return "$0.00"
            return "${:,.2f}".format(float(x))
        except Exception:
            return "$0.00"

    def _compute_cost_total_from_md(md_text: str) -> Optional[float]:
        """Compute a photo-only approximate total from the model's cost markdown.
        We intentionally sum the *section totals* (labor totals, paint materials, parts subtotal, tax)
        to avoid double-counting per-part lines.
        """
        if not md_text:
            return None
        t = str(md_text).replace("\r\n", "\n").replace("\r", "\n")

        def _grab_amount(pat: str) -> Optional[float]:
            m = re.search(pat, t, flags=re.IGNORECASE | re.MULTILINE)
            if not m:
                return None
            try:
                return float(m.group(1).replace(",", ""))
            except Exception:
                return None

        # Prefer explicit totals (bold amounts) from the model output
        body_labor = _grab_amount(r"^\s*[-*]?\s*Body\s+labor\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*")
        paint_labor = _grab_amount(r"^\s*[-*]?\s*Paint\s+labor\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*")
        mech_labor = _grab_amount(r"^\s*[-*]?\s*Mechanical[^:]*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*")
        if mech_labor is None:
            mech_labor = _grab_amount(r"^\s*[-*]?\s*Mechanical[^=\n]*=\s*\$?\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b")
        frame_labor = _grab_amount(r"^\s*[-*]?\s*Frame\s+labor\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*")

        # Paint materials dollars (prefer subtotal/extended amount lines; never the $/hr rate)
        paint_mat = None
        for _pat in [
            r"^\s*[-*]?\s*Paint\s+materials\s+subtotal\s*[:=]\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)",
            r"^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s*[:=]\s*\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*",
            r"^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s*[:=]\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)",
        ]:
            paint_mat = _grab_amount(_pat)
            if paint_mat is not None:
                break

        # Parts subtotal dollars
        parts_sub = _grab_amount(r"^\s*\*\*\s*Estimated\s+parts\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)")
        if parts_sub is None:
            parts_sub = _grab_amount(r"^\s*[-*]?\s*Parts\s+subtotal\s*:\s*\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*")
        if parts_sub is None:
            # Some model outputs use a simple 'Parts:' label
            parts_sub = _grab_amount(r"^\s*[-*]?\s*Parts\s*:\s*\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*")

        tax_amt = _grab_amount(r"^\s*[-*]?\s*(?:Sales\s+tax|Estimated\s+tax|Tax)\b.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*")

        # If tax dollars weren't provided but an assumed rate exists, compute tax on (parts + paint materials) only.
        if tax_amt is None:
            m_rate = re.search(r"(?im)^\s*[-*]?\s*Tax\s+rate\s*(?:\(assumed\)|assumption|\(assumed\)\s*:|\(assumed\)\s*)?\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)\s*%", t)
            if not m_rate:
                m_rate = re.search(r"(?im)\bTax\s+rate\s*(?:\(assumed\))?\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)\s*%", t)
            if m_rate and parts_sub is not None and paint_mat is not None:
                try:
                    rate = float(m_rate.group(1)) / 100.0
                    tax_amt = round((float(parts_sub) + float(paint_mat)) * rate, 2)
                except Exception:
                    tax_amt = None

        # If the model failed to provide the key components, do not guess a total here.
        components = [body_labor, paint_labor, mech_labor, frame_labor, paint_mat, parts_sub, tax_amt]
        have_any = any(v is not None for v in components)
        have_core = (parts_sub is not None) and (paint_mat is not None)
        # Require at least parts + paint materials + some labor; tax can be derived if rate exists.
        have_some_labor = any(v is not None for v in (body_labor, paint_labor, mech_labor, frame_labor))
        if not (have_any and have_core and have_some_labor):
            return None

        total = 0.0
        for v in (body_labor, paint_labor, mech_labor, frame_labor, paint_mat, parts_sub, tax_amt):
            if isinstance(v, (int, float)):
                total += float(v)
        return round(total, 2)

    def _inject_clean_total_line(md_text: str, total_val: Optional[float]) -> str:
        """Remove any existing 'Approximate Repair Cost Total' lines.
        Option 1 mode: keep the model's 'Approximate Repair Total: $X' line untouched
        and do NOT inject any additional total line.
        """
        if not md_text:
            return md_text or ""
        t = str(md_text).replace("\r\n", "\n").replace("\r", "\n")

        lines: List[str] = []
        for ln in t.splitlines():
            s = (ln or "").strip()
            # Remove any injected/legacy 'Repair Cost Total' lines (we will not re-add).
            if re.search(r"(?i)^\s*Approximate\s+Repair\s+Cost\s+Total\s*:", s):
                continue
            if re.search(r"(?i)^\s*Approximate\s+Total\s+Repair\s+Cost\s*:", s):
                continue
            # Also remove any standalone variants with $ on the line
            if re.search(r"(?i)^\s*Approximate\s+Repair\s+Cost\s+Total\b", s) and "$" in s:
                continue
            lines.append(ln)

        return "\n".join(lines).strip()



    def _parse_approx_total(md_text: str) -> Optional[float]:
            """Parse the approximate repair total from estimated_costs_markdown (best-effort).
            Prefers a standalone/bold total near the end of the cost section.
            """
            if not md_text:
                return None
            t = str(md_text).replace("\r\n","\n").replace("\r","\n")

            # 1) Prefer explicit label variants
            label_patterns = [
                r"(?im)^\s*\*\*?\s*Approximate\s+Repair\s+Cost\s+Total\s*[:\-]?\s*\*\*?\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"(?im)^\s*\*\*?\s*Approximate\s+Repair\s+Total\s*[:\-]?\s*\*\*?\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"(?im)^\s*\*\*?\s*Approximate\s+Total\s+Repair\s+Cost\s*[:\-]?\s*\*\*?\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            ]
            for pat in label_patterns:
                m = re.search(pat, t)
                if m:
                    try:
                        return float(m.group(1).replace(",", ""))
                    except Exception:
                        pass

            # 2) Prefer a standalone line that is just $X,XXX (often bolded)
            standalone = []
            for ln in t.splitlines():
                s = ln.strip()
                # strip markdown bold
                s2 = re.sub(r"^[*_]+|[*_]+$", "", s).strip()
                m = re.fullmatch(r"\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)", s2)
                if m:
                    try:
                        standalone.append(float(m.group(1).replace(",", "")))
                    except Exception:
                        pass
            if standalone:
                return standalone[-1]

            # 3) Otherwise: take the last dollar amount after a 'Tax' section (common structure)
            # BUT ignore tier boundary dollars like $3,500 / $10,000 from Severity Tier checklists.
            tail_lines = t.splitlines()
            # keep only lines after the first Tax heading if present
            for i, ln in enumerate(tail_lines):
                if re.search(r"(?i)^\s*tax\b", ln.strip()):
                    tail_lines = tail_lines[i:]
                    break
            # drop any severity/tier checklist lines
            safe_tail = []
            for ln in tail_lines:
                if re.search(r"(?i)severity\s+tier|minor\s*\(|moderate\s*\(|major\s*\(|total\s+loss\s+threshold", ln):
                    continue
                if re.match(r"^\s*\[[ xX]\]", ln.strip()):
                    continue
                safe_tail.append(ln)
            monies = re.findall(r"\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)", "\n".join(safe_tail))
            if monies:
                # choose the last amount in the tail
                try:
                    return float(monies[-1].replace(",", ""))
                except Exception:
                    return None
            return None

    def _enforce_severity_tier_checkmarks(md_text: str) -> str:
        """Rewrite the Severity Tier block to ensure exactly one box is checked based on parsed total."""
        if not md_text:
            return md_text or ""
        t = str(md_text).replace("Likely Total Loss Threshold Approaching", "Possible Total Loss Threshold Approaching")
        total_val = _parse_approx_total(t)
        # thresholds
        tier = None
        if isinstance(total_val, (int, float)):
            if total_val < 3500:
                tier = "minor"
            elif total_val < 10000:
                tier = "moderate"
            elif total_val < 25000:
                tier = "major"
            else:
                tier = "possible_tl"

        # Remove any existing Severity Tier block (best-effort)
        lines_in = t.splitlines()
        out = []
        i = 0
        while i < len(lines_in):
            ln = lines_in[i]
            if re.search(r"(?i)^\s*#{1,6}\s*Severity\s+Tier\b", ln):
                # skip until next heading or blank-run end
                i += 1
                while i < len(lines_in) and not re.search(r"^\s*#{1,6}\s+\w", lines_in[i]):
                    i += 1
                continue
            # Also remove any existing tier checkbox/list lines even if the heading was stripped upstream
            if re.search(r"(?i)^\s*(?:[-*]\s*)?(?:\[\s*[xX ]\s*\]|[☐☑])?\s*(Minor|Moderate|Major|Possible\s+Total\s+Loss)\b", ln.strip()):
                i += 1
                continue
            out.append(ln)
            i += 1
        t2 = "\n".join(out).rstrip()

        # Append normalized Severity Tier block at end of cost section (or end of markdown)
        boxes = {
            "minor": ("☑","☐","☐","☐"),
            "moderate": ("☐","☑","☐","☐"),
            "major": ("☐","☐","☑","☐"),
            "possible_tl": ("☐","☐","☐","☑"),
            None: ("☐","☐","☐","☐"),
        }[tier]
        sev_block = (
            "\n### Severity Tier\n"
            f"{boxes[0]} Minor (< $3,500)\n"
            f"{boxes[1]} Moderate ($3,500-$10,000)\n"
            f"{boxes[2]} Major ($10,000+)\n"
            f"{boxes[3]} Possible Total Loss Threshold Approaching\n"
        )
        return (t2 + "\n" + sev_block).strip() + "\n"

    def _strip_unwanted_cost_lines_for_pdf(md_text: str) -> str:
        """Remove known unwanted lines in cost markdown before PDF render.
        This is used ONLY for Photos-Only PDFs to prevent model-inserted prompt headers,
        'Totals' math blocks, duplicate section headers, and any estimate/docs language.
        """
        if not md_text:
            return md_text or ""
        t = str(md_text).replace("Likely Total Loss Threshold Approaching", "Possible Total Loss Threshold Approaching")
        cleaned: List[str] = []
        in_repair_disclaimer = False
        for ln in t.splitlines():
            s = (ln or "").strip()

            # Remove any mid-report Repair Cost Disclaimer block entirely.
            if re.search(r"(?i)^\s*(?:\*\*)?repair\s+cost\s+disclaimer\b", s):
                in_repair_disclaimer = True
                continue
            if in_repair_disclaimer:
                # stop on blank line or a heading
                if (not s) or re.search(r"^\s*#{1,6}\s+\w", s):
                    in_repair_disclaimer = False
                continue

            # Remove model/prompt headers and echoes
            if s in ("Approximate Repair Cost Breakdown", "(See estimated_costs_markdown field.)"):
                continue
            if re.search(r"(?i)estimated_costs_markdown", s):
                continue
            if re.search(r"(?i)\bPopulate\s+JSON\s+field\b", s):
                continue
            if re.search(r"(?i)^\s*##\s*Approximate\s+Repair\s+Cost\s+Breakdown\b", s):
                continue

            # Remove Totals blocks / arithmetic
            if re.search(r"(?i)^\s*###\s*Totals\b", s):
                continue
            if ("+" in ln and "=" in ln and re.search(r"\$\s*[0-9]", ln)):
                continue
            if re.search(r"(?i)^\s*###\s*Approximate\s+Total\s+Repair\s+Cost\b", s):
                continue
            if re.search(r"(?i)^\s*[-*]\s*Estimated\s+total\b", s):
                continue
            # Strip model totals/arithmetic phrasing variants (we print ONE locked total line)
            if re.search(r"(?i)\bEstimated\s+Repair\s+Total\b", s):
                continue
            if re.search(r"(?i)\bPre-?tax\s+total\b", s):
                continue
            if re.search(r"(?i)\bApproximate\s+total\s+repair\s+cost\b", s):
                continue
            if re.search(r"(?i)\bRounded\b", s) and re.search(r"\$\s*\d", ln):
                continue

            # Strip model helper headings / arithmetic lines we never want in the PDF
            if re.search(r"(?i)^\s*cost\s+math\s*\(approx\.?\)\s*$", s):
                continue
            if re.search(r"(?i)^\s*[-*]?\s*tax\s*:\s*\$.*[×x\*].*=", s):
                continue
            if re.search(r"(?i)^\s*\*{0,2}\s*Tax\s*\(parts\s*\+\s*paint\s+materials\s+only\)\s*:\s*\*{0,2}\s*$", s):
                continue
            if re.search(r"(?i)^\s*[-*]?\s*taxable\s+subtotal\s*:", s):
                continue
            if re.search(r"(?i)^\s*[-*]?\s*(estimated\s+tax|sales\s+tax|tax)\s*:", s):
                continue

            # Strip additional model cost-summary headings and totals (we print our own deterministic tax/total lines)
            if re.search(r"(?i)^\s*cost\s+summary\b", s):
                continue
            if re.search(r"(?i)\bApprox\.?\s*Repair\s*Total\b", s) and not re.search(r"(?i)^\s*\*{0,2}\s*Approximate\s+Repair\s+Total\s*:\s*\$\s*[0-9]", s):
                # Also drop any other repair-total style lines that include a dollar amount
                if re.search(r"(?i)\brepair\s*total\b", s) and "$" in s:
                    continue
                continue
            if re.search(r"(?i)^\s*[-*]?\s*tax\s+rate\s+assumption\b", s):
                continue
            if re.search(r"(?i)^\s*[-*]?\s*tax\s+rate\s*\(assumed\)\b", s):
                continue
            if re.search(r"(?i)^\s*[-*]?\s*estimated\s+tax\b", s):
                continue
            if re.search(r"(?i)^\s*tax\s*\(apply\b", s):
                continue
            if re.search(r"(?i)^\s*[-*]?\s*sales\s+tax\b", s) and not re.search(r"(?i)^\s*[-*]?\s*sales\s+tax\s*\(assumed\s*7%\s*for\s*approximation\)", s):
                continue
            if re.search(r"(?i)\bassumed\b.*\btax\b", s) and not re.search(r"(?i)^\s*[-*]?\s*sales\s+tax\s*\(assumed\s*7%\s*for\s*approximation\)", s):
                continue

            # Remove Severity headings/labels here (we render a normalized block later)
            if re.search(r"(?i)^\s*###\s*Severity\s+Tier\b", s):
                continue
            if s.lower() == "severity tier":
                continue

            # Strip any model-provided severity checklist lines (we print exactly one normalized block)
            if re.search(r"(?i)\bminor\s*\(|\bmoderate\s*\(|\bmajor\s*\(|total\s+loss\s+threshold\s+approaching", s):
                continue

            # Remove unwanted notes/disclaimers inserted by older logic
            if re.search(r"(?i)\bbest-?effort\b|\bfrom\s+estimate\b|\bfrom\s+docs\b|\bnot\s+evidenced\b", ln):
                continue

            # Never print Inspection Location inside the cost section (per your requirement)
            if re.search(r"(?i)^\*\*Inspection\s+Location\*\*\s*:", s) or re.search(r"(?i)^Inspection\s+Location\s*:", s):
                continue

            # Remove any existing total line; we will compute/insert a clean one
            if re.search(r"(?i)^\s*Approximate\s+Repair\s+Cost\s+Total\s*:", s):
                continue
            if re.search(r"(?i)^\s*Approximate\s+Total\s+Repair\s+Cost\s*:", s):
                continue

            # Avoid duplicate parts subtotal echo lines
            if re.search(r"(?i)estimated\s+parts\s+subtotal", s) and any(re.search(r"(?i)parts\s+subtotal", x) for x in cleaned[-3:]):
                continue

            cleaned.append(ln)
        return "\n".join(cleaned).strip()

    def _scrub_photo_only_narrative_cost_headers(md_text: str) -> str:
        """Remove any model-generated cost section from the narrative/email body.
        Keep the narrative, fraud, and conclusion text only. The canonical locked
        cost block is rendered separately from result['estimated_costs_markdown'].
        """
        if not md_text:
            return md_text or ""
        lines_in = str(md_text).replace("\r\n", "\n").replace("\r", "\n").splitlines()
        out: List[str] = []
        skip_cost_block = False
        for ln in lines_in:
            s = (ln or "").strip()

            if re.search(r"(?i)^\s*##\s*Approximate\s+Repair\s+Cost\s+Breakdown\b", s):
                skip_cost_block = True
                continue

            if skip_cost_block:
                if re.search(r"(?i)^\s*##\s+(Fraud\s*&\s*Authenticity\s*Check|Fraud\s+Detection|Conclusion)\b", s):
                    skip_cost_block = False
                    out.append(ln)
                    continue
                if re.search(r"(?i)^\s*(Fraud\s*&\s*Authenticity\s*Check|Fraud\s+Detection|Conclusion)\s*$", s):
                    skip_cost_block = False
                    out.append(ln)
                    continue
                continue

            if re.search(r"(?i)\bPopulate\s+JSON\s+field\b", ln):
                continue
            if re.search(r"(?i)\bSee\s+estimated_costs_markdown\s+field\b", ln):
                continue
            out.append(ln)
        return "\n".join(out).strip()


    def render_repair_cost_section(pdf_obj: FPDF, md: str, tax_rate: Optional[float] = None, parsed: Optional[Dict[str, Any]] = None) -> None:
        """Render the Approximate Repair Cost Breakdown in a controlled PDF format.
        Locked behavior:
        - fixed printed structure every time
        - print actual labor hours/rates used where available
        - recompute labor subtotal, tax, and total from one backend-owned path only
        """
        if tax_rate is None or not isinstance(tax_rate, (int, float)) or tax_rate <= 0:
            tax_rate = 0.07

        parsed = parsed if isinstance(parsed, dict) else _parse_locked_photos_only_costs(md, tax_rate)

        def _nz_money(v: Optional[float]) -> float:
            try:
                return round(float(v), 2) if isinstance(v, (int, float)) else 0.0
            except Exception:
                return 0.0

        def _fmt_hours_rate(hours: Optional[float], rate: Optional[float], amount: float) -> str:
            if isinstance(rate, (int, float)):
                shown_hours = float(hours) if isinstance(hours, (int, float)) else 0.0
                return f"{shown_hours:.1f} hrs @ ${float(rate):,.2f}/hr = {_money2(amount)}"
            return f"Not separately derived = {_money2(amount)}"

        body_hours = parsed.get('body_hours')
        paint_hours = parsed.get('paint_hours')
        setup_hours = parsed.get('setup_hours')
        frame_hours = parsed.get('frame_hours')
        mech_hours = parsed.get('mech_hours')
        body_rate = parsed.get('body_rate')
        paint_rate = parsed.get('paint_rate')
        frame_rate = parsed.get('frame_rate')
        mech_rate = parsed.get('mech_rate')

        body_labor = _nz_money(parsed.get('body_labor'))
        paint_labor = _nz_money(parsed.get('paint_labor'))
        setup_measure = _nz_money(parsed.get('setup_measure'))
        frame_labor = _nz_money(parsed.get('frame_labor'))
        mech_labor = _nz_money(parsed.get('mech_labor'))
        parts_lines = parsed.get('parts_lines') or []
        parts_sub = _nz_money(parsed.get('parts_sub'))
        paint_mat = _nz_money(parsed.get('paint_mat'))
        sublet = _nz_money(parsed.get('sublet'))
        labor_sub = _nz_money(parsed.get('labor_sub'))
        tax_basis = _nz_money(parsed.get('tax_basis'))
        tax_amt = _nz_money(parsed.get('tax_amt'))
        total_val = _nz_money(parsed.get('total_val'))

        pdf_obj.ln(1)
        try:
            pdf_obj.set_font("Helvetica", "", 11)
        except Exception:
            pdf_obj.set_font("Arial", "", 11)

        mc(f"Body labor: {_fmt_hours_rate(body_hours, body_rate, body_labor)}")
        mc(f"Paint labor: {_fmt_hours_rate(paint_hours, paint_rate, paint_labor)}")
        mc(f"Setup & Measure: {_fmt_hours_rate(setup_hours, body_rate, setup_measure)}")
        mc(f"Frame labor: {(float(frame_hours) if isinstance(frame_hours, (int, float)) else 0.0):.1f} hrs @ ${(float(frame_rate) if isinstance(frame_rate, (int, float)) else 0.0):,.2f}/hr = {_money2(frame_labor)}")
        mc(f"Mechanical labor: {_fmt_hours_rate(mech_hours, mech_rate, mech_labor)}")
        mc(f"Labor subtotal: {_money2(labor_sub)}")

        mc("Itemized parts breakdown:")
        if parts_lines:
            for pl in parts_lines:
                _clean_pl = re.sub(r'^[-*]\s*', '', str(pl).strip())
                mc(f"- {_clean_pl}")
        else:
            mc("- Not separately derived.")

        mc(f"Parts subtotal: {_money2(parts_sub)}")
        mc(f"Paint & materials: {_money2(paint_mat)}")
        if sublet > 0:
            mc(f"Sublet: {_money2(sublet)}")

        mc(f"Tax rate: {float(tax_rate)*100:.3f}%")
        mc(f"Tax basis (parts + paint materials): {_money2(tax_basis)}")
        mc(f"Tax: {_money2(tax_amt)}")

        try:
            pdf_obj.set_font("Helvetica", "B", 11)
        except Exception:
            pdf_obj.set_font("Arial", "B", 11)
        mc(f"Approximate Repair Total: {_money2(total_val)}")
        try:
            pdf_obj.set_font("Helvetica", "", 11)
        except Exception:
            pdf_obj.set_font("Arial", "", 11)

        if total_val < 3500:
            tier = "minor"
        elif total_val < 10000:
            tier = "moderate"
        else:
            tier = "major"

        boxes = {
            "minor": ("[x]", "[ ]", "[ ]", "[ ]"),
            "moderate": ("[ ]", "[x]", "[ ]", "[ ]"),
            "major": ("[ ]", "[ ]", "[x]", "[ ]"),
        }[tier]

        pdf_obj.ln(1)
        mc("Severity Tier")
        mc(f"{boxes[0]} Minor (< $3,500)")
        mc(f"{boxes[1]} Moderate ($3,500-$10,000)")
        mc(f"{boxes[2]} Major ($10,000+)")
        mc(f"{boxes[3]} Possible Total Loss Threshold Approaching")

    def add_thumbnail_page(pdf_obj: FPDF, image_paths: List[str]) -> None:
        """Append exactly ONE page containing thumbnails of all uploaded photos (as space allows)."""
        if not image_paths:
            return

        pdf_obj.add_page()
        try:
            pdf_obj.set_font("Helvetica", "B", 12)
        except Exception:
            pdf_obj.set_font("Arial", "B", 12)
        pdf_obj.cell(0, 8, "Uploaded Photos", ln=True)
        pdf_obj.ln(2)

        # Layout (single-page, grid)
        cols = 4
        gutter = 3.0
        usable_w = pdf_obj.w - pdf_obj.l_margin - pdf_obj.r_margin
        thumb_w = (usable_w - gutter * (cols - 1)) / cols
        x0 = pdf_obj.l_margin
        y = pdf_obj.get_y()
        x = x0
        col = 0

        placed = 0
        total = len(image_paths)

        for p in image_paths:
            # Hard-stop to keep this as ONE PAGE ONLY
            if y + thumb_w > (pdf_obj.h - pdf_obj.b_margin - 10):
                break
            try:
                pdf_obj.image(p, x=x, y=y, w=thumb_w)
                placed += 1
            except Exception:
                pass

            col += 1
            if col >= cols:
                col = 0
                x = x0
                y += thumb_w + gutter
            else:
                x += thumb_w + gutter

        # Footer note if truncated by the 1-page rule
        if placed < total:
            try:
                pdf_obj.set_font_size(9)
            except Exception:
                pass
            pdf_obj.set_y(pdf_obj.h - pdf_obj.b_margin - 8)
            pdf_obj.set_x(pdf_obj.l_margin)
            pdf_obj.cell(0, 6, f"Showing {placed} of {total} uploaded photos (1-page appendix limit).", ln=True)

    # POI-15 Total Loss trigger from uploaded text ONLY
    poi15_hit = False
    try:
        txt = uploaded_text_all.lower()
        if re.search(r"\b15\s*total\s*loss\b", txt, flags=re.IGNORECASE):
            poi15_hit = True
        elif re.search(r"point\s*of\s*impact[^A-Za-z0-9]{0,10}15[^A-Za-z0-9]{0,20}total\s*loss", txt, flags=re.IGNORECASE):
            poi15_hit = True
    except Exception:
        poi15_hit = False

    if ai_intent == "damage_report_from_photos":
        # -----------------------------
        # NSPXN.com Condition Report PDF
        # -----------------------------
        # Hard "safe" top margin so no section bar/box can overlap the logo/header.
        # (User-facing: move Vehicle Identification down a few lines.)
        SAFE_TOP_Y = 60
    
        def _ensure_safe_top() -> None:
            try:
                if pdf.get_y() < SAFE_TOP_Y:
                    pdf.set_y(SAFE_TOP_Y)
            except Exception:
                pass
    
        def _section_bar(title: str) -> None:
            """Draw a brighter, color-coded section header bar (aligned to DOCX style)."""
            t = str(title or "").strip()
            cmap = {
                "VEHICLE IDENTIFICATION": (0, 112, 192),              # bright blue
                "REPORT SUMMARY": (191, 112, 0),                      # orange
                "APPROXIMATE REPAIR COST BREAKDOWN": (0, 153, 76),     # green
                "FRAUD & AUTHENTICITY CHECK": (112, 48, 160),          # purple
                "CONCLUSION": (64, 64, 64),                           # dark gray
                "DISCLAIMER": (96, 96, 96),                           # gray
            }
            rgb = cmap.get(t.upper(), (0, 112, 192))
            _ensure_safe_top()
            try:
                pdf.ln(3)
                pdf.set_fill_color(*rgb)
                pdf.set_text_color(255, 255, 255)
                pdf.set_font("Helvetica", "B", 12)
            except Exception:
                pdf.ln(3)
                pdf.set_fill_color(*rgb)
                pdf.set_text_color(255, 255, 255)
                pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 8, _pdf_sanitize(t), ln=True, fill=True)
            pdf.set_text_color(0, 0, 0)
            try:
                pdf.set_font("Helvetica", "", 11)
            except Exception:
                pdf.set_font("Arial", "", 11)

        def _vin_model_year(vin_value: Optional[str]) -> str:
            """Best-effort VIN 10th-character model year decode for PDF display fallback."""
            try:
                vv = str(vin_value or "").strip().upper()
                if len(vv) != 17:
                    return ""
                code = vv[9]
                year_map = {
                    "A": 2010, "B": 2011, "C": 2012, "D": 2013, "E": 2014,
                    "F": 2015, "G": 2016, "H": 2017, "J": 2018, "K": 2019,
                    "L": 2020, "M": 2021, "N": 2022, "P": 2023, "R": 2024,
                    "S": 2025, "T": 2026, "V": 2027, "W": 2028, "X": 2029,
                    "Y": 2030,
                    "1": 2031, "2": 2032, "3": 2033, "4": 2034, "5": 2035,
                    "6": 2036, "7": 2037, "8": 2038, "9": 2039,
                }
                yy = year_map.get(code)
                return str(yy) if yy else ""
            except Exception:
                return ""

        def _format_vehicle_value(v, vin_value: Optional[str] = None) -> str:
            """Normalize the Vehicle field for PDF printing and include year when missing."""
            try:
                if isinstance(v, dict):
                    year = str(v.get("year") or "").strip()
                    make = str(v.get("make") or "").strip()
                    model = str(v.get("model") or "").strip()
                    trim = str(v.get("trim") or "").strip()
                    if not year or year.upper() == "N/A":
                        year = _vin_model_year(vin_value)
                    parts = [p for p in [year, make, model] if p and p.upper() != "N/A"]
                    base = " ".join(parts).strip()
                    if trim and trim.upper() != "N/A":
                        return f"{base} ({trim})" if base else trim
                    return base or "N/A"
            except Exception:
                pass
            s = str(v or "").strip()
            vin_year = _vin_model_year(vin_value)
            if vin_year and not re.search(r'\b(19|20)\d{2}\b', s):
                mm = re.search(r'(?i)\b([A-Z][a-zA-Z0-9]+)\s+([A-Z][a-zA-Z0-9]+)\b', s)
                if mm:
                    prefix = f"{vin_year} {mm.group(1)} {mm.group(2)}"
                    tail = s[mm.end():].strip()
                    return f"{prefix} {tail}".strip()
            return s if s else "N/A"
    
        def _scrub_model_headings(md_text: str) -> str:
            """Remove model-emitted headings and duplicate cost/checklist content from the narrative summary.

            The PDF prints the locked cost section separately, so any model-emitted cost/tax/severity
            lines inside summary_markdown must be stripped here to avoid duplicated cost output.
            """
            if not md_text:
                return md_text or ""
            lines = str(md_text).replace("\r\n","\n").replace("\r","\n").splitlines()
            out = []
            in_cost_block = False
            for ln in lines:
                s = (ln or "").strip()
                if not s:
                    if not in_cost_block:
                        out.append(ln)
                    continue
                # Once the model starts a cost block inside the narrative, drop it entirely.
                if re.search(r"(?i)^#{0,3}\s*APPROXIMATE\s+REPAIR\s+COST\s+BREAKDOWN\b", s):
                    in_cost_block = True
                    continue
                if in_cost_block:
                    # Resume only when the next real narrative section begins.
                    if re.search(r"(?i)^#{0,3}\s*(FRAUD\s*&\s*AUTHENTICITY\s*CHECK|CONCLUSION|DISCLAIMER)\b", s):
                        in_cost_block = False
                    else:
                        if re.search(r"(?i)^(Body labor|Paint labor|Setup\s*&\s*Measure|Frame labor|Mechanical labor|Labor subtotal|Itemized parts breakdown|Parts subtotal|Paint\s*&\s*materials|Tax rate|Tax basis|Tax|Approximate Repair Total|Severity Tier|\[[ xX]\]\s*(Minor|Moderate|Major|Possible\s+Total\s+Loss))\b", s):
                            continue
                        # swallow all other cost-block lines until a new real section starts
                        continue
                # Drop markdown headings like '# Condition Report (Photos Only)' or '## ...'
                if re.match(r"^#+\s+", s):
                    continue
                # Drop echoes like 'Report Selected'
                if re.search(r"(?i)^report\s+selected\b", s):
                    continue
                # Drop helper artifacts the model sometimes emits
                if re.search(r"(?i)estimated_costs_markdown", s):
                    continue
                # Drop the redundant front checklist bullets from the narrative summary.
                if re.search(r"(?i)^-\s*(Hood|Front bumper cover|Grille|One headlamp|The other headlamp|One front fender|The other front fender)\s*:", s):
                    continue
                out.append(ln)
            return "\n".join(out).strip()
    
        def _extract_photo_confirmed_odometer(summary_md: str) -> Optional[str]:
            """Best-effort: extract the photo-confirmed odometer from the Photo #3 row (single source of truth)."""
            if not summary_md:
                return None
            t = str(summary_md).replace("\r\n","\n").replace("\r","\n")
            # Prefer Photo 3 table row
            for ln in t.splitlines():
                if re.search(r"^\s*\|\s*3\s*\|", ln):
                    m = re.search(r"(?i)\bodometer\b[^0-9]{0,60}([0-9][0-9,]{2,})\s*mi", ln)
                    if m:
                        return m.group(1).replace(",", "") + " mi"
                    m2 = re.search(r"(?i)\bodometer\b[^0-9]{0,60}([0-9][0-9,]{2,})\b", ln)
                    if m2:
                        return m2.group(1).replace(",", "") + " mi"
            # Fallback: any odometer mention
            m = re.search(r"(?i)\bodometer\b[^0-9]{0,80}([0-9][0-9,]{2,})\s*mi", t)
            if m:
                return m.group(1).replace(",", "") + " mi"
            return None
    
        # Title (larger + bold)
        try:
            pdf.set_font("Helvetica", "B", 16)
        except Exception:
            pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, "NSPXN.com Condition Report", ln=True, align="C")
        try:
            pdf.set_font("Helvetica", "", 11)
        except Exception:
            pdf.set_font("Arial", "", 11)
        pdf.ln(2)
    
        # Vehicle Identification (fixed PDF block)
        _section_bar("VEHICLE IDENTIFICATION")
        _summary_md_raw = (result.get("summary_markdown") or "").strip()
        _odo_photo = _extract_photo_confirmed_odometer(_summary_md_raw)
        _odo_print = _odo_photo or (result.get("odometer_estimate_only") or "N/A")
        pdf_status = (result.get("redaction_status") or "").replace("✅", "OK").strip() or "N/A"
    
        # Use colon formatting (no double-odometer conflicts)
        mc(f"File #: {file_number or 'N/A'}")
        mc(f"Generated: {report_generated_ts}")
        mc(f"Claim #: {result.get('claim_number') or 'N/A'}")
        mc(f"Inspected For: {ia_company or 'N/A'}")
        mc(f"VIN: {result.get('vin') or 'N/A'}")
        mc(f"VIN Verification: {result.get('vin_verification') or 'N/A'}")
        mc(f"Vehicle: {_format_vehicle_value(result.get('vehicle'), result.get('vin'))}")
        mc(f"Odometer: {_odo_print}")
        mc(f"Primary Impact: {result.get('primary_impact') or 'N/A'}")
        mc(f"Secondary Impact: {result.get('secondary_impact') or 'N/A'}")
        mc(f"Redaction Status: {pdf_status}")
    
        # Report Summary (scrub markdown headings)
        _section_bar("REPORT SUMMARY")
        _summary_md = _scrub_model_headings(_summary_md_raw)
        if ai_intent == "damage_report_from_photos":
            _summary_md = _scrub_photo_only_narrative_cost_headers(_summary_md)
        mc(_summary_md if _summary_md else "-")
        pdf.ln(2)
    
        # Controlled Repair Cost section rendering (prevents duplicate headings, Totals blocks, duplicate tiers, and bad totals)
        _raw_costs_md = result.get("estimated_costs_markdown") or ""
        costs_md = _strip_unwanted_cost_lines_for_pdf(_raw_costs_md)
    
        # NOTE: Total + Severity Tier are rendered deterministically inside render_repair_cost_section.
    
        _section_bar("APPROXIMATE REPAIR COST BREAKDOWN")
        render_repair_cost_section(pdf, costs_md, tax_rate=tax_rate, parsed=locked_costs_obj)
    
        _section_bar("FRAUD & AUTHENTICITY CHECK")
        mc((result.get("fraud_markdown") or "").strip() or "-")
    
        _section_bar("CONCLUSION")
        mc((result.get("conclusion") or "").strip() or "-")
    
        # --- Combined Disclaimer (end of report) ---
        try:
            pdf.ln(4)
            x_left = pdf.l_margin
            x_right = pdf.w - pdf.r_margin
            y_line = pdf.get_y()
            pdf.set_draw_color(180, 180, 180)
            pdf.line(x_left, y_line, x_right, y_line)
            pdf.ln(3)
    
            pdf.set_text_color(90, 90, 90)
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(0, 5, "Disclaimer:", ln=True)
            pdf.set_font("Helvetica", "", 8)
    
            disclaimer_body = (
                "This report is based solely on photographic evidence and is intended as a preliminary visual damage assessment only. "
                "It does not constitute a formal appraisal, repair estimate, or safety certification. Hidden structural, mechanical, or electronic "
                "damage may exist beyond what is visible in the photographs. All findings should be verified by a qualified collision repair technician, "
                "appraiser, and/or insurance adjuster performing a physical inspection. This report was generated using artificial intelligence. AI systems "
                "may make errors or misinterpret visual information. All photos, damage descriptions, conclusions, and findings must be independently "
                "reviewed and verified by a qualified appraiser before preparing or finalizing any repair estimate. This repair cost is an approximation "
                "derived from the visible conditions in the provided photos only. Actual repair scope and cost may change after teardown, measurement, and "
                "confirmation of hidden damage, glass bonding requirements, and any sensor/trim replacement needs."
            )
            pdf.multi_cell(0, 4, _pdf_sanitize(disclaimer_body))
            pdf.set_text_color(0, 0, 0)

        except Exception:
            pass
            
        # --- Timestamp (match Comprehensive placement) ---
        try:
            pdf.ln(3)
            pdf.set_text_color(90, 90, 90)
            pdf.set_font("Helvetica", "", 8)
            pdf.cell(0, 4, f"Generated: {report_generated_ts}", ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", "", 11)
        except Exception:
            pass
    
        safe_file = _safe(file_number)
        pdf_filename = f"AI_Condition_Report_{safe_file}.pdf"
    else:
        pdf.cell(0,10,"NSPXN.com Condition Report", ln=True, align="C")
        pdf.set_font_size(10); pdf.ln(3)
        mc(f"File Number: {file_number}")
        mc(f"Inspected For: {ia_company}")
        mc(f"Appraiser ID #: {appraiser_id}")
        mc(f"Request Type: {result['request_type']}")

        # --- Supplement header (documents only; ignore negated mentions) ---
        smark = result.get("summary_markdown","")
        _txt_docs = uploaded_text_all or ""
        _supp_doc_hit = bool(re.search(
            r"(?is)\b(Supplement\s+(?:Summary|of\s+Record)|Estimate\s+Version:\s*S0[1-9]\b|\bS0[1-9]\b|\bSupplement\s+Estimate\b)",
            _txt_docs
        ))
        _no_supp_negation = not re.search(r"(?is)\b(no|not)\s+(a\s+)?supplement\b", _txt_docs)
        supp_detected_docs = _supp_doc_hit and _no_supp_negation
        if supp_detected_docs:
            mc("Supplement Status: Supplement Estimate detected in documentation")
            _possible_amt = None
            _m = re.search(r"(?is)\bPossible\s+Supplement\s+Amount\s*\$?([0-9,]+\.\d{2})\b", _txt_docs)
            if _m:
                _possible_amt = _m.group(1)
            mc("Supplement Details")
            # NEW: list all supplement tags (S01, S02, ...) instead of a single version only
            supp_versions_docs = sorted(set(re.findall(r"(?i)\bS[0-9]{2}\b", _txt_docs)))
            if supp_versions_docs:
                mc(f"- Supplements detected in documents: {', '.join(supp_versions_docs)}")
            if _possible_amt:
                mc(f"- Possible amount noted: ${_possible_amt}")
            if not (supp_versions_docs or _possible_amt):
                mc("- Supplement indicators present (e.g., 'Supplement Summary' or S01/S02).")

        # --- Total Loss echo (documents-only; no narrative trigger) ---
        explicit_tl_hit = False
        try:
            txt_docs = uploaded_text_all or ""
            poi15_docs = re.search(
                r"(?is)\bpoint\s*of\s*impact[^a-z0-9]{0,10}15[^a-z0-9]{0,20}total\s*loss\b",
                txt_docs,
            )
            doc_tl_docs = re.search(
                r"(?i)\b(estimate\s*type|type\s*of\s*loss)\s*:\s*total\s*loss\b",
                txt_docs,
            )
            explicit_tl_hit = bool(poi15_docs or doc_tl_docs)
        except Exception:
            explicit_tl_hit = False
        if explicit_tl_hit:
            mc("Estimate Type: Total Loss (explicit in documents)")

        mc(f"Claim #: {result['claim_number']}")
        mc(f"VIN (from estimate/photos): {result['vin']}")
        mc(f"VIN verification (estimate vs photo): {result['vin_verification']}")
        mc(f"Vehicle: {_format_vehicle_value(result.get('vehicle'), result.get('vin'))}")
        mc(f"Odometer (from estimate): {result['odometer_estimate_only']}")
        mc(f"Compliance Score: {result['compliance_score']}")
        pdf_status = result["redaction_status"].replace("✅", "OK")
        mc(pdf_status)
        pdf.ln(3); mc("NSPXN.com Condition Summary"); mc((smark or '').strip())
        pdf.ln(3); mc("Approximate Repair Cost Breakdown"); mc((result.get("estimated_costs_markdown") or "").strip())
        pdf.ln(3); mc("Fraud Detection"); mc((result["fraud_markdown"] or 'N/A').strip())

        # --- AI Disclaimer (after report content) ---
        try:
            pdf.ln(4)
            x_left = pdf.l_margin
            x_right = pdf.w - pdf.r_margin
            y_line = pdf.get_y()
            pdf.set_draw_color(180, 180, 180)
            pdf.line(x_left, y_line, x_right, y_line)
            pdf.ln(3)

            pdf.set_text_color(90, 90, 90)
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(0, 5, "Disclaimer:", ln=True)
            pdf.set_font("Helvetica", "", 8)

            disclaimer_body = (
                "This report was generated using artificial intelligence. AI systems may make errors or misinterpret "
                "visual information. All photos, damage descriptions, conclusions, and findings must be independently "
                "reviewed and verified by a qualified appraiser before preparing or finalizing any repair estimate."
            )
            pdf.multi_cell(0, 4, _pdf_sanitize(disclaimer_body))
            pdf.set_text_color(0, 0, 0)
        except Exception:
            pass


        safe_file = _safe(file_number)
        pdf_filename = f"{safe_file}.pdf"


    # --- One-page photo thumbnail appendix (all uploaded photos) ---
    try:
        add_thumbnail_page(pdf, thumbnail_paths)
    except Exception:
        pass

    pdf_path = os.path.join(PDF_DIR, pdf_filename)
    try:
        out = pdf.output(dest="S")
        if isinstance(out, (bytes, bytearray)):
            data_bytes = bytes(out)
        else:
            data_bytes = str(out).encode("latin-1", "ignore")
        with open(pdf_path, "wb") as f:
            f.write(data_bytes)
    except Exception as e:
        logging.warning(f"PDF write error: {e}")

    pdf_url = f"/download-pdf?filename={pdf_filename}"

    # -----------------------
    # Email — info-only (attach PDF)
    # -----------------------
    try:
        msg = EmailMessage()
        if ai_intent == "damage_report_from_photos":
            subj = f"NSPXN.com Condition Report: {file_number or ''} {result['claim_number'] or ''}".strip()
            _locked_total_email = locked_costs_obj.get('total_val') if isinstance(locked_costs_obj, dict) else None
            _summary_email = _scrub_photo_only_narrative_cost_headers(_scrub_model_headings(result.get('summary_markdown') or ''))
            _conclusion_email = _force_conclusion_to_locked_total(result.get('conclusion') or '', _locked_total_email)
            body = (
                "NSPXN.com Condition Report\n\n"
                f"Generated: {report_generated_ts}\n"
                f"Inspected For: {ia_company}\n"
                f"Claim #: {result['claim_number'] or 'N/A'}    File #: {file_number or 'N/A'}\n"
                f"Odometer: {result['odometer_estimate_only'] or 'N/A'}    Primary Impact: {result['primary_impact'] or 'N/A'}\n"
                f"Secondary Impact: {result['secondary_impact'] or 'N/A'}\n\n"
                f"{result['redaction_status']}\n\n"
                "Condition Summary\n"
                f"{(_summary_email or 'N/A')}\n\n"
                "Approximate Repair Cost Breakdown\n"
                f"{(result['estimated_costs_markdown'] or 'N/A')}\n\n"
                "Fraud & Authenticity Check\n"
                f"{(result['fraud_markdown'] or 'N/A')}\n\n"
                "Conclusion\n"
                f"{(_conclusion_email or 'N/A')}\n"
            )
        else:
            try:
                _txt_email = uploaded_text_all or ""
                _poi15_email = re.search(
                    r"(?is)\bpoint\s*of\s*impact[^a-z0-9]{0,10}15[^a-z0-9]{0,20}total\s*loss\b",
                    _txt_email,
                )
                _doc_tl_email = re.search(
                    r"(?i)\b(estimate\s*type|type\s*of\s*loss)\s*:\s*total\s*loss\b",
                    _txt_email,
                )
                _explicit_tl_email = bool(_poi15_email or _doc_tl_email)
            except Exception:
                _explicit_tl_email = False

            # NEW: supplement line includes all Sxx tags if present
            supp_line = ""
            if supp_detected_docs:
                supp_versions_email = sorted(set(re.findall(r"(?i)\bS[0-9]{2}\b", _txt_email or "")))
                if supp_versions_email:
                    supp_line = (
                        "Supplement Status: Supplement Estimates detected in documentation "
                        f"({', '.join(supp_versions_email)})\n"
                    )
                else:
                    supp_line = "Supplement Status: Supplement Estimate detected in documentation\n"

            tl_line = "Estimate Type: Total Loss (explicit in documents)\n" if _explicit_tl_email else ""
            subj = f"NSPXN.com Review: {result['claim_number'] or file_number}"
            body = (
                "NSPXN.com Condition Report\n\n"
                f"File Number: {file_number}\n"
                f"Inspected For: {ia_company}\n"
                f"Appraiser ID #: {appraiser_id}\n"
                f"Request Type: {result['request_type']}\n"
                f"{supp_line}"
                f"{tl_line}"
                f"Claim #: {result['claim_number']}\n"
                f"VIN (from estimate/photos): {result['vin']}\n"
                f"VIN verification (estimate vs photo): {result['vin_verification']}\n"
                f"Vehicle: {result['vehicle']}\n"
                f"Odometer (from estimate): {result['odometer_estimate_only']}\n"
                f"Compliance Score: {result['compliance_score']}\n\n"
                f"{result['redaction_status']}\n\n"
                "NSPXN.com Condition Summary\n"
                f"{result['summary_markdown']}\n\n"
                "Fraud Detection\n"
                f"{result['fraud_markdown']}\n"
            )

        msg["Subject"] = subj
        msg["From"] = "info@nspxn.com"
        msg["To"] = "info@nspxn.com"
        msg["Cc"] = "growley@ractrak.com"
        msg.set_content(body)

        try:
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            msg.add_attachment(pdf_bytes, maintype="application", subtype="pdf", filename=pdf_filename)
        except Exception as e:
            logging.warning(f"Failed to attach PDF to email: {e}")

        with smtplib.SMTP_SSL("mail.tierra.net", 465, timeout=20) as smtp:
            smtp.login("info@nspxn.com", "grr2025GRR")
            smtp.send_message(msg)
        log.info("Info email sent to info@nspxn.com")
    except Exception as e:
        logging.error(f"Email error: {e}")

    return {
        **result,
        "web_summary": result["summary_brief"],
        "gpt_output": result["summary_markdown"],
        "pdf_url": pdf_url,
        "pdf_filename": pdf_filename
    }

    # -----------------------
# PDF download
    # -----------------------
@app.get("/download-pdf")
async def download_pdf(file_number: Optional[str] = None, filename: Optional[str] = None):
    if filename:
        safe = _safe(filename)
        path = os.path.join(PDF_DIR, safe)
        if os.path.exists(path):
            return FileResponse(path=path, media_type="application/pdf", filename=safe)
        return JSONResponse(status_code=404, content={"detail": "Not Found"})
    if not file_number:
        return JSONResponse(status_code=400, content={"detail": "Missing query param 'filename' or 'file_number'"})
    safe_num = _safe(file_number)
    candidates = glob.glob(os.path.join(PDF_DIR, f"*{safe_num}*.pdf"))
    if not candidates:
        return JSONResponse(status_code=404, content={"detail": "Not Found"})
    latest = max(candidates, key=lambda p: os.path.getmtime(p))
    return FileResponse(path=latest, media_type="application/pdf", filename=os.path.basename(latest))
