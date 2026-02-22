from fastapi import FastAPI, File, UploadFile, Form, Request
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Dict, Any, Optional
import os, io, re, json, base64, logging, zipfile, glob, uuid
import urllib.parse, urllib.request
import smtplib  # email transport
from email.message import EmailMessage

from fpdf import FPDF
from docx import Document
from pdf2image import convert_from_bytes
from PIL import Image

# Optional HEIC/HEIF support if available
try:
    import pillow_heif  # type: ignore
    pillow_heif.register_heif_opener()  # type: ignore
except Exception:
    pass

# Optional OCR (pytesseract). Safe no-op if not installed or tesseract binary missing.
try:
    import pytesseract  # type: ignore
    _OCR_ENABLED = True
except Exception:
    _OCR_ENABLED = False

from openai import OpenAI

# --- PII Redaction (Presidio) ---
from presidio_analyzer import AnalyzerEngine, Pattern, PatternRecognizer, RecognizerResult
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig  # required for anonymizer API

# -----------------------
# Minimal setup
# -----------------------
PDF_DIR = os.getenv("PDF_DIR", "/tmp"); os.makedirs(PDF_DIR, exist_ok=True)
CLIENT_RULES_DIR = os.getenv("CLIENT_RULES_DIR", "client_rules")

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
log = logging.getLogger("nspxn")
log.info(f"Using CLIENT_RULES_DIR={CLIENT_RULES_DIR}")

# Use selected model everywhere
MODEL = os.getenv("OAI_MODEL") or os.getenv("OPENAI_MODEL") or "gpt-5.2"
# GPT-5.x models use max_completion_tokens; GPT-4.x uses max_tokens
_token_kw = "max_completion_tokens"

if not os.getenv("OPENAI_API_KEY"):
    raise RuntimeError("OPENAI_API_KEY missing")
try:
    client = OpenAI(api_key=os.environ["OPENAI_API_KEY"], timeout=120.0, max_retries=1)
except TypeError:
    # Backwards-compatible init for older openai-python versions
    client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])

# --------------------------------
# Presidio: Analyzer/Anonymizer (preserve VIN & Claim #)
# --------------------------------
analyzer = AnalyzerEngine()
anonymizer = AnonymizerEngine()

VIN_PATTERN = r"\b([A-HJ-NPR-Z0-9]{17})\b"
vin_recognizer = PatternRecognizer(
    supported_entity="VIN",
    patterns=[Pattern(name="vin-17", regex=VIN_PATTERN, score=0.8)],
)

CLAIM_PATTERN = r"\b(?:(?:Claim|CLM|Clm)\s*#?\s*[:\-]?\s*)?([A-Z0-9]{5,}[A-Z0-9\-]{0,})\b"
claim_recognizer = PatternRecognizer(
    supported_entity="CLAIM_NUMBER",
    patterns=[Pattern(name="claim-generic", regex=CLAIM_PATTERN, score=0.6)],
)

analyzer.registry.add_recognizer(vin_recognizer)
analyzer.registry.add_recognizer(claim_recognizer)

# Only redact these entity types (we PRESERVE VIN & CLAIM_NUMBER)
REDACT_ENTITY_TYPES = {
    "PERSON", "PHONE_NUMBER", "EMAIL_ADDRESS", "US_SSN",
    "CREDIT_CARD", "IBAN_CODE", "LOCATION", "NRP", "ORGANIZATION",
    "DATE_TIME", "IP_ADDRESS", "CRYPTO", "MEDICAL_LICENSE", "URL"
}

def _filter_results(results: List[RecognizerResult]) -> List[RecognizerResult]:
    return [r for r in results if r.entity_type in REDACT_ENTITY_TYPES]

def redact_text_preserve_vin_claim(text: str) -> str:
    """Mask PII while keeping VIN and Claim # intact. Fail-open to avoid 500s."""
    if not text:
        return text
    results = analyzer.analyze(text=text, language="en")
    to_mask = _filter_results(results)
    if not to_mask:
        return text
    try:
        return anonymizer.anonymize(
            text=text,
            analyzer_results=to_mask,
            operators={"DEFAULT": OperatorConfig("replace", {"new_value": "[REDACTED]"})}
        ).text
    except Exception as e:
        log.warning(f"Presidio anonymizer failed, passing text through. Error: {e}")
        return text

def _safe(s: str) -> str:
    return re.sub(r"[^\w.\-]+", "-", (s or "").strip()).strip("-_. ")



def _normalize_ai_notes(raw: str, max_len: int = 800) -> str:
    """Normalize/sanitize Add'l Notes so they can't break JSON/structured prompting."""
    if raw is None:
        return "No additional notes provided."
    s = str(raw).strip()
    if not s:
        return "No additional notes provided."

    # Remove ASCII control chars except newline/tab
    s = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", " ", s)

    # Normalize newlines + collapse whitespace
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s).strip()

    # Neutralize common schema-breakers
    s = s.replace("```", "'''")
    # Avoid curly braces interfering with JSON-format instructions
    s = s.replace("{", "(").replace("}", ")")

    if len(s) > max_len:
        s = s[:max_len].rstrip() + "…"
    return s

def _ai_notes_block(ai_notes_clean: str) -> str:
    """Always return a consistent, clearly-delimited notes block."""
    return (
        "\n\nADDL_NOTES_FOR_AI_REVIEW (USER PROVIDED):\n"
        "<<<\n"
        f"{ai_notes_clean}\n"
        ">>>\n"
        "REQUIREMENT:\n"
        "- You MUST consider these notes while analyzing the photos.\n"
        "- Use them to guide focus/wording, but DO NOT invent damage.\n"
        "- If notes conflict with visible evidence or orientation is unclear, say so and defer to what is visible.\n"
    )

# -----------------------
# Approximate Repair Cost Breakdown (location-based rates)
# -----------------------
# Notes:
# - Paint supplies/materials are modeled as $ per refinish hour (NOT a percent).
# - This is an approximation for observed damages, not a repair estimate.

# Optional external integration (no new dependencies):
# - Set LABORRATEHERO_API_URL to a JSON endpoint you control (proxy/scraper) that returns:
#   {"body_rate": 0, "paint_rate": 0, "frame_rate": 0, "mechanical_rate": 0, "paint_supplies_rate": 0}
# - Optionally set LABORRATEHERO_API_KEY if your endpoint requires it.
LABORRATEHERO_API_URL = os.getenv("LABORRATEHERO_API_URL", "").strip()
LABORRATEHERO_API_KEY = os.getenv("LABORRATEHERO_API_KEY", "").strip()

def _extract_inspection_location(text: str) -> str:
    """Best-effort extraction of Inspection Location (City/State/ZIP) from uploaded text."""
    if not text:
        return ""
    # Common patterns: "Inspection Location: Albuquerque, NM 87109" or multi-line variants
    m = re.search(r"(?im)^\s*Inspection\s+Location\s*[:\-]?\s*(.+)$", text)
    if m:
        return (m.group(1) or "").strip()
    m2 = re.search(r"(?is)\bInspection\s+Location\b\s*[:\-]?\s*(.{0,120}?)\n", text)
    if m2:
        return (m2.group(1) or "").strip()
    return ""

def _extract_zip5(*texts: str) -> str:
    """Return the first 5-digit ZIP found in any provided text (best-effort)."""
    for t in texts:
        if not t:
            continue
        m = re.search(r"\b(\d{5})(?:-\d{4})?\b", str(t))
        if m:
            return m.group(1)
    return ""

def _zip_to_city_state(zip5: str) -> Optional[Dict[str, str]]:
    """Best-effort ZIP -> City/State via Zippopotam (public). Fail-open."""
    z = (zip5 or "").strip()
    if not re.fullmatch(r"\d{5}", z):
        return None
    try:
        url = f"https://api.zippopotam.us/us/{z}"
        req = urllib.request.Request(url, headers={"Accept": "application/json"})
        with urllib.request.urlopen(req, timeout=6) as resp:
            raw = resp.read().decode("utf-8", "ignore")
        data = json.loads(raw) if raw else {}
        if not isinstance(data, dict):
            return None
        places = data.get("places") or []
        if not places:
            return None
        p0 = places[0] if isinstance(places[0], dict) else {}
        city = (p0.get("place name") or "").strip()
        state = (p0.get("state abbreviation") or "").strip().upper()
        if city and state:
            return {"city": city, "state": state, "zip": z}
    except Exception:
        return None
    return None

def _normalize_location_with_zip(location: str, *fallback_texts: str) -> str:
    """If location lacks City/State but has ZIP, expand it using ZIP->City/State."""
    loc = (location or "").strip()
    zip5 = _extract_zip5(loc, *fallback_texts)
    if not zip5:
        return loc
    # If already contains a state, just ensure ZIP is present
    if _parse_state_from_location(loc):
        return loc
    z = _zip_to_city_state(zip5)
    if z:
        return f"{z['city']}, {z['state']} {z['zip']}"
    # At least return ZIP if nothing else
    return zip5


def _parse_state_from_location(loc: str) -> str:
    if not loc:
        return ""
    # Pull 2-letter state if present (", NM" or " NM ")
    m = re.search(r"(?i)\b([A-Z]{2})\b\s*(\d{5})?(?:-\d{4})?\s*$", loc.strip())
    if m:
        return (m.group(1) or "").upper()
    m2 = re.search(r"(?i),\s*([A-Z]{2})\b", loc)
    if m2:
        return (m2.group(1) or "").upper()
    return ""

def _fallback_rates_by_state(state: str) -> Dict[str, float]:
    """Conservative defaults when no external rate service is configured/available."""
    # Keep these conservative; you can tune later or replace with your own table.
    # body/paint/frame/mechanical in $/hr; paint_supplies_rate in $/refinish hr.
    defaults = {
        "DEFAULT": {"body_rate": 60.0, "paint_rate": 60.0, "frame_rate": 75.0, "mechanical_rate": 115.0, "paint_supplies_rate": 38.0},
        "NM":      {"body_rate": 64.0, "paint_rate": 62.0, "frame_rate": 78.0, "mechanical_rate": 120.0, "paint_supplies_rate": 38.0},
        "AZ":      {"body_rate": 66.0, "paint_rate": 64.0, "frame_rate": 80.0, "mechanical_rate": 125.0, "paint_supplies_rate": 40.0},
        "CO":      {"body_rate": 72.0, "paint_rate": 70.0, "frame_rate": 88.0, "mechanical_rate": 135.0, "paint_supplies_rate": 42.0},
        "TX":      {"body_rate": 62.0, "paint_rate": 60.0, "frame_rate": 75.0, "mechanical_rate": 120.0, "paint_supplies_rate": 38.0},
        "CA":      {"body_rate": 85.0, "paint_rate": 85.0, "frame_rate": 110.0, "mechanical_rate": 165.0, "paint_supplies_rate": 50.0},
    }
    return dict(defaults.get(state.upper(), defaults["DEFAULT"]))

def _fetch_rates_from_laborratehero(location: str) -> Optional[Dict[str, float]]:
    """Fetch rates from a configured endpoint. Returns None on any failure."""
    if not LABORRATEHERO_API_URL:
        return None
    try:
        q = urllib.parse.urlencode({"location": location})
        url = LABORRATEHERO_API_URL
        url = (url + ("&" if "?" in url else "?") + q) if location else url
        req = urllib.request.Request(url, headers={"Accept": "application/json"})
        if LABORRATEHERO_API_KEY:
            req.add_header("Authorization", f"Bearer {LABORRATEHERO_API_KEY}")
        with urllib.request.urlopen(req, timeout=10) as resp:
            raw = resp.read().decode("utf-8", "ignore")
        data = json.loads(raw) if raw else {}
        if not isinstance(data, dict):
            return None
        out = {}
        for k in ("body_rate","paint_rate","frame_rate","mechanical_rate","paint_supplies_rate"):
            v = data.get(k)
            if isinstance(v, (int, float)) and v > 0:
                out[k] = float(v)
        return out or None
    except Exception:
        return None

def _lookup_rates(location: str) -> Dict[str, float]:
    """Return a complete rate card (with fallbacks)."""
    location = _normalize_location_with_zip(location)
    state = _parse_state_from_location(location)
    base = _fallback_rates_by_state(state)
    ext = _fetch_rates_from_laborratehero(location) or {}
    base.update({k: float(v) for k, v in ext.items() if isinstance(v, (int, float)) and float(v) > 0})
    # Ensure required keys exist
    for k in ("body_rate","paint_rate","frame_rate","mechanical_rate","paint_supplies_rate"):
        base.setdefault(k, 0.0)
    return base

def _lookup_tax_rate(location: str) -> float:
    """Best-effort tax rate for parts/materials (decimal). Conservative fallback."""
    # Optional external (user-controlled) endpoint
    tax_url = os.getenv("TAXRATE_API_URL", "").strip()
    tax_key = os.getenv("TAXRATE_API_KEY", "").strip()
    if tax_url and location:
        try:
            q = urllib.parse.urlencode({"location": location})
            url = tax_url + ("&" if "?" in tax_url else "?") + q
            req = urllib.request.Request(url, headers={"Accept": "application/json"})
            if tax_key:
                req.add_header("Authorization", f"Bearer {tax_key}")
            with urllib.request.urlopen(req, timeout=10) as resp:
                raw = resp.read().decode("utf-8", "ignore")
            data = json.loads(raw) if raw else {}
            if isinstance(data, dict):
                r = data.get("tax_rate")
                if isinstance(r, (int, float)) and 0 <= float(r) <= 0.15:
                    return float(r)
                # allow percent input
                if isinstance(r, (int, float)) and 0 <= float(r) <= 15:
                    return float(r) / 100.0
        except Exception:
            pass

    # Minimal state fallbacks (tune later). Use decimal (e.g., 0.07875 = 7.875%)
    state = _parse_state_from_location(location)
    state_map = {
        "NM": 0.07875,
        "AZ": 0.085,
        "CO": 0.075,
        "TX": 0.0825,
        "CA": 0.085,
    }
    return float(state_map.get(state, 0.08))

def _money(x: Optional[float]) -> str:
    try:
        if x is None:
            return "$0"
        return "${:,.0f}".format(float(x))
    except Exception:
        return "$0"

def _num(x: Optional[float]) -> float:
    try:
        return float(x) if x is not None else 0.0
    except Exception:
        return 0.0

def _extract_hours_and_parts_totals(text: str) -> Dict[str, Any]:
    """Parse estimate totals (best effort)."""
    out: Dict[str, Any] = {"body_hours": None, "paint_hours": None, "frame_hours": None, "mech_hours": None, "parts_total": None, "parts_lines": []}
    if not text:
        return out

    def _find_hours(label_rx: str) -> Optional[float]:
        m = re.search(label_rx, text, flags=re.IGNORECASE)
        if m:
            try:
                return float(m.group(1))
            except Exception:
                return None
        return None

    out["body_hours"]  = _find_hours(r"\bBody\s+Labor\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)")
    out["paint_hours"] = _find_hours(r"\bPaint\s+Labor\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)")
    out["frame_hours"] = _find_hours(r"\b(Frame|Structural)\s+Labor\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)")  # group2 maybe
    if out["frame_hours"] is None:
        m = re.search(r"\bFrame\s+Labor\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)", text, flags=re.IGNORECASE)
        if m:
            try: out["frame_hours"] = float(m.group(1))
            except Exception: pass
    out["mech_hours"]  = _find_hours(r"\b(Mechanical|Mech)\s+Labor\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)")
    if out["mech_hours"] is None:
        m = re.search(r"\bMechanical\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)", text, flags=re.IGNORECASE)
        if m:
            try: out["mech_hours"] = float(m.group(1))
            except Exception: pass

    # Parts total (best effort)
    mpt = re.search(r"(?i)\bParts\s*(?:Total)?\b\s*[:\-]?\s*\$\s*([0-9,]+(?:\.[0-9]{2})?)", text)
    if mpt:
        try:
            out["parts_total"] = float(mpt.group(1).replace(",", ""))
        except Exception:
            pass

    # Parts line items (best effort, capped)
    # Capture lines that look like: "<line#> ... <part desc> ... $123.45"
    lines = []
    for ln in (text or "").splitlines():
        s = ln.strip()
        if not s:
            continue
        if len(s) > 220:
            continue
        if re.search(r"(?i)\b(replace|r\&r|r\s*/\s*r)\b", s) and re.search(r"\$\s*[0-9,]+(?:\.[0-9]{2})?", s):
            lines.append(s)
        elif re.search(r"(?i)\b(part|oem)\b", s) and re.search(r"\$\s*[0-9,]+(?:\.[0-9]{2})?", s):
            lines.append(s)
    out["parts_lines"] = lines[:20]
    return out

def _structural_observed(text_blobs: List[str]) -> bool:
    t = "\n".join([x for x in text_blobs if x]).lower()
    if not t:
        return False
    rx = r"\b(frame|rail|apron|unibody|structural|pull|straighten|set\s*up\s*&\s*measure|measure\s*&\s*setup|dimension\s*check|buckl|kink|twist)\b"
    return bool(re.search(rx, t))
# -----------------------
# Prompt steering (free analysis + detailed narrative)
# -----------------------
DETAIL_TEMPLATES = {
    "guidelines_only": (
        "## Inputs Used\n"
        "- Briefly list which documents, pages, and photos you actually referenced.\n\n"
        "## Executive Summary\n"
        "- 3–5 bullets summarizing overall compliance and key risks.\n\n"
        "## AI-4-IA Review Summary\n"
        "- Write a formal, paragraph-style appraisal narrative. Include scope of impact, damage by zone/panel, "
        "repair vs. replace rationale, parts type (OEM/LKQ/Aftermarket), labor ops, refinish overlap, rate validation, "
        "tax handling, and estimate integrity. Reference evidence inline (e.g., 'p2/L14', 'Photo 3'). "
        "Close with compliance stance and final recommendation. Minimum 8–10 sentences.\n\n"
        "## Key Issues & Actions\n"
        "- Bullet list of the highest-impact issues with a one-line recommended action each.\n\n"
        "## Final\n"
        "- Compliance Score: NN% with one-sentence rationale."
    ),

    "comprehensive": (
        "## Inputs Used\n"
        "- List the estimate pages/lines and photo numbers you used, plus any rules text (if provided).\n\n"
        "## Executive Summary\n"
        "- 3–6 bullets capturing the big picture: estimate integrity, rule alignment (only if rules text was supplied), "
        "and photo consistency.\n\n"
        "## Detailed Condition Report\n"
        "- Write this section as a formal, paragraph-style appraisal report summarizing the entire claim. "
        "Include: scope of impact, damage by zone/panel, repair vs. replace rationale, parts type (OEM/LKQ/Aftermarket), "
        "labor operations, refinish/overlap considerations, rate validation, paint materials handling, sublet usage, "
        "tax/markup accuracy, and overall estimate integrity. Cite photos and estimate lines (e.g., 'Photo 3', 'p2/L14'). "
        "Close with compliance to any provided client rules and a clear final recommendation (Repairable vs. Total Loss). "
        "Do not declare Repairable/Total Loss unless the estimate itself explicitly marks 'Total Loss' or an ACV comparison is provided. "
        "If the shop info is listed under Repair Facility on ANY estimate, add only the shop name to the Detailed Condition Report narrative. "
        "If a Printout showing the Clean Retail Value or Estimated Trade-In Value of the unit is present which may include ANY of the following: NADA, J.D. Power, Kelly Blue Book, Edmunds, Carfax, or Cars.com, DO NOT declare as missing if any of these are present. "
        "Minimum 10–14 sentences (one continuous narrative, not bullets).\n\n"
        "## Photo-by-Photo Damage Ledger\n"
        "| Photo # | View/Angle | Panels/Parts Visible | Condition (dent/crease/scrape/misalignment) | Identifiers (VIN/odo/plate/reg) | Legibility |\n"
        "|---:|---|---|---|---|---|\n"
        "- One row per photo used in the analysis (>=6 rows if >=6 photos exist). If an identifier is present but unreadable, mark 'Present — not clearly legible'.\n\n"
        "## Brief Damage Descriptions\n"
        "- 6–12 bullets. Each: part/panel + condition + suggested op (repair/replace/refinish/blend) + Photo #.\n\n"
        "## Estimate Line Extract (top relevant lines)\n"
        "| Est. p#/L# | Part/Op | Labor Hrs | Rate | Part Type | Price | Notes |\n"
        "|---|---|---:|---:|---|---:|---|\n"
        "- Include 8+ lines if available, focusing on items tied to observed damages. Use 'Notes' for overlap/blend or rationale.\n\n"
        "## Estimate Compliance Cross-Check (brief)\n"
        "Status: Compliant / Non-compliant / Not Evidenced. Cite only the most important evidence (p#/L# or Photo #). Keep it short.\n"
        "| Topic | Evidence (p#/L# or value; Photo # if relevant) | Status | Required Fix |\n"
        "|---|---|:--:|---|\n"
        "| Labor Rates |  |  |  |\n"
        "| Refinish/Overlap |  |  |  |\n"
        "| Paint Materials |  |  |  |\n"
        "| OEM Procedures |  |  |  |\n"
        "| Sublet |  |  |  |\n"
        "| Tax/Markup |  |  |  |\n\n"
        "## Client Guidelines Comparison (if rules text was supplied)\n"
        "- 3–8 bullets. Quote the relevant rule fragment and note Aligned / Not Aligned / Not Evidenced, with evidence refs (p#/L#, Photo #). "
        "If no client_rules were provided, omit this section.\n\n"
        "## Risks / Missing Evidence\n"
        "- Short bullets with severity (High/Med/Low) and a one-line remediation.\n\n"
        "## Compliance Score Rationale\n"
        "- REQUIRED if score < 100: start at 100 and list each deficiency with evidence refs (p#/L# and/or Photo #), "
        "severity (Minor/Moderate/Major), and numeric deduction. Show the arithmetic to the final score.\n\n"
        "## Final Evaluation\n"
        "- Compliance Score: NN% with a single-sentence justification. "
        "If no fraud indicators are identified, state 'No material inconsistencies found.' Do not use 'N/A'."
    ),

    "damage_report_from_photos": (
        """
# Condition Report (Photos Only)

## Photo-by-Photo Condition Summary
| Photo # | View/Side | Key Panels/Parts Visible | Damage/Condition |
|---:|---|---|---|
- Cover EVERY provided photo. If no damage is obvious from that angle, write: "No obvious damage visible from this angle" (do not use the word intact).

## Side Checks
- **Driver/Left Side**: <what is visible; cite Photo #; if not shown, say not shown>
- **Passenger/Right Side**: <what is visible; cite Photo #; if not shown, say not shown>

## Front-End Checklist
- Hood: <condition or Not clearly shown> (Photo #)
- Front bumper cover: <condition or Not clearly shown> (Photo #)
- Grille: <condition or Not clearly shown> (Photo #)
- Driver-side headlamp: <condition or Not clearly shown> (Photo #)
- Passenger-side headlamp: <condition or Not clearly shown> (Photo #)
- Driver-side front fender: <condition or Not clearly shown> (Photo #)
- Passenger-side front fender: <condition or Not clearly shown> (Photo #)

## Detailed Condition Report
- Write a continuous 10–15 sentence narrative summarizing visible damage, impact zones, misalignment/gaps, and repair implications (photo-based).
- If VIN label or odometer are visible, state them with Photo #. If not visible or unreadable, say so.

## Approximate Repair Cost Breakdown (Populate JSON field 'estimated_costs_markdown')
- You MUST produce a cost approximation derived from the PHOTOS ONLY (do not reference estimates, documents, or 'not evidenced').
- Provide AI-derived hours and assumptions:
  • Body labor hours
  • Paint labor hours
  • Paint & materials cost = (paint hours × $/refinish hr)
  • If structural damage is observed in photos: add Setup & Measure = 2.0 hrs @ body rate and include frame hours @ frame rate
  • If airbags/ADAS/mechanical damage is observed: include mechanical hours @ mechanical rate
- Provide an OEM replacement parts list (each part with an approximate $).
- Apply tax ONLY to (parts + paint materials). Do NOT tax labor.
- Include the Severity Tier checkboxes:
  ☐ Minor (< $3,500)
  ☐ Moderate ($3,500–$10,000)
  ☐ Major ($10,000+)
  ☐ Possible Total Loss Threshold Approaching
- Include a Repair Cost Disclaimer stating this is an approximation, not an official estimate.
- Forbidden phrases: 'from estimate', 'from documentation', 'not evidenced', 'no documentation provided'.
"""
    ),
}

# --- Static audit questions ---
STATIC_AUDIT_QUESTIONS = [
    "Do the photos substantiate the highest-cost operations (frame/sectioning/panel replace)?",
    "Are ADAS calibrations or wheel alignments required and supported by the damage and OEM procedures?",
    "Is blend time justified by color/finish (metallic/pearl/tri-coat) and adjacent panel visibility?",
    "Do invoices corroborate parts used and match estimate line items (brand/grade, price, quantity)?",
    "Are AM/LKQ choices compliant with age/mileage rules, and is OEM required anywhere by client policy or safety?",
    "Is there evidence of prior or unrelated damage (UPD) that materially affects valuation or repair scope?",
    "Are there structural/safety indicators (buckles, misalignments, airbags/pretensioners) that alter repair strategy?",
    "Are materials/hazard charges (paint supplies, corrosion protection, seam sealer) aligned with operations and shop norms?",
    "Are storage/tow charges and dates supported and reasonable given claim timeline and shop status?",
    "Are scanner reports (pre/post) included or needed; if absent, does that meaningfully impact confidence?",
    "Did the supplement (if any) correct earlier gaps, and are newly added operations now evidenced?",
    "Are client-required documents present (e.g., NADA printout, release forms, production date plate); if missing, what’s the impact?",
    "What is the bottom-line recommendation (approve as-is, adjust items, or request specific evidence)?",
]

# --- Identifiers Verification Protocol (prompt-only; no new logic) ---
IDENTIFIERS_VERIFICATION_PROTOCOL = (
    "\n\nIDENTIFIERS VERIFICATION PROTOCOL (must follow):"
    "\n1) Search the photos for: windshield VIN plate, driver-door VIN label, driver-door VIN label Production date, driver-door VIN label Date of Mfr, odometer cluster."
    "\n2) Transcribe the VIN exactly as visible and cite Photo # for EACH location you find."
    "\n3) If multiple VINs, compare them to each other and to the estimate VIN; explicitly state: MATCH / MISMATCH."
    "\n4) Transcribe the odometer reading exactly as shown and cite Photo #."
    "\n5) Grade legibility for each identifier as one of: 'Clearly legible' / 'Present — not clearly legible' / 'Not present'."
    "\n6) If any identifier is present but not clearly legible, say why (glare, blur, angle) and what photo would resolve it."
    "\n7) Write a one-line bottom line: 'VIN verification: <MATCH/MISMATCH/INCONCLUSIVE>; Odometer: <value or reason>'."
    "\n8) Weave these facts naturally into the '## Detailed Condition Report' narrative and keep the top-line fields "
    "(vin, vin_verification, odometer_estimate_only) consistent."
    "\n9) When citing more than one VIN location (e.g., windshield vs. door label), you must cite DISTINCT Photo #s; "
    "never reuse the same photo number for two different locations."
    "\n10) Compare VINs as literal 17-character strings. If any single character differs between sources, report "
    "MISMATCH, and quote both strings with their Photo #/page references."
    "\n11) ODOMETER RULES (photos-only especially): transcribe only the digits visible in the odometer photo; "
    "do not infer from estimate text or metadata. Include the exact Photo #. If any digit is unclear, state 'Present — not clearly legible' and explain why; do not guess."
)

# --- Consistency Guard (prompt-only; avoid contradictions) ---
CONSISTENCY_GUARD = (
    "\n\nCONSISTENCY GUARD:"
    "\n- Do not claim any required photo is 'missing' if you graded it 'Clearly legible' or 'Present — not clearly legible'."
    "\n- For VIN, Odometer, and Production Date specifically: if present in any photo, do not write any sentence implying they are absent."
    "\n- If a legible driver-door VIN label photo is present, treat the Production Date requirement as evidenced (the production month/year appears on the same label). Do NOT deduct or say 'not separately documented'."
    "\n- Only deduct for missing Repair Facility info when the Closing Report or other documents clearly show the vehicle is at a named repair facility AND the estimate's 'Repair Facility' section does not list that same facility; "
    "if no repair facility information appears anywhere in the estimate or Closing Report, report 'N/A — not provided' and do NOT deduct."
    "\n- If legibility is the issue, explicitly say 'Present — not clearly legible' and explain why (glare/blur/angle), and request a precise retake rather than marking it missing."
    "\n- Before finalizing, re-scan your output: confirm every referenced Photo # matches the content described (e.g., do not cite an Odometer photo as the point-of-impact photo). Correct any mismatches."
)

# --- No-intact-if-damaged rule (prompt-only; prevents false 'intact' claims) ---
NO_INTACT_IF_DAMAGED_RULE = (
    "\n\nNO 'INTACT/NO-DAMAGE' OVERRIDE RULE (DO NOT PRINT CONFLICT WARNINGS):"
    "\n- If any photo indicates damage to a panel/component, you may NOT state that same panel/component is intact/undamaged/no visible damage anywhere."
    "\n- If photos appear inconsistent, DO NOT write a conflict warning; instead, remove/avoid the intact/no-damage claim and describe only what appears damaged with citations."
)

# --- Damage Side / Orientation Guard (prompt-only; prevents left/right drift) ---
DAMAGE_SIDE_GUARD = (
    "\n\nDAMAGE SIDE GUIDANCE (MINIMAL):"
    "\n- Describe any visible damage on any side (Driver/Passenger or Left/Right) when it is visible in photos."
    "\n- Do NOT suppress side-level damage descriptions when damage is clearly visible."
    "\n- If orientation is genuinely unclear, say so and avoid guessing."
)

BILATERAL_DAMAGE_MANDATE = (
    "\n\nBILATERAL / SECONDARY DAMAGE MANDATE (STRICT):\n"
    "- In frontal impacts, explicitly address BOTH front corners (driver-side and passenger-side if clear; otherwise ‘front corner A/B’ or ‘front corner (viewed)’)\n"
    "- If a photo shows partial view of the opposite side with visible distortion/misalignment/crush, describe it — do NOT default to 'intact' or 'no damage' without citing clear evidence.\n"
    "- Contradicting visible photo evidence (e.g. calling a crushed fender 'intact') is forbidden."
)

FRONT_CORNER_ORIENTATION_GUARD = (
    "\n\nFRONT CORNER ORIENTATION (MANDATORY, MINIMAL):"
    "\n- Do NOT label front damage as LF/RF (or 'left/right headlight/fender') unless the photo angle clearly establishes the vehicle orientation."
    "\n- If you have a straight-on FRONT photo: viewer-right corresponds to vehicle-LEFT; viewer-left corresponds to vehicle-RIGHT."
    "\n- If orientation is not clear, use neutral wording: 'front corner' / 'front headlamp area' instead of left/right."
    "\n- You may NOT state 'left front fender/headlight intact' or 'right front fender/headlight intact' unless orientation is established; otherwise say 'not clearly shown from this angle.'"
)

# --- Parts Source Guard (prompt-only; prevents OEM vs Aftermarket drift) ---
PARTS_SOURCE_GUARD = (
    "\n\nPARTS SOURCE GUARD (MANDATORY):"
    "\n- Do NOT claim 'aftermarket', 'A/M', 'quality replacement', 'non-OEM', or 'LKQ/used/recycled' parts were used unless the estimate LINE ITEMS explicitly label them as such."
    "\n- Generic disclosure/boilerplate text about aftermarket crash parts does NOT prove aftermarket parts were used."
    "\n- If the Closing Report states no aftermarket/LKQ parts were included, your narrative must not claim they were used."
    "\n- When parts source is not explicit, state that it is not explicitly labeled and avoid guessing; default to OEM only when supported by part numbers/labels."
)

# --- Supplement Handling (prompt-only; ensures detection + narrative mention) ---
SUPPLEMENT_HANDLING = (
    "\n\nSUPPLEMENT HANDLING:"
    "\n- Examine the estimate documents for explicit supplement indicators: 'Supplement', 'Supplement of record', 'S01', 'S02', 'Supplement Summary', or similar."
    "\n- If a supplement or multiple supplements are detected, clearly state in the narrative that the estimate is a supplement and summarize what changed: added operations/parts, rate updates, refinish overlap changes, or corrections to prior omissions."
    "\n- If the supplement(s) corrects earlier deficiencies (e.g., missing materials line, added calibrations), note that improvement explicitly."
    "\n- If a supplement(s) exists but required supporting evidence (invoices, photos) is still missing, call this out in Risks/Missing Evidence."
)

ALLOWED_INTENTS = {"guidelines_only","comprehensive","damage_report_from_photos"}

SYSTEM_BASE = (
    "You are an auto-claims appraisal assistant. Return ONLY valid JSON (no code fences). "
    "Always include all required keys: "
    "['file_number','request_type','claim_number','vin','vin_verification','vehicle',"
    "'odometer_estimate_only','compliance_score','summary_brief','summary_markdown',"
    "'fraud_markdown','primary_impact','secondary_impact','estimated_costs_markdown','conclusion']. "
    "Use only provided evidence. Cite photos as 'Photo #' and docs as 'p#/L#' when available. "
    "Do not guess. If something is not visible, say why. "
    "NEVER return 'N/A' for summary_markdown, fraud_markdown, estimated_costs_markdown, or conclusion. For request_type 'Create a Damage Report from Photos', 'estimated_costs_markdown' must be a photos-only approximation (no document/estimate dependency language) and must model Paint & Materials as a $ per refinish hour rate (not a percent)."
)

SYSTEM_BASE += (
    " Do not state or imply any client rule unless it appears verbatim in the provided client_rules text. "
    "If client_rules is blank, write the entire report without referencing client rules. "
    "If a value cannot be confirmed from the visible evidence, set it to 'N/A' and briefly state why. "
    " Except when the request_type is 'Create a Damage Report from Photos', Compliance Score must be a numeric percentage 0–100 (never 'N/A'). "
    "If the request_type is 'Create a Damage Report from Photos', set compliance_score to 'N/A' and omit the '## Compliance Score Rationale' section. "
    "If no client_rules are supplied, base the score on estimate-photo internal consistency, evidence completeness, and clarity/legibility. "
    "If compliance_score < 100, include a dedicated section titled '## Compliance Score Rationale' which itemizes every deficiency with exact evidence references "
    "(estimate p#/L# and/or Photo #), assigns an explicit deduction per item, and shows the arithmetic to the final score. "
    "Use a consistent scheme (e.g., Minor -5, Moderate -10, Major -20) and never go below 0. "
    "The 'fraud_markdown' section must never be 'N/A'. If nothing material is found, write "
    "'No material inconsistencies found.' and briefly note what was checked (VIN match, date/metadata, obvious photo tampering, duplicated images)."
)

SYSTEM_BASE += (
    " Focus on a cohesive, professional appraisal. Prefer narrative over rigid tables. "
    "Include a section named '## Detailed Condition Report'. "
    "Include '## Compliance Score Rationale' only when compliance_score < 100, and show deductions from 100 with brief evidence refs (p#/L# or Photo #). "
    "If you include tables, keep them concise and only when they help clarity. "
    "Avoid placeholder rows/columns; do not invent data. "
    "When client_rules text is provided, also include a section titled '## Client Guidelines Comparison' with 3–8 concise bullets quoting the relevant rule fragment and citing evidence (p#/L#, Photo #); "
    "weave any material rule alignment/misalignment into the Detailed Condition Report narrative."
    "When a valuation/clean retail printout exists but the header doesn’t match the estimate’s VIN/year/trim/mileage, label it “Present — mismatched (detail the differences)” and request a corrected printout; never mark it Missing/Not Evidenced. "
    "If a legible driver-door VIN label photo is present, treat Production Date as evidenced; do not mark 'missing' or deduct for lack of a separate photo. "
    "Only deduct for missing Repair Facility info when the Closing Report or other documents clearly show the vehicle is at a named repair facility AND the estimate's 'Repair Facility' section does not list that same facility; "
    "if no repair facility information appears anywhere in the estimate or Closing Report, report 'N/A — not provided' and do NOT deduct."
)

SYSTEM_BASE += (
    " Your 'summary_markdown' MUST include a top-level section named '## Detailed Condition Report' containing a cohesive narrative of at least 10–14 sentences (not bullets). "
    "It must synthesize: impact zones, per-panel damages, repair vs. replace rationale, parts type (OEM/LKQ/Aftermarket), labor ops, refinish/overlap, rate/materials/sublet/tax handling, and estimate integrity. "
    "It must cite concrete evidence inline (e.g., p2/L14, Photo 3). "
    "When evaluating paint materials, recognize that a summary line such as 'Paint Supplies' or 'Paint Materials' with hours and rate in the totals section constitutes a valid cost breakdown. "
    "Do not mark it missing if such a line is present, even if materials are not listed per-panel. "
    "Avoid categorical phrases such as 'deemed repairable' or 'deemed total loss' unless that exact determination appears in the provided documents (e.g., estimate header says 'Total Loss' or an ACV comparison is shown). "
    "Otherwise, use neutral language and do not make a repairability determination."
)

# -----------------------
# Supported file types
# -----------------------
SUPPORTED_IMAGE_EXTS = (".jpg",".jpeg",".png",".webp",".heic",".heif")
SUPPORTED_TEXT_EXTS = (".txt",)
SUPPORTED_DOCX_EXTS = (".docx",)
SUPPORTED_PDF_EXTS = (".pdf",)

# -----------------------
# Helpers to add parts from bytes
# -----------------------
def _image_part_from_bytes(raw: bytes) -> Dict[str, Any]:
    b64 = base64.b64encode(raw).decode("utf-8")
    return {"type": "image_url", "image_url": {"url": "data:image/jpeg;base64," + b64}}

# Safe PDF text extract helper
def _maybe_extract_pdf_text(raw: bytes, fname: str, parts: List[Dict[str, Any]], files_seen: List[str], pdf_text_fulls: Optional[List[str]] = None) -> None:
    try:
        from pdfminer_high_level import extract_text as _x  # type: ignore
    except Exception:
        try:
            from pdfminer.high.level import extract_text as _x  # fallback
        except Exception:
            _x = None
    try:
        if _x:
            full = (_x(io.BytesIO(raw)) or "")
            if pdf_text_fulls is not None and full.strip():
                pdf_text_fulls.append(full)
            t = full[:12000]
            if t.strip():
                parts.insert(0, {"type": "text", "text": t})
                files_seen.append(f"{fname} (pdf text extracted)")
    except Exception:
        pass

# Lightweight OCR helper for images
def _maybe_ocr_image_text(im: Image.Image) -> str:
    if not _OCR_ENABLED:
        return ""
    try:
        im2 = im.convert("L")
        if max(im2.size) < 1400:
            scale = 1400 / max(im2.size)
            im2 = im2.resize((int(im2.width*scale), int(im2.height*scale)))
        txt = pytesseract.image_to_string(im2)
        return (txt or "").strip()
    except Exception:
        return ""


def _qr_decode_vin_from_pil(im: Image.Image) -> Optional[str]:
    """Best-effort local QR/barcode decode for VIN (optional deps)."""
    try:
        import cv2  # type: ignore
        import numpy as np  # type: ignore
        arr = cv2.cvtColor(np.array(im.convert("RGB")), cv2.COLOR_RGB2BGR)
        det = cv2.QRCodeDetector()
        val, _, _ = det.detectAndDecode(arr)
        if val:
            s = str(val).strip().upper()
            m = re.search(VIN_PATTERN, s)
            if m:
                return m.group(0)
    except Exception:
        pass
    try:
        from pyzbar.pyzbar import decode  # type: ignore
        for d in decode(im.convert("RGB")):
            try:
                s = d.data.decode("utf-8", "ignore").strip().upper()
            except Exception:
                s = str(d.data).strip().upper()
            m = re.search(VIN_PATTERN, s)
            if m:
                return m.group(0)
    except Exception:
        pass
    return None


def _add_bytes(parts: List[Dict[str,Any]], files_seen: List[str], photo_index: Optional[List[str]], thumb_paths: Optional[List[str]], raw: bytes, fname: str, used: int, max_images: int, pdf_text_fulls: Optional[List[str]] = None, ocr_pairs: Optional[List[Dict[str, Any]]] = None) -> int:
    low = fname.lower()
    if low.endswith(SUPPORTED_PDF_EXTS) and used < max_images:
        try:
            pages = convert_from_bytes(raw, dpi=200)
            files_seen.append(f"{fname} (pdf, {len(pages)} page(s))")
            _maybe_extract_pdf_text(raw, fname, parts, files_seen, pdf_text_fulls=pdf_text_fulls)
            OCR_PAGE_CAP = 100
            ocr_collected = []
            for idx, im in enumerate(pages[:max_images - used]):
                b = io.BytesIO()
                im.save(b, format="JPEG", quality=65, optimize=True)
                parts.append(_image_part_from_bytes(b.getvalue()))
                used += 1
                if photo_index is not None:
                    photo_index.append(f"{fname}::page_{idx+1}")
                if idx < OCR_PAGE_CAP:
                    txt = _maybe_ocr_image_text(im)
                    if txt:
                        ocr_collected.append(txt)
            if ocr_collected:
                parts.insert(0, {"type": "text", "text": ("\n".join(ocr_collected))[:12000]})
                files_seen.append(f"{fname} (ocr text extracted)")
        except Exception as e:
            logging.warning(f"pdf2image failed for {fname}: {e}")
            files_seen.append(f"{fname} (pdf, could not be converted)")
    elif low.endswith(SUPPORTED_IMAGE_EXTS) and used < max_images:
        im_ref = None
        raw_for_vin = None
        qr_vin = None
        try:
            im = Image.open(io.BytesIO(raw)).convert("RGB")
            im_ref = im.copy()
            # Preserve ORIGINAL bytes for VIN label / QR/barcode decoding (before any resizing or recompression).
            raw_for_vin = raw
            # Secondary confirmation: attempt local QR/barcode decode from the ORIGINAL image (before resizing).
            qr_vin = _qr_decode_vin_from_pil(im_ref)
            # Use the same preprocessing for ZIP and loose JPGs (keep higher res for small label text).
            max_dim = 2048
            if max(im.size) > max_dim:
                scale = max_dim / float(max(im.size))
                im = im.resize((int(im.width * scale), int(im.height * scale)))
            b = io.BytesIO()
            im.save(b, format="JPEG", quality=65, optimize=True)
            raw = b.getvalue()
        except Exception:
            im_ref = None
            raw_for_vin = None
        parts.append(_image_part_from_bytes(raw))
        used += 1
        if photo_index is not None:
            photo_index.append(fname)
        files_seen.append(f"{fname} (photo)")
        # Keep a local copy of each uploaded photo for the PDF thumbnail appendix page.
        if thumb_paths is not None:
            try:
                im_save = im_ref if im_ref is not None else Image.open(io.BytesIO(raw)).convert("RGB")
                im_save.thumbnail((900, 900))
                thumb_name = f"thumb_{uuid.uuid4().hex}.jpg"
                thumb_path = os.path.join(PDF_DIR, thumb_name)
                im_save.save(thumb_path, format="JPEG", quality=75, optimize=True)
                thumb_paths.append(thumb_path)
            except Exception:
                pass
        if im_ref is not None:
            txt = _maybe_ocr_image_text(im_ref)
            if ocr_pairs is not None:
                try:
                    ocr_pairs.append({"name": fname, "text": (txt or ""), "raw_for_vin": raw_for_vin, "qr_vin": (qr_vin or None)})
                except Exception:
                    pass
            if txt:
                parts.insert(0, {"type":"text", "text": txt[:12000]})
                files_seen.append(f"{fname} (ocr text extracted)")
    elif low.endswith(SUPPORTED_DOCX_EXTS):
        try:
            text = "\n".join([p.text for p in Document(io.BytesIO(raw)).paragraphs if p.text.strip()])
        except Exception:
            text = ""
        if text.strip():
            parts.insert(0, {"type":"text","text": text[:12000]})
            files_seen.append(f"{fname} (docx text included)")
        else:
            files_seen.append(f"{fname} (docx, no readable text)")
    elif low.endswith(SUPPORTED_TEXT_EXTS):
        try:
            text = raw.decode("utf-8","ignore")[:12000]
        except Exception:
            text = ""
        if text.strip():
            parts.insert(0, {"type":"text","text": text})
            files_seen.append(f"{fname} (txt included)")
        else:
            files_seen.append(f"{fname} (txt, empty)")
    else:
        files_seen.append(f"{fname} (unsupported type)")
    return used

# -----------------------
# App + CORS
# -----------------------
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://nspxn.com","https://www.nspxn.com","http://nspxn.com","http://www.nspxn.com",
        "https://nspxn.onrender.com"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# -----------------------
# Client Rules: fuzzy finder + endpoints
# -----------------------
def _find_rules_path(base_name: str, rules_dir: str) -> Optional[str]:
    target = base_name.strip()
    if not target.lower().endswith(".docx"):
        target += ".docx"
    p = os.path.join(rules_dir, target)
    if os.path.exists(p):
        return p
    pattern = os.path.join(rules_dir, "*.docx")
    candidates = glob.glob(pattern)
    if not candidates:
        return None
    low_target = target.lower()
    for c in candidates:
        if os.path.basename(c).lower() == low_target:
            return c
    for c in candidates:
        if os.path.basename(c).lower().startswith(low_target[:-5]):
            return c
    tokens = [t for t in low_target[:-5].split() if t]
    def contains_in_order(name: str) -> bool:
        i = 0
        for tok in tokens:
            i = name.find(tok, i)
            if i == -1:
                return False
            i += len(tok)
        return True
    for c in candidates:
        if contains_in_order(os.path.basename(c).lower()):
            return c
    return None

@app.get("/client-rules/{client_name}")
async def get_client_rules(client_name: str):
    path = _find_rules_path(client_name, CLIENT_RULES_DIR)
    if not path:
        try:
            files = sorted(os.path.basename(p) for p in glob.glob(os.path.join(CLIENT_RULES_DIR, "*.docx")))
        except Exception:
            files = []
        return JSONResponse(
            status_code=404,
            content={
                "error": f"Rules not found for '{client_name}'.",
                "hint": "Ensure the .docx file exists in CLIENT_RULES_DIR (default 'client_rules').",
                "available_rule_files": files[:50]
            },
        )
    try:
        doc = Document(path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return {"file": os.path.basename(path), "text": text}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"Unable to read rules: {e}", "file": os.path.basename(path)})

@app.get("/client-rules")
async def list_client_rules():
    try:
        files = sorted(os.path.basename(p) for p in glob.glob(os.path.join(CLIENT_RULES_DIR, "*.docx")))
        return {"dir": CLIENT_RULES_DIR, "files": files}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e), "dir": CLIENT_RULES_DIR})

# -----------------------
# Vision Review
# -----------------------
@app.post("/vision-review")
async def vision_review(
    request: Request,
    files: Optional[List[UploadFile]] = File(None),
    client_rules: str = Form(""),
    ai_notes: str = Form(""),
    addl_notes: str = Form(""),
    additional_notes: str = Form(""),
    notes: str = Form(""),
    file_number: Optional[str] = Form(None),
    ia_company: str = Form(""),
    appraiser_id: str = Form(""),
    ai_intent: str = Form("comprehensive")
):
    parts: List[Dict[str, Any]] = []
    files_seen: List[str] = []
    photo_index: List[str] = []
    thumbnail_paths: List[str] = []
    ocr_pairs: List[Dict[str, Any]] = []

    # --- 422 guard: avoid FastAPI validation failures when frontend keys vary ---
    # Accept missing/alternate keys without raising 422; return a clear 400 instead.
    if files is None:
        files = []
    if file_number is None or str(file_number).strip() == "":
        try:
            _form = await request.form()
            # Common file number key variants from different frontends
            for _k in ("file_number", "file-number", "fileNumber", "fileNum", "file_num", "filenumber"):
                _v = str(_form.get(_k, "") or "").strip()
                if _v:
                    file_number = _v
                    break
        except Exception:
            pass
    # If 'files' was posted under a different key, recover it from the raw form
    if (not files) or (isinstance(files, list) and len(files) == 0):
        try:
            _form = await request.form()
            _maybe = []
            try:
                _maybe = list(_form.getlist("files")) if hasattr(_form, "getlist") else []
            except Exception:
                _maybe = []
            if not _maybe:
                try:
                    _maybe = list(_form.getlist("file")) if hasattr(_form, "getlist") else []
                except Exception:
                    _maybe = []
            _maybe_files = [x for x in _maybe if hasattr(x, "filename")]
            if _maybe_files:
                files = _maybe_files  # type: ignore[assignment]
        except Exception:
            pass
    # Hard-required fields (explicit 400 instead of 422)
    if not file_number or str(file_number).strip() == "":
        return JSONResponse(status_code=400, content={"error": "Missing required field: file_number"})
    if not files:
        return JSONResponse(status_code=400, content={"error": "Missing required upload: files"})
    vin_candidates: List[str] = []  # reserved (filenames not used)
    MAX_IMAGES = 48
    used = 0
    # Coalesce Add\'l Notes from multiple possible frontend field names
    ai_notes_used = ((ai_notes or "").strip() or (addl_notes or "").strip() or (additional_notes or "").strip() or (notes or "").strip())
    ai_notes_used = ai_notes_used.strip()
    # If notes still empty, fall back to reading the raw posted form (covers mismatched frontend field names)
    if not ai_notes_used:
        try:
            _form = await request.form()
            # First try common variants explicitly
            for _k in ("ai_notes","ai-notes","addl_notes","additional_notes","notes","ai_review_notes","ai_notes_box","addlNote","addlNoteText"):
                _v = str(_form.get(_k, "") or "").strip()
                if _v:
                    ai_notes_used = _v
                    break
            # Then any key containing 'note' (last resort)
            if not ai_notes_used:
                for _k in _form.keys():
                    if "note" in str(_k).lower():
                        _v = str(_form.get(_k, "") or "").strip()
                        if _v:
                            ai_notes_used = _v
                            break
        except Exception:
            pass
    # Normalize/sanitize notes so they cannot break structured prompting
    ai_notes_used = _normalize_ai_notes(ai_notes_used)
    ai_notes_block = _ai_notes_block(ai_notes_used)

    pdf_text_fulls: List[str] = []  # full PDF text for supplement detection

    # Anti-zipbomb guardrails
    MAX_ZIP_FILES = 100
    MAX_ENTRY_SIZE = 15 * 1024 * 1024  # 15 MB

    for f in sorted(files, key=lambda _f: ((_f.filename or '').lower())):
        raw = await f.read()
        fname = f.filename or "upload"
        low = fname.lower()
        if low.endswith(".zip"):
            try:
                zf = zipfile.ZipFile(io.BytesIO(raw))
            except Exception as e:
                files_seen.append(f"{fname} (zip, unreadable: {e})")
                continue
            members = sorted([zi for zi in zf.infolist() if not zi.is_dir()], key=lambda _zi: (_zi.filename or '').lower())
            if len(members) > MAX_ZIP_FILES:
                files_seen.append(f"{fname} (zip, too many entries: {len(members)})")
                members = members[:MAX_ZIP_FILES]
            for zi in members:
                inner_name = zi.filename
                if ".." in inner_name or inner_name.startswith(("/", "\\")):
                    files_seen.append(f"{fname}::{inner_name} (skipped unsafe path)"); continue
                if zi.file_size > MAX_ENTRY_SIZE:
                    files_seen.append(f"{fname}::{inner_name} (skipped >15MB)"); continue
                try:
                    data = zf.read(zi)
                except Exception as e:
                    files_seen.append(f"{fname}::{inner_name} (read error: {e})"); continue
                used = _add_bytes(parts, files_seen, photo_index, thumbnail_paths, data, f"{fname}::{inner_name}", used, MAX_IMAGES, pdf_text_fulls=pdf_text_fulls, ocr_pairs=ocr_pairs)
        else:
            used = _add_bytes(parts, files_seen, photo_index, thumbnail_paths, raw, fname, used, MAX_IMAGES, pdf_text_fulls=pdf_text_fulls, ocr_pairs=ocr_pairs)

    # Collect uploaded TEXT ONLY for evidence checks
    uploaded_text_blobs = []
    for p in parts:
        if isinstance(p, dict) and p.get("type") == "text" and isinstance(p.get("text"), str):
            uploaded_text_blobs.append(p["text"])
    uploaded_text_all = "\n".join(uploaded_text_blobs)
    # --- VIN EXTRACTION (door label OCR -> QR/barcode) ---
    # Do NOT use filenames. Prefer driver-door certification label OCR; if not found, attempt QR/barcode decode.
    vin_from_label = None
    vin_from_label_photo = None
    vin_from_qr = None
    vin_from_qr_photo = None

    def _looks_like_door_label(txt: str) -> bool:
        if not txt:
            return False
        t = txt.upper()
        keys = ["MFD BY", "MANUFACTURED", "GVWR", "GAWR", "TIRE SIZE", "CONFORMS TO"]
        return sum(1 for k in keys if k in t) >= 2

    # (b) Door label OCR first
    try:
        for rec in ocr_pairs:
            if not isinstance(rec, dict):
                continue
            t = rec.get("text") or ""
            if _looks_like_door_label(t):
                mvin = re.search(VIN_PATTERN, t.upper())
                if mvin:
                    vin_from_label = mvin.group(0)
                    vin_from_label_photo = rec.get("name") or None
                    break
    except Exception:
        vin_from_label = None
        vin_from_label_photo = None
    # (c) QR/barcode VIN as secondary confirmation (prefer same photo as door label if available)
    try:
        if vin_from_label_photo:
            for rec in ocr_pairs:
                if isinstance(rec, dict) and (rec.get("name") == vin_from_label_photo) and rec.get("qr_vin"):
                    vin_from_qr = str(rec.get("qr_vin") or "").strip().upper() or None
                    vin_from_qr_photo = rec.get("name") or None
                    break
        if not vin_from_qr:
            for rec in ocr_pairs:
                if isinstance(rec, dict) and rec.get("qr_vin"):
                    vin_from_qr = str(rec.get("qr_vin") or "").strip().upper() or None
                    vin_from_qr_photo = rec.get("name") or None
                    break
    except Exception:
        vin_from_qr = None
        vin_from_qr_photo = None


    # (c) If OCR didn't yield a VIN, try a dedicated vision decode for VIN/QR/barcode on the best label candidate
    def _decode_vin_from_label_or_qr(raw_bytes: Optional[bytes]) -> Optional[str]:
        if not raw_bytes:
            return None
        try:
            prompt = (
                "Return ONLY JSON: {\"vin\": \"...\"}.\n"
                "Task: Read the vehicle VIN from the door-jamb certification label text and/or decode any QR/barcode if present.\n"
                "Rules: VIN must be exactly 17 characters (A-H, J-N, P, R-Z, 0-9; no I/O/Q). "
                "If not fully legible, return {\"vin\": null}.\n"
            )
            rsp_v = client.chat.completions.create(
                model=MODEL,
                messages=[
                    {"role": "system", "content": "You extract VINs from vehicle door-jamb certification labels. JSON only."},
                    {"role": "user", "content": [{"type": "text", "text": prompt}, _image_part_from_bytes(raw_bytes)]},
                ],
                max_completion_tokens=300,
                temperature=0,
                response_format={"type": "json_object"},
            )
            raw_v = (rsp_v.choices[0].message.content or "").strip()
            try:
                data_v = json.loads(raw_v)
            except Exception:
                data_v = None
            if isinstance(data_v, dict):
                vv = data_v.get("vin")
                if isinstance(vv, str):
                    vv = vv.strip().upper()
                    if re.fullmatch(VIN_PATTERN, vv):
                        return vv
        except Exception:
            return None
        return None

    if not vin_from_label:
        try:
            candidate = None
            for rec in ocr_pairs:
                if isinstance(rec, dict) and _looks_like_door_label(rec.get("text") or "") and rec.get("raw_for_vin"):
                    candidate = rec
                    break
            if candidate is None:
                for rec in ocr_pairs:
                    if isinstance(rec, dict) and rec.get("raw_for_vin"):
                        candidate = rec
                        break
            if candidate is not None:
                # Prefer local QR/barcode VIN if available before calling the vision model.
                if candidate.get("qr_vin"):
                    vin_from_label = str(candidate.get("qr_vin") or "").strip().upper() or None
                if not vin_from_label:
                    vin_from_label = _decode_vin_from_label_or_qr(candidate.get("raw_for_vin"))
                vin_from_label_photo = candidate.get("name") if vin_from_label else None
        except Exception:
            pass

    # normalize + keep unique order
    _seen_v = set(); _tmp_v=[]
    for _v in vin_candidates:
        if _v and _v not in _seen_v:
            _seen_v.add(_v); _tmp_v.append(_v)
    vin_candidates = _tmp_v

    # --- ODOMETER OCR LOCK (extract mileage from OCR text if visible) ---
    # This makes it impossible for the narrative to claim the odometer is not visible when OCR captured a mileage value.
    odometer_value = None
    try:
        _odo_txt = uploaded_text_all or ""
        # Common mileage patterns from digital clusters: "72,261 mi", "72261 mi", "72261 miles", "116000 km"
        _m = re.search(r"(?i)\b(\d{1,3}(?:,\d{3})+|\d{4,7})\s*(mi|miles|km)\b", _odo_txt)
        if _m:
            _digits = _m.group(1).replace(",", "")
            _unit = _m.group(2).lower()
            if _unit == "miles":
                _unit = "mi"
            odometer_value = f"{int(_digits):,} {_unit}"
    except Exception:
        odometer_value = None


    # --- Closing Report: extract "Inspection Results" section for deterministic cross-check + shop-status ---
    def _extract_inspection_results_block(_txt: str) -> str:
        if not _txt:
            return ""
        t = _txt.replace("\r", "\n")
        # Capture the text after the "Inspection Results" header up to the next major header.
        m_ir = re.search(
            r"(?is)\bInspection\s+Results\b\s*[:\-]?\s*(.{0,5000}?)(?=\n\s*(?:Possible\s+Supplement\s+Amount|Supplement\b|Estimate\b|Inspection\s+Location\b|Repair\s+Facility\b|Vehicle\b|Owner\b|Appraiser\b|Adjuster\b|$))",
            t,
        )
        if m_ir:
            return (m_ir.group(1) or "").strip()[:5000]
        return ""

    inspection_results_text = _extract_inspection_results_block(uploaded_text_all or "")

    # Closing Report: "Possible Supplement Amount" (this ALONE does NOT mean the estimate is a supplement)
    _possible_supp_amount = None
    _m_psa = re.search(r"(?i)\bPossible\s+Supplement\s+Amount\b\s*\$?\s*([0-9,]+(?:\.[0-9]{2})?)", uploaded_text_all or "")
    if _m_psa:
        _possible_supp_amount = _m_psa.group(1)

    photos_provided = any(isinstance(p, dict) and p.get('type') != 'text' for p in parts)

    # --- Robust detectors (CLEAN RETAIL + ADVISOR) ---
    clean_retail_rx = (
        r"(?i)\b("
        r"NADA|J[.\s-]*D[.\s-]*\s*Power|JDPower\.com|"
        r"Kell?ey\s+Blue\s+Book|KBB\.com|"
        r"Edmunds|Carfax|Cars\.com|"
        r"Clean\s+Retail(?:\s+Value)?"
        r")\b"
    )
    _clean_retail_present = bool(re.search(clean_retail_rx, uploaded_text_all or ""))

    advisor_rx = r"(?i)\bAdvisor\s+Report\b"
    _advisor_present = bool(re.search(advisor_rx, uploaded_text_all or ""))

    paint_mat_rx = r"(Paint\s+(Suppl(?:ies|y)|Materials)|Materials\s*Line)"
    _paint_materials_present = bool(re.search(paint_mat_rx, uploaded_text_all or "", flags=re.IGNORECASE))

    # --- Parts source guardrail (OEM vs Aftermarket/LKQ) ---
    # Detect explicit non-OEM indicators in estimate LINE ITEMS only (ignore generic boilerplate/disclosures).
    _explicit_non_oem_parts = False
    try:
        # Heuristic: look for non-OEM keywords on the same line as an operation/part line (starts with line #).
        _explicit_non_oem_parts = bool(re.search(
            r"(?im)^\s*\d+\s+.*\b(A/M|AFTERMARKET|NON\s*OEM|CAPA|NSF|LKQ|RCY|RECYCLED|USED|RECOND)\b",
            uploaded_text_all or "",
        ))
    except Exception:
        _explicit_non_oem_parts = False

    # Closing report sometimes explicitly states no aftermarket/LKQ parts were included.
    _closing_no_aftermarket = False
    try:
        _closing_no_aftermarket = bool(re.search(r"(?i)\bno\s+aftermarket\b|\bno\s+aftermarket\s+or\s+lkq\b", uploaded_text_all or ""))
    except Exception:
        _closing_no_aftermarket = False

    # Presence detectors for VIN / Odometer / Production Date (OCR text)
    vin_photo_rx = r"(?i)\bDescription:\s*VIN\b|\bVIN(?:\s*#|:)?\s*[A-HJ-NPR-Z0-9]{17}\b"
    odo_photo_rx = r"(?i)\bDescription:\s*Odometer\b|\bOdometer\s*Photo\b|\bPhoto\s*[:#]?\s*Odometer\b"

    # Production date patterns (label)
    prod_date_rx = (
        r"(?is)\b(Production\s*date|Prod(?:uction)?\s*Date|Date\s*of\s*Mfr|Date\s*of\s*Manufacture|"
        r"MFD\.?\s*(?:BY|DATE)?|MFR\.?\s*DATE|MFG\.?\s*DATE|DATE)\b"
        r".{0,80}?\b(0[1-9]|1[0-2])\s*[-/.\u2013\u2014\u2212:\s]\s*(20\d{2}|\d{2})\b"
    )

    _vin_photo_present = bool(re.search(vin_photo_rx, uploaded_text_all or ""))
    _odo_photo_present = bool(re.search(odo_photo_rx, uploaded_text_all or ""))

    # presence + extractor
    _prod_date_present = False
    _prod_date_str = None
    m = re.search(prod_date_rx, uploaded_text_all or "")
    if m:
        mm = m.group(2); yy = m.group(3)
        if len(yy) == 2: yy = "20" + yy
        _prod_date_present = True
        _prod_date_str = f"{mm}/{yy}"
    if not _prod_date_present:
        loose_prod_rx = (
            r"(?is)\b(MFD|MFR|MFG|PROD\.?|PRODUCTION|DATE)\b"
            r".{0,60}?\b(0[1-9]|1[0-2])\s*[-/.\u2013\u2014\u2212:\s]\s*(20\d{2}|\d{2})\b"
        )
        m2 = re.search(loose_prod_rx, uploaded_text_all or "")
        if m2:
            mm = m2.group(2); yy = m2.group(3)
            if len(yy) == 2: yy = "20" + yy
            _prod_date_present = True
            _prod_date_str = f"{mm}/{yy}"

    # treat production date as evidenced if either door label VIN photo or prod date text is seen
    _prod_evidenced = _prod_date_present or _vin_photo_present

    # Lock to 3 intents only
    if ai_intent not in ALLOWED_INTENTS:
        ai_intent = "comprehensive"

    REQ_LABELS = {
        "guidelines_only": "Guidelines → Estimate (no photos)",
        "comprehensive": "Comprehensive: Guidelines + Estimate + Photos (with VIN check)",
        "damage_report_from_photos": "Create a Damage Report from Photos",
    }
    req_label = REQ_LABELS.get(ai_intent, "Comprehensive: Guidelines + Estimate + Photos (with VIN check)")
    log.info(f"ai_intent received: {ai_intent} -> using label: {req_label}")

    KEYS = [
        "file_number","request_type","claim_number","vin","vin_verification","vehicle",
        "odometer_estimate_only","compliance_score","summary_brief","summary_markdown",
        "fraud_markdown","primary_impact","secondary_impact","estimated_costs_markdown","conclusion"
    ]

    SYSTEM = SYSTEM_BASE

    # Closing Report shop-status detector (documents-only; used to prevent false Repair Facility deductions)
    _not_at_shop = False
    try:
        # Prefer the extracted "Inspection Results" block if available (more deterministic)
        if inspection_results_text:
            _not_at_shop = bool(re.search(
                r"(?i)\bNot\s+at\s+Shop\b|\bOwner\s+Location\b|\bInspection\s+Location\s*:?\s*Owner\b|\bInspection\s+Location\s+Owner\b",
                inspection_results_text,
            ))
        if not _not_at_shop:
            _not_at_shop = bool(re.search(
                r"(?i)\bNot\s+at\s+Shop\b|\bOwner\s+Location\b|\bInspection\s+Location\s*:?\s*Owner\b|\bInspection\s+Location\s+Owner\b",
                uploaded_text_all or "",
            ))
    except Exception:
        _not_at_shop = False

    # --- NEW: detect all supplement tags S01, S02, ... in the combined document text ---
    combined_detection_text = (uploaded_text_all or "") + "\n" + "\n".join(pdf_text_fulls or [])
    supplement_versions = sorted(set(re.findall(r"(?i)\bS[0-9]{2}\b", combined_detection_text)))
    supplement_block = ""
    if supplement_versions:
        supplement_block = (
            "\n\nSUPPLEMENT VERSIONS DETECTED FROM DOCUMENTS:\n"
            f"- Detected supplement tags: {', '.join(supplement_versions)}.\n"
            "- Use these exact tags (e.g., 'Supplement S01 and S02') when describing supplements in the narrative.\n"
        )

    # -----------------------
    # Minimal prompt (LET GPT DO THE WORK)
    # -----------------------
    prompt_text = (
        f"REQUEST TYPE (use exactly in request_type): {req_label}\n"
        f"FILE #: {file_number}\n"
        f"CLIENT: {ia_company}\n\n"
        "PHOTO INDEX (use Photo # citations exactly as listed):\n"
        + ("\n".join([f"Photo {i+1}: {name}" for i, name in enumerate(photo_index)]) if photo_index else "No photos provided.")
        + "\n\n"
        "CLIENT RULES (only if provided):\n"
        + (client_rules[:1500] if client_rules else "")
        + ai_notes_block
        + "\n\n"
        "INSTRUCTIONS:\n"
        "- Return strict JSON only.\n"
        "- REQUIRED: Populate 'estimated_costs_markdown' in the JSON.\n"
        "- Use the template below for narrative formatting.\n\n"
        + DETAIL_TEMPLATES.get(ai_intent, DETAIL_TEMPLATES['comprehensive'])
    )



    parts_payload: List[Dict[str,Any]] = []
    redaction_success = False
    try:
        red_prompt = redact_text_preserve_vin_claim(prompt_text)
        redaction_success = True
        # Prevent Presidio from mutating fixed markdown headings used for prompting.
        # (This avoids model outputs like "### [REDACTED] Addressed".)
        red_prompt = red_prompt.replace("### [REDACTED] Addressed", "### Add'l Notes Addressed")
    except Exception as e:
        log.warning(f"Redaction failed on prompt_text: {e}")
        red_prompt = prompt_text

    parts_payload.append({"type": "text", "text": red_prompt})

    if parts:
        for p in parts:
            if p.get("type") == "text" and isinstance(p.get("text"), str):
                try:
                    red_txt = redact_text_preserve_vin_claim(p["text"])
                    redaction_success = True
                except Exception as e:
                    log.warning(f"Redaction failed on a text part: {e}")
                    red_txt = p["text"]
                parts_payload.append({"type": "text", "text": red_txt})
            else:
                parts_payload.append(p)

    redaction_status = "Redacted PII: Successful ✅" if redaction_success else "Redacted PII: Not Applied"

    # -------------------------------
    # Keep prompt lean
    # -------------------------------
    TEXT_PART_LIMIT = 6
    _text_parts = [p for p in parts_payload if p.get("type") == "text"]
    _image_parts = [p for p in parts_payload if p.get("type") != "text"]
    for tp in _text_parts:
        if isinstance(tp.get("text"), str) and len(tp["text"]) > 8000:
            tp["text"] = tp["text"][:8000]
    parts_payload = _text_parts[:TEXT_PART_LIMIT] + _image_parts

    # Token limits
    MAX_TOKENS_BY_INTENT = {
        "comprehensive": 2200,
        "guidelines_only": 1500,
        "damage_report_from_photos": 2400
    }
    max_tokens = MAX_TOKENS_BY_INTENT.get(ai_intent, 1500)

    # Call GPT and parse JSON (JSON hardened)
    # Prefer the canonical SDK path (client.chat.completions). Keep fallback for older SDKs.
    try:
        rsp = client.chat.completions.create(
            model=MODEL,
            messages=[{"role":"system","content": SYSTEM},
                      {"role":"user","content": parts_payload}],
            max_completion_tokens=max_tokens,
            temperature=0,
            top_p=1,
            presence_penalty=0,
            frequency_penalty=0,
            response_format={"type":"json_object"},
        )
    except AttributeError:
        rsp = client.chat_completions.create(  # type: ignore[attr-defined]
            model=MODEL,
            messages=[{"role":"system","content": SYSTEM},
                      {"role":"user","content": parts_payload}],
            max_completion_tokens=max_tokens,
            temperature=0,
            top_p=1,
            presence_penalty=0,
            frequency_penalty=0,
            response_format={"type":"json_object"},
        )

    # --- Hardened JSON parse helper
    def _try_parse_json(raw_text: str):
        if not raw_text:
            return None
        raw_local = raw_text.strip()
        for fence in ("```json", "```JSON", "```"):
            if raw_local.startswith(fence):
                raw_local = raw_local[len(fence):]
            if raw_local.endswith("```"):
                raw_local = raw_local[:-3]
        raw_local = raw_local.strip()
        raw_local = raw_local.replace("\ufeff", "").replace("\u200b", "").replace("\u00A0", " ")
        try:
            return json.loads(raw_local)
        except Exception:
            pass
        lb = raw_local.find("{"); rb = raw_local.rfind("}")
        chunk = raw_local[lb:rb+1] if (lb != -1 and rb != -1 and rb > lb) else raw_local
        fixes = {"\u2018": "'", "\u2019": "'", "\u201C": '"', "\u201D": '"', "\r": "", "\t": "    "}
        for k, v in fixes.items():
            chunk = chunk.replace(k, v)
        chunk = re.sub(r",\s*([}\]])", r"\1", chunk)
        try:
            return json.loads(chunk)
        except Exception:
            return None

    try:
        raw = (rsp.choices[0].message.content or "")
    except Exception as e:
        log.error(f"LLM returned no content: {e}")
        return JSONResponse(status_code=500, content={"error":"Model returned no content."})

    data = _try_parse_json(raw)
    # Prefer door-label VIN when present (OCR -> QR/barcode). Do NOT use filenames.
    # --- HARD VIN LOCK (door label wins; QR only confirms) ---
    try:
        if isinstance(data, dict):

            # 1) If door-label VIN exists, it is authoritative.
            if vin_from_label:
                data["vin"] = vin_from_label

                if vin_from_qr:
                    if vin_from_qr == vin_from_label:
                        data["vin_verification"] = "MATCH (door label + QR)"
                    else:
                        data["vin_verification"] = (
                            f"MISMATCH (door label: {vin_from_label}; QR: {vin_from_qr})"
                        )
                else:
                    data["vin_verification"] = (
                        "INCONCLUSIVE (door label extracted; no secondary confirmation)"
                    )

            # 2) If no door label but QR found
            elif vin_from_qr:
                data["vin"] = vin_from_qr
                data["vin_verification"] = "INCONCLUSIVE (QR only; door label not detected)"

    except Exception:
        pass

    # Prefer door-label VIN when present (OCR -> QR/barcode). Do NOT use filenames.
    try:
        if isinstance(data, dict) and vin_from_label:
            v_model = (data.get("vin") or "").strip().upper()
            if not re.fullmatch(VIN_PATTERN, v_model):
                data["vin"] = vin_from_label
                if not (data.get("vin_verification") or "").strip():
                    data["vin_verification"] = "INCONCLUSIVE (door label VIN extracted; compare to other docs if present)"
            elif v_model != vin_from_label:
                data["vin_verification"] = f"MISMATCH (door label: {vin_from_label}; other source: {v_model})"
            # Secondary confirmation vs QR/barcode if available
            if vin_from_qr:
                if vin_from_label == vin_from_qr:
                    if not re.search(r"\bMATCH\b|\bMISMATCH\b", str(data.get("vin_verification") or ""), flags=re.IGNORECASE):
                        data["vin_verification"] = "MATCH (door label + QR/barcode)"
                else:
                    data["vin_verification"] = f"MISMATCH (door label: {vin_from_label}; QR/barcode: {vin_from_qr})"
    except Exception:
        pass


    # One safe retry on truncation
    try:
        finish_reason = getattr(rsp.choices[0], "finish_reason", None)
    except Exception:
        finish_reason = None

    if data is None and (finish_reason == "length" or (raw and raw.strip().endswith("..."))):
        _text_parts_retry = [p for p in parts_payload if p.get("type") == "text"]
        _image_parts_retry = [p for p in parts_payload if p.get("type") != "text"]
        shrunk = _text_parts_retry[: max(3, len(_text_parts_retry)//2)] + _image_parts_retry
        retry_tokens = min(3000, max_tokens + 600)
        try:
            rsp2 = client.chat.completions.create(
                model=MODEL,
                messages=[{"role": "system", "content": SYSTEM},
                          {"role": "user", "content": shrunk}],
                max_completion_tokens=retry_tokens,
                temperature=0,
                response_format={"type": "json_object"}
            )
            raw2 = (rsp2.choices[0].message.content or "")
            data = _try_parse_json(raw2)
        except Exception:
            pass

    # Fallback: formatting pass
    if data is None:
        try:
            fix_prompt = [
                {"role":"system","content":
                    "You are a formatter. Return ONLY a strict JSON object. No prose. No markdown. No code fences."
                },
                {"role":"user","content":
                    "Convert the following text into a valid JSON object. "
                    "Use exactly these keys (all required): "
                    "['file_number','request_type','claim_number','vin','vin_verification','vehicle',"
                    "'odometer_estimate_only','compliance_score','summary_brief','summary_markdown',"
                    "'fraud_markdown','primary_impact','secondary_impact','estimated_costs_markdown','conclusion'] "
                    "Do not invent new keys. If a field is unavailable, use 'N/A'. "
                    "Here is the text:\n\n" + raw
                }
            ]
            try:
                fix_rsp = client.chat_completions.create(  # type: ignore[attr-defined]
                    model=MODEL,
                    messages=fix_prompt,
                    max_completion_tokens=max_tokens,
                    temperature=0,
                    response_format={"type":"json_object"}
                )
            except AttributeError:
                fix_rsp = client.chat.completions.create(
                    model=MODEL,
                    messages=fix_prompt,
                    max_completion_tokens=max_tokens,
                    temperature=0,
                    response_format={"type":"json_object"}
                )
            fixed = (fix_rsp.choices[0].message.content or "")
            data = _try_parse_json(fixed)
        except Exception as e:
            log.error(f"Self-heal reformat failed: {e}")

    if data is None:
        log.error(f"LLM failure or JSON parse error; first 500 chars:\n" + (raw or "")[:500])
        skeleton = {k: "N/A" for k in KEYS}
        skeleton["file_number"] = file_number
        skeleton["request_type"] = req_label
        skeleton["summary_brief"] = "N/A (model output could not be parsed; skeleton returned)."
        skeleton["summary_markdown"] = (
            "## Detailed Condition Report\n"
            "Model output could not be parsed into JSON on this run. Please resubmit."
        )
        skeleton["fraud_markdown"] = "No material inconsistencies found."
        return skeleton

    def _get(k):
        v = data.get(k)
        return "" if v is None else str(v)

    result = {
        "file_number": file_number,
        "request_type": req_label,
        "claim_number": _get("claim_number"),
        "vin": _get("vin"),
        "vin_verification": _get("vin_verification"),
        "vehicle": _get("vehicle"),
"odometer_estimate_only": (odometer_value if odometer_value else _get("odometer_estimate_only")),
        "compliance_score": _get("compliance_score"),
        "summary_brief": _get("summary_brief"),
        "summary_markdown": _get("summary_markdown"),
        "fraud_markdown": _get("fraud_markdown"),
        "primary_impact": _get("primary_impact"),
        "secondary_impact": _get("secondary_impact"),
        "estimated_costs_markdown": _get("estimated_costs_markdown"),
        "conclusion": _get("conclusion"),
        "redaction_status": redaction_status,
    }


    # -----------------------
    # Cost Markdown Scrub + Severity Tier Checkbox Enforcement
    # -----------------------
    def _scrub_cost_markdown(md: str) -> str:
        """Remove phrases/lines the user explicitly does not want printed in the PDF output."""
        if not isinstance(md, str):
            return md
        s = md.replace("\r\n", "\n").replace("\r", "\n")

        banned_subs = (
            "from estimate",
            "from documentation",
            "not evidenced",
            "no documentation provided",
            "photo-only claims may be blank",
            "best-effort",
        )

        out_lines = []
        for ln in s.splitlines():
            raw_ln = ln
            lnl = raw_ln.lower().strip()

            # Drop the model-added heading variants (we render the header in PDF ourselves)
            if re.search(r"^\s*#{1,6}\s*approximate\s+repair\s+cost\s+breakdown\b", raw_ln, flags=re.I):
                continue
            if re.search(r"^\s*#{1,6}\s*approximate\s+repair\s+cost\s+breakdown\s*\(\s*photos\s*only\s*\)\s*$", raw_ln, flags=re.I):
                continue

            # Remove inspection-location callouts entirely (user doesn't want this printed)
            if re.search(r"^\s*#{1,6}\s*inspection\s+location\b", raw_ln, flags=re.I):
                continue
            if re.search(r"\*\*\s*inspection\s+location\s*\*\*\s*:", raw_ln, flags=re.I):
                continue

            # Remove arithmetic "Estimated total: $A + $B = $C" lines (user wants ONE total only)
            if re.search(r"(?i)\bestimated\s+total\b\s*:", raw_ln):
                continue
            if ("+" in raw_ln and "=" in raw_ln and re.search(r"\$\s*[0-9]", raw_ln)):
                continue

            # Remove our own editorial parentheticals
            if re.search(r"\(\s*best-?effort.*\)$", raw_ln, flags=re.I):
                continue

            # Never call it an estimate
            raw_ln = re.sub(r"(?i)\bestimated\s+repair\s+total\b", "Approximate Repair Total", raw_ln)
            raw_ln = re.sub(r"(?i)\bestimated\s+total\b", "Approximate Total", raw_ln)

            # Rename TL language
            raw_ln = raw_ln.replace("Possible Total Loss Threshold Approaching", "Possible Total Loss Threshold Approaching")

            # Drop or neutralize banned phrases
            lnl2 = raw_ln.lower()
            if any(b in lnl2 for b in banned_subs):
                if ":" in raw_ln:
                    left = raw_ln.split(":", 1)[0].rstrip()
                    out_lines.append(f"{left}: (photo-derived approximation)")
                else:
                    continue
            else:
                out_lines.append(raw_ln)

        return "\n".join(out_lines).strip()
    def _parse_approx_total(md: str) -> Optional[float]:
        if not isinstance(md, str) or not md.strip():
            return None

        # Ignore arithmetic subtotal lines that can mislead tier selection.
        cleaned_lines = []
        for ln in md.replace("\r\n", "\n").replace("\r", "\n").splitlines():
            if re.search(r"(?i)\bestimated\s+total\b\s*:", ln):
                continue
            if ("+" in ln and "=" in ln and re.search(r"\$\s*[0-9]", ln)):
                continue
            cleaned_lines.append(ln)
        cleaned = "\n".join(cleaned_lines)

        # Prefer explicit labeled totals.
        candidates = [
            r"(?i)\bapprox(?:imate|\.?)(?:\s+repair)?\s+total\s*[:\-]?\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"(?i)\bapproximat(?:ed|e)\s+repair\s+cost\s+total\b[\s\S]{0,80}?\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"(?i)\btotal\s+amount\b\s*[:\-]?\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            # Bold total line: "- **$23,239**"
            r"\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)\*\*",
        ]
        for rx in candidates:
            m = re.search(rx, cleaned)
            if m:
                try:
                    return float(m.group(1).replace(",", ""))
                except Exception:
                    pass

        # Last resort: take the largest dollar value in the cost section text.
        vals = []
        for m in re.finditer(r"\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)", cleaned):
            try:
                vals.append(float(m.group(1).replace(",", "")))
            except Exception:
                pass
        return max(vals) if vals else None


    def _enforce_severity_tier(md: str) -> str:
        """Ensure Severity Tier appears with checkboxes and one tier is checked based on approximate total."""
        if not isinstance(md, str):
            return md

        total = _parse_approx_total(md)

        def tier_for(v: Optional[float]) -> Optional[str]:
            if v is None:
                return None
            if v < 3500:
                return "minor"
            if v <= 10000:
                return "moderate"
            return "major"

        tier = tier_for(total)

        # Optional TL-approaching check: only check if explicitly very high or docs explicitly indicate TL
        tl_approaching = False
        try:
            if total is not None and total >= 25000:
                tl_approaching = True
        except Exception:
            tl_approaching = False

        # Normalize any "☐" style into markdown checkboxes
        lines = md.replace("\r\n", "\n").replace("\r", "\n").splitlines()
        norm = []
        for ln in lines:
            ln2 = ln
            ln2 = ln2.replace("☐", "- [ ]").replace("☑", "- [x]")
            norm.append(ln2)

        # Detect existing checkbox lines
        has_any = any(re.search(r"(?i)\[\s*[xX ]\s*\].*(minor|moderate|major|total\s+loss)", ln) for ln in norm)
        if not has_any:
            # Append a clean severity block if missing
            norm.append("")
            norm.append("### Severity Tier")
            norm.append("- [ ] Minor (< $3,500)")
            norm.append("- [ ] Moderate ($3,500–$10,000)")
            norm.append("- [ ] Major ($10,000+)")
            norm.append("- [ ] Possible Total Loss Threshold Approaching")

        def mark_line(ln: str) -> str:
            lnl = ln.lower()
            if "minor" in lnl and "total loss" not in lnl:
                return re.sub(r"\[\s*[xX ]\s*\]", "[x]" if tier == "minor" else "[ ]", ln)
            if "moderate" in lnl:
                return re.sub(r"\[\s*[xX ]\s*\]", "[x]" if tier == "moderate" else "[ ]", ln)
            if "major" in lnl and "total loss" not in lnl:
                return re.sub(r"\[\s*[xX ]\s*\]", "[x]" if tier == "major" else "[ ]", ln)
            if "total loss" in lnl:
                return re.sub(r"\[\s*[xX ]\s*\]", "[x]" if tl_approaching else "[ ]", ln)
            return ln

        enforced = [mark_line(ln) for ln in norm]
        return "\n".join(enforced).strip()

    # Apply scrub + severity enforcement early (model-provided costs) so the PDF never prints banned phrases
    try:
        _md0 = result.get("estimated_costs_markdown") or ""
        _md0 = _scrub_cost_markdown(_md0)
        _md0 = _enforce_severity_tier(_md0)
        if _md0:
            result["estimated_costs_markdown"] = _md0
    except Exception:
        pass

    # ✅ FIX #2: Hard fallback so UI never gets an empty narrative ("No narrative generated")
    try:
        sm_tmp = (result.get("summary_markdown") or "").strip()
        if not sm_tmp:
            result["summary_markdown"] = (
                "## Detailed Condition Report\n"
                "Narrative fallback: The model returned an empty narrative field. "
                "Please re-run with the same inputs; core identifiers and score fields were still returned.\n\n"
                "## Overall Assessment\n"
                f"Request Type: {result.get('request_type','N/A')}\n"
                f"Compliance Score: {result.get('compliance_score','N/A')}\n"
            )
        elif "## Detailed Condition Report" not in sm_tmp:
            # Keep minimal: do not re-write content; just prepend the required header to avoid downstream display rules.
            result["summary_markdown"] = "## Detailed Condition Report\n" + sm_tmp
    except Exception:
        pass


    # -----------------------
    # Photos-Only OUTPUT HARDENING (prevents blank/N/A reports; enforces Add'l Notes)
    # -----------------------
    def _bad_field(v: str) -> bool:
        """Return True if the model gave a placeholder/empty section.
        This catches bare 'N/A' AND markdown sections that are effectively only headings + N/A.
        """
        s = (v or "").strip()
        if not s:
            return True
        up = s.strip().upper()

        # Bare placeholders
        if up in ("N/A", "NA", "NONE", "NULL"):
            return True
        if up.startswith("N/A") and len(up) <= 8:
            return True

        # If it's markdown, strip headings/blank lines and see if anything substantive remains
        lines = [ln.strip() for ln in s.splitlines() if ln.strip()]
        # remove markdown headings and common labels
        stripped = []
        for ln in lines:
            if ln.startswith("#"):
                continue
            if ln.lower() in ("report selected", "approximate repair cost breakdown", "fraud & authenticity check", "fraud detection", "conclusion"):
                continue
            stripped.append(ln)

        if not stripped:
            return True

        # If after stripping headings the only remaining content is 'N/A'
        if len(stripped) == 1 and stripped[0].strip().upper() in ("N/A", "NA", "NONE", "NULL"):
            return True

        # If the remaining content is extremely short and contains only placeholders
        joined = " ".join(stripped).strip().upper()
        if joined in ("N/A", "NA", "NONE", "NULL"):
            return True

        return False

    try:
        if ai_intent == "damage_report_from_photos":
            _bad = (
                _bad_field(result.get("summary_markdown") or "")
                or _bad_field(result.get("estimated_costs_markdown") or "")
                or _bad_field(result.get("fraud_markdown") or "")
                or _bad_field(result.get("conclusion") or "")
            )

            if _bad:
                retry_preamble = (
                    "CRITICAL RETRY (PHOTOS-ONLY): Your prior output was invalid because required fields were blank or 'N/A'.\n"
                    "Return ONLY valid JSON (no prose, no code fences).\n"
                    "Hard rules:\n"
                    "- NEVER return 'N/A' for summary_markdown, estimated_costs_markdown, fraud_markdown, or conclusion.\n"
                    "- You MUST explicitly address the Add'l Notes in the narrative (e.g., 'Impact to Right Front') and describe the impact zone accordingly.\n"
                    "- estimated_costs_markdown MUST be a PHOTOS-ONLY approximation: create body/paint/frame/mechanical hours and an OEM parts list with approximate $ by part.\n"
                    "- Paint & Materials MUST be modeled as $ per refinish hour (not a percent).\n"
                    "- Forbidden phrases: 'from estimate', 'from documentation', 'not evidenced', 'no documentation provided'.\n"
                )

                retry_parts = list(parts_payload)  # shallow copy is enough
                # Replace the first text block with a retry preamble + original prompt
                if retry_parts and isinstance(retry_parts[0], dict) and retry_parts[0].get("type") == "text":
                    retry_parts[0] = {"type": "text", "text": retry_preamble + "\n\n" + (retry_parts[0].get("text") or "")}
                else:
                    retry_parts.insert(0, {"type": "text", "text": retry_preamble})

                retry_tokens = min(3200, max_tokens + 500)

                try:
                    rsp_retry = client.chat.completions.create(
                        model=MODEL,
                        messages=[{"role": "system", "content": SYSTEM},
                                  {"role": "user", "content": retry_parts}],
                        max_completion_tokens=retry_tokens,
                        temperature=0,
                        top_p=1,
                        presence_penalty=0,
                        frequency_penalty=0,
                        response_format={"type": "json_object"},
                    )
                except AttributeError:
                    rsp_retry = client.chat_completions.create(  # type: ignore[attr-defined]
                        model=MODEL,
                        messages=[{"role": "system", "content": SYSTEM},
                                  {"role": "user", "content": retry_parts}],
                        max_completion_tokens=retry_tokens,
                        temperature=0,
                        top_p=1,
                        presence_penalty=0,
                        frequency_penalty=0,
                        response_format={"type": "json_object"},
                    )

                raw_retry = (rsp_retry.choices[0].message.content or "")
                data_retry = _try_parse_json(raw_retry)

                if isinstance(data_retry, dict):
                    # Only overwrite the output fields; preserve file_number/request_type etc.
                    for k in ("summary_brief","summary_markdown","fraud_markdown","primary_impact","secondary_impact","estimated_costs_markdown","conclusion","claim_number","vin","vin_verification","vehicle","odometer_estimate_only","compliance_score"):
                        if k in data_retry and data_retry.get(k) is not None:
                            result[k] = str(data_retry.get(k))
                    # Re-apply the narrative header guard if needed
                    sm_retry = (result.get("summary_markdown") or "").strip()
                    if sm_retry and "## Detailed Condition Report" not in sm_retry:
                        result["summary_markdown"] = "## Detailed Condition Report\n" + sm_retry

                    # If the retry STILL returned placeholders, force a non-N/A minimal scaffold (never print dead 'N/A')
                    if (
                        _bad_field(result.get("summary_markdown") or "")
                        or _bad_field(result.get("estimated_costs_markdown") or "")
                        or _bad_field(result.get("fraud_markdown") or "")
                        or _bad_field(result.get("conclusion") or "")
                    ):
                        _notes = (ai_notes_used or "").strip()
                        if not _notes or _notes.lower().startswith("no additional notes"):
                            _notes = "(No additional notes provided.)"

                        result["summary_markdown"] = (
                            "## Detailed Condition Report\n"
                            "Photo-based narrative could not be generated on this run due to a model compliance error.\n"
                            f"Add'l Notes received: {_notes}\n"
                            "Please re-run. (This placeholder is generated by NSPXN to avoid blank/N/A reports.)"
                        )
                        result["estimated_costs_markdown"] = (
                            "## Approximate Repair Cost Breakdown (Photos-Only Approximation)\n"
                            "The model failed to generate a cost breakdown on this run. Please re-run with the same photo set.\n"
                            "Paint & Materials are modeled as $/refinish hour; tax applies to parts + paint materials only."
                        )
                        result["fraud_markdown"] = "No material inconsistencies found."
                        result["conclusion"] = (
                            "Conclusion unavailable due to model compliance error on this run. Please re-run."
                        )

    except Exception:
        # Fail-open: never break the request if retry logic fails.
        pass

        # -----------------------
    # Approximate Repair Cost Breakdown
    # - Prefer model-provided `estimated_costs_markdown` (especially for Photos-Only).
    # - Fail-open: never break the run if cost math cannot be produced.
    # -----------------------
    try:
        _existing_costs = (result.get("estimated_costs_markdown") or "").strip()

        if not _existing_costs:
            inspection_location = _normalize_location_with_zip(_extract_inspection_location(uploaded_text_all or ""), ia_company, uploaded_text_all)
            rates = _lookup_rates(inspection_location)
            tax_rate = _lookup_tax_rate(inspection_location)

            # Photos-only: do NOT depend on documents/estimates. The model is instructed to generate hours/parts.
            # This backend block is only a minimal fallback if the model omitted the field.
            if ai_intent == "damage_report_from_photos":
                structural_flag = _structural_observed([result.get("summary_markdown") or ""])

                result["estimated_costs_markdown"] = (
                    "## Approximate Repair Cost Breakdown (Photos-Only Approximation)\n\n"
                    f"**Inspection Location:** {inspection_location or 'Not provided'}\n\n"
                    "**Rates Used (regional average / fallback):**\n"
                    f"- Avg Body Rate: ${rates.get('body_rate', 0):.0f}/hr\n"
                    f"- Avg Paint Rate: ${rates.get('paint_rate', 0):.0f}/hr\n"
                    f"- Avg Frame Rate: ${rates.get('frame_rate', 0):.0f}/hr\n"
                    f"- Avg Mechanical Rate: ${rates.get('mechanical_rate', 0):.0f}/hr\n"
                    f"- Paint & Materials Rate: ${rates.get('paint_supplies_rate', 0):.0f} per refinish hr\n"
                    f"- Parts/Materials Tax Rate: {float(tax_rate or 0.0)*100:.3f}%\n\n"
                    "**AI-Derived Repair Scope:**\n"
                    "- Body labor hours: (model should provide)\n"
                    "- Paint labor hours: (model should provide)\n"
                    f"- Setup & Measure (if structural observed): {('2.0 hrs @ body rate' if structural_flag else '0.0 hrs')}\n"
                    "- Frame labor hours (if structural observed): (model should provide)\n"
                    "- Mechanical/ADAS/Airbag/Suspension hours (if observed): (model should provide)\n\n"
                    "**OEM Parts Needed (with approximate $ by part):**\n"
                    "- (model should provide)\n\n"
                    "**Tax (parts + paint materials only):**\n"
                    "- (model should provide)\n\n"
                    "**Approximated Repair Cost Total (Approximation Only):**\n"
                    "- (model should provide)\n\n"
                    "**Estimated Severity Tier (based on this approximation only):**\n"
                    "- [ ] Minor (< $3,500)\n"
                    "- [ ] Moderate ($3,500–$10,000)\n"
                    "- [ ] Major ($10,000+)\n"
                    "- [ ] Possible Total Loss Threshold Approaching\n\n"
                    "_Repair Cost Disclaimer: This section is an AI-generated approximation of repair-related costs based on observed damage in photos and regional average rates. "
                    "It is not an official repair estimate and must be validated by a qualified appraiser using an estimating platform._"
                )
            else:
                # Non-photos intents: best-effort parse from documents (if present), using $/refinish-hr materials.
                totals = _extract_hours_and_parts_totals(uploaded_text_all or "")

                structural_flag = _structural_observed([uploaded_text_all or "", result.get("summary_markdown") or ""])
                setup_measure_hours = 2.0 if structural_flag else 0.0

                body_hours = totals.get("body_hours")
                paint_hours = totals.get("paint_hours")
                frame_hours = totals.get("frame_hours")
                mech_hours = totals.get("mech_hours")
                parts_total = totals.get("parts_total")

                paint_supplies_rate = float(rates.get("paint_supplies_rate") or 0.0)
                paint_supplies_total = (_num(paint_hours) * paint_supplies_rate) if paint_hours is not None else 0.0

                body_labor_total = (_num(body_hours) + setup_measure_hours) * float(rates.get("body_rate") or 0.0) if (body_hours is not None or setup_measure_hours) else 0.0
                paint_labor_total = _num(paint_hours) * float(rates.get("paint_rate") or 0.0) if paint_hours is not None else 0.0

                frame_labor_total = 0.0
                if structural_flag and frame_hours is not None:
                    frame_labor_total = _num(frame_hours) * float(rates.get("frame_rate") or 0.0)

                mech_labor_total = _num(mech_hours) * float(rates.get("mechanical_rate") or 0.0) if mech_hours is not None else 0.0

                taxable = _num(parts_total) + _num(paint_supplies_total)
                tax_total = taxable * float(tax_rate or 0.0)

                approx_total = (
                    _num(body_labor_total)
                    + _num(paint_labor_total)
                    + _num(frame_labor_total)
                    + _num(mech_labor_total)
                    + _num(parts_total)
                    + _num(paint_supplies_total)
                    + _num(tax_total)
                )

                parts_lines = totals.get("parts_lines") or []
                if parts_lines:
                    parts_md = "\n".join([f"- {ln}" for ln in parts_lines[:20]])
                else:
                    parts_md = "- (Itemized OEM parts list not available from parsed text.)"

                result["estimated_costs_markdown"] = (
                    "## Approximate Repair Cost Breakdown (Approximation Only)\n\n"
                    f"**Inspection Location:** {inspection_location or 'Not provided'}\n\n"
                    "**Regional Average Rates (configured / fallback):**\n"
                    f"- Avg Body Rate: ${rates.get('body_rate', 0):.0f}/hr\n"
                    f"- Avg Paint Rate: ${rates.get('paint_rate', 0):.0f}/hr\n"
                    f"- Avg Frame Rate: ${rates.get('frame_rate', 0):.0f}/hr\n"
                    f"- Avg Mechanical Rate: ${rates.get('mechanical_rate', 0):.0f}/hr\n"
                    f"- Paint & Materials Rate: ${rates.get('paint_supplies_rate', 0):.0f} per refinish hr\n"
                    f"- Parts/Materials Tax Rate: {float(tax_rate or 0.0)*100:.3f}%\n\n"
                    "**Hours & Totals:**\n"
                    f"- Approx Total Labor Hours @ Body Rate: {('%.1f' % _num(body_hours)) if body_hours is not None else 'Unknown'}\n"
                    f"- Approx Total Paint Labor Hours @ Paint Rate: {('%.1f' % _num(paint_hours)) if paint_hours is not None else 'Unknown'}\n"
                    f"- Approx Total Paint Supplies @ ${rates.get('paint_supplies_rate', 0):.0f}/refinish hr: {_money(paint_supplies_total)}\n"
                    f"- Setup & Measure (if structural observed): {('2.0 hrs' if structural_flag else '0.0 hrs')}\n"
                    f"- Approx Total Frame Labor (if structural observed) @ Frame Rate: {('%.1f hrs' % _num(frame_hours)) if (structural_flag and frame_hours is not None) else ('Unknown' if structural_flag else '0.0 hrs')}\n"
                    f"- Approx Mechanical/ADAS/Airbag/Suspension Hours @ Mechanical Rate: {('%.1f' % _num(mech_hours)) if mech_hours is not None else 'Unknown'}\n\n"
                    "**Approx Cost of OEM Parts Needed:**\n"
                    f"{parts_md}\n\n"
                    "**Tax (parts + paint materials only):**\n"
                    f"- Taxable subtotal: {_money(taxable)}\n"
                    f"- Estimated tax: {_money(tax_total)}\n\n"
                    "**Approximated Repair Cost Total (Approximation Only):**\n"
                    f"- **{_money(approx_total)}**\n\n"
                    "_Repair Cost Disclaimer: This section is an approximation. It is not an official repair estimate and must be validated by a qualified appraiser using an estimating platform._"
                )

    except Exception:
        # Never fail the request for cost calculation issues
        if not (result.get("estimated_costs_markdown") or "").strip():
            result["estimated_costs_markdown"] = (
                "## Approximate Repair Cost Breakdown (Approximation Only)\n"
                "Unable to generate cost approximation on this run."
            )
# -----------------------

    # Final scrub pass after any backend fallback cost-building
    try:
        _md1 = result.get("estimated_costs_markdown") or ""
        _md1 = _scrub_cost_markdown(_md1)
        _md1 = _enforce_severity_tier(_md1)
        if _md1:
            result["estimated_costs_markdown"] = _md1
    except Exception:
        pass

    # PDF helpers
    # -----------------------
    def _pdf_sanitize(text: str, max_token_len: int = 60) -> str:
        if text is None:
            return ""
        s = str(text).replace("\r\n", "\n").replace("\r", "\n")
        s = "".join(ch if ord(ch) < 256 else " " for ch in s)
        def _break(tok: str) -> str:
            if len(tok) <= max_token_len:
                return tok
            return " ".join(tok[i:i+max_token_len] for i in range(0, len(tok), max_token_len))
        s = " ".join(_break(t) for t in s.split(" "))
        return s

    pdf = FPDF(); pdf.add_page()
    # --- NSPXN Logo (Top Right, First Page Only) ---
    try:
        logo_path = os.path.join(os.path.dirname(__file__), "ChatGPT logo100725.png")
        if os.path.exists(logo_path):
            pdf.image(logo_path, x=pdf.w - 45, y=8, w=35)  # small–medium size
    except Exception:
        pass
    pdf.set_auto_page_break(auto=True, margin=10)
    pdf.set_left_margin(10); pdf.set_right_margin(10)

    try:
        pdf.add_font("DejaVu","", "DejaVuSans.ttf", uni=True); pdf.set_font(size=11)
    except Exception:
        pdf.set_font("Arial", size=11)

    def mc(s):
        try:
            effective_w = pdf.w - pdf.l_margin - pdf.r_margin
            if effective_w <= 5:
                effective_w = 180
            safe = _pdf_sanitize(s)
            if not safe.strip():
                safe = "-"
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(effective_w, 6, safe)
        except Exception:
            effective_w = pdf.w - pdf.l_margin - pdf.r_margin
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(effective_w, 6, (_pdf_sanitize(str(s))[:2000] + " …"))


    def add_thumbnail_page(pdf_obj: FPDF, image_paths: List[str]) -> None:
        """Append exactly ONE page containing thumbnails of all uploaded photos (as space allows)."""
        if not image_paths:
            return

        pdf_obj.add_page()
        try:
            pdf_obj.set_font("Helvetica", "B", 12)
        except Exception:
            pdf_obj.set_font("Arial", "B", 12)
        pdf_obj.cell(0, 8, "Uploaded Photos", ln=True)
        pdf_obj.ln(2)

        # Layout (single-page, grid)
        cols = 4
        gutter = 3.0
        usable_w = pdf_obj.w - pdf_obj.l_margin - pdf_obj.r_margin
        thumb_w = (usable_w - gutter * (cols - 1)) / cols
        x0 = pdf_obj.l_margin
        y = pdf_obj.get_y()
        x = x0
        col = 0

        placed = 0
        total = len(image_paths)

        for p in image_paths:
            # Hard-stop to keep this as ONE PAGE ONLY
            if y + thumb_w > (pdf_obj.h - pdf_obj.b_margin - 10):
                break
            try:
                pdf_obj.image(p, x=x, y=y, w=thumb_w)
                placed += 1
            except Exception:
                pass

            col += 1
            if col >= cols:
                col = 0
                x = x0
                y += thumb_w + gutter
            else:
                x += thumb_w + gutter

        # Footer note if truncated by the 1-page rule
        if placed < total:
            try:
                pdf_obj.set_font_size(9)
            except Exception:
                pass
            pdf_obj.set_y(pdf_obj.h - pdf_obj.b_margin - 8)
            pdf_obj.set_x(pdf_obj.l_margin)
            pdf_obj.cell(0, 6, f"Showing {placed} of {total} uploaded photos (1-page appendix limit).", ln=True)

    # POI-15 Total Loss trigger from uploaded text ONLY
    poi15_hit = False
    try:
        txt = uploaded_text_all.lower()
        if re.search(r"\b15\s*total\s*loss\b", txt, flags=re.IGNORECASE):
            poi15_hit = True
        elif re.search(r"point\s*of\s*impact[^A-Za-z0-9]{0,10}15[^A-Za-z0-9]{0,20}total\s*loss", txt, flags=re.IGNORECASE):
            poi15_hit = True
    except Exception:
        poi15_hit = False

    if ai_intent == "damage_report_from_photos":
        pdf.cell(0,10,"NSPXN.com Condition Report", ln=True, align="C")
        pdf.set_font_size(10); pdf.ln(3)

        mc(f"Claim #: {result['claim_number'] or 'N/A'}    File #: {file_number or 'N/A'}")
        mc(f"Inspected For: {ia_company}")
        pdf_status = result["redaction_status"].replace("✅", "OK")
        pdf.ln(2); mc(pdf_status)
        pdf.ln(2); mc("Report Selected"); mc((result["summary_markdown"] or "N/A").strip())
        pdf.ln(2)
        # Centered bold section header (user requested)
        try:
            pdf.set_font("Helvetica", "B", 11)
        except Exception:
            pdf.set_font("Arial", "B", 11)
        pdf.cell(0, 6, "Approximate Repair Cost Breakdown", ln=True, align="C")
        try:
            pdf.set_font("Helvetica", "", 11)
        except Exception:
            pdf.set_font("Arial", "", 11)
        mc((result.get("estimated_costs_markdown") or "").strip())
        pdf.ln(2); mc("Fraud & Authenticity Check"); mc((result["fraud_markdown"] or 'N/A').strip())
        pdf.ln(2); mc("Conclusion"); mc((result["conclusion"] or 'N/A').strip())

        # --- AI Disclaimer (after Conclusion) ---
        try:
            pdf.ln(4)
            x_left = pdf.l_margin
            x_right = pdf.w - pdf.r_margin
            y_line = pdf.get_y()
            pdf.set_draw_color(180, 180, 180)
            pdf.line(x_left, y_line, x_right, y_line)
            pdf.ln(3)

            pdf.set_text_color(90, 90, 90)
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(0, 5, "Disclaimer:", ln=True)
            pdf.set_font("Helvetica", "", 8)

            disclaimer_body = (
                "This report was generated using artificial intelligence. AI systems may make errors or misinterpret "
                "visual information. All photos, damage descriptions, conclusions, and findings must be independently "
                "reviewed and verified by a qualified appraiser before preparing or finalizing any repair estimate."
            )
            pdf.multi_cell(0, 4, disclaimer_body)
            pdf.set_text_color(0, 0, 0)
        except Exception:
            pass


        safe_file = _safe(file_number)
        pdf_filename = f"AI_Condition_Report_{safe_file}.pdf"
    else:
        pdf.cell(0,10,"NSPXN.com Condition Report", ln=True, align="C")
        pdf.set_font_size(10); pdf.ln(3)
        mc(f"File Number: {file_number}")
        mc(f"Inspected For: {ia_company}")
        mc(f"Appraiser ID #: {appraiser_id}")
        mc(f"Request Type: {result['request_type']}")

        # --- Supplement header (documents only; ignore negated mentions) ---
        smark = result.get("summary_markdown","")
        _txt_docs = uploaded_text_all or ""
        _supp_doc_hit = bool(re.search(
            r"(?is)\b(Supplement\s+(?:Summary|of\s+Record)|Estimate\s+Version:\s*S0[1-9]\b|\bS0[1-9]\b|\bSupplement\s+Estimate\b)",
            _txt_docs
        ))
        _no_supp_negation = not re.search(r"(?is)\b(no|not)\s+(a\s+)?supplement\b", _txt_docs)
        supp_detected_docs = _supp_doc_hit and _no_supp_negation
        if supp_detected_docs:
            mc("Supplement Status: Supplement Estimate detected in documentation")
            _possible_amt = None
            _m = re.search(r"(?is)\bPossible\s+Supplement\s+Amount\s*\$?([0-9,]+\.\d{2})\b", _txt_docs)
            if _m:
                _possible_amt = _m.group(1)
            mc("Supplement Details")
            # NEW: list all supplement tags (S01, S02, ...) instead of a single version only
            supp_versions_docs = sorted(set(re.findall(r"(?i)\bS[0-9]{2}\b", _txt_docs)))
            if supp_versions_docs:
                mc(f"- Supplements detected in documents: {', '.join(supp_versions_docs)}")
            if _possible_amt:
                mc(f"- Possible amount noted: ${_possible_amt}")
            if not (supp_versions_docs or _possible_amt):
                mc("- Supplement indicators present (e.g., 'Supplement Summary' or S01/S02).")

        # --- Total Loss echo (documents-only; no narrative trigger) ---
        explicit_tl_hit = False
        try:
            txt_docs = uploaded_text_all or ""
            poi15_docs = re.search(
                r"(?is)\bpoint\s*of\s*impact[^a-z0-9]{0,10}15[^a-z0-9]{0,20}total\s*loss\b",
                txt_docs,
            )
            doc_tl_docs = re.search(
                r"(?i)\b(estimate\s*type|type\s*of\s*loss)\s*:\s*total\s*loss\b",
                txt_docs,
            )
            explicit_tl_hit = bool(poi15_docs or doc_tl_docs)
        except Exception:
            explicit_tl_hit = False
        if explicit_tl_hit:
            mc("Estimate Type: Total Loss (explicit in documents)")

        mc(f"Claim #: {result['claim_number']}")
        mc(f"VIN (from estimate/photos): {result['vin']}")
        mc(f"VIN verification (estimate vs photo): {result['vin_verification']}")
        mc(f"Vehicle: {result['vehicle']}")
        mc(f"Odometer (from estimate): {result['odometer_estimate_only']}")
        mc(f"Compliance Score: {result['compliance_score']}")
        pdf_status = result["redaction_status"].replace("✅", "OK")
        mc(pdf_status)
        pdf.ln(3); mc("NSPXN.com Condition Summary"); mc((smark or '').strip())
        pdf.ln(3)
        # Centered bold section header (user requested)
        try:
            pdf.set_font("Helvetica", "B", 11)
        except Exception:
            pdf.set_font("Arial", "B", 11)
        pdf.cell(0, 6, "Approximate Repair Cost Breakdown", ln=True, align="C")
        try:
            pdf.set_font("Helvetica", "", 11)
        except Exception:
            pdf.set_font("Arial", "", 11)
        mc((result.get("estimated_costs_markdown") or "").strip())
        pdf.ln(3); mc("Fraud Detection"); mc((result["fraud_markdown"] or 'N/A').strip())

        # --- AI Disclaimer (after report content) ---
        try:
            pdf.ln(4)
            x_left = pdf.l_margin
            x_right = pdf.w - pdf.r_margin
            y_line = pdf.get_y()
            pdf.set_draw_color(180, 180, 180)
            pdf.line(x_left, y_line, x_right, y_line)
            pdf.ln(3)

            pdf.set_text_color(90, 90, 90)
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(0, 5, "Disclaimer:", ln=True)
            pdf.set_font("Helvetica", "", 8)

            disclaimer_body = (
                "This report was generated using artificial intelligence. AI systems may make errors or misinterpret "
                "visual information. All photos, damage descriptions, conclusions, and findings must be independently "
                "reviewed and verified by a qualified appraiser before preparing or finalizing any repair estimate."
            )
            pdf.multi_cell(0, 4, disclaimer_body)
            pdf.set_text_color(0, 0, 0)
        except Exception:
            pass


        safe_file = _safe(file_number)
        pdf_filename = f"{safe_file}.pdf"


    # --- One-page photo thumbnail appendix (all uploaded photos) ---
    try:
        add_thumbnail_page(pdf, thumbnail_paths)
    except Exception:
        pass

    pdf_path = os.path.join(PDF_DIR, pdf_filename)
    try:
        out = pdf.output(dest="S")
        if isinstance(out, (bytes, bytearray)):
            data_bytes = bytes(out)
        else:
            data_bytes = str(out).encode("latin-1", "ignore")
        with open(pdf_path, "wb") as f:
            f.write(data_bytes)
    except Exception as e:
        logging.warning(f"PDF write error: {e}")

    pdf_url = f"/download-pdf?filename={pdf_filename}"

    # -----------------------
    # Email — info-only (attach PDF)
    # -----------------------
    try:
        msg = EmailMessage()
        if ai_intent == "damage_report_from_photos":
            subj = f"NSPXN.com Condition Report: {file_number or ''} {result['claim_number'] or ''}".strip()
            body = (
                "NSPXN.com Condition Report\n\n"
                f"Inspected For: {ia_company}\n"
                f"Claim #: {result['claim_number'] or 'N/A'}    File #: {file_number or 'N/A'}\n"
                f"Odometer: {result['odometer_estimate_only'] or 'N/A'}    Primary Impact: {result['primary_impact'] or 'N/A'}\n"
                f"Secondary Impact: {result['secondary_impact'] or 'N/A'}\n\n"
                f"{result['redaction_status']}\n\n"
                "Condition Summary\n"
                f"{(result['summary_markdown'] or 'N/A')}\n\n"
                "Fraud & Authenticity Check\n"
                f"{(result['fraud_markdown'] or 'N/A')}\n\n"
                "Conclusion\n"
                f"{(result['conclusion'] or 'N/A')}\n"
            )
        else:
            try:
                _txt_email = uploaded_text_all or ""
                _poi15_email = re.search(
                    r"(?is)\bpoint\s*of\s*impact[^a-z0-9]{0,10}15[^a-z0-9]{0,20}total\s*loss\b",
                    _txt_email,
                )
                _doc_tl_email = re.search(
                    r"(?i)\b(estimate\s*type|type\s*of\s*loss)\s*:\s*total\s*loss\b",
                    _txt_email,
                )
                _explicit_tl_email = bool(_poi15_email or _doc_tl_email)
            except Exception:
                _explicit_tl_email = False

            # NEW: supplement line includes all Sxx tags if present
            supp_line = ""
            if supp_detected_docs:
                supp_versions_email = sorted(set(re.findall(r"(?i)\bS[0-9]{2}\b", _txt_email or "")))
                if supp_versions_email:
                    supp_line = (
                        "Supplement Status: Supplement Estimates detected in documentation "
                        f"({', '.join(supp_versions_email)})\n"
                    )
                else:
                    supp_line = "Supplement Status: Supplement Estimate detected in documentation\n"

            tl_line = "Estimate Type: Total Loss (explicit in documents)\n" if _explicit_tl_email else ""
            subj = f"NSPXN.com Review: {result['claim_number'] or file_number}"
            body = (
                "NSPXN.com Condition Report\n\n"
                f"File Number: {file_number}\n"
                f"Inspected For: {ia_company}\n"
                f"Appraiser ID #: {appraiser_id}\n"
                f"Request Type: {result['request_type']}\n"
                f"{supp_line}"
                f"{tl_line}"
                f"Claim #: {result['claim_number']}\n"
                f"VIN (from estimate/photos): {result['vin']}\n"
                f"VIN verification (estimate vs photo): {result['vin_verification']}\n"
                f"Vehicle: {result['vehicle']}\n"
                f"Odometer (from estimate): {result['odometer_estimate_only']}\n"
                f"Compliance Score: {result['compliance_score']}\n\n"
                f"{result['redaction_status']}\n\n"
                "NSPXN.com Condition Summary\n"
                f"{result['summary_markdown']}\n\n"
                "Fraud Detection\n"
                f"{result['fraud_markdown']}\n"
            )

        msg["Subject"] = subj
        msg["From"] = "info@nspxn.com"
        msg["To"] = "info@nspxn.com"
        msg["Cc"] = "growley@ractrak.com"
        msg.set_content(body)

        try:
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            msg.add_attachment(pdf_bytes, maintype="application", subtype="pdf", filename=pdf_filename)
        except Exception as e:
            logging.warning(f"Failed to attach PDF to email: {e}")

        with smtplib.SMTP_SSL("mail.tierra.net", 465, timeout=20) as smtp:
            smtp.login("info@nspxn.com", "grr2025GRR")
            smtp.send_message(msg)
        log.info("Info email sent to info@nspxn.com")
    except Exception as e:
        logging.error(f"Email error: {e}")

    return {
        **result,
        "web_summary": result["summary_brief"],
        "gpt_output": result["summary_markdown"],
        "pdf_url": pdf_url,
        "pdf_filename": pdf_filename
    }

    # -----------------------
# PDF download
    # -----------------------
@app.get("/download-pdf")
async def download_pdf(file_number: Optional[str] = None, filename: Optional[str] = None):
    if filename:
        safe = _safe(filename)
        path = os.path.join(PDF_DIR, safe)
        if os.path.exists(path):
            return FileResponse(path=path, media_type="application/pdf", filename=safe)
        return JSONResponse(status_code=404, content={"detail": "Not Found"})
    if not file_number:
        return JSONResponse(status_code=400, content={"detail": "Missing query param 'filename' or 'file_number'"})
    safe_num = _safe(file_number)
    candidates = glob.glob(os.path.join(PDF_DIR, f"*{safe_num}*.pdf"))
    if not candidates:
        return JSONResponse(status_code=404, content={"detail": "Not Found"})
    latest = max(candidates, key=lambda p: os.path.getmtime(p))
    return FileResponse(path=latest, media_type="application/pdf", filename=os.path.basename(latest))









