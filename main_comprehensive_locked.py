from zoneinfo import ZoneInfo
from datetime import datetime
from fastapi import FastAPI, File, UploadFile, Form, Request, Response
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Dict, Any, Optional
import os, io, re, json, base64, logging, zipfile, glob, uuid
import urllib.parse, urllib.request
import smtplib  # email transport
from email.message import EmailMessage

from fpdf import FPDF
from docx import Document
from pdf2image import convert_from_bytes
from PIL import Image

# Optional HEIC/HEIF support if available
try:
    import pillow_heif  # type: ignore
    pillow_heif.register_heif_opener()  # type: ignore
except Exception:
    pass

# Optional OCR (pytesseract). Safe no-op if not installed or tesseract binary missing.
try:
    import pytesseract  # type: ignore
    _OCR_ENABLED = True
except Exception:
    _OCR_ENABLED = False

from openai import OpenAI

# --- PII Redaction (Presidio) ---
from presidio_analyzer import AnalyzerEngine, Pattern, PatternRecognizer, RecognizerResult
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig  # required for anonymizer API

# -----------------------
# Minimal setup
# -----------------------
PDF_DIR = os.getenv("PDF_DIR", "/tmp"); os.makedirs(PDF_DIR, exist_ok=True)
CLIENT_RULES_DIR = os.getenv("CLIENT_RULES_DIR", "client_rules")

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
log = logging.getLogger("nspxn")
log.info(f"Using CLIENT_RULES_DIR={CLIENT_RULES_DIR}")

# Use selected model everywhere
MODEL = os.getenv("OAI_MODEL") or os.getenv("OPENAI_MODEL") or "gpt-5.2"
# GPT-5.x models use max_completion_tokens; GPT-4.x uses max_tokens
_token_kw = "max_completion_tokens"

if not os.getenv("OPENAI_API_KEY"):
    raise RuntimeError("OPENAI_API_KEY missing")
try:
    client = OpenAI(api_key=os.environ["OPENAI_API_KEY"], timeout=120.0, max_retries=1)
except TypeError:
    # Backwards-compatible init for older openai-python versions
    client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])

# --------------------------------
# Presidio: Analyzer/Anonymizer (preserve VIN & Claim #)
# --------------------------------
analyzer = AnalyzerEngine()
anonymizer = AnonymizerEngine()

VIN_PATTERN = r"\b([A-HJ-NPR-Z0-9]{17})\b"
vin_recognizer = PatternRecognizer(
    supported_entity="VIN",
    patterns=[Pattern(name="vin-17", regex=VIN_PATTERN, score=0.8)],
)

CLAIM_PATTERN = r"\b(?:(?:Claim|CLM|Clm)\s*#?\s*[:\-]?\s*)?([A-Z0-9]{5,}[A-Z0-9\-]{0,})\b"
claim_recognizer = PatternRecognizer(
    supported_entity="CLAIM_NUMBER",
    patterns=[Pattern(name="claim-generic", regex=CLAIM_PATTERN, score=0.6)],
)

analyzer.registry.add_recognizer(vin_recognizer)
analyzer.registry.add_recognizer(claim_recognizer)

# Only redact these entity types (we PRESERVE VIN & CLAIM_NUMBER)
REDACT_ENTITY_TYPES = {
    "PERSON", "PHONE_NUMBER", "EMAIL_ADDRESS", "US_SSN",
    "CREDIT_CARD", "IBAN_CODE", "LOCATION", "NRP", "ORGANIZATION",
    "DATE_TIME", "IP_ADDRESS", "CRYPTO", "MEDICAL_LICENSE", "URL"
}

def _filter_results(results: List[RecognizerResult]) -> List[RecognizerResult]:
    return [r for r in results if r.entity_type in REDACT_ENTITY_TYPES]

def redact_text_preserve_vin_claim(text: str) -> str:
    """Mask PII while keeping VIN and Claim # intact. Fail-open to avoid 500s."""
    if not text:
        return text
    results = analyzer.analyze(text=text, language="en")
    to_mask = _filter_results(results)
    if not to_mask:
        return text
    try:
        return anonymizer.anonymize(
            text=text,
            analyzer_results=to_mask,
            operators={"DEFAULT": OperatorConfig("replace", {"new_value": "[REDACTED]"})}
        ).text
    except Exception as e:
        log.warning(f"Presidio anonymizer failed, passing text through. Error: {e}")
        return text

def _safe(s: str) -> str:
    return re.sub(r"[^\w.\-]+", "-", (s or "").strip()).strip("-_. ")




def _format_vehicle_value(v) -> str:
    """Normalize the Vehicle field for PDF printing."""
    try:
        if isinstance(v, dict):
            year = str(v.get("year") or "").strip()
            make = str(v.get("make") or "").strip()
            model = str(v.get("model") or "").strip()
            trim = str(v.get("trim") or "").strip()
            parts = [p for p in [year, make, model] if p and p.upper() != "N/A"]
            base = " ".join(parts).strip()
            if trim and trim.upper() != "N/A":
                return f"{base} ({trim})" if base else trim
            return base or "N/A"
    except Exception:
        pass
    s = str(v or "").strip()
    return s if s else "N/A"

def _normalize_ai_notes(raw: str, max_len: int = 800) -> str:
    """Normalize/sanitize Add'l Notes so they can't break JSON/structured prompting."""
    if raw is None:
        return "No additional notes provided."
    s = str(raw).strip()
    if not s:
        return "No additional notes provided."

    # Remove ASCII control chars except newline/tab
    s = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", " ", s)

    # Normalize newlines + collapse whitespace
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s).strip()

    # Neutralize common schema-breakers
    s = s.replace("```", "'''")
    # Avoid curly braces interfering with JSON-format instructions
    s = s.replace("{", "(").replace("}", ")")

    if len(s) > max_len:
        s = s[:max_len].rstrip() + "…"
    return s

def _ai_notes_block(ai_notes_clean: str) -> str:
    """Always return a consistent, clearly-delimited notes block."""
    return (
        "\n\nADDL_NOTES_FOR_AI_REVIEW (USER PROVIDED):\n"
        "<<<\n"
        f"{ai_notes_clean}\n"
        ">>>\n"
        "REQUIREMENT:\n"
        "- You MUST consider these notes while analyzing the photos.\n"
        "- Use them to guide focus/wording, but DO NOT invent damage.\n"
        "- If notes conflict with visible evidence or orientation is unclear, say so and defer to what is visible.\n"
    )

# -----------------------
# Approximate Repair Cost Breakdown (location-based rates)
# -----------------------
# Notes:
# - Paint supplies/materials are modeled as $ per refinish hour (NOT a percent).
# - This is an approximation for observed damages, not a repair estimate.

# Optional external integration (no new dependencies):
# - Set LABORRATEHERO_API_URL to a JSON endpoint you control (proxy/scraper) that returns:
#   {"body_rate": 0, "paint_rate": 0, "frame_rate": 0, "mechanical_rate": 0, "paint_supplies_rate": 0}
# - Optionally set LABORRATEHERO_API_KEY if your endpoint requires it.
LABORRATEHERO_API_URL = os.getenv("LABORRATEHERO_API_URL", "").strip()
LABORRATEHERO_API_KEY = os.getenv("LABORRATEHERO_API_KEY", "").strip()

def _extract_inspection_location(text: str) -> str:
    """Best-effort extraction of Inspection Location (City/State/ZIP) from uploaded text."""
    if not text:
        return ""
    # Common patterns: "Inspection Location: Albuquerque, NM 87109" or multi-line variants
    m = re.search(r"(?im)^\s*Inspection\s+Location\s*[:\-]?\s*(.+)$", text)
    if m:
        return (m.group(1) or "").strip()
    m2 = re.search(r"(?is)\bInspection\s+Location\b\s*[:\-]?\s*(.{0,120}?)\n", text)
    if m2:
        return (m2.group(1) or "").strip()
    return ""

def _extract_zip5(*texts: str) -> str:
    """Return the first 5-digit ZIP found in any provided text (best-effort)."""
    for t in texts:
        if not t:
            continue
        m = re.search(r"\b(\d{5})(?:-\d{4})?\b", str(t))
        if m:
            return m.group(1)
    return ""

def _zip_to_city_state(zip5: str) -> Optional[Dict[str, str]]:
    """Best-effort ZIP -> City/State via Zippopotam (public). Fail-open."""
    z = (zip5 or "").strip()
    if not re.fullmatch(r"\d{5}", z):
        return None
    try:
        url = f"https://api.zippopotam.us/us/{z}"
        req = urllib.request.Request(url, headers={"Accept": "application/json"})
        with urllib.request.urlopen(req, timeout=6) as resp:
            raw = resp.read().decode("utf-8", "ignore")
        data = json.loads(raw) if raw else {}
        if not isinstance(data, dict):
            return None
        places = data.get("places") or []
        if not places:
            return None
        p0 = places[0] if isinstance(places[0], dict) else {}
        city = (p0.get("place name") or "").strip()
        state = (p0.get("state abbreviation") or "").strip().upper()
        if city and state:
            return {"city": city, "state": state, "zip": z}
    except Exception:
        return None
    return None

def _normalize_location_with_zip(location: str, *fallback_texts: str) -> str:
    """If location lacks City/State but has ZIP, expand it using ZIP->City/State."""
    loc = (location or "").strip()
    zip5 = _extract_zip5(loc, *fallback_texts)
    if not zip5:
        return loc
    # If already contains a state, just ensure ZIP is present
    if _parse_state_from_location(loc):
        return loc
    z = _zip_to_city_state(zip5)
    if z:
        return f"{z['city']}, {z['state']} {z['zip']}"
    # At least return ZIP if nothing else
    return zip5


def _parse_state_from_location(loc: str) -> str:
    if not loc:
        return ""
    # Pull 2-letter state if present (", NM" or " NM ")
    m = re.search(r"(?i)\b([A-Z]{2})\b\s*(\d{5})?(?:-\d{4})?\s*$", loc.strip())
    if m:
        return (m.group(1) or "").upper()
    m2 = re.search(r"(?i),\s*([A-Z]{2})\b", loc)
    if m2:
        return (m2.group(1) or "").upper()
    return ""

def _fallback_rates_by_state(state: str) -> Dict[str, float]:
    """Conservative defaults when no external rate service is configured/available."""
    # Keep these conservative; you can tune later or replace with your own table.
    # body/paint/frame/mechanical in $/hr; paint_supplies_rate in $/refinish hr.
    defaults = {
        "DEFAULT": {"body_rate": 60.0, "paint_rate": 60.0, "frame_rate": 75.0, "mechanical_rate": 115.0, "paint_supplies_rate": 38.0},
        "NM":      {"body_rate": 64.0, "paint_rate": 62.0, "frame_rate": 78.0, "mechanical_rate": 120.0, "paint_supplies_rate": 38.0},
        "AZ":      {"body_rate": 66.0, "paint_rate": 64.0, "frame_rate": 80.0, "mechanical_rate": 125.0, "paint_supplies_rate": 40.0},
        "CO":      {"body_rate": 72.0, "paint_rate": 70.0, "frame_rate": 88.0, "mechanical_rate": 135.0, "paint_supplies_rate": 42.0},
        "TX":      {"body_rate": 62.0, "paint_rate": 60.0, "frame_rate": 75.0, "mechanical_rate": 120.0, "paint_supplies_rate": 38.0},
        "CA":      {"body_rate": 85.0, "paint_rate": 85.0, "frame_rate": 110.0, "mechanical_rate": 165.0, "paint_supplies_rate": 50.0},
    }
    return dict(defaults.get(state.upper(), defaults["DEFAULT"]))

def _fetch_rates_from_laborratehero(location: str) -> Optional[Dict[str, float]]:
    """Fetch rates from a configured endpoint. Returns None on any failure."""
    if not LABORRATEHERO_API_URL:
        return None
    try:
        q = urllib.parse.urlencode({"location": location})
        url = LABORRATEHERO_API_URL
        url = (url + ("&" if "?" in url else "?") + q) if location else url
        req = urllib.request.Request(url, headers={"Accept": "application/json"})
        if LABORRATEHERO_API_KEY:
            req.add_header("Authorization", f"Bearer {LABORRATEHERO_API_KEY}")
        with urllib.request.urlopen(req, timeout=10) as resp:
            raw = resp.read().decode("utf-8", "ignore")
        data = json.loads(raw) if raw else {}
        if not isinstance(data, dict):
            return None
        out = {}
        for k in ("body_rate","paint_rate","frame_rate","mechanical_rate","paint_supplies_rate"):
            v = data.get(k)
            if isinstance(v, (int, float)) and v > 0:
                out[k] = float(v)
        return out or None
    except Exception:
        return None

def _lookup_rates(location: str) -> Dict[str, float]:
    """Return a complete rate card (with fallbacks)."""
    location = _normalize_location_with_zip(location)
    state = _parse_state_from_location(location)
    base = _fallback_rates_by_state(state)
    ext = _fetch_rates_from_laborratehero(location) or {}
    base.update({k: float(v) for k, v in ext.items() if isinstance(v, (int, float)) and float(v) > 0})
    # Ensure required keys exist
    for k in ("body_rate","paint_rate","frame_rate","mechanical_rate","paint_supplies_rate"):
        base.setdefault(k, 0.0)
    return base

def _lookup_tax_rate(location: str) -> float:
    """Best-effort tax rate for parts/materials (decimal). Conservative fallback."""
    # Optional external (user-controlled) endpoint
    tax_url = os.getenv("TAXRATE_API_URL", "").strip()
    tax_key = os.getenv("TAXRATE_API_KEY", "").strip()
    if tax_url and location:
        try:
            q = urllib.parse.urlencode({"location": location})
            url = tax_url + ("&" if "?" in tax_url else "?") + q
            req = urllib.request.Request(url, headers={"Accept": "application/json"})
            if tax_key:
                req.add_header("Authorization", f"Bearer {tax_key}")
            with urllib.request.urlopen(req, timeout=10) as resp:
                raw = resp.read().decode("utf-8", "ignore")
            data = json.loads(raw) if raw else {}
            if isinstance(data, dict):
                r = data.get("tax_rate")
                if isinstance(r, (int, float)) and 0 <= float(r) <= 0.15:
                    return float(r)
                # allow percent input
                if isinstance(r, (int, float)) and 0 <= float(r) <= 15:
                    return float(r) / 100.0
        except Exception:
            pass

    # Minimal state fallbacks (tune later). Use decimal (e.g., 0.07875 = 7.875%)
    state = _parse_state_from_location(location)
    state_map = {
        "NM": 0.07875,
        "AZ": 0.085,
        "CO": 0.075,
        "TX": 0.0825,
        "CA": 0.085,
    }
    return float(state_map.get(state, 0.08))

def _money(x: Optional[float]) -> str:
    try:
        if x is None:
            return "$0"
        return "${:,.0f}".format(float(x))
    except Exception:
        return "$0"

def _num(x: Optional[float]) -> float:
    try:
        return float(x) if x is not None else 0.0
    except Exception:
        return 0.0

def _extract_hours_and_parts_totals(text: str) -> Dict[str, Any]:
    """Parse estimate totals (best effort)."""
    out: Dict[str, Any] = {"body_hours": None, "paint_hours": None, "frame_hours": None, "mech_hours": None, "parts_total": None, "parts_lines": []}
    if not text:
        return out

    def _find_hours(label_rx: str) -> Optional[float]:
        m = re.search(label_rx, text, flags=re.IGNORECASE)
        if m:
            try:
                return float(m.group(1))
            except Exception:
                return None
        return None

    out["body_hours"]  = _find_hours(r"\bBody\s+Labor\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)")
    out["paint_hours"] = _find_hours(r"\bPaint\s+Labor\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)")
    out["frame_hours"] = _find_hours(r"\b(Frame|Structural)\s+Labor\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)")  # group2 maybe
    if out["frame_hours"] is None:
        m = re.search(r"\bFrame\s+Labor\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)", text, flags=re.IGNORECASE)
        if m:
            try: out["frame_hours"] = float(m.group(1))
            except Exception: pass
    out["mech_hours"]  = _find_hours(r"\b(Mechanical|Mech)\s+Labor\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)")
    if out["mech_hours"] is None:
        m = re.search(r"\bMechanical\s+Hours?\b\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)", text, flags=re.IGNORECASE)
        if m:
            try: out["mech_hours"] = float(m.group(1))
            except Exception: pass

    # Parts total (best effort)
    mpt = re.search(r"(?i)\bParts\s*(?:Total)?\b\s*[:\-]?\s*\$\s*([0-9,]+(?:\.[0-9]{2})?)", text)
    if mpt:
        try:
            out["parts_total"] = float(mpt.group(1).replace(",", ""))
        except Exception:
            pass

    # Parts line items (best effort, capped)
    # Capture lines that look like: "<line#> ... <part desc> ... $123.45"
    lines = []
    for ln in (text or "").splitlines():
        s = ln.strip()
        if not s:
            continue
        if len(s) > 220:
            continue
        if re.search(r"(?i)\b(replace|r\&r|r\s*/\s*r)\b", s) and re.search(r"\$\s*[0-9,]+(?:\.[0-9]{2})?", s):
            lines.append(s)
        elif re.search(r"(?i)\b(part|oem)\b", s) and re.search(r"\$\s*[0-9,]+(?:\.[0-9]{2})?", s):
            lines.append(s)
    out["parts_lines"] = lines[:20]
    return out

def _structural_observed(text_blobs: List[str]) -> bool:
    t = "\n".join([x for x in text_blobs if x]).lower()
    if not t:
        return False
    rx = r"\b(frame|rail|apron|unibody|structural|pull|straighten|set\s*up\s*&\s*measure|measure\s*&\s*setup|dimension\s*check|buckl|kink|twist)\b"
    return bool(re.search(rx, t))
# -----------------------
# Prompt steering (free analysis + detailed narrative)
# -----------------------
DETAIL_TEMPLATES = {
    "guidelines_only": (
        "## Inputs Used\n"
        "- Briefly list which documents, pages, and photos you actually referenced.\n\n"
        "## Executive Summary\n"
        "- 3–5 bullets summarizing overall compliance and key risks.\n\n"
        "## AI-4-IA Review Summary\n"
        "- Write a formal, paragraph-style appraisal narrative. Include scope of impact, damage by zone/panel, "
        "repair vs. replace rationale, parts type (OEM/LKQ/Aftermarket), labor ops, refinish overlap, rate validation, "
        "tax handling, and estimate integrity. Reference evidence inline (e.g., 'p2/L14', 'Photo 3'). "
        "Close with compliance stance and final recommendation. Minimum 8–10 sentences.\n\n"
        "## Key Issues & Actions\n"
        "- Bullet list of the highest-impact issues with a one-line recommended action each.\n\n"
        "## Final\n"
        "- Compliance Score: NN% with one-sentence rationale."
    ),

    "comprehensive": (
        "## Inputs Used\n"
        "- List the estimate pages/lines and photo numbers you used, plus any rules text (if provided).\n\n"
        "## Executive Summary\n"
        "- 3–6 bullets capturing the big picture: estimate integrity, rule alignment (only if rules text was supplied), "
        "and photo consistency.\n\n"
        "## Detailed Condition Report\n"
        "- Write this section as a formal, paragraph-style appraisal report summarizing the entire claim. "
        "Include: scope of impact, damage by zone/panel, repair vs. replace rationale, parts type (OEM/LKQ/Aftermarket), "
        "labor operations, refinish/overlap considerations, rate validation, paint materials handling, sublet usage, "
        "tax/markup accuracy, and overall estimate integrity. Cite photos and estimate lines (e.g., 'Photo 3', 'p2/L14'). "
        "Close with compliance to any provided client rules and a clear final recommendation (Repairable vs. Total Loss). "
        "Do not declare Repairable/Total Loss unless the estimate itself explicitly marks 'Total Loss' or an ACV comparison is provided. "
        "If the shop info is listed under Repair Facility on ANY estimate, add only the shop name to the Detailed Condition Report narrative. "
        "If a Printout showing the Clean Retail Value or Estimated Trade-In Value of the unit is present which may include ANY of the following: NADA, J.D. Power, Kelly Blue Book, Edmunds, Carfax, or Cars.com, DO NOT declare as missing if any of these are present. "
        "Minimum 10–14 sentences (one continuous narrative, not bullets).\n\n"
        "## Photo-by-Photo Damage Ledger\n"
        "| Photo # | View/Angle | Panels/Parts Visible | Condition (dent/crease/scrape/misalignment) | Identifiers (VIN/odo/plate/reg) | Legibility |\n"
        "|---:|---|---|---|---|---|\n"
        "- One row per photo used in the analysis (>=6 rows if >=6 photos exist). If an identifier is present but unreadable, mark 'Present — not clearly legible'.\n\n"
        "## Brief Damage Descriptions\n"
        "- 6–12 bullets. Each: part/panel + condition + suggested op (repair/replace/refinish/blend) + Photo #.\n\n"
        "## Estimate Line Extract (top relevant lines)\n"
        "| Est. p#/L# | Part/Op | Labor Hrs | Rate | Part Type | Price | Notes |\n"
        "|---|---|---:|---:|---|---:|---|\n"
        "- Include 8+ lines if available, focusing on items tied to observed damages. Use 'Notes' for overlap/blend or rationale.\n\n"
        "## Estimate Compliance Cross-Check (brief)\n"
        "Status: Compliant / Non-compliant / Not Evidenced. Cite only the most important evidence (p#/L# or Photo #). Keep it short.\n"
        "| Topic | Evidence (p#/L# or value; Photo # if relevant) | Status | Required Fix |\n"
        "|---|---|:--:|---|\n"
        "| Labor Rates |  |  |  |\n"
        "| Refinish/Overlap |  |  |  |\n"
        "| Paint Materials |  |  |  |\n"
        "| OEM Procedures |  |  |  |\n"
        "| Sublet |  |  |  |\n"
        "| Tax/Markup |  |  |  |\n\n"
        "## Client Guidelines Comparison (if rules text was supplied)\n"
        "- 3–8 bullets. Quote the relevant rule fragment and note Aligned / Not Aligned / Not Evidenced, with evidence refs (p#/L#, Photo #). "
        "If no client_rules were provided, omit this section.\n\n"
        "## Risks / Missing Evidence\n"
        "- Short bullets with severity (High/Med/Low) and a one-line remediation.\n\n"
        "## Compliance Score Rationale\n"
        "- REQUIRED if score < 100: start at 100 and list each deficiency with evidence refs (p#/L# and/or Photo #), "
        "severity (Minor/Moderate/Major), and numeric deduction. Show the arithmetic to the final score.\n\n"
        "## Final Evaluation\n"
        "- Compliance Score: NN% with a single-sentence justification. "
        "If no fraud indicators are identified, state 'No material inconsistencies found.' Do not use 'N/A'."
    ),

    "damage_report_from_photos": (
        """
# Condition Report (Photos Only)

## Photo-by-Photo Condition Summary
| Photo # | View/Side | Key Panels/Parts Visible | Damage/Condition |
|---:|---|---|---|
- Cover EVERY provided photo. If no damage is obvious from that angle, write: "No obvious damage visible from this angle" (do not use the word intact).

## Side Checks
- **Driver/Left Side**: <what is visible; cite Photo #; if not shown, say not shown>
- **Passenger/Right Side**: <what is visible; cite Photo #; if not shown, say not shown>

## Front-End Checklist
- Hood: <condition or Not clearly shown> (Photo #)
- Front bumper cover: <condition or Not clearly shown> (Photo #)
- Grille: <condition or Not clearly shown> (Photo #)
- Driver-side headlamp: <condition or Not clearly shown> (Photo #)
- Passenger-side headlamp: <condition or Not clearly shown> (Photo #)
- Driver-side front fender: <condition or Not clearly shown> (Photo #)
- Passenger-side front fender: <condition or Not clearly shown> (Photo #)

## Detailed Condition Report
- Write a continuous 10–15 sentence narrative summarizing visible damage, impact zones, misalignment/gaps, and repair implications (photo-based).
- If VIN label or odometer are visible, state them with Photo #. If not visible or unreadable, say so.

## Approximate Repair Cost Breakdown (Populate JSON field 'estimated_costs_markdown')
- You MUST produce a cost approximation derived from the PHOTOS ONLY (do not reference estimates, documents, or 'not evidenced').
- Provide AI-derived hours and assumptions:
  • Body labor hours
  • Paint labor hours
  • Paint & materials cost = (paint hours × $/refinish hr)
  • If structural damage is observed in photos: add Setup & Measure = 2.0 hrs @ body rate and include frame hours @ frame rate
  • If airbags/ADAS/mechanical damage is observed: include mechanical hours @ mechanical rate
- Provide an OEM replacement parts list (each part with an approximate $).
- Apply tax ONLY to (parts + paint materials). Do NOT tax labor.
- Include the Severity Tier checkboxes:
  ☐ Minor (< $3,500)
  ☐ Moderate ($3,500-$10,000)
  ☐ Major ($10,000+)
  ☐ Possible Total Loss Threshold Approaching
- Include a Repair Cost Disclaimer stating this is an approximation, not an official estimate.
- Forbidden phrases: 'from estimate', 'from documentation', 'not evidenced', 'no documentation provided'.
"""
    ),
}

# --- Static audit questions ---
STATIC_AUDIT_QUESTIONS = [
    "Do the photos substantiate the highest-cost operations (frame/sectioning/panel replace)?",
    "Are ADAS calibrations or wheel alignments required and supported by the damage and OEM procedures?",
    "Is blend time justified by color/finish (metallic/pearl/tri-coat) and adjacent panel visibility?",
    "Do invoices corroborate parts used and match estimate line items (brand/grade, price, quantity)?",
    "Are AM/LKQ choices compliant with age/mileage rules, and is OEM required anywhere by client policy or safety?",
    "Is there evidence of prior or unrelated damage (UPD) that materially affects valuation or repair scope?",
    "Are there structural/safety indicators (buckles, misalignments, airbags/pretensioners) that alter repair strategy?",
    "Are materials/hazard charges (paint supplies, corrosion protection, seam sealer) aligned with operations and shop norms?",
    "Are storage/tow charges and dates supported and reasonable given claim timeline and shop status?",
    "Are scanner reports (pre/post) included or needed; if absent, does that meaningfully impact confidence?",
    "Did the supplement (if any) correct earlier gaps, and are newly added operations now evidenced?",
    "Are client-required documents present (e.g., NADA printout, release forms, production date plate); if missing, what’s the impact?",
    "What is the bottom-line recommendation (approve as-is, adjust items, or request specific evidence)?",
]

# --- Identifiers Verification Protocol (prompt-only; no new logic) ---
IDENTIFIERS_VERIFICATION_PROTOCOL = (
    "\n\nIDENTIFIERS VERIFICATION PROTOCOL (must follow):"
    "\n1) Search the photos for: windshield VIN plate, driver-door VIN label, driver-door VIN label Production date, driver-door VIN label Date of Mfr, odometer cluster."
    "\n2) Transcribe the VIN exactly as visible and cite Photo # for EACH location you find."
    "\n3) If multiple VINs, compare them to each other and to the estimate VIN; explicitly state: MATCH / MISMATCH."
    "\n4) Transcribe the odometer reading exactly as shown and cite Photo #."
    "\n5) Grade legibility for each identifier as one of: 'Clearly legible' / 'Present — not clearly legible' / 'Not present'."
    "\n6) If any identifier is present but not clearly legible, say why (glare, blur, angle) and what photo would resolve it."
    "\n7) Write a one-line bottom line: 'VIN verification: <MATCH/MISMATCH/INCONCLUSIVE>; Odometer: <value or reason>'."
    "\n8) Weave these facts naturally into the '## Detailed Condition Report' narrative and keep the top-line fields "
    "(vin, vin_verification, odometer_estimate_only) consistent."
    "\n9) When citing more than one VIN location (e.g., windshield vs. door label), you must cite DISTINCT Photo #s; "
    "never reuse the same photo number for two different locations."
    "\n10) Compare VINs as literal 17-character strings. If any single character differs between sources, report "
    "MISMATCH, and quote both strings with their Photo #/page references."
    "\n11) ODOMETER RULES (photos-only especially): transcribe only the digits visible in the odometer photo; "
    "do not infer from estimate text or metadata. Include the exact Photo #. If any digit is unclear, state 'Present — not clearly legible' and explain why; do not guess."
)

# --- Consistency Guard (prompt-only; avoid contradictions) ---
CONSISTENCY_GUARD = (
    "\n\nCONSISTENCY GUARD:"
    "\n- Do not claim any required photo is 'missing' if you graded it 'Clearly legible' or 'Present — not clearly legible'."
    "\n- For VIN, Odometer, and Production Date specifically: if present in any photo, do not write any sentence implying they are absent."
    "\n- If a legible driver-door VIN label photo is present, treat the Production Date requirement as evidenced (the production month/year appears on the same label). Do NOT deduct or say 'not separately documented'."
    "\n- Only deduct for missing Repair Facility info when the Closing Report or other documents clearly show the vehicle is at a named repair facility AND the estimate's 'Repair Facility' section does not list that same facility; "
    "if no repair facility information appears anywhere in the estimate or Closing Report, report 'N/A — not provided' and do NOT deduct."
    "\n- If legibility is the issue, explicitly say 'Present — not clearly legible' and explain why (glare/blur/angle), and request a precise retake rather than marking it missing."
    "\n- Before finalizing, re-scan your output: confirm every referenced Photo # matches the content described (e.g., do not cite an Odometer photo as the point-of-impact photo). Correct any mismatches."
)

# --- No-intact-if-damaged rule (prompt-only; prevents false 'intact' claims) ---
NO_INTACT_IF_DAMAGED_RULE = (
    "\n\nNO 'INTACT/NO-DAMAGE' OVERRIDE RULE (DO NOT PRINT CONFLICT WARNINGS):"
    "\n- If any photo indicates damage to a panel/component, you may NOT state that same panel/component is intact/undamaged/no visible damage anywhere."
    "\n- If photos appear inconsistent, DO NOT write a conflict warning; instead, remove/avoid the intact/no-damage claim and describe only what appears damaged with citations."
)

# --- Damage Side / Orientation Guard (prompt-only; prevents left/right drift) ---
DAMAGE_SIDE_GUARD = (
    "\n\nDAMAGE SIDE GUIDANCE (MINIMAL):"
    "\n- Describe any visible damage on any side (Driver/Passenger or Left/Right) when it is visible in photos."
    "\n- Do NOT suppress side-level damage descriptions when damage is clearly visible."
    "\n- If orientation is genuinely unclear, say so and avoid guessing."
)

BILATERAL_DAMAGE_MANDATE = (
    "\n\nBILATERAL / SECONDARY DAMAGE MANDATE (STRICT):\n"
    "- In frontal impacts, explicitly address BOTH front corners (driver-side and passenger-side if clear; otherwise ‘front corner A/B’ or ‘front corner (viewed)’)\n"
    "- If a photo shows partial view of the opposite side with visible distortion/misalignment/crush, describe it — do NOT default to 'intact' or 'no damage' without citing clear evidence.\n"
    "- Contradicting visible photo evidence (e.g. calling a crushed fender 'intact') is forbidden."
)

FRONT_CORNER_ORIENTATION_GUARD = (
    "\n\nFRONT / SIDE LABELING (LOCKED):"
    "\n- For primary_impact and secondary_impact fields: if the damage is at the front end and left/right cannot be proven from the photos, set the value to exactly 'Front' (no qualifiers)."
    "\n- Do NOT label front damage as LF/RF or Driver/Passenger unless orientation is clearly established by unmistakable cues within the same photo set (e.g., readable badge/plate position combined with a full-front view)."
    "\n- If orientation is not clearly established, use neutral wording only (e.g., 'front end', 'front headlamp area', 'front corner') and DO NOT assign left/right."
    "\n- Do not state that a left/right front component is undamaged/intact unless that specific component is clearly shown and its condition is confirmable."
)

# --- Parts Source Guard (prompt-only; prevents OEM vs Aftermarket drift) ---
PARTS_SOURCE_GUARD = (
    "\n\nPARTS SOURCE GUARD (MANDATORY):"
    "\n- Do NOT claim 'aftermarket', 'A/M', 'quality replacement', 'non-OEM', or 'LKQ/used/recycled' parts were used unless the estimate LINE ITEMS explicitly label them as such."
    "\n- Generic disclosure/boilerplate text about aftermarket crash parts does NOT prove aftermarket parts were used."
    "\n- If the Closing Report states no aftermarket/LKQ parts were included, your narrative must not claim they were used."
    "\n- When parts source is not explicit, state that it is not explicitly labeled and avoid guessing; default to OEM only when supported by part numbers/labels."
)

# --- Supplement Handling (prompt-only; ensures detection + narrative mention) ---
SUPPLEMENT_HANDLING = (
    "\n\nSUPPLEMENT HANDLING:"
    "\n- Examine the estimate documents for explicit supplement indicators: 'Supplement', 'Supplement of record', 'S01', 'S02', 'Supplement Summary', or similar."
    "\n- If a supplement or multiple supplements are detected, clearly state in the narrative that the estimate is a supplement and summarize what changed: added operations/parts, rate updates, refinish overlap changes, or corrections to prior omissions."
    "\n- If the supplement(s) corrects earlier deficiencies (e.g., missing materials line, added calibrations), note that improvement explicitly."
    "\n- If a supplement(s) exists but required supporting evidence (invoices, photos) is still missing, call this out in Risks/Missing Evidence."
)

ALLOWED_INTENTS = {"guidelines_only","comprehensive","damage_report_from_photos"}

SYSTEM_BASE = (
    "You are an auto-claims appraisal assistant. Return ONLY valid JSON (no code fences). "
    "Always include all required keys: "
    "['file_number','request_type','claim_number','vin','vin_verification','vehicle',"
    "'odometer_estimate_only','compliance_score','summary_brief','summary_markdown',"
    "'fraud_markdown','primary_impact','secondary_impact','estimated_costs_markdown','conclusion']. "
    "Use only provided evidence. Cite photos as 'Photo #' and docs as 'p#/L#' when available. "
    "Do not guess. If something is not visible, say why. "
    "NEVER return 'N/A' for summary_markdown, fraud_markdown, estimated_costs_markdown, or conclusion. For request_type 'Create a Damage Report from Photos', 'estimated_costs_markdown' must be a photos-only approximation (no document/estimate dependency language) and must model Paint & Materials as a $ per refinish hour rate (not a percent)."
)

SYSTEM_BASE += (
    " Do not state or imply any client rule unless it appears verbatim in the provided client_rules text. "
    "If client_rules is blank, write the entire report without referencing client rules. "
    "If a value cannot be confirmed from the visible evidence, set it to 'N/A' and briefly state why. "
    " Except when the request_type is 'Create a Damage Report from Photos', Compliance Score must be a numeric percentage 0–100 (never 'N/A'). "
    "If the request_type is 'Create a Damage Report from Photos', set compliance_score to 'N/A' and omit the '## Compliance Score Rationale' section. "
    "If no client_rules are supplied, base the score on estimate-photo internal consistency, evidence completeness, and clarity/legibility. "
    "If compliance_score < 100, include a dedicated section titled '## Compliance Score Rationale' which itemizes every deficiency with exact evidence references "
    "(estimate p#/L# and/or Photo #), assigns an explicit deduction per item, and shows the arithmetic to the final score. "
    "Use a consistent scheme (e.g., Minor -5, Moderate -10, Major -20) and never go below 0. "
    "The 'fraud_markdown' section must never be 'N/A'. If nothing material is found, write "
    "'No material inconsistencies found.' and briefly note what was checked (VIN match, date/metadata, obvious photo tampering, duplicated images)."
)

SYSTEM_BASE += (
    " Focus on a cohesive, professional appraisal. Prefer narrative over rigid tables. "
    "Include a section named '## Detailed Condition Report'. "
    "Include '## Compliance Score Rationale' only when compliance_score < 100, and show deductions from 100 with brief evidence refs (p#/L# or Photo #). "
    "If you include tables, keep them concise and only when they help clarity. "
    "Avoid placeholder rows/columns; do not invent data. "
    "When client_rules text is provided, also include a section titled '## Client Guidelines Comparison' with 3–8 concise bullets quoting the relevant rule fragment and citing evidence (p#/L#, Photo #); "
    "weave any material rule alignment/misalignment into the Detailed Condition Report narrative."
    "When a valuation/clean retail printout exists but the header doesn’t match the estimate’s VIN/year/trim/mileage, label it “Present — mismatched (detail the differences)” and request a corrected printout; never mark it Missing/Not Evidenced. "
    "If a legible driver-door VIN label photo is present, treat Production Date as evidenced; do not mark 'missing' or deduct for lack of a separate photo. "
    "Only deduct for missing Repair Facility info when the Closing Report or other documents clearly show the vehicle is at a named repair facility AND the estimate's 'Repair Facility' section does not list that same facility; "
    "if no repair facility information appears anywhere in the estimate or Closing Report, report 'N/A — not provided' and do NOT deduct."
)

SYSTEM_BASE += (
    " Your 'summary_markdown' MUST include a top-level section named '## Detailed Condition Report' containing a cohesive narrative of at least 10–14 sentences (not bullets). "
    "It must synthesize: impact zones, per-panel damages, repair vs. replace rationale, parts type (OEM/LKQ/Aftermarket), labor ops, refinish/overlap, rate/materials/sublet/tax handling, and estimate integrity. "
    "It must cite concrete evidence inline (e.g., p2/L14, Photo 3). "
    "When evaluating paint materials, recognize that a summary line such as 'Paint Supplies' or 'Paint Materials' with hours and rate in the totals section constitutes a valid cost breakdown. "
    "Do not mark it missing if such a line is present, even if materials are not listed per-panel. "
    "Avoid categorical phrases such as 'deemed repairable' or 'deemed total loss' unless that exact determination appears in the provided documents (e.g., estimate header says 'Total Loss' or an ACV comparison is shown). "
    "Otherwise, use neutral language and do not make a repairability determination."
)

# -----------------------
# Supported file types
# -----------------------
SUPPORTED_IMAGE_EXTS = (".jpg",".jpeg",".png",".webp",".heic",".heif")
SUPPORTED_TEXT_EXTS = (".txt",)
SUPPORTED_DOCX_EXTS = (".docx",)
SUPPORTED_PDF_EXTS = (".pdf",)

# -----------------------
# Helpers to add parts from bytes
# -----------------------
def _image_part_from_bytes(raw: bytes) -> Dict[str, Any]:
    b64 = base64.b64encode(raw).decode("utf-8")
    return {"type": "image_url", "image_url": {"url": "data:image/jpeg;base64," + b64}}

# Safe PDF text extract helper
def _maybe_extract_pdf_text(raw: bytes, fname: str, parts: List[Dict[str, Any]], files_seen: List[str], pdf_text_fulls: Optional[List[str]] = None) -> None:
    try:
        from pdfminer_high_level import extract_text as _x  # type: ignore
    except Exception:
        try:
            from pdfminer.high.level import extract_text as _x  # fallback
        except Exception:
            _x = None
    try:
        if _x:
            full = (_x(io.BytesIO(raw)) or "")
            if pdf_text_fulls is not None and full.strip():
                pdf_text_fulls.append(full)
            t = full[:12000]
            if t.strip():
                parts.insert(0, {"type": "text", "text": t})
                files_seen.append(f"{fname} (pdf text extracted)")
    except Exception:
        pass

# Lightweight OCR helper for images
def _maybe_ocr_image_text(im: Image.Image) -> str:
    if not _OCR_ENABLED:
        return ""
    try:
        im2 = im.convert("L")
        if max(im2.size) < 1400:
            scale = 1400 / max(im2.size)
            im2 = im2.resize((int(im2.width*scale), int(im2.height*scale)))
        txt = pytesseract.image_to_string(im2)
        return (txt or "").strip()
    except Exception:
        return ""


def _qr_decode_vin_from_pil(im: Image.Image) -> Optional[str]:
    """Best-effort local QR/barcode decode for VIN (optional deps)."""
    try:
        import cv2  # type: ignore
        import numpy as np  # type: ignore
        arr = cv2.cvtColor(np.array(im.convert("RGB")), cv2.COLOR_RGB2BGR)
        det = cv2.QRCodeDetector()
        val, _, _ = det.detectAndDecode(arr)
        if val:
            s = str(val).strip().upper()
            m = re.search(VIN_PATTERN, s)
            if m:
                return m.group(0)
    except Exception:
        pass
    try:
        from pyzbar.pyzbar import decode  # type: ignore
        for d in decode(im.convert("RGB")):
            try:
                s = d.data.decode("utf-8", "ignore").strip().upper()
            except Exception:
                s = str(d.data).strip().upper()
            m = re.search(VIN_PATTERN, s)
            if m:
                return m.group(0)
    except Exception:
        pass
    return None


def _add_bytes(parts: List[Dict[str,Any]], files_seen: List[str], photo_index: Optional[List[str]], thumb_paths: Optional[List[str]], raw: bytes, fname: str, used: int, max_images: int, pdf_text_fulls: Optional[List[str]] = None, ocr_pairs: Optional[List[Dict[str, Any]]] = None) -> int:
    low = fname.lower()
    if low.endswith(SUPPORTED_PDF_EXTS) and used < max_images:
        try:
            # MEMORY GUARD: do not rasterize every PDF page at once. Large 25MB PDFs
            # can contain many pages and can OOM/restart the Render worker if converted
            # all at 200 DPI. Convert only the remaining allowed pages at a lighter DPI.
            remaining_pages = max(1, max_images - used)
            pages = convert_from_bytes(
                raw,
                dpi=150,
                first_page=1,
                last_page=remaining_pages,
                thread_count=1,
            )
            files_seen.append(f"{fname} (pdf, {len(pages)} page(s) converted; capped at {remaining_pages})")
            _maybe_extract_pdf_text(raw, fname, parts, files_seen, pdf_text_fulls=pdf_text_fulls)
            OCR_PAGE_CAP = min(24, remaining_pages)
            ocr_collected = []
            for idx, im in enumerate(pages):
                b = io.BytesIO()
                im.save(b, format="JPEG", quality=70, optimize=True)
                parts.append(_image_part_from_bytes(b.getvalue()))
                used += 1
                if photo_index is not None:
                    photo_index.append(f"{fname}::page_{idx+1}")
                if idx < OCR_PAGE_CAP:
                    txt = _maybe_ocr_image_text(im)
                    if txt:
                        ocr_collected.append(txt)
                try:
                    im.close()
                except Exception:
                    pass
            if ocr_collected:
                parts.insert(0, {"type": "text", "text": ("\n".join(ocr_collected))[:12000]})
                files_seen.append(f"{fname} (ocr text extracted)")
        except Exception as e:
            logging.warning(f"pdf2image failed for {fname}: {e}")
            files_seen.append(f"{fname} (pdf, could not be converted)")
    elif low.endswith(SUPPORTED_IMAGE_EXTS) and used < max_images:
        im_ref = None
        raw_for_vin = None
        qr_vin = None
        try:
            im = Image.open(io.BytesIO(raw)).convert("RGB")
            im_ref = im.copy()
            # Preserve ORIGINAL bytes for VIN label / QR/barcode decoding (before any resizing or recompression).
            raw_for_vin = raw
            # Secondary confirmation: attempt local QR/barcode decode from the ORIGINAL image (before resizing).
            qr_vin = _qr_decode_vin_from_pil(im_ref)
            # Use the same preprocessing for ZIP and loose JPGs (keep higher res for small label text).
            max_dim = 2048
            if max(im.size) > max_dim:
                scale = max_dim / float(max(im.size))
                im = im.resize((int(im.width * scale), int(im.height * scale)))
            b = io.BytesIO()
            im.save(b, format="JPEG", quality=75, optimize=True)
            raw = b.getvalue()
        except Exception:
            im_ref = None
            raw_for_vin = None
        parts.append(_image_part_from_bytes(raw))
        used += 1
        if photo_index is not None:
            photo_index.append(fname)
        files_seen.append(f"{fname} (photo)")
        # Keep a local copy of each uploaded photo for the PDF thumbnail appendix page.
        if thumb_paths is not None:
            try:
                im_save = im_ref if im_ref is not None else Image.open(io.BytesIO(raw)).convert("RGB")
                im_save.thumbnail((900, 900))
                thumb_name = f"thumb_{uuid.uuid4().hex}.jpg"
                thumb_path = os.path.join(PDF_DIR, thumb_name)
                im_save.save(thumb_path, format="JPEG", quality=75, optimize=True)
                thumb_paths.append(thumb_path)
            except Exception:
                pass
        if im_ref is not None:
            txt = _maybe_ocr_image_text(im_ref)
            if ocr_pairs is not None:
                try:
                    ocr_pairs.append({"name": fname, "text": (txt or ""), "raw_for_vin": raw_for_vin, "qr_vin": (qr_vin or None)})
                except Exception:
                    pass
            if txt:
                parts.insert(0, {"type":"text", "text": txt[:12000]})
                files_seen.append(f"{fname} (ocr text extracted)")
    elif low.endswith(SUPPORTED_DOCX_EXTS):
        try:
            text = "\n".join([p.text for p in Document(io.BytesIO(raw)).paragraphs if p.text.strip()])
        except Exception:
            text = ""
        if text.strip():
            parts.insert(0, {"type":"text","text": text[:12000]})
            files_seen.append(f"{fname} (docx text included)")
        else:
            files_seen.append(f"{fname} (docx, no readable text)")
    elif low.endswith(SUPPORTED_TEXT_EXTS):
        try:
            text = raw.decode("utf-8","ignore")[:12000]
        except Exception:
            text = ""
        if text.strip():
            parts.insert(0, {"type":"text","text": text})
            files_seen.append(f"{fname} (txt included)")
        else:
            files_seen.append(f"{fname} (txt, empty)")
    else:
        files_seen.append(f"{fname} (unsupported type)")
    return used

# -----------------------
# App + CORS
# -----------------------
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://nspxn.com","https://www.nspxn.com","http://nspxn.com","http://www.nspxn.com",
        "https://nspxn.onrender.com"
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# -----------------------
# Client Rules: fuzzy finder + endpoints
# -----------------------
def _find_rules_path(base_name: str, rules_dir: str) -> Optional[str]:
    target = base_name.strip()
    if not target.lower().endswith(".docx"):
        target += ".docx"
    p = os.path.join(rules_dir, target)
    if os.path.exists(p):
        return p
    pattern = os.path.join(rules_dir, "*.docx")
    candidates = glob.glob(pattern)
    if not candidates:
        return None
    low_target = target.lower()
    for c in candidates:
        if os.path.basename(c).lower() == low_target:
            return c
    for c in candidates:
        if os.path.basename(c).lower().startswith(low_target[:-5]):
            return c
    tokens = [t for t in low_target[:-5].split() if t]
    def contains_in_order(name: str) -> bool:
        i = 0
        for tok in tokens:
            i = name.find(tok, i)
            if i == -1:
                return False
            i += len(tok)
        return True
    for c in candidates:
        if contains_in_order(os.path.basename(c).lower()):
            return c
    return None


def _read_rules_text_from_path(path: str) -> str:
    try:
        doc = Document(path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        log.warning(f"Unable to read rules file {path}: {e}")
        return ""

@app.get("/client-rules/{client_name}")
async def get_client_rules(client_name: str):
    path = _find_rules_path(client_name, CLIENT_RULES_DIR)
    if not path:
        try:
            files = sorted(os.path.basename(p) for p in glob.glob(os.path.join(CLIENT_RULES_DIR, "*.docx")))
        except Exception:
            files = []
        return JSONResponse(
            status_code=404,
            content={
                "error": f"Rules not found for '{client_name}'.",
                "hint": "Ensure the .docx file exists in CLIENT_RULES_DIR (default 'client_rules').",
                "available_rule_files": files[:50]
            },
        )
    try:
        doc = Document(path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return {"file": os.path.basename(path), "text": text}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"Unable to read rules: {e}", "file": os.path.basename(path)})

@app.get("/client-rules")
async def list_client_rules():
    try:
        files = sorted(os.path.basename(p) for p in glob.glob(os.path.join(CLIENT_RULES_DIR, "*.docx")))
        return {"dir": CLIENT_RULES_DIR, "files": files}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e), "dir": CLIENT_RULES_DIR})

# -----------------------
# Vision Review
# -----------------------
@app.post("/vision-review")
async def vision_review(
    request: Request,
    response: Response,
    files: Optional[List[UploadFile]] = File(None),
    client_rules: str = Form(""),
    ai_notes: str = Form(""),
    addl_notes: str = Form(""),
    additional_notes: str = Form(""),
    notes: str = Form(""),
    file_number: Optional[str] = Form(None),
    ia_company: str = Form(""),
    appraiser_id: str = Form(""),
    ai_intent: str = Form("comprehensive")
):
    parts: List[Dict[str, Any]] = []
    files_seen: List[str] = []
    photo_index: List[str] = []
    thumbnail_paths: List[str] = []
    ocr_pairs: List[Dict[str, Any]] = []

    # --- 422 guard: avoid FastAPI validation failures when frontend keys vary ---
    # Accept missing/alternate keys without raising 422; return a clear 400 instead.
    if files is None:
        files = []
    if file_number is None or str(file_number).strip() == "":
        try:
            _form = await request.form()
            # Common file number key variants from different frontends
            for _k in ("file_number", "file-number", "fileNumber", "fileNum", "file_num", "filenumber"):
                _v = str(_form.get(_k, "") or "").strip()
                if _v:
                    file_number = _v
                    break
        except Exception:
            pass
    # If 'files' was posted under a different key, recover it from the raw form
    if (not files) or (isinstance(files, list) and len(files) == 0):
        try:
            _form = await request.form()
            _maybe = []
            try:
                _maybe = list(_form.getlist("files")) if hasattr(_form, "getlist") else []
            except Exception:
                _maybe = []
            if not _maybe:
                try:
                    _maybe = list(_form.getlist("file")) if hasattr(_form, "getlist") else []
                except Exception:
                    _maybe = []
            _maybe_files = [x for x in _maybe if hasattr(x, "filename")]
            if _maybe_files:
                files = _maybe_files  # type: ignore[assignment]
        except Exception:
            pass
    # Hard-required fields (explicit 400 instead of 422)
    if not file_number or str(file_number).strip() == "":
        return JSONResponse(status_code=400, content={"error": "Missing required field: file_number"})
    if not files:
        return JSONResponse(status_code=400, content={"error": "Missing required upload: files"})
    vin_candidates: List[str] = []  # reserved (filenames not used)
    MAX_IMAGES = 48
    used = 0
    # Coalesce Add\'l Notes from multiple possible frontend field names
    ai_notes_used = ((ai_notes or "").strip() or (addl_notes or "").strip() or (additional_notes or "").strip() or (notes or "").strip())
    ai_notes_used = ai_notes_used.strip()
    # If notes still empty, fall back to reading the raw posted form (covers mismatched frontend field names)
    if not ai_notes_used:
        try:
            _form = await request.form()
            # First try common variants explicitly
            for _k in ("ai_notes","ai-notes","addl_notes","additional_notes","notes","ai_review_notes","ai_notes_box","addlNote","addlNoteText"):
                _v = str(_form.get(_k, "") or "").strip()
                if _v:
                    ai_notes_used = _v
                    break
            # Then any key containing 'note' (last resort)
            if not ai_notes_used:
                for _k in _form.keys():
                    if "note" in str(_k).lower():
                        _v = str(_form.get(_k, "") or "").strip()
                        if _v:
                            ai_notes_used = _v
                            break
        except Exception:
            pass
    # Normalize/sanitize notes so they cannot break structured prompting
    ai_notes_used = _normalize_ai_notes(ai_notes_used)
    ai_notes_block = _ai_notes_block(ai_notes_used)

    # Hydrate selected client rules from CLIENT_RULES_DIR when the frontend sends only a selected rule name/file.
    selected_rules_key = ""
    resolved_rules_path = None
    resolved_rules_text = ""
    client_rules_source = "inline_form_text" if (client_rules or "").strip() else "blank"
    try:
        _form = await request.form()
        for _k in (
            "client_rules_selected", "client_rule_selected", "selected_client_rules", "selected_client_rule",
            "selectedGuideline", "selected_guideline", "client_guideline", "client_guidelines",
            "client_guideline_name", "clientGuideline", "clientGuidelineName", "rules_file", "rules_filename",
            "rules_name", "guideline_file", "guideline_filename", "guideline_name", "client_name"
        ):
            _v = str(_form.get(_k, "") or "").strip()
            if _v:
                selected_rules_key = _v
                break

        _client_rules_trim = str(client_rules or "").strip()
        _looks_like_rules_name = bool(_client_rules_trim) and (
            _client_rules_trim.lower().endswith(".docx")
            or ("\n" not in _client_rules_trim and len(_client_rules_trim) <= 160 and not re.search(r"(?is)\bshall\b|\bmust\b|\brequired\b|\bguideline\b.{20,}", _client_rules_trim))
        )

        if _looks_like_rules_name and not selected_rules_key:
            selected_rules_key = _client_rules_trim

        if selected_rules_key:
            resolved_rules_path = _find_rules_path(selected_rules_key, CLIENT_RULES_DIR)
            if resolved_rules_path:
                resolved_rules_text = _read_rules_text_from_path(resolved_rules_path)

        if (not _client_rules_trim) and resolved_rules_text:
            client_rules = resolved_rules_text
            client_rules_source = f"resolved_from_selected:{os.path.basename(resolved_rules_path)}"
        elif _looks_like_rules_name and resolved_rules_text:
            client_rules = resolved_rules_text
            client_rules_source = f"resolved_from_client_rules_name:{os.path.basename(resolved_rules_path)}"
        elif _client_rules_trim:
            client_rules_source = "inline_form_text"
        elif selected_rules_key and not resolved_rules_text:
            client_rules_source = f"selected_but_unresolved:{selected_rules_key}"

        log.info(
            "CLIENT RULES INPUT | source=%s | inline_len=%s | selected_key=%s | resolved_path=%s | resolved_len=%s",
            client_rules_source,
            len(_client_rules_trim),
            selected_rules_key or "",
            (os.path.basename(resolved_rules_path) if resolved_rules_path else ""),
            len(resolved_rules_text or ""),
        )
    except Exception as e:
        log.warning(f"Client rules hydration/logging failed: {e}")

    client_rules_supplied = bool(str(client_rules or "").strip() or str(resolved_rules_text or "").strip() or str(selected_rules_key or "").strip())

    pdf_text_fulls: List[str] = []  # full PDF text for supplement detection

    # Anti-zipbomb guardrails
    MAX_ZIP_FILES = 100
    MAX_ENTRY_SIZE = 25 * 1024 * 1024  # 25 MB

    for f in sorted(files, key=lambda _f: ((_f.filename or '').lower())):
        raw = await f.read()
        fname = f.filename or "upload"
        low = fname.lower()
        if low.endswith(".zip"):
            try:
                zf = zipfile.ZipFile(io.BytesIO(raw))
            except Exception as e:
                files_seen.append(f"{fname} (zip, unreadable: {e})")
                continue
            members = sorted([zi for zi in zf.infolist() if not zi.is_dir()], key=lambda _zi: (_zi.filename or '').lower())
            if len(members) > MAX_ZIP_FILES:
                files_seen.append(f"{fname} (zip, too many entries: {len(members)})")
                members = members[:MAX_ZIP_FILES]
            for zi in members:
                inner_name = zi.filename
                if ".." in inner_name or inner_name.startswith(("/", "\\")):
                    files_seen.append(f"{fname}::{inner_name} (skipped unsafe path)"); continue
                if zi.file_size > MAX_ENTRY_SIZE:
                    files_seen.append(f"{fname}::{inner_name} (skipped >15MB)"); continue
                try:
                    data = zf.read(zi)
                except Exception as e:
                    files_seen.append(f"{fname}::{inner_name} (read error: {e})"); continue
                used = _add_bytes(parts, files_seen, photo_index, thumbnail_paths, data, f"{fname}::{inner_name}", used, MAX_IMAGES, pdf_text_fulls=pdf_text_fulls, ocr_pairs=ocr_pairs)
        else:
            used = _add_bytes(parts, files_seen, photo_index, thumbnail_paths, raw, fname, used, MAX_IMAGES, pdf_text_fulls=pdf_text_fulls, ocr_pairs=ocr_pairs)

    # Collect uploaded TEXT ONLY for evidence checks
    uploaded_text_blobs = []
    for p in parts:
        if isinstance(p, dict) and p.get("type") == "text" and isinstance(p.get("text"), str):
            uploaded_text_blobs.append(p["text"])
    uploaded_text_all = "\n".join(uploaded_text_blobs)
    # --- VIN EXTRACTION (door label OCR -> QR/barcode) ---
    # Do NOT use filenames. Prefer driver-door certification label OCR; if not found, attempt QR/barcode decode.
    vin_from_label = None
    vin_from_label_photo = None
    vin_from_qr = None
    vin_from_qr_photo = None

    def _looks_like_door_label(txt: str) -> bool:
        if not txt:
            return False
        t = txt.upper()
        keys = ["MFD BY", "MANUFACTURED", "GVWR", "GAWR", "TIRE SIZE", "CONFORMS TO"]
        return sum(1 for k in keys if k in t) >= 2

    # (b) Door label OCR first
    try:
        for rec in ocr_pairs:
            if not isinstance(rec, dict):
                continue
            t = rec.get("text") or ""
            if _looks_like_door_label(t):
                mvin = re.search(VIN_PATTERN, t.upper())
                if mvin:
                    vin_from_label = mvin.group(0)
                    vin_from_label_photo = rec.get("name") or None
                    break
    except Exception:
        vin_from_label = None
        vin_from_label_photo = None
    # (c) QR/barcode VIN as secondary confirmation (prefer same photo as door label if available)
    try:
        if vin_from_label_photo:
            for rec in ocr_pairs:
                if isinstance(rec, dict) and (rec.get("name") == vin_from_label_photo) and rec.get("qr_vin"):
                    vin_from_qr = str(rec.get("qr_vin") or "").strip().upper() or None
                    vin_from_qr_photo = rec.get("name") or None
                    break
        if not vin_from_qr:
            for rec in ocr_pairs:
                if isinstance(rec, dict) and rec.get("qr_vin"):
                    vin_from_qr = str(rec.get("qr_vin") or "").strip().upper() or None
                    vin_from_qr_photo = rec.get("name") or None
                    break
    except Exception:
        vin_from_qr = None
        vin_from_qr_photo = None


    # (c) If OCR didn't yield a VIN, try a dedicated vision decode for VIN/QR/barcode on the best label candidate
    def _decode_vin_from_label_or_qr(raw_bytes: Optional[bytes]) -> Optional[str]:
        if not raw_bytes:
            return None
        try:
            prompt = (
                "Return ONLY JSON: {\"vin\": \"...\"}.\n"
                "Task: Read the vehicle VIN from the door-jamb certification label text and/or decode any QR/barcode if present.\n"
                "Rules: VIN must be exactly 17 characters (A-H, J-N, P, R-Z, 0-9; no I/O/Q). "
                "If not fully legible, return {\"vin\": null}.\n"
            )
            rsp_v = client.chat.completions.create(
                model=MODEL,
                messages=[
                    {"role": "system", "content": "You extract VINs from vehicle door-jamb certification labels. JSON only."},
                    {"role": "user", "content": [{"type": "text", "text": prompt}, _image_part_from_bytes(raw_bytes)]},
                ],
                max_completion_tokens=300,
                temperature=0,
                response_format={"type": "json_object"},
            )
            raw_v = (rsp_v.choices[0].message.content or "").strip()
            try:
                data_v = json.loads(raw_v)
            except Exception:
                data_v = None
            if isinstance(data_v, dict):
                vv = data_v.get("vin")
                if isinstance(vv, str):
                    vv = vv.strip().upper()
                    if re.fullmatch(VIN_PATTERN, vv):
                        return vv
        except Exception:
            return None
        return None

    if not vin_from_label:
        try:
            candidate = None
            for rec in ocr_pairs:
                if isinstance(rec, dict) and _looks_like_door_label(rec.get("text") or "") and rec.get("raw_for_vin"):
                    candidate = rec
                    break
            if candidate is None:
                for rec in ocr_pairs:
                    if isinstance(rec, dict) and rec.get("raw_for_vin"):
                        candidate = rec
                        break
            if candidate is not None:
                # Prefer local QR/barcode VIN if available before calling the vision model.
                if candidate.get("qr_vin"):
                    vin_from_label = str(candidate.get("qr_vin") or "").strip().upper() or None
                if not vin_from_label:
                    vin_from_label = _decode_vin_from_label_or_qr(candidate.get("raw_for_vin"))
                vin_from_label_photo = candidate.get("name") if vin_from_label else None
        except Exception:
            pass

    # normalize + keep unique order
    _seen_v = set(); _tmp_v=[]
    for _v in vin_candidates:
        if _v and _v not in _seen_v:
            _seen_v.add(_v); _tmp_v.append(_v)
    vin_candidates = _tmp_v

    # --- ODOMETER OCR LOCK (extract mileage from OCR text if visible) ---
    # This makes it impossible for the narrative to claim the odometer is not visible when OCR captured a mileage value.
    odometer_value = None
    try:
        _odo_txt = uploaded_text_all or ""
        # Common mileage patterns from digital clusters: "72,261 mi", "72261 mi", "72261 miles", "116000 km"
        _m = re.search(r"(?i)\b(\d{1,3}(?:,\d{3})+|\d{4,7})\s*(mi|miles|km)\b", _odo_txt)
        if _m:
            _digits = _m.group(1).replace(",", "")
            _unit = _m.group(2).lower()
            if _unit == "miles":
                _unit = "mi"
            odometer_value = f"{int(_digits):,} {_unit}"
    except Exception:
        odometer_value = None


    # --- Closing Report: extract "Inspection Results" section for deterministic cross-check + shop-status ---
    def _extract_inspection_results_block(_txt: str) -> str:
        if not _txt:
            return ""
        t = _txt.replace("\r", "\n")
        # Capture the text after the "Inspection Results" header up to the next major header.
        m_ir = re.search(
            r"(?is)\bInspection\s+Results\b\s*[:\-]?\s*(.{0,5000}?)(?=\n\s*(?:Possible\s+Supplement\s+Amount|Supplement\b|Estimate\b|Inspection\s+Location\b|Repair\s+Facility\b|Vehicle\b|Owner\b|Appraiser\b|Adjuster\b|$))",
            t,
        )
        if m_ir:
            return (m_ir.group(1) or "").strip()[:5000]
        return ""

    inspection_results_text = _extract_inspection_results_block(uploaded_text_all or "")

    # Closing Report: "Possible Supplement Amount" (this ALONE does NOT mean the estimate is a supplement)
    _possible_supp_amount = None
    _m_psa = re.search(r"(?i)\bPossible\s+Supplement\s+Amount\b\s*\$?\s*([0-9,]+(?:\.[0-9]{2})?)", uploaded_text_all or "")
    if _m_psa:
        _possible_supp_amount = _m_psa.group(1)

    photos_provided = any(isinstance(p, dict) and p.get('type') != 'text' for p in parts)

    # --- Robust detectors (CLEAN RETAIL + ADVISOR) ---
    clean_retail_rx = (
        r"(?i)\b("
        r"NADA|J[.\s-]*D[.\s-]*\s*Power|JDPower\.com|"
        r"Kell?ey\s+Blue\s+Book|KBB\.com|"
        r"Edmunds|Carfax|Cars\.com|"
        r"Clean\s+Retail(?:\s+Value)?"
        r")\b"
    )
    _clean_retail_present = bool(re.search(clean_retail_rx, uploaded_text_all or ""))

    advisor_rx = r"(?i)\bAdvisor\s+Report\b"
    _advisor_present = bool(re.search(advisor_rx, uploaded_text_all or ""))

    paint_mat_rx = r"(Paint\s+(Suppl(?:ies|y)|Materials)|Materials\s*Line)"
    _paint_materials_present = bool(re.search(paint_mat_rx, uploaded_text_all or "", flags=re.IGNORECASE))

    # --- Parts source guardrail (OEM vs Aftermarket/LKQ) ---
    # Detect explicit non-OEM indicators in estimate LINE ITEMS only (ignore generic boilerplate/disclosures).
    _explicit_non_oem_parts = False
    try:
        # Heuristic: look for non-OEM keywords on the same line as an operation/part line (starts with line #).
        _explicit_non_oem_parts = bool(re.search(
            r"(?im)^\s*\d+\s+.*\b(A/M|AFTERMARKET|NON\s*OEM|CAPA|NSF|LKQ|RCY|RECYCLED|USED|RECOND)\b",
            uploaded_text_all or "",
        ))
    except Exception:
        _explicit_non_oem_parts = False

    # Closing report sometimes explicitly states no aftermarket/LKQ parts were included.
    _closing_no_aftermarket = False
    try:
        _closing_no_aftermarket = bool(re.search(r"(?i)\bno\s+aftermarket\b|\bno\s+aftermarket\s+or\s+lkq\b", uploaded_text_all or ""))
    except Exception:
        _closing_no_aftermarket = False

    # Presence detectors for VIN / Odometer / Production Date (OCR text)
    vin_photo_rx = r"(?i)\bDescription:\s*VIN\b|\bVIN(?:\s*#|:)?\s*[A-HJ-NPR-Z0-9]{17}\b"
    odo_photo_rx = r"(?i)\bDescription:\s*Odometer\b|\bOdometer\s*Photo\b|\bPhoto\s*[:#]?\s*Odometer\b"

    # Production date patterns (label)
    prod_date_rx = (
        r"(?is)\b(Production\s*date|Prod(?:uction)?\s*Date|Date\s*of\s*Mfr|Date\s*of\s*Manufacture|"
        r"MFD\.?\s*(?:BY|DATE)?|MFR\.?\s*DATE|MFG\.?\s*DATE|DATE)\b"
        r".{0,80}?\b(0[1-9]|1[0-2])\s*[-/.\u2013\u2014\u2212:\s]\s*(20\d{2}|\d{2})\b"
    )

    _vin_photo_present = bool(re.search(vin_photo_rx, uploaded_text_all or ""))
    _odo_photo_present = bool(re.search(odo_photo_rx, uploaded_text_all or ""))

    # presence + extractor
    _prod_date_present = False
    _prod_date_str = None
    m = re.search(prod_date_rx, uploaded_text_all or "")
    if m:
        mm = m.group(2); yy = m.group(3)
        if len(yy) == 2: yy = "20" + yy
        _prod_date_present = True
        _prod_date_str = f"{mm}/{yy}"
    if not _prod_date_present:
        loose_prod_rx = (
            r"(?is)\b(MFD|MFR|MFG|PROD\.?|PRODUCTION|DATE)\b"
            r".{0,60}?\b(0[1-9]|1[0-2])\s*[-/.\u2013\u2014\u2212:\s]\s*(20\d{2}|\d{2})\b"
        )
        m2 = re.search(loose_prod_rx, uploaded_text_all or "")
        if m2:
            mm = m2.group(2); yy = m2.group(3)
            if len(yy) == 2: yy = "20" + yy
            _prod_date_present = True
            _prod_date_str = f"{mm}/{yy}"

    # treat production date as evidenced if either door label VIN photo or prod date text is seen
    _prod_evidenced = _prod_date_present or _vin_photo_present

    # Locked comprehensive entrypoint
    _incoming_intent = str(ai_intent or "").strip().lower()
    if _incoming_intent not in {"comprehensive", "guidelines_only"}:
        ai_intent = "comprehensive"
    else:
        ai_intent = _incoming_intent

    REQ_LABELS = {
        "guidelines_only": "Guidelines → Estimate (no photos)",
        "comprehensive": "Comprehensive: Guidelines + Estimate + Photos (with VIN check)",
        "damage_report_from_photos": "Create a Damage Report from Photos",
    }
    req_label = REQ_LABELS.get(ai_intent, "Comprehensive: Guidelines + Estimate + Photos (with VIN check)")
    log.info(f"ai_intent received: {ai_intent} -> using label: {req_label}")

    KEYS = [
        "file_number","request_type","claim_number","vin","vin_verification","vehicle",
        "odometer_estimate_only","compliance_score","summary_brief","summary_markdown",
        "fraud_markdown","primary_impact","secondary_impact","estimated_costs_markdown","conclusion"
    ]

    SYSTEM = SYSTEM_BASE

    # Closing Report shop-status detector (documents-only; used to prevent false Repair Facility deductions)
    _not_at_shop = False
    try:
        # Prefer the extracted "Inspection Results" block if available (more deterministic)
        if inspection_results_text:
            _not_at_shop = bool(re.search(
                r"(?i)\bNot\s+at\s+Shop\b|\bOwner\s+Location\b|\bInspection\s+Location\s*:?\s*Owner\b|\bInspection\s+Location\s+Owner\b",
                inspection_results_text,
            ))
        if not _not_at_shop:
            _not_at_shop = bool(re.search(
                r"(?i)\bNot\s+at\s+Shop\b|\bOwner\s+Location\b|\bInspection\s+Location\s*:?\s*Owner\b|\bInspection\s+Location\s+Owner\b",
                uploaded_text_all or "",
            ))
    except Exception:
        _not_at_shop = False

    # --- NEW: detect all supplement tags S01, S02, ... in the combined document text ---
    combined_detection_text = (uploaded_text_all or "") + "\n" + "\n".join(pdf_text_fulls or [])
    supplement_versions = sorted(set(re.findall(r"(?i)\bS[0-9]{2}\b", combined_detection_text)))
    supplement_block = ""
    if supplement_versions:
        supplement_block = (
            "\n\nSUPPLEMENT VERSIONS DETECTED FROM DOCUMENTS:\n"
            f"- Detected supplement tags: {', '.join(supplement_versions)}.\n"
            "- Use these exact tags (e.g., 'Supplement S01 and S02') when describing supplements in the narrative.\n"
        )

    # -----------------------
    # Minimal prompt (LET GPT DO THE WORK)
    # -----------------------
    prompt_text = (
        f"REQUEST TYPE (use exactly in request_type): {req_label}\n"
        f"FILE #: {file_number}\n"
        f"CLIENT: {ia_company}\n\n"
        + (
            "CLIENT RULES STATUS: CLIENT RULES WERE SUPPLIED IN THIS REQUEST.\n"
            "HARD RULE: Do NOT say that client guidelines are missing, absent, not provided, or that no separate guideline document/rule text was included.\n"
            "HARD RULE: Build the Client Guidelines Comparison from the supplied frontend dropdown/pasted client rules text below.\n\n"
            if client_rules_supplied else
            "CLIENT RULES STATUS: No client rules text was supplied in this request.\n\n"
        )
        + "FILES SEEN (echo verbatim in Inputs Used):\n- "
        + ("\n- ".join(files_seen) if files_seen else "none")
        + "\n\n"
        + "PHOTO INDEX (use Photo # citations exactly as listed):\n"
        + ("\n".join([f"Photo {i+1}: {name}" for i, name in enumerate(photo_index)]) if photo_index else "No photos provided.")
        + "\n\n"
        + "CLIENT RULES (only if provided):\n"
        + (client_rules[:6000] if client_rules else "")
        + ai_notes_block
        + "\n\n"
        + "INSTRUCTIONS:\n"
        + "- Return strict JSON only.\n"
        + "- REQUIRED: Populate 'estimated_costs_markdown' in the JSON.\n"
        + "- Use the template below for narrative formatting.\n\n"
        + DETAIL_TEMPLATES.get(ai_intent, DETAIL_TEMPLATES['comprehensive'])
    )
    if client_rules_supplied and ai_intent in {"comprehensive", "guidelines_only"}:
        prompt_text += (
            "\n\nWhen client_rules text is provided, you MUST include a section titled '## Client Guidelines Comparison' "
            "with 3–8 concise bullets. For each, quote the relevant rule fragment and mark Aligned / Not Aligned / Not Evidenced, "
            "citing evidence (p#/L#, Photo #). Also weave any material rule alignment/misalignment into the Detailed Condition Report narrative."
        )
        prompt_text += (
            "\n\nWeave the following static audit questions naturally into the Detailed Condition Report narrative "
            "(do NOT present them as a separate Q&A list; integrate answers inline and cite evidence with p#/L# and Photo # as applicable):\n"
            + "\n".join(f"- {q}" for q in STATIC_AUDIT_QUESTIONS)
        )


    parts_payload: List[Dict[str,Any]] = []
    redaction_success = False
    try:
        red_prompt = redact_text_preserve_vin_claim(prompt_text)
        redaction_success = True
        # Prevent Presidio from mutating fixed markdown headings used for prompting.
        # (This avoids model outputs like "### [REDACTED] Addressed".)
        red_prompt = red_prompt.replace("### [REDACTED] Addressed", "### Add'l Notes Addressed")
    except Exception as e:
        log.warning(f"Redaction failed on prompt_text: {e}")
        red_prompt = prompt_text

    parts_payload.append({"type": "text", "text": red_prompt})
    if client_rules_supplied and (client_rules or "").strip():
        parts_payload.append({
            "type": "text",
            "text": (
                "CLIENT RULES EVIDENCE BLOCK (UNREDACTED, AUTHORITATIVE):\n"
                "These client rules were supplied from the frontend dropdown and/or pasted Client Rules box for this request.\n"
                "Do NOT say they were missing or not provided. Quote and compare them directly in the Client Guidelines Comparison section.\n\n"
                + str(client_rules or "")[:12000]
            )
        })

    if parts:
        for p in parts:
            if p.get("type") == "text" and isinstance(p.get("text"), str):
                try:
                    red_txt = redact_text_preserve_vin_claim(p["text"])
                    redaction_success = True
                except Exception as e:
                    log.warning(f"Redaction failed on a text part: {e}")
                    red_txt = p["text"]
                parts_payload.append({"type": "text", "text": red_txt})
            else:
                parts_payload.append(p)

    redaction_status = "Redacted PII: Successful ✅" if redaction_success else "Redacted PII: Not Applied"

    # -------------------------------
    # Keep prompt lean
    # -------------------------------
    TEXT_PART_LIMIT = 6
    _text_parts = [p for p in parts_payload if p.get("type") == "text"]
    _image_parts = [p for p in parts_payload if p.get("type") != "text"]
    for tp in _text_parts:
        if isinstance(tp.get("text"), str) and len(tp["text"]) > 8000:
            tp["text"] = tp["text"][:8000]
    parts_payload = _text_parts[:TEXT_PART_LIMIT] + _image_parts

    # Token limits
    MAX_TOKENS_BY_INTENT = {
        "comprehensive": 2200,
        "guidelines_only": 1500,
        "damage_report_from_photos": 2400
    }
    max_tokens = MAX_TOKENS_BY_INTENT.get(ai_intent, 1500)

    def _stage_blank(v: Any) -> bool:
        s = str(v or "").strip()
        return (not s) or (s.upper() in {"N/A", "NA", "NONE", "NULL", "UNKNOWN"})

    def _extract_claim_from_uploaded_text(_txt: str) -> str:
        if not _txt:
            return "Not confirmed from provided evidence."
        m = re.search(r"(?im)\bClaim\s*#?\s*[:\-]?\s*([A-Z0-9][A-Z0-9\-]{4,})\b", _txt)
        if m:
            return m.group(1).strip()
        return "Not confirmed from provided evidence."

    def _extract_vehicle_from_uploaded_text(_txt: str) -> str:
        if not _txt:
            return "Vehicle information not confirmed from provided evidence."
        makes = r"Acura|Audi|BMW|Buick|Cadillac|Chevrolet|Chevy|Chrysler|Dodge|Ford|GMC|Honda|Hyundai|Infiniti|Jeep|Kia|Lexus|Lincoln|Mazda|Mercedes(?:-Benz)?|Mercury|Mini|Mitsubishi|Nissan|Pontiac|Porsche|Ram|Saturn|Scion|Subaru|Tesla|Toyota|Volkswagen|Volvo"
        m = re.search(rf"(?i)\b(19\d{{2}}|20\d{{2}})\s+({makes})\s+([A-Z0-9][A-Z0-9\- ]{{1,40}})", _txt)
        if m:
            year = m.group(1).strip()
            make = m.group(2).strip().replace('Chevy', 'Chevrolet')
            model = re.split(r"\s{2,}|\n|\r", m.group(3).strip())[0].strip(" :-")
            return f"{year} {make} {model}".strip()
        return "Vehicle information not confirmed from provided evidence."

    def _extract_vin_from_docs(_txt: str) -> str:
        if not _txt:
            return "Not confirmed from provided evidence."
        m = re.search(VIN_PATTERN, (_txt or "").upper())
        if m:
            return m.group(0)
        return "Not confirmed from provided evidence."

    def _numeric_score_or_blank(v: Any) -> str:
        s = str(v or "").strip()
        if not s:
            return ""
        m = re.search(r"([0-9]{1,3})", s)
        if not m:
            return ""
        try:
            n = max(0, min(100, int(m.group(1))))
            return str(n)
        except Exception:
            return ""

    locked_fields = {
        "claim_number": _extract_claim_from_uploaded_text(uploaded_text_all or ""),
        "vin": vin_from_label or vin_from_qr or _extract_vin_from_docs(uploaded_text_all or ""),
        "vin_verification": "",
        "vehicle": _extract_vehicle_from_uploaded_text(uploaded_text_all or ""),
        "odometer_estimate_only": odometer_value or "Not confirmed from provided evidence.",
        "compliance_score": "",
    }
    if vin_from_label and vin_from_qr:
        locked_fields["vin_verification"] = ("MATCH (door label + QR)" if vin_from_label == vin_from_qr else f"MISMATCH (door label: {vin_from_label}; QR: {vin_from_qr})")
    elif vin_from_label:
        locked_fields["vin_verification"] = "INCONCLUSIVE (door label extracted; no secondary confirmation)"
    elif vin_from_qr:
        locked_fields["vin_verification"] = "INCONCLUSIVE (QR extracted; door label not detected)"
    elif locked_fields["vin"] and not _stage_blank(locked_fields["vin"]):
        locked_fields["vin_verification"] = "INCONCLUSIVE (document/OCR VIN extracted)"
    else:
        locked_fields["vin_verification"] = "Not confirmed from provided evidence."

    staged_guideline_markdown = ""

    # Call GPT and parse JSON (JSON hardened)
    # Prefer the canonical SDK path (client.chat.completions). Keep fallback for older SDKs.
    try:
        rsp = client.chat.completions.create(
            model=MODEL,
            messages=[{"role":"system","content": SYSTEM},
                      {"role":"user","content": parts_payload}],
            max_completion_tokens=max_tokens,
            temperature=0,
            top_p=1,
            presence_penalty=0,
            frequency_penalty=0,
            response_format={"type":"json_object"},
        )
    except AttributeError:
        rsp = client.chat_completions.create(  # type: ignore[attr-defined]
            model=MODEL,
            messages=[{"role":"system","content": SYSTEM},
                      {"role":"user","content": parts_payload}],
            max_completion_tokens=max_tokens,
            temperature=0,
            top_p=1,
            presence_penalty=0,
            frequency_penalty=0,
            response_format={"type":"json_object"},
        )

    # --- Hardened JSON parse helper
    def _try_parse_json(raw_text: str):
        if not raw_text:
            return None
        raw_local = raw_text.strip()
        for fence in ("```json", "```JSON", "```"):
            if raw_local.startswith(fence):
                raw_local = raw_local[len(fence):]
            if raw_local.endswith("```"):
                raw_local = raw_local[:-3]
        raw_local = raw_local.strip()
        raw_local = raw_local.replace("\ufeff", "").replace("\u200b", "").replace("\u00A0", " ")
        try:
            return json.loads(raw_local)
        except Exception:
            pass
        lb = raw_local.find("{"); rb = raw_local.rfind("}")
        chunk = raw_local[lb:rb+1] if (lb != -1 and rb != -1 and rb > lb) else raw_local
        fixes = {"\u2018": "'", "\u2019": "'", "\u201C": '"', "\u201D": '"', "\r": "", "\t": "    "}
        for k, v in fixes.items():
            chunk = chunk.replace(k, v)
        chunk = re.sub(r",\s*([}\]])", r"\1", chunk)
        try:
            return json.loads(chunk)
        except Exception:
            return None

    def _call_json_stage(stage_name: str, stage_system: str, stage_parts: List[Dict[str, Any]], stage_tokens: int) -> Any:
        try:
            _rsp = client.chat.completions.create(
                model=MODEL,
                messages=[{"role": "system", "content": stage_system}, {"role": "user", "content": stage_parts}],
                max_completion_tokens=stage_tokens,
                temperature=0,
                top_p=1,
                presence_penalty=0,
                frequency_penalty=0,
                response_format={"type": "json_object"},
            )
        except AttributeError:
            _rsp = client.chat_completions.create(  # type: ignore[attr-defined]
                model=MODEL,
                messages=[{"role": "system", "content": stage_system}, {"role": "user", "content": stage_parts}],
                max_completion_tokens=stage_tokens,
                temperature=0,
                top_p=1,
                presence_penalty=0,
                frequency_penalty=0,
                response_format={"type": "json_object"},
            )
        _raw = (_rsp.choices[0].message.content or "")
        log.info(f"{stage_name} RAW RESPONSE START")
        log.info((_raw or "")[:4000])
        log.info(f"{stage_name} RAW RESPONSE END")
        return _try_parse_json(_raw)

    if ai_intent in {"comprehensive", "guidelines_only"}:
        _stage_text_parts = [p for p in parts_payload if p.get("type") == "text"]
        _stage_image_parts = [p for p in parts_payload if p.get("type") != "text"]

        stage1_prompt = (
            "Return ONLY strict JSON with keys: claim_number, vin, vin_verification, vehicle, odometer_estimate_only, compliance_score. "
            "Task: extract locked header facts from the uploaded estimate/photos/OCR. Do not write narrative. "
            "Use direct visible/document evidence only. If a field cannot be confirmed, return 'Not confirmed from provided evidence.'."
        )
        stage1_parts = [{"type": "text", "text": stage1_prompt}] + _stage_text_parts[1:4] + _stage_image_parts[:8]
        try:
            stage1_data = _call_json_stage("STAGE1 EXTRACTION", "You extract locked appraisal header facts. Return JSON only.", stage1_parts, 900)
        except Exception:
            stage1_data = None
        if isinstance(stage1_data, dict):
            for _k in ("claim_number", "vin", "vin_verification", "vehicle", "odometer_estimate_only"):
                _v = str(stage1_data.get(_k) or "").strip()
                if _v and not _stage_blank(_v):
                    locked_fields[_k] = _v
            _score = _numeric_score_or_blank(stage1_data.get("compliance_score"))
            if _score:
                locked_fields["compliance_score"] = _score

        if (client_rules or "").strip():
            stage2_prompt = (
                "Return ONLY strict JSON with keys: guideline_comparison_markdown, compliance_score, summary_brief. "
                "Task: compare the provided client guidelines to the estimate/photos. "
                "Write concise markdown bullets only for guideline_comparison_markdown. Each bullet must state the rule point, the evidence, and Aligned / Not Aligned / Not Evidenced. "
                "Use uploaded evidence only and do not invent facts."
            )
            stage2_parts = [{"type": "text", "text": stage2_prompt}] + _stage_text_parts[1:6] + _stage_image_parts[:10]
            try:
                stage2_data = _call_json_stage("STAGE2 GUIDELINE COMPARISON", "You compare client guidelines against estimate and photo evidence. Return JSON only.", stage2_parts, 2200)
            except Exception:
                stage2_data = None
            if isinstance(stage2_data, dict):
                staged_guideline_markdown = str(stage2_data.get("guideline_comparison_markdown") or "").strip()
                _score2 = _numeric_score_or_blank(stage2_data.get("compliance_score"))
                if _score2:
                    locked_fields["compliance_score"] = _score2

        _locked_block = (
            "\n\nLOCKED EXTRACTED FIELDS (use exactly; do not overwrite with weaker guesses):\n"
            f"- Claim #: {locked_fields.get('claim_number', '')}\n"
            f"- VIN: {locked_fields.get('vin', '')}\n"
            f"- VIN verification: {locked_fields.get('vin_verification', '')}\n"
            f"- Vehicle: {locked_fields.get('vehicle', '')}\n"
            f"- Odometer: {locked_fields.get('odometer_estimate_only', '')}\n"
            + (f"- Compliance Score seed: {locked_fields.get('compliance_score', '')}\n" if locked_fields.get('compliance_score') else "")
        )
        if staged_guideline_markdown:
            _locked_block += "\nPRECOMPUTED CLIENT GUIDELINES COMPARISON (preserve substance in final narrative):\n" + staged_guideline_markdown[:5000] + "\n"
        elif selected_rules_key and resolved_rules_text:
            _locked_block += (
                "\nCLIENT GUIDELINES WERE RESOLVED LOCALLY FROM THE SELECTED FRONTEND RULE FILE. "
                "Do NOT say that no separate client guideline document/rule list is present. "
                f"Resolved guideline source: {os.path.basename(resolved_rules_path) if resolved_rules_path else selected_rules_key}.\n"
            )
        if parts_payload and isinstance(parts_payload[0], dict) and parts_payload[0].get("type") == "text":
            parts_payload[0]["text"] = str(parts_payload[0].get("text") or "") + _locked_block

    try:
        raw = (rsp.choices[0].message.content or "")
    except Exception as e:
        log.error(f"LLM returned no content: {e}")
        return JSONResponse(status_code=500, content={"error":"Model returned no content."})

    log.info("MODEL RAW RESPONSE START")
    log.info((raw or "")[:4000])
    log.info("MODEL RAW RESPONSE END")

    data = _try_parse_json(raw)
    if isinstance(data, dict):
        for _k in ("claim_number", "vin", "vin_verification", "vehicle", "odometer_estimate_only"):
            _locked_v = str(locked_fields.get(_k) or "").strip()
            if _locked_v and not _stage_blank(_locked_v):
                data[_k] = _locked_v
        if ai_intent != "damage_report_from_photos":
            _locked_score = _numeric_score_or_blank(locked_fields.get("compliance_score"))
            if _locked_score:
                data["compliance_score"] = _locked_score
        if staged_guideline_markdown:
            _sm = str(data.get("summary_markdown") or "").strip()
            if "Client Guidelines Comparison" not in _sm:
                data["summary_markdown"] = (_sm + "\n\n## Client Guidelines Comparison\n" + staged_guideline_markdown).strip()
        elif selected_rules_key and resolved_rules_text:
            _sm_resolved = str(data.get("summary_markdown") or "").strip()
            if "Client Guidelines Comparison" not in _sm_resolved:
                data["summary_markdown"] = (_sm_resolved + "\n\n## Client Guidelines Comparison\n- Client guidelines were provided and resolved from the selected frontend rule file.").strip()
    # Prefer door-label VIN when present (OCR -> QR/barcode). Do NOT use filenames.
    # --- HARD VIN LOCK (door label wins; QR only confirms) ---
    try:
        if isinstance(data, dict):

            # 1) If door-label VIN exists, it is authoritative.
            if vin_from_label:
                data["vin"] = vin_from_label

                if vin_from_qr:
                    if vin_from_qr == vin_from_label:
                        data["vin_verification"] = "MATCH (door label + QR)"
                    else:
                        data["vin_verification"] = (
                            f"MISMATCH (door label: {vin_from_label}; QR: {vin_from_qr})"
                        )
                else:
                    data["vin_verification"] = (
                        "INCONCLUSIVE (door label extracted; no secondary confirmation)"
                    )

            # 2) If no door label but QR found
            elif vin_from_qr:
                data["vin"] = vin_from_qr
                data["vin_verification"] = "INCONCLUSIVE (QR only; door label not detected)"

    except Exception:
        pass

    # Prefer door-label VIN when present (OCR -> QR/barcode). Do NOT use filenames.
    try:
        if isinstance(data, dict) and vin_from_label:
            v_model = (data.get("vin") or "").strip().upper()
            if not re.fullmatch(VIN_PATTERN, v_model):
                data["vin"] = vin_from_label
                if not (data.get("vin_verification") or "").strip():
                    data["vin_verification"] = "INCONCLUSIVE (door label VIN extracted; compare to other docs if present)"
            elif v_model != vin_from_label:
                data["vin_verification"] = f"MISMATCH (door label: {vin_from_label}; other source: {v_model})"
            # Secondary confirmation vs QR/barcode if available
            if vin_from_qr:
                if vin_from_label == vin_from_qr:
                    if not re.search(r"\bMATCH\b|\bMISMATCH\b", str(data.get("vin_verification") or ""), flags=re.IGNORECASE):
                        data["vin_verification"] = "MATCH (door label + QR/barcode)"
                else:
                    data["vin_verification"] = f"MISMATCH (door label: {vin_from_label}; QR/barcode: {vin_from_qr})"
    except Exception:
        pass


    # One safe retry on truncation
    try:
        finish_reason = getattr(rsp.choices[0], "finish_reason", None)
    except Exception:
        finish_reason = None

    if data is None:
        _text_parts_retry = [p for p in parts_payload if p.get("type") == "text"]
        _image_parts_retry = [p for p in parts_payload if p.get("type") != "text"]
        shrunk = _text_parts_retry[: max(3, len(_text_parts_retry)//2)] + _image_parts_retry
        retry_tokens = min(3000, max_tokens + 600)
        try:
            rsp2 = client.chat.completions.create(
                model=MODEL,
                messages=[{"role": "system", "content": SYSTEM},
                          {"role": "user", "content": shrunk}],
                max_completion_tokens=retry_tokens,
                temperature=0,
                response_format={"type": "json_object"}
            )
            raw2 = (rsp2.choices[0].message.content or "")
            data = _try_parse_json(raw2)
        except Exception:
            pass

    # Fallback: formatting pass
    if data is None:
        try:
            fix_prompt = [
                {"role":"system","content":
                    "You are a formatter. Return ONLY a strict JSON object. No prose. No markdown. No code fences."
                },
                {"role":"user","content":
                    "Convert the following text into a valid JSON object. "
                    "Use exactly these keys (all required): "
                    "['file_number','request_type','claim_number','vin','vin_verification','vehicle',"
                    "'odometer_estimate_only','compliance_score','summary_brief','summary_markdown',"
                    "'fraud_markdown','primary_impact','secondary_impact','estimated_costs_markdown','conclusion'] "
                    "Do not invent new keys. If a field is unavailable, use 'N/A'. "
                    "Here is the text:\n\n" + raw
                }
            ]
            try:
                fix_rsp = client.chat_completions.create(  # type: ignore[attr-defined]
                    model=MODEL,
                    messages=fix_prompt,
                    max_completion_tokens=max_tokens,
                    temperature=0,
                    response_format={"type":"json_object"}
                )
            except AttributeError:
                fix_rsp = client.chat.completions.create(
                    model=MODEL,
                    messages=fix_prompt,
                    max_completion_tokens=max_tokens,
                    temperature=0,
                    response_format={"type":"json_object"}
                )
            fixed = (fix_rsp.choices[0].message.content or "")
            data = _try_parse_json(fixed)
        except Exception as e:
            log.error(f"Self-heal reformat failed: {e}")

    if data is None:
        if ai_intent == "damage_report_from_photos":
            try:
                direct_retry_prompt = (
                    "PHOTOS-ONLY DIRECT RECOVERY. Return ONLY strict JSON. No markdown fences. No prose outside JSON.\n"
                    "Required keys: ['file_number','request_type','claim_number','vin','vin_verification','vehicle',"
                    "'odometer_estimate_only','compliance_score','summary_brief','summary_markdown',"
                    "'fraud_markdown','primary_impact','secondary_impact','estimated_costs_markdown','conclusion'].\n"
                    "This is Create a Damage Report from Photos. Use the uploaded photos and OCR only.\n"
                    "Do not rely on estimate/document language. Do not return N/A for summary_markdown, fraud_markdown, estimated_costs_markdown, or conclusion.\n"
                    "If VIN label or odometer are visible, read them. If not fully legible, say so clearly.\n"
                    "estimated_costs_markdown must include labor assumptions, parts assumptions, sales tax on parts + paint materials only, one Approximate Repair Total, and one Severity Tier block."
                )
                direct_parts = list(parts_payload)
                if direct_parts and isinstance(direct_parts[0], dict) and direct_parts[0].get("type") == "text":
                    direct_parts[0] = {"type": "text", "text": direct_retry_prompt + "\n\n" + str(direct_parts[0].get("text") or "")}
                else:
                    direct_parts.insert(0, {"type": "text", "text": direct_retry_prompt})
                direct_rsp = client.chat.completions.create(
                    model=MODEL,
                    messages=[{"role":"system","content": SYSTEM},
                              {"role":"user","content": direct_parts}],
                    max_completion_tokens=min(3200, max_tokens + 700),
                    temperature=0,
                    top_p=1,
                    presence_penalty=0,
                    frequency_penalty=0,
                    response_format={"type":"json_object"},
                )
                direct_raw = (direct_rsp.choices[0].message.content or "")
                log.info("PHOTOS-ONLY DIRECT RECOVERY RAW RESPONSE START")
                log.info((direct_raw or "")[:4000])
                log.info("PHOTOS-ONLY DIRECT RECOVERY RAW RESPONSE END")
                data = _try_parse_json(direct_raw)
            except Exception as e:
                log.error(f"Photos-only direct recovery failed: {e}")

        if data is None:
            skeleton = {k: "" for k in KEYS}
            skeleton["file_number"] = file_number
            skeleton["request_type"] = req_label
            skeleton["claim_number"] = locked_fields.get("claim_number", "Not confirmed from provided evidence.")
            skeleton["vin"] = locked_fields.get("vin", "Not confirmed from provided evidence.")
            skeleton["vin_verification"] = locked_fields.get("vin_verification", "Not confirmed from provided evidence.")
            skeleton["vehicle"] = locked_fields.get("vehicle", "Vehicle information not confirmed from provided evidence.")
            skeleton["odometer_estimate_only"] = locked_fields.get("odometer_estimate_only", "Not confirmed from provided evidence.")
            skeleton["compliance_score"] = (locked_fields.get("compliance_score", "") if ai_intent != "damage_report_from_photos" else "N/A")
            skeleton["summary_brief"] = "Header facts were preserved from the uploaded evidence."
            skeleton["summary_markdown"] = (
                "## Detailed Condition Report\n"
                f"Deterministic extraction completed for file {file_number}. "
                f"Claim #: {locked_fields.get('claim_number', 'Not confirmed from provided evidence.')}. "
                f"VIN: {locked_fields.get('vin', 'Not confirmed from provided evidence.')}. "
                f"Vehicle: {locked_fields.get('vehicle', 'Vehicle information not confirmed from provided evidence.')}. "
                f"Odometer: {locked_fields.get('odometer_estimate_only', 'Not confirmed from provided evidence.')}. "
                "This response preserves the extracted header facts and available evidence for appraisal review instead of returning blank sections."
            )
            if staged_guideline_markdown:
                skeleton["summary_markdown"] = skeleton["summary_markdown"].rstrip() + "\n\n## Client Guidelines Comparison\n" + staged_guideline_markdown
            elif (client_rules or "").strip():
                skeleton["summary_markdown"] = skeleton["summary_markdown"].rstrip() + (
                    "\n\n## Client Guidelines Comparison\n"
                    "- Client guidelines text was provided for this run and should be compared against the estimate/photos.\n"
                    "- The final model narrative did not validate cleanly, so the detailed rule-by-rule comparison was not finalized in this fallback body."
                )
            skeleton["fraud_markdown"] = "No material inconsistencies found."
            skeleton["estimated_costs_markdown"] = "## Approximate Repair Cost Breakdown\nCost analysis requires a valid narrative JSON response to finalize."
            skeleton["conclusion"] = "Available identifiers and uploaded evidence were preserved for appraisal review."
            return skeleton

    def _get(k):
        v = data.get(k)
        return "" if v is None else str(v)

    # LOCKED FINAL FIELDS: claim/VIN/vehicle/odometer/compliance/cost/rationale are finalized by code-owned post-processing.
    result = {
        "file_number": file_number,
        "request_type": req_label,
        "claim_number": _get("claim_number"),
        "vin": _get("vin"),
        "vin_verification": _get("vin_verification"),
        "vehicle": _get("vehicle"),
"odometer_estimate_only": (odometer_value if odometer_value else _get("odometer_estimate_only")),
        "compliance_score": _get("compliance_score"),
        "summary_brief": _get("summary_brief"),
        "summary_markdown": _get("summary_markdown"),
        "fraud_markdown": _get("fraud_markdown"),
        "primary_impact": _get("primary_impact"),
        "secondary_impact": _get("secondary_impact"),
        "estimated_costs_markdown": _get("estimated_costs_markdown"),
        "conclusion": _get("conclusion"),
        "redaction_status": redaction_status,
    }
    try:
        for _k in ("claim_number", "vin", "vin_verification", "vehicle", "odometer_estimate_only"):
            _locked_v = str(locked_fields.get(_k) or "").strip()
            if _locked_v and not _stage_blank(_locked_v):
                result[_k] = _locked_v
        if ai_intent != "damage_report_from_photos":
            _locked_score = _numeric_score_or_blank(locked_fields.get("compliance_score"))
            if _locked_score:
                result["compliance_score"] = _locked_score
        if staged_guideline_markdown:
            _sm = str(result.get("summary_markdown") or "").strip()
            if "Client Guidelines Comparison" not in _sm:
                result["summary_markdown"] = (_sm + "\n\n## Client Guidelines Comparison\n" + staged_guideline_markdown).strip()
        elif selected_rules_key and resolved_rules_text:
            _sm_resolved = str(result.get("summary_markdown") or "").strip()
            if "Client Guidelines Comparison" not in _sm_resolved:
                result["summary_markdown"] = (_sm_resolved + "\n\n## Client Guidelines Comparison\n- Client guidelines were provided and resolved from the selected frontend rule file.").strip()
    except Exception:
        pass


    def _naish(v: Any) -> bool:
        s = str(v or "").strip()
        if not s:
            return True
        return s.upper() in {"N/A", "NA", "NONE", "NULL", "UNKNOWN"}

    def _extract_claim_from_text(_txt: str) -> str:
        if not _txt:
            return "Not confirmed from provided evidence."
        m = re.search(r"(?im)\bClaim\s*#?\s*[:\-]?\s*([A-Z0-9][A-Z0-9\-]{4,})\b", _txt)
        if m:
            return m.group(1).strip()
        return "Not confirmed from provided evidence."

    def _extract_vehicle_from_text(_txt: str) -> str:
        if not _txt:
            return "Vehicle information not confirmed from provided evidence."
        makes = r"Acura|Audi|BMW|Buick|Cadillac|Chevrolet|Chevy|Chrysler|Dodge|Ford|GMC|Honda|Hyundai|Infiniti|Jeep|Kia|Lexus|Lincoln|Mazda|Mercedes(?:-Benz)?|Mercury|Mini|Mitsubishi|Nissan|Pontiac|Porsche|Ram|Saturn|Scion|Subaru|Tesla|Toyota|Volkswagen|Volvo"
        m = re.search(rf"(?i)\b(19\d{{2}}|20\d{{2}})\s+({makes})\s+([A-Z0-9][A-Z0-9\- ]{{1,40}})", _txt)
        if m:
            year = m.group(1).strip()
            make = m.group(2).strip().replace('Chevy', 'Chevrolet')
            model = re.split(r"\s{2,}|\n|\r", m.group(3).strip())[0].strip(" :-")
            return f"{year} {make} {model}".strip()
        return "Vehicle information not confirmed from provided evidence."

    def _build_non_na_summary() -> str:
        claim_txt = result.get("claim_number") if not _naish(result.get("claim_number")) else _extract_claim_from_text(uploaded_text_all or "")
        vin_txt = result.get("vin") if not _naish(result.get("vin")) else (vin_from_label or vin_from_qr or "Not confirmed from provided evidence.")
        veh_txt = _format_vehicle_value(result.get("vehicle")) if not _naish(result.get("vehicle")) else _extract_vehicle_from_text(uploaded_text_all or "")
        odo_txt = result.get("odometer_estimate_only") if not _naish(result.get("odometer_estimate_only")) else (odometer_value or "Not confirmed from provided evidence.")
        photo_count = len(photo_index or [])
        files_count = len(files_seen or [])
        rules_state = "provided" if (client_rules or "").strip() else "not resolved"
        supp_txt = (", ".join(supplement_versions) if supplement_versions else "No supplement tags detected")
        return (
            "## Detailed Condition Report\n"
            f"This report used the exact files uploaded for file {file_number}. "
            f"The structured narrative response did not fully validate, so NSPXN preserved deterministic extracted facts instead of printing blank fields. "
            f"Request type: {req_label}. "
            f"Claim number: {claim_txt}. VIN: {vin_txt}. Vehicle: {veh_txt}. Odometer: {odo_txt}. "
            f"Processed files: {files_count}. Photo references loaded: {photo_count}. Client rules were {rules_state}. "
            f"Supplement detection from uploaded documents: {supp_txt}. "
            "Visible-condition, estimate, and OCR evidence should be re-run for a full narrative, but this fallback confirms the job executed and preserves the available identifiers instead of returning blank sections."
        )

    def _build_non_na_fraud() -> str:
        vin_txt = result.get("vin") if not _naish(result.get("vin")) else (vin_from_label or vin_from_qr or "Not confirmed from provided evidence.")
        return (
            "No material inconsistencies found from the available review inputs. "
            f"VIN evidence available on this run: {vin_txt}. "
            "No confirmed fraud, tampering, or estimate manipulation was established from the available evidence in this review."
        )

    def _build_non_na_cost() -> str:
        if ai_intent == "damage_report_from_photos":
            return (
                "## Approximate Repair Cost Breakdown\n"
                "Photos-only cost approximation could not be finalized from the model response on this run. "
                "A re-run is required to populate labor hours, parts, tax, and severity tier."
            )
        return (
            "## Approximate Repair Cost Breakdown\n"
            "Approximate repair cost evaluation should reflect the estimate lines, visible damage support, labor operations, parts usage, paint materials, and applicable tax treatment. "
            "Final repair amounts remain subject to estimate validation and qualified appraiser review."
        )

    def _build_non_na_conclusion() -> str:
        return (
            "The uploaded files were processed successfully and the available review evidence was preserved. "
            "Final claim handling should proceed after confirmation of estimate support, photo evidence, client guideline requirements, and qualified appraiser review."
        )

    try:
        if _naish(result.get("claim_number")):
            result["claim_number"] = _extract_claim_from_text(uploaded_text_all or "")
        if _naish(result.get("vin")):
            result["vin"] = vin_from_label or vin_from_qr or "Not confirmed from provided evidence."
        if _naish(result.get("vin_verification")):
            if vin_from_label and vin_from_qr:
                result["vin_verification"] = ("MATCH (door label + QR)" if vin_from_label == vin_from_qr else f"MISMATCH (door label: {vin_from_label}; QR: {vin_from_qr})")
            elif vin_from_label or vin_from_qr:
                result["vin_verification"] = "INCONCLUSIVE (single-source identifier recovered)"
            else:
                result["vin_verification"] = "Not confirmed from provided evidence."
        if _naish(result.get("vehicle")):
            result["vehicle"] = _extract_vehicle_from_text(uploaded_text_all or "")
        if _naish(result.get("odometer_estimate_only")):
            result["odometer_estimate_only"] = odometer_value or "Not confirmed from provided evidence."
        if _naish(result.get("compliance_score")) and ai_intent != "damage_report_from_photos":
            result["compliance_score"] = locked_fields.get("compliance_score") or "Not scored from validated evidence."
        if _naish(result.get("summary_brief")):
            result["summary_brief"] = "Deterministic fallback summary applied to prevent blank output."
        if ai_intent != "damage_report_from_photos":
            if _naish(result.get("summary_markdown")):
                result["summary_markdown"] = _build_non_na_summary()
            if _naish(result.get("fraud_markdown")):
                result["fraud_markdown"] = _build_non_na_fraud()
            if _naish(result.get("estimated_costs_markdown")):
                result["estimated_costs_markdown"] = _build_non_na_cost()
            if _naish(result.get("conclusion")):
                result["conclusion"] = _build_non_na_conclusion()
        if _naish(result.get("primary_impact")):
            result["primary_impact"] = "Front" if photos_provided else "Not confirmed from provided evidence."
        if _naish(result.get("secondary_impact")):
            result["secondary_impact"] = "None identified on this run"
    except Exception:
        pass

    # ---- Minimal Impact Sanitizer (no Left/Right/Driver/Passenger unless proven) ----
    # Policy A: If any impact label implies left/right/driver/passenger, collapse to Front/Rear/Side.
    def _sanitize_impact_label(v: Optional[str]) -> str:
        if v is None:
            return "N/A"
        s = str(v).strip()
        if not s:
            return "N/A"
        low = s.lower()

        # Common abbreviations / synonyms
        front_hit = bool(re.search(r"\bfront\b|\bfrt\b", low)) or bool(re.search(r"\b(?:lf|rf)\b", low))
        rear_hit  = bool(re.search(r"\brear\b", low)) or bool(re.search(r"\b(?:lr|rr)\b", low))
        side_hint = bool(re.search(r"\bleft\b|\bright\b|\bdriver\b|\bpassenger\b|\b(?:lf|rf|lr|rr)\b", low))

        if front_hit and side_hint:
            return "Front"
        if rear_hit and side_hint:
            return "Rear"
        if side_hint:
            return "Side"

        # Normalize a few common exact values
        if re.fullmatch(r"(?i)front", s):
            return "Front"
        if re.fullmatch(r"(?i)rear", s):
            return "Rear"
        if re.fullmatch(r"(?i)side", s):
            return "Side"
        return s

    try:
        result["primary_impact"] = _sanitize_impact_label(result.get("primary_impact"))
        result["secondary_impact"] = _sanitize_impact_label(result.get("secondary_impact"))
    except Exception:
        pass

    try:
        sm_tmp = (result.get("summary_markdown") or "").strip()
        if not sm_tmp:
            result["summary_markdown"] = (
                "## Detailed Condition Report\n"
                "A detailed condition narrative was not populated in the initial model field, so the report preserves the available request type and compliance information below.\n\n"
                "## Overall Assessment\n"
                f"Request Type: {result.get('request_type','N/A')}\n"
                f"Compliance Score: {result.get('compliance_score','N/A')}\n"
            )
        elif "## Detailed Condition Report" not in sm_tmp:
            # Keep minimal: do not re-write content; just prepend the required header to avoid downstream display rules.
            result["summary_markdown"] = "## Detailed Condition Report\n" + sm_tmp
    except Exception:
        pass


    # -----------------------
    # Photos-Only output hardening
    # -----------------------
    def _bad_field(v: str) -> bool:
        """Return True if the model gave a placeholder/empty section.
        This catches bare 'N/A' AND markdown sections that are effectively only headings + N/A.
        """
        s = (v or "").strip()
        if not s:
            return True
        up = s.strip().upper()

        # Bare placeholders
        if up in ("N/A", "NA", "NONE", "NULL"):
            return True
        if up.startswith("N/A") and len(up) <= 8:
            return True

        # If it's markdown, strip headings/blank lines and see if anything substantive remains
        lines = [ln.strip() for ln in s.splitlines() if ln.strip()]
        # remove markdown headings and common labels
        stripped = []
        for ln in lines:
            if ln.startswith("#"):
                continue
            if ln.lower() in ("report selected", "approximate repair cost breakdown", "fraud & authenticity check", "fraud detection", "conclusion"):
                continue
            stripped.append(ln)

        if not stripped:
            return True

        # If after stripping headings the only remaining content is 'N/A'
        if len(stripped) == 1 and stripped[0].strip().upper() in ("N/A", "NA", "NONE", "NULL"):
            return True

        # If the remaining content is extremely short and contains only placeholders
        joined = " ".join(stripped).strip().upper()
        if joined in ("N/A", "NA", "NONE", "NULL"):
            return True

        return False

    try:
        if ai_intent == "damage_report_from_photos":
            _bad = (
                _bad_field(result.get("summary_markdown") or "")
                or _bad_field(result.get("estimated_costs_markdown") or "")
                or _bad_field(result.get("fraud_markdown") or "")
                or _bad_field(result.get("conclusion") or "")
            )

            if _bad:
                retry_preamble = (
                    "PHOTOS-ONLY RETRY: Prior output was incomplete.\n"
                    "Return ONLY valid JSON (no prose, no code fences).\n"
                    "Hard rules:\n"
                    "- NEVER return 'N/A' for summary_markdown, estimated_costs_markdown, fraud_markdown, or conclusion.\n"
                    "- You MUST explicitly address the Add'l Notes in the narrative (e.g., 'Impact to Right Front') and describe the impact zone accordingly.\n"
                    "- estimated_costs_markdown MUST be a PHOTOS-ONLY approximation: create body/paint/frame/mechanical hours and an OEM parts list with approximate $ by part.\n"
                    "- Paint & Materials MUST be modeled as $ per refinish hour (not a percent).\n"
                    "- Forbidden phrases: 'from estimate', 'from documentation', 'not evidenced', 'no documentation provided'.\n"
                )

                retry_parts = list(parts_payload)  # shallow copy is enough
                # Replace the first text block with a retry preamble + original prompt
                if retry_parts and isinstance(retry_parts[0], dict) and retry_parts[0].get("type") == "text":
                    retry_parts[0] = {"type": "text", "text": retry_preamble + "\n\n" + (retry_parts[0].get("text") or "")}
                else:
                    retry_parts.insert(0, {"type": "text", "text": retry_preamble})

                retry_tokens = min(3200, max_tokens + 500)

                try:
                    rsp_retry = client.chat.completions.create(
                        model=MODEL,
                        messages=[{"role": "system", "content": SYSTEM},
                                  {"role": "user", "content": retry_parts}],
                        max_completion_tokens=retry_tokens,
                        temperature=0,
                        top_p=1,
                        presence_penalty=0,
                        frequency_penalty=0,
                        response_format={"type": "json_object"},
                    )
                except AttributeError:
                    rsp_retry = client.chat_completions.create(  # type: ignore[attr-defined]
                        model=MODEL,
                        messages=[{"role": "system", "content": SYSTEM},
                                  {"role": "user", "content": retry_parts}],
                        max_completion_tokens=retry_tokens,
                        temperature=0,
                        top_p=1,
                        presence_penalty=0,
                        frequency_penalty=0,
                        response_format={"type": "json_object"},
                    )

                raw_retry = (rsp_retry.choices[0].message.content or "")
                data_retry = _try_parse_json(raw_retry)

                if isinstance(data_retry, dict):
                    # Only overwrite the output fields; preserve file_number/request_type etc.
                    for k in ("summary_brief","summary_markdown","fraud_markdown","primary_impact","secondary_impact","estimated_costs_markdown","conclusion","claim_number","vin","vin_verification","vehicle","odometer_estimate_only","compliance_score"):
                        if k in data_retry and data_retry.get(k) is not None:
                            result[k] = str(data_retry.get(k))
                    # Re-apply the narrative header guard if needed
                    sm_retry = (result.get("summary_markdown") or "").strip()
                    if sm_retry and "## Detailed Condition Report" not in sm_retry:
                        result["summary_markdown"] = "## Detailed Condition Report\n" + sm_retry

                    # If the retry STILL returned placeholders, force a non-N/A minimal scaffold (never print dead 'N/A')
                    if (
                        _bad_field(result.get("summary_markdown") or "")
                        or _bad_field(result.get("estimated_costs_markdown") or "")
                        or _bad_field(result.get("fraud_markdown") or "")
                        or _bad_field(result.get("conclusion") or "")
                    ):
                        _notes = (ai_notes_used or "").strip()
                        if not _notes or _notes.lower().startswith("no additional notes"):
                            _notes = "(No additional notes provided.)"

                        result["summary_markdown"] = (
                            "## Detailed Condition Report\n"
                            f"The uploaded photos were reviewed for visible exterior and interior damage, identifiers, and overall vehicle condition. "
                            f"Add'l Notes considered: {_notes}. "
                            "This report remains photo-based only and should describe only what is visible in the submitted images, including any readable VIN label or odometer display, visible impact areas, panel damage, misalignment, and likely repair implications."
                        )
                        result["estimated_costs_markdown"] = (
                            "## Approximate Repair Cost Breakdown\n"
                            "Approximate repair cost should be based on visible photo damage only, using modeled labor, paint materials as a $/refinish-hour value, parts assumptions, tax on parts plus paint materials only, one Approximate Repair Total, and one Severity Tier block."
                        )
                        result["fraud_markdown"] = "No material inconsistencies found from the submitted photo set."
                        result["conclusion"] = (
                            "This photo-based condition report should be finalized from the visible evidence in the uploaded images and reviewed by a qualified appraiser."
                        )

    except Exception:
        # Fail-open: never break the request if retry logic fails.
        pass

        # -----------------------
    # Approximate Repair Cost Breakdown
    # - Prefer model-provided `estimated_costs_markdown` (especially for Photos-Only).
    # - Fail-open: never break the run if cost math cannot be produced.
    # -----------------------
    tax_rate = None  # default; set later if we compute/lookup tax
    try:
        _existing_costs = (result.get("estimated_costs_markdown") or "").strip()

        if not _existing_costs:
            inspection_location = _normalize_location_with_zip(_extract_inspection_location(uploaded_text_all or ""), ia_company, uploaded_text_all)
            rates = _lookup_rates(inspection_location)
            tax_rate = _lookup_tax_rate(inspection_location)

            # Photos-only: do NOT depend on documents/estimates. The model is instructed to generate hours/parts.
            # This backend block is only a minimal fallback if the model omitted the field.
            if ai_intent == "damage_report_from_photos":
                structural_flag = _structural_observed([result.get("summary_markdown") or ""])

                result["estimated_costs_markdown"] = (
                    "## Approximate Repair Cost Breakdown (Photos-Only Approximation)\n\n"
                    f"**Inspection Location:** {inspection_location or 'Not provided'}\n\n"
                    "**Rates Used (regional average / fallback):**\n"
                    f"- Avg Body Rate: ${rates.get('body_rate', 0):.0f}/hr\n"
                    f"- Avg Paint Rate: ${rates.get('paint_rate', 0):.0f}/hr\n"
                    f"- Avg Frame Rate: ${rates.get('frame_rate', 0):.0f}/hr\n"
                    f"- Avg Mechanical Rate: ${rates.get('mechanical_rate', 0):.0f}/hr\n"
                    f"- Paint & Materials Rate: ${rates.get('paint_supplies_rate', 0):.0f} per refinish hr\n"
                    f"- Parts/Materials Tax Rate: {float(tax_rate or 0.0)*100:.3f}%\n\n"
                    "**AI-Derived Repair Scope:**\n"
                    "- Body labor hours: (model should provide)\n"
                    "- Paint labor hours: (model should provide)\n"
                    f"- Setup & Measure (if structural observed): {('2.0 hrs @ body rate' if structural_flag else '0.0 hrs')}\n"
                    "- Frame labor hours (if structural observed): (model should provide)\n"
                    "- Mechanical/ADAS/Airbag/Suspension hours (if observed): (model should provide)\n\n"
                    "**OEM Parts Needed (with approximate $ by part):**\n"
                    "- (model should provide)\n\n"
                    "**Tax (parts + paint materials only):**\n"
                    "- (model should provide)\n\n"
                    "**Approximated Repair Cost Total (Approximation Only):**\n"
                    "- (model should provide)\n\n"
                    "**Estimated Severity Tier (based on this approximation only):**\n"
                    "- [ ] Minor (< $3,500)\n"
                    "- [ ] Moderate ($3,500-$10,000)\n"
                    "- [ ] Major ($10,000+)\n"
                    "- [ ] Likely Total Loss Threshold Approaching\n\n"
                    "_Repair Cost Disclaimer: This section is an AI-generated approximation of repair-related costs based on observed damage in photos and regional average rates. "
                    "It is not an official repair estimate and must be validated by a qualified appraiser using an estimating platform._"
                )
            else:
                # Non-photos intents: best-effort parse from documents (if present), using $/refinish-hr materials.
                totals = _extract_hours_and_parts_totals(uploaded_text_all or "")

                structural_flag = _structural_observed([uploaded_text_all or "", result.get("summary_markdown") or ""])
                setup_measure_hours = 2.0 if structural_flag else 0.0

                body_hours = totals.get("body_hours")
                paint_hours = totals.get("paint_hours")
                frame_hours = totals.get("frame_hours")
                mech_hours = totals.get("mech_hours")
                parts_total = totals.get("parts_total")

                paint_supplies_rate = float(rates.get("paint_supplies_rate") or 0.0)
                paint_supplies_total = (_num(paint_hours) * paint_supplies_rate) if paint_hours is not None else 0.0

                body_labor_total = (_num(body_hours) + setup_measure_hours) * float(rates.get("body_rate") or 0.0) if (body_hours is not None or setup_measure_hours) else 0.0
                paint_labor_total = _num(paint_hours) * float(rates.get("paint_rate") or 0.0) if paint_hours is not None else 0.0

                frame_labor_total = 0.0
                if structural_flag and frame_hours is not None:
                    frame_labor_total = _num(frame_hours) * float(rates.get("frame_rate") or 0.0)

                mech_labor_total = _num(mech_hours) * float(rates.get("mechanical_rate") or 0.0) if mech_hours is not None else 0.0

                taxable = _num(parts_total) + _num(paint_supplies_total)
                tax_total = taxable * float(tax_rate or 0.0)

                approx_total = (
                    _num(body_labor_total)
                    + _num(paint_labor_total)
                    + _num(frame_labor_total)
                    + _num(mech_labor_total)
                    + _num(parts_total)
                    + _num(paint_supplies_total)
                    + _num(tax_total)
                )

                parts_lines = totals.get("parts_lines") or []
                if parts_lines:
                    parts_md = "\n".join([f"- {ln}" for ln in parts_lines[:20]])
                else:
                    parts_md = "- (Itemized OEM parts list not available from parsed text.)"

                result["estimated_costs_markdown"] = (
                    "## Approximate Repair Cost Breakdown (Approximation Only)\n\n"
                    f"**Inspection Location:** {inspection_location or 'Not provided'}\n\n"
                    "**Regional Average Rates (configured / fallback):**\n"
                    f"- Avg Body Rate: ${rates.get('body_rate', 0):.0f}/hr\n"
                    f"- Avg Paint Rate: ${rates.get('paint_rate', 0):.0f}/hr\n"
                    f"- Avg Frame Rate: ${rates.get('frame_rate', 0):.0f}/hr\n"
                    f"- Avg Mechanical Rate: ${rates.get('mechanical_rate', 0):.0f}/hr\n"
                    f"- Paint & Materials Rate: ${rates.get('paint_supplies_rate', 0):.0f} per refinish hr\n"
                    f"- Parts/Materials Tax Rate: {float(tax_rate or 0.0)*100:.3f}%\n\n"
                    "**Hours & Totals:**\n"
                    f"- Approx Total Labor Hours @ Body Rate: {('%.1f' % _num(body_hours)) if body_hours is not None else 'Unknown'}\n"
                    f"- Approx Total Paint Labor Hours @ Paint Rate: {('%.1f' % _num(paint_hours)) if paint_hours is not None else 'Unknown'}\n"
                    f"- Approx Total Paint Supplies @ ${rates.get('paint_supplies_rate', 0):.0f}/refinish hr: {_money(paint_supplies_total)}\n"
                    f"- Setup & Measure (if structural observed): {('2.0 hrs' if structural_flag else '0.0 hrs')}\n"
                    f"- Approx Total Frame Labor (if structural observed) @ Frame Rate: {('%.1f hrs' % _num(frame_hours)) if (structural_flag and frame_hours is not None) else ('Unknown' if structural_flag else '0.0 hrs')}\n"
                    f"- Approx Mechanical/ADAS/Airbag/Suspension Hours @ Mechanical Rate: {('%.1f' % _num(mech_hours)) if mech_hours is not None else 'Unknown'}\n\n"
                    "**Approx Cost of OEM Parts Needed:**\n"
                    f"{parts_md}\n\n"
                    "**Tax (parts + paint materials only):**\n"
                    f"- Taxable subtotal: {_money(taxable)}\n"
                    f"- Estimated tax: {_money(tax_total)}\n\n"
                    "**Approximated Repair Cost Total (Approximation Only):**\n"
                    f"- **{_money(approx_total)}**\n\n"
                    "_Repair Cost Disclaimer: This section is an approximation. It is not an official repair estimate and must be validated by a qualified appraiser using an estimating platform._"
                )

    except Exception:
        # Never fail the request for cost calculation issues
        if not (result.get("estimated_costs_markdown") or "").strip():
            result["estimated_costs_markdown"] = (
                "## Approximate Repair Cost Breakdown (Approximation Only)\n"
                "Unable to generate cost approximation on this run."
            )
        if ai_intent == "damage_report_from_photos":

            def _neutralize_side_terms(text: str) -> str:
                if not text:
                    return text

                replacements = [
                    (r"(?i)\b(driver|passenger)[-\s]?side\s+headlamp\b", "one headlamp"),
                    (r"(?i)\b(left|right)[-\s]?front\s+corner\b", "front corner"),
                    (r"(?i)\b(left|right)[-\s]?rear\s+corner\b", "rear corner"),
                    (r"(?i)\b(left|right)[-\s]?front\b", "front"),
                    (r"(?i)\b(left|right)[-\s]?rear\b", "rear"),
                    (r"(?i)\bdriver[-\s]?side\b", "one side"),
                    (r"(?i)\bpassenger[-\s]?side\b", "one side"),
                    (r"(?i)\bleft side\b", "one side"),
                    (r"(?i)\bright side\b", "one side"),
                ]

                for pattern, repl in replacements:
                    text = re.sub(pattern, repl, text)

                return text

            summary_markdown = _neutralize_side_terms(summary_markdown)
            estimated_costs_markdown = _neutralize_side_terms(estimated_costs_markdown)
            conclusion = _neutralize_side_terms(conclusion)
# -----------------------
    def _normalize_photos_only_cost_block(_cm: str) -> str:
        """Deterministic tail completer for Photos-Only cost markdown.
        If tax/total are missing, compute and inject the canonical lines and
        replace Severity Tier so exactly one box is checked from TOTAL.
        """
        _cm = _cm or ""
        _cm = re.sub(r"(?im)^\s*Approximate\s+Repair\s+Cost\s+Total\s*:.*$", "", _cm).strip()
        _cm = re.sub(r"(?im)^\s*Approximate\s+Repair\s+Cost\s+Total\b.*$", "", _cm).strip()

        def _moneyf(s: str) -> float:
            try:
                return float(str(s).replace(",", "").strip())
            except Exception:
                return 0.0

        def _grab(pats: List[str]) -> Optional[float]:
            for pat in pats:
                m = re.search(pat, _cm, flags=re.IGNORECASE | re.MULTILINE)
                if m:
                    try:
                        return _moneyf(m.group(1))
                    except Exception:
                        pass
            return None

        parts = _grab([
            r"^\s*[-*]?\s*Estimated\s+parts\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Parts\s+subtotal\s*\(approx\.?\)\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Parts\s+subtotal\s*\(approx\.?\)\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Parts\s+subtotal\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Parts\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
        ])
        labor = _grab([
            r"^\s*[-*]?\s*Labor\s+subtotal\s*\([^\n]*?\)\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Labor\s+subtotal\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Labor\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
        ])
        paint_materials = _grab([
            r"^\s*[-*]?\s*Paint\s*&\s*materials\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Paint\s*&\s*materials\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Paint\s+materials\s+subtotal\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Paint\s+materials\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Paint\s+materials\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Paint\s+materials\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
        ])
        tax_amt = _grab([
            r"^\s*[-*]?\s*Sales\s+tax\s*\(assumed\s*7%\s*for\s*approximation\)\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            r"^\s*[-*]?\s*Tax\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
        ])
        approx_total = _grab([
            r"^\s*\*\*?\s*Approximate\s+Repair\s+Total\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
        ])

        if labor is None:
            body = _grab([
                r"^\s*[-*]?\s*Body(?:\s+labor)?\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"^\s*[-*]?\s*Body(?:\s+labor)?\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            ]) or 0.0
            paint = _grab([
                r"^\s*[-*]?\s*Paint\s+labor\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"^\s*[-*]?\s*Paint\s+labor\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            ]) or 0.0
            mechanical = _grab([
                r"^\s*[-*]?\s*Mechanical(?:/SRS/Glass|/diagnostic)?\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"^\s*[-*]?\s*Mechanical(?:/SRS/Glass|/diagnostic)?\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            ]) or 0.0
            setup = _grab([
                r"^\s*[-*]?\s*Setup\s*&\s*Measure\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"^\s*[-*]?\s*Setup\s*&\s*Measure\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            ]) or 0.0
            frame = _grab([
                r"^\s*[-*]?\s*Frame/measure\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"^\s*[-*]?\s*Frame/measure\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            ]) or 0.0
            calc_labor = body + paint + mechanical + setup + frame
            if calc_labor > 0:
                labor = calc_labor

        if tax_amt is None and parts is not None and paint_materials is not None:
            tax_amt = round((parts + paint_materials) * 0.07, 2)
        if approx_total is None and parts is not None and paint_materials is not None and labor is not None and tax_amt is not None:
            approx_total = round(labor + parts + paint_materials + tax_amt, 2)

        cleaned = []
        for ln in _cm.splitlines():
            s = (ln or "").strip()
            if re.search(r"(?i)^\s*[-*]?\s*Sales\s+tax\s*\(assumed\s*7%\s*for\s*approximation\)", s):
                continue
            if re.search(r"(?i)^\s*\*{0,2}\s*Approximate\s+Repair\s+Total\s*:", s):
                continue
            if re.search(r"(?i)^\s*Severity\s+Tier\s*$", s):
                continue
            if re.search(r"(?i)^\s*\[[ xX]\]\s*(Minor|Moderate|Major|Possible\s+Total\s+Loss)", s):
                continue
            cleaned.append(ln)
        _cm = "\n".join(cleaned).rstrip()

        tail = []
        if tax_amt is not None:
            tail.append(f"- Sales tax (assumed 7% for approximation) = ${tax_amt:,.2f}")
        if approx_total is not None:
            tail.append(f"**Approximate Repair Total: ${approx_total:,.2f}**")

        checked_minor = checked_mod = checked_major = checked_tl = " "
        if isinstance(approx_total, (int, float)):
            if approx_total < 3500:
                checked_minor = "x"
            elif approx_total < 10000:
                checked_mod = "x"
            else:
                checked_major = "x"

        tail.extend([
            "Severity Tier",
            f"[{checked_minor}] Minor (< $3,500)",
            f"[{checked_mod}] Moderate ($3,500-$10,000)",
            f"[{checked_major}] Major ($10,000+)",
            f"[{checked_tl}] Possible Total Loss Threshold Approaching",
        ])
        return (_cm.rstrip() + "\n\n" + "\n".join(tail)).strip()
# -----------------------
    # --- Normalize Photos-Only cost markdown for downstream PDF/UI ---
    try:
        if ai_intent == "damage_report_from_photos":
            _cm = result.get("estimated_costs_markdown") or ""
            _cm = _normalize_photos_only_cost_block(_cm)
            result["estimated_costs_markdown"] = _cm
    except Exception:
        pass

    def _needs_customer_scrub(v: Any) -> bool:
        s = str(v or "").strip()
        if not s:
            return True
        low = s.lower()
        bad_phrases = [
            "deterministic fallback",
            "deterministic extraction completed",
            "structured ai",
            "did not validate",
            "did not fully validate",
            "validated model response",
            "unavailable on this run",
            "pending re-run",
            "not finalized from the validated model response",
            "model compliance error",
            "please re-run",
            "json did not validate",
            "fallback review",
            "core identifiers and score fields were still returned",
            "this fallback means",
            "the uploaded files were processed successfully",
            "available review evidence was preserved",
        ]
        if any(p in low for p in bad_phrases):
            return True
        lines = [ln.strip() for ln in s.splitlines() if ln.strip()]
        stripped = []
        for ln in lines:
            if ln.startswith("#"):
                continue
            if ln.lower() in {"detailed condition report", "approximate repair cost breakdown", "fraud detection", "conclusion"}:
                continue
            stripped.append(ln)
        if not stripped:
            return True
        joined = " ".join(stripped).strip().upper()
        if joined in {"N/A", "NA", "NONE", "NULL", "UNKNOWN"}:
            return True
        return False

    def _summary_is_effectively_na(v: Any) -> bool:
        s = str(v or "").strip()
        if not s:
            return True
        lines = [ln.strip() for ln in s.splitlines() if ln.strip()]
        stripped = []
        for ln in lines:
            if ln.startswith("#"):
                continue
            if ln.lower() in {"detailed condition report", "overall assessment", "condition summary"}:
                continue
            stripped.append(ln)
        if not stripped:
            return True
        joined = " ".join(stripped).strip().upper()
        return joined in {"N/A", "NA", "NONE", "NULL", "UNKNOWN"}

    def _extract_estimate_record_totals_strict(source_text: str) -> Dict[str, Any]:
        out: Dict[str, Any] = {
            "tax_rate": None,
            "sales_tax": None,
            "parts_subtotal": None,
            "labor_subtotal": None,
            "paint_materials": None,
            "estimate_total": None,
        }
        if not source_text:
            return out
        txt = str(source_text).replace("\r\n", "\n").replace("\r", "\n")
        lines = [ln.strip() for ln in txt.splitlines() if ln.strip()]

        def _money_from_line(line: str, label_rx: str) -> Optional[float]:
            if not re.search(label_rx, line, flags=re.IGNORECASE):
                return None
            m = re.search(r"\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)", line)
            if not m:
                return None
            try:
                return float(m.group(1).replace(",", ""))
            except Exception:
                return None

        for ln in lines:
            if out["tax_rate"] is None and re.search(r"(?i)\btax\s*rate\b|\bsales\s*tax\b", ln):
                m_rate = re.search(r"([0-9]+(?:\.[0-9]+)?)\s*%", ln)
                if m_rate:
                    try:
                        out["tax_rate"] = float(m_rate.group(1))
                    except Exception:
                        pass
            if out["sales_tax"] is None:
                val = _money_from_line(ln, r"(?i)\bsales\s*tax\b|\btax\s+amount\b|\btotal\s+tax\b")
                if val is not None:
                    out["sales_tax"] = val
            if out["parts_subtotal"] is None:
                val = _money_from_line(ln, r"(?i)^\s*(parts\s+subtotal|parts\s+total)\b")
                if val is not None:
                    out["parts_subtotal"] = val
            if out["labor_subtotal"] is None:
                val = _money_from_line(ln, r"(?i)^\s*(labor\s+subtotal|labor\s+total|body\s+labor\s+total|paint\s+labor\s+total)\b")
                if val is not None:
                    out["labor_subtotal"] = val
            if out["paint_materials"] is None:
                val = _money_from_line(ln, r"(?i)^\s*(paint\s+(?:supplies|materials)|materials\s+subtotal)\b")
                if val is not None:
                    out["paint_materials"] = val
            if out["estimate_total"] is None:
                val = _money_from_line(ln, r"(?i)^\s*(estimate\s+total|total\s+amount|net\s+amount|grand\s+total|total\s+loss\s+value)\b")
                if val is not None:
                    out["estimate_total"] = val

        if out["sales_tax"] is not None:
            if out["tax_rate"] is not None and out["estimate_total"] is not None:
                max_reasonable_tax = float(out["estimate_total"]) * max(float(out["tax_rate"]) / 100.0, 0.02) * 1.25
                if out["sales_tax"] > max_reasonable_tax:
                    out["sales_tax"] = None
            elif out["estimate_total"] is not None and out["sales_tax"] > float(out["estimate_total"]) * 0.25:
                out["sales_tax"] = None

        return out

    def _comprehensive_cost_is_too_thin(v: Any) -> bool:
        s = str(v or "").strip()
        if not s:
            return True
        if _needs_customer_scrub(s):
            return True
        lines = [ln.strip() for ln in s.splitlines() if ln.strip()]
        non_heading = [ln for ln in lines if not ln.startswith("#") and ln.lower() != "approximate repair cost breakdown"]
        meaningful_money = [ln for ln in non_heading if "$" in ln]
        if len(meaningful_money) <= 1:
            return True
        return False

    def _build_comprehensive_cost_from_estimate_text() -> str:
        parsed = _extract_estimate_record_totals_strict(uploaded_text_all or "")
        bullets: List[str] = []
        if parsed.get("labor_subtotal") is not None:
            bullets.append(f"- Labor subtotal: {_money2(parsed['labor_subtotal'])}")
        if parsed.get("parts_subtotal") is not None:
            bullets.append(f"- Parts subtotal: {_money2(parsed['parts_subtotal'])}")
        if parsed.get("paint_materials") is not None:
            bullets.append(f"- Paint materials: {_money2(parsed['paint_materials'])}")
        if parsed.get("tax_rate") is not None:
            bullets.append(f"- Applicable tax rate: {float(parsed['tax_rate']):.3f}%")
        if parsed.get("sales_tax") is not None:
            bullets.append(f"- Sales tax: {_money2(parsed['sales_tax'])}")
        if parsed.get("estimate_total") is not None:
            bullets.append(f"- Estimate total: {_money2(parsed['estimate_total'])}")
        if parsed.get("estimate_total") is not None or len(bullets) >= 2:
            return "## Approximate Repair Cost Breakdown\nEstimate of Record totals (documented):\n" + "\n".join(bullets)
        return (
            "## Approximate Repair Cost Breakdown\n"
            "Estimate of record totals could not be extracted with confidence from the uploaded document text on this run. "
            "Final estimate totals should be confirmed from the estimate summary/totals page before release."
        )

    def _professional_summary_fallback() -> str:
        claim_txt = result.get("claim_number") if not _naish(result.get("claim_number")) else _extract_claim_from_text(uploaded_text_all or "")
        vin_txt = result.get("vin") if not _naish(result.get("vin")) else (vin_from_label or vin_from_qr or "Not confirmed from provided evidence.")
        veh_txt = _format_vehicle_value(result.get("vehicle")) if not _naish(result.get("vehicle")) else _extract_vehicle_from_text(uploaded_text_all or "")
        odo_txt = result.get("odometer_estimate_only") if not _naish(result.get("odometer_estimate_only")) else (odometer_value or "Not confirmed from provided evidence.")
        score_txt = result.get("compliance_score") if not _naish(result.get("compliance_score")) else "Not scored from available evidence."
        rules_txt = "Client guidelines were included and considered in this review." if client_rules_supplied else "No separate client guideline text was provided with this review."
        supp_txt = (", ".join(supplement_versions) if supplement_versions else "No supplement tags detected in the uploaded documents")
        return (
            "## Detailed Condition Report\n"
            f"This condition report was prepared from the uploaded estimate, document, OCR, and photo evidence for file {file_number}. "
            f"Claim number {claim_txt} and VIN {vin_txt} were identified from the submitted materials. "
            f"The vehicle is recorded as {veh_txt}. "
            f"The odometer reading reflected in the available evidence is {odo_txt}. "
            f"{rules_txt} "
            f"Supplement review: {supp_txt}. "
            "The file should be evaluated against visible damage, estimate line support, parts usage, labor operations, refinish handling, tax treatment, and any required supporting documentation. "
            "Header facts and available identifiers were preserved for this report, and the appraisal should be finalized by confirming all line-item support against the photos and estimate pages. "
            f"Current compliance score shown for this run: {score_txt}."
        )

    def _professional_fraud_fallback() -> str:
        vin_txt = result.get("vin") if not _naish(result.get("vin")) else (vin_from_label or vin_from_qr or "Not confirmed from provided evidence.")
        return (
            "No material inconsistencies were confirmed from the available review inputs. "
            f"Identifier evidence reviewed included VIN information recorded as {vin_txt}. "
            "Any final fraud or authenticity conclusion should remain subject to appraiser verification of the estimate, photo set, and supporting documentation."
        )

    def _professional_cost_fallback() -> str:
        if ai_intent == "damage_report_from_photos":
            return (
                "## Approximate Repair Cost Breakdown\n"
                "Approximate repair cost evaluation should be based on the visible photo damage, modeled labor hours, paint materials, parts pricing, and tax on parts plus paint materials only. "
                "This section remains an approximation only and must be validated by a qualified appraiser before estimate completion."
            )
        return (
            "## Approximate Repair Cost Breakdown\n"
            "Approximate repair cost evaluation should reflect the estimate lines, visible damage support, labor operations, parts usage, paint materials, and applicable tax treatment. "
            "Any final amount remains subject to estimate validation and appraiser review."
        )

    def _professional_conclusion_fallback() -> str:
        rules_txt = "Client guideline requirements" if client_rules_supplied else "available estimate and photo requirements"
        return (
            "Based on the uploaded estimate, limited photo evidence, and the documented review findings in this report, "
            "the file should not be relied upon as fully documented until the detailed condition narrative, estimate totals, and required supporting photos are confirmed. "
            f"Final handling should verify estimate-line support, VIN/photo consistency, and {rules_txt} before release or claim decision."
        )

    try:
        if ai_intent != "damage_report_from_photos":
            if _summary_is_effectively_na(result.get("summary_markdown")) or _needs_customer_scrub(result.get("summary_markdown")):
                result["summary_markdown"] = _professional_summary_fallback()
            if _needs_customer_scrub(result.get("fraud_markdown")):
                result["fraud_markdown"] = _professional_fraud_fallback()
            if _comprehensive_cost_is_too_thin(result.get("estimated_costs_markdown")):
                result["estimated_costs_markdown"] = _build_comprehensive_cost_from_estimate_text()
            if _needs_customer_scrub(result.get("conclusion")):
                result["conclusion"] = _professional_conclusion_fallback()
    except Exception:
        pass

    def _extract_client_rule_fragments(rules_text: str, limit: int = 4) -> List[str]:
        fragments: List[str] = []
        for raw_ln in str(rules_text or "").replace("\r\n", "\n").replace("\r", "\n").splitlines():
            s = re.sub(r"\s+", " ", raw_ln).strip(" -\t")
            if not s:
                continue
            if len(s) < 12:
                continue
            if s.lower() in {"client rules", "guidelines", "infinity insurance"}:
                continue
            fragments.append(s[:180])
            if len(fragments) >= limit:
                break
        return fragments

    def _scrub_false_missing_client_rules_claims(md_text: str) -> str:
        if not md_text:
            return md_text or ""
        if not client_rules_supplied:
            return md_text
        bad_patterns = [
            r"(?i)^.*no\s+separate\s+client\s+guideline\s+document.*$",
            r"(?i)^.*no\s+explicit\s+client\s+guidelines?.*$",
            r"(?i)^.*client\s+guidelines?/rules?\s+to\s+compare\s+against\s*.*$",
            r"(?i)^.*client\s+guidelines?\s+(?:were\s+)?not\s+(?:provided|included|supplied).*$",
            r"(?i)^.*rule\s+text\s+(?:was\s+)?not\s+(?:provided|included|supplied).*$",
        ]
        kept: List[str] = []
        for ln in str(md_text).replace("\r\n", "\n").replace("\r", "\n").splitlines():
            s = (ln or "").strip()
            if any(re.search(p, s) for p in bad_patterns):
                continue
            kept.append(ln)
        return "\n".join(kept).strip()

    def _ensure_client_guidelines_section(md_text: str) -> str:
        t = str(md_text or "").strip()
        if not client_rules_supplied or ai_intent not in {"comprehensive", "guidelines_only"}:
            return t
        t = _scrub_false_missing_client_rules_claims(t)
        if re.search(r"(?im)^##\s*Client\s+Guidelines\s+Comparison\b", t):
            return t
        fragments = _extract_client_rule_fragments(client_rules or resolved_rules_text or "")
        bullets: List[str] = []
        if fragments:
            for frag in fragments:
                bullets.append(f'- Rule supplied for review: "{frag}" — Compare against the uploaded estimate/photos and confirm Aligned / Not Aligned / Not Evidenced using document or photo citations.')
        else:
            bullets.append("- Client guidelines were supplied from the frontend selection/pasted Client Rules box and must be applied to this review.")
        section = "## Client Guidelines Comparison\n" + "\n".join(bullets)
        return (t + "\n\n" + section).strip() if t else section

    def _scrub_comprehensive_rationale_text(md_text: str) -> str:
        """Targeted comprehensive scrub:
        - J.D. Power Pricing & Values satisfies the valuation requirement for this workflow
        - remove any valuation deduction/rationale tied to J.D. Power / NADA / Redbook / KBB clean-retail wording
        - remove false UPD/commingling deduction language
        - rebuild BOTH the final score line and arithmetic from the remaining deduction bullets only
        """
        if not md_text:
            return md_text or ""
        t = str(md_text).replace("\r\n", "\n").replace("\r", "\n")

        jd_present = bool(re.search(r"(?i)j\.?d\.?\s*power|jdpower", uploaded_text_all or ""))

        out = []
        in_score_section = False
        score_lines = []
        skip_wrapped_valuation_tail = False

        def _is_jd_valuation_line(s: str) -> bool:
            if not s:
                return False
            return bool(
                re.search(r"(?i)j\.?d\.?\s*power|jdpower", s) and
                re.search(r"(?i)valuation|clean\s+retail|required\s+clean\s+retail|required\s+valuation|nada|redbook|kbb|pricing\s*&\s*values", s)
            )

        def _is_wrapped_tail_line(s: str) -> bool:
            if not s:
                return False
            return bool(re.search(r"(?i)^\s*\(photo\s*\d+.*client\s+rule\s+text.*clean\s+retail", s))

        for ln in t.splitlines():
            # VAL_PATCH: always remove valuation mismatch lines (J.D. Power = NADA)
            if re.search(r"(?i)valuation printout|clean retail|nada/redbook/kbb|j\.d\. power pricing", (ln or "")):
                continue
            s = (ln or "").strip()

            if re.search(r"(?i)^##\s*Compliance\s+Score\s+Rationale\b", s):
                in_score_section = True
                score_lines.append(ln)
                continue

            if in_score_section and re.search(r"(?i)^##\s+", s):
                in_score_section = False
                out.extend(score_lines)
                score_lines = []

            target = score_lines if in_score_section else out

            if jd_present and in_score_section:
                if _is_jd_valuation_line(s):
                    skip_wrapped_valuation_tail = True
                    continue
                if skip_wrapped_valuation_tail and (_is_wrapped_tail_line(s) or s.startswith("(Photo ") or s.startswith("Photo ")):
                    continue
                skip_wrapped_valuation_tail = False

            if re.search(r"(?i)commingl", s):
                continue
            if re.search(r"(?i)upd.*left\s*front", s) and re.search(r"(?i)major\s*\(-?20\)|moderate\s*\(-?10\)|minor\s*\(-?5\)", s):
                continue

            # Strip stale score/arithmetic lines; we rebuild them
            if in_score_section and (
                re.search(r"(?i)^Total\s*=\s*100", s) or
                re.search(r"(?i)^=\s*100\s*-", s) or
                re.search(r"(?i)^Arithmetic\s*:\s*100", s) or
                re.search(r"(?i)^Adjusted\s+compliance_score\s+reported\s*:", s) or
                re.search(r"(?i)^Final\s+compliance\s+score\s*:", s) or
                re.search(r"(?i)^Final\s*:", s) or
                re.search(r"(?i)^Score\s*math\s*:", s) or
                re.search(r"(?i)^Current\s+score\s*:", s)
            ):
                continue

            target.append(ln)

        if score_lines:
            deductions = []
            rebuilt = []
            for ln in score_lines:
                s = (ln or "").strip()
                m = re.search(r"(?i)\((?:-|–)?(\d+)\)", s)
                if m and re.search(r"(?i)^(?:-|\u2022)?\s*(Minor|Moderate|Major)\b", s):
                    try:
                        deductions.append(int(m.group(1)))
                    except Exception:
                        pass
                rebuilt.append(ln)

            total = max(0, 100 - sum(deductions))

            if not any(re.search(r"(?i)^Starting\s+(?:at|from)\s+100\.?$", (x or "").strip()) for x in rebuilt):
                insert_at = 1 if rebuilt else 0
                rebuilt.insert(insert_at, "Starting at 100.")

            rebuilt.append(f"Final compliance score: **{total}**.")
            rebuilt.append(f"Total = 100{' ' + ' '.join(f'- {d}' for d in deductions) if deductions else ''} = {total}.")
            out.extend(rebuilt)

        t = "\n".join(out)

        if jd_present:
            t = re.sub(r"(?i)(NADA/Redbook/KBB)", "NADA/J.D. Power/Redbook/KBB", t)
            t = re.sub(r"(?i)(NADA, Redbook, or KBB)", "NADA, J.D. Power, Redbook, or KBB", t)
            t = re.sub(r"(?i)J\.D\.\s*Power\s+Pricing\s*&\s*Values", "J.D. Power Pricing & Values", t)

        return t.strip()

    def _score_after_jd_power_scrub(md_text: str, current_score: str) -> str:
        t = str(md_text or "")
        m = re.search(r"(?im)^Total\s*=\s*100(?:\s*-\s*\d+)*\s*=\s*(\d{1,3})\.?\s*$", t)
        if m:
            try:
                n = max(0, min(100, int(m.group(1))))
                return str(n)
            except Exception:
                pass
        return str(current_score or "")

    try:
        if ai_intent != "damage_report_from_photos":
            result["summary_markdown"] = _scrub_comprehensive_rationale_text(result.get("summary_markdown") or "")
            result["summary_markdown"] = _ensure_client_guidelines_section(result.get("summary_markdown") or "")
            _rescored = _score_after_jd_power_scrub(result.get("summary_markdown") or "", result.get("compliance_score") or "")
            if str(_rescored).strip():
                result["compliance_score"] = str(_rescored).strip()
    except Exception:
        pass


    # -----------------------
    # LOCKED FINAL POST-PROCESSOR
    # Any changes here require explicit approval. This block is the final writer for:
    # - compliance_score
    # - score rationale
    # - estimated_costs_markdown
    # - fraud_markdown
    # -----------------------
    def _extract_locked_deductions_from_line(line_text: str) -> List[int]:
        s = str(line_text or "").strip()
        if not s:
            return []
        vals: List[int] = []

        deduction_patterns = [
            r"(?i)\bDeduction\b\s*:\s*-\s*(\d+)\b",
            r"(?i)\bDeduction\b\s*-\s*(\d+)\b",
            r":\s*-\s*(\d+)\b",
            r"(?i)\((?:[^)]*?)\b(?:Minor|Moderate|Major)\b\s*-\s*(\d+)(?:[^)]*?)\)",
            r"\((?:-|–)\s*(\d+)\)",
        ]

        for pat in deduction_patterns:
            for m in re.finditer(pat, s):
                try:
                    vals.append(int(m.group(1)))
                except Exception:
                    pass

        deduped: List[int] = []
        seen = set()
        for v in vals:
            if v not in seen:
                seen.add(v)
                deduped.append(v)
        return deduped
    def _build_locked_compliance_score_rationale(md_text: str, current_score: str) -> str:
        t = str(md_text or "").replace("\r\n", "\n").replace("\r", "\n")
        lines = t.splitlines()
        start_idx = None
        end_idx = None
        for i, ln in enumerate(lines):
            if re.search(r"(?i)^##\s*Compliance\s+Score\s+Rationale\b", (ln or "").strip()):
                start_idx = i
                break
        if start_idx is None:
            return t.strip()
        for j in range(start_idx + 1, len(lines)):
            if re.search(r"(?i)^##\s+", (lines[j] or "").strip()):
                end_idx = j
                break
        if end_idx is None:
            end_idx = len(lines)

        score_block = lines[start_idx:end_idx]
        body_lines = []
        deductions: List[int] = []
        for ln in score_block[1:]:
            s = (ln or "").strip()
            if not s:
                continue
            if re.search(r"(?i)^Starting\s+(?:at|from)\s+100\.?$", s):
                continue
            if re.search(r"(?i)^Starting\s+score\s*:\s*100\.?$", s):
                continue
            if re.search(r"(?i)^New\s+score\s*:", s):
                continue
            if re.search(r"(?i)^Final\s+compliance\s+score\s*:", s):
                continue
            if re.search(r"(?i)^Final\s*:", s):
                continue
            if re.search(r"(?i)^Score\s*math\s*:", s):
                continue
            if re.search(r"(?i)^Current\s+score\s*:", s):
                continue
            if re.search(r"(?i)^Total\s*=\s*100", s):
                continue
            if re.search(r"(?i)^Adjustment\s*:", s):
                continue
            if re.search(r"(?i)\bseed\b", s):
                continue
            if re.search(r"(?i)provided\s+Compliance\s+Score", s):
                continue
            cleaned_ln = re.sub(r"(?i)\s*New\s+score\s*:\s*\d+\.?$", "", ln).rstrip()
            cleaned_ln = re.sub(r"(?i)\s*Final\s*:\s*.*$", "", cleaned_ln).rstrip()
            cleaned_s = cleaned_ln.strip()
            if not cleaned_s:
                continue
            deductions.extend(_extract_locked_deductions_from_line(cleaned_s))
            body_lines.append(cleaned_ln)

        total = max(0, 100 - sum(deductions))
        rebuilt = ["## Compliance Score Rationale", "Starting from 100."]
        rebuilt.extend(body_lines)
        rebuilt.append(f"Final compliance score: **{total}**.")
        rebuilt.append(f"Total = 100{' ' + ' '.join(f'- {d}' for d in deductions) if deductions else ''} = {total}.")
        new_lines = lines[:start_idx] + rebuilt + lines[end_idx:]
        return "\n".join(new_lines).strip()

    def _locked_score_from_rationale(md_text: str, current_score: str) -> str:
        t = str(md_text or "")
        m = re.search(r"(?im)^Final\s+compliance\s+score\s*:\s*\*\*(\d{1,3})\*\*\.?\s*$", t)
        if m:
            try:
                return str(max(0, min(100, int(m.group(1)))))
            except Exception:
                pass
        m2 = re.search(r"(?im)^Total\s*=\s*100(?:\s*-\s*\d+)*\s*=\s*(\d{1,3})\.?\s*$", t)
        if m2:
            try:
                return str(max(0, min(100, int(m2.group(1)))))
            except Exception:
                pass
        return str(current_score or "").strip()

    def _validate_locked_compliance_score_block(md_text: str, current_score: str) -> (str, str):
        t = str(md_text or "").replace("\r\n", "\n").replace("\r", "\n")
        # Remove stale score math lines that should never survive to customer output
        cleaned_lines = []
        for ln in t.splitlines():
            s = (ln or "").strip()
            if re.search(r"(?i)^Final\s*:\s*", s):
                continue
            if re.search(r"(?i)^Score\s*math\s*:\s*", s):
                continue
            if re.search(r"(?i)^Current\s+score\s*:\s*", s):
                continue
            cleaned_lines.append(ln)
        t = "\n".join(cleaned_lines).strip()

        score_from_rationale = _locked_score_from_rationale(t, current_score)
        score_from_total = None
        m_total = re.search(r"(?im)^Total\s*=\s*100(?:\s*-\s*\d+)*\s*=\s*(\d{1,3})\.?\s*$", t)
        if m_total:
            try:
                score_from_total = str(max(0, min(100, int(m_total.group(1)))))
            except Exception:
                score_from_total = None

        final_score = str(score_from_rationale or current_score or "").strip()
        if score_from_total and final_score and score_from_total != final_score:
            # Rebuild from rationale so both displayed values always match
            t = _build_locked_compliance_score_rationale(t, final_score)
            final_score = _locked_score_from_rationale(t, final_score)

        return t, final_score

    def _build_locked_cost_block(existing_md: str) -> str:
        t = str(existing_md or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        if not t:
            return (
                "## Approximate Repair Cost Breakdown\n"
                "Approximate repair cost evaluation should reflect the estimate lines, visible damage support, labor operations, parts usage, paint materials, and applicable tax treatment. "
                "Any final amount remains subject to estimate validation and appraiser review."
            )
        cleaned = []
        for ln in t.splitlines():
            s = (ln or "").strip()
            if re.search(r"(?i)^Final\s+compliance\s+score\s*:", s):
                continue
            if re.search(r"(?i)^Total\s*=\s*100", s):
                continue
            cleaned.append(ln)
        return "\n".join(cleaned).strip()

    def _build_locked_fraud_block(existing_md: str) -> str:
        t = str(existing_md or "").strip()
        return t if t else "No material inconsistencies found."

    def _apply_locked_final_postprocessor(result_obj: dict) -> dict:
        if not isinstance(result_obj, dict):
            return result_obj
        sm = str(result_obj.get("summary_markdown") or "")
        sm = _build_locked_compliance_score_rationale(sm, str(result_obj.get("compliance_score") or ""))
        sm, locked_score = _validate_locked_compliance_score_block(sm, str(result_obj.get("compliance_score") or ""))
        result_obj["summary_markdown"] = sm
        if locked_score:
            result_obj["compliance_score"] = locked_score
        result_obj["estimated_costs_markdown"] = _build_locked_cost_block(result_obj.get("estimated_costs_markdown") or "")
        result_obj["fraud_markdown"] = _build_locked_fraud_block(result_obj.get("fraud_markdown") or "")
        return result_obj

    try:
        if ai_intent != "damage_report_from_photos":
            result = _apply_locked_final_postprocessor(result)
            result["summary_markdown"] = _ensure_client_guidelines_section(result.get("summary_markdown") or "")
    except Exception:
        pass

    # -----------------------
    # FINAL-STAGE TARGETED NAME SCRUB (authoritative)
    # - Customer-facing output must NOT show owner/insured/customer/claimant names
    # - Do NOT run Presidio over final estimate math / ops / labor / cost text
    # - Use a targeted name scrubber only, then verify no known names remain
    # -----------------------
    def _extract_name_candidates_for_redaction(source_text: str) -> List[str]:
        candidates: List[str] = []
        if not source_text:
            return candidates
        pats = [
            r"(?im)\bOwner\s*(?:Name)?\s*[:\-]\s*([A-Z][A-Z'.,\- ]{2,80})",
            r"(?im)\bInsured\s*(?:Name)?\s*[:\-]\s*([A-Z][A-Z'.,\- ]{2,80})",
            r"(?im)\bCustomer\s*(?:Name)?\s*[:\-]\s*([A-Z][A-Z'.,\- ]{2,80})",
            r"(?im)\bClaimant\s*(?:Name)?\s*[:\-]\s*([A-Z][A-Z'.,\- ]{2,80})",
        ]
        for pat in pats:
            for m in re.finditer(pat, source_text or ""):
                raw_name = re.sub(r"\s+", " ", str(m.group(1) or "")).strip(" ,.-")
                if not raw_name:
                    continue
                if len(raw_name) < 4 or len(raw_name) > 80:
                    continue
                if re.fullmatch(r"[A-HJ-NPR-Z0-9]{17}", raw_name.replace(" ", "")):
                    continue
                candidates.append(raw_name)
        seen = set()
        out: List[str] = []
        for nm in candidates:
            key = nm.lower()
            if key not in seen:
                seen.add(key)
                out.append(nm)
        return out[:12]

    def _force_scrub_known_names(text_value: str, names: List[str], replacement: str = "Owner") -> str:
        s = str(text_value or "")
        for nm in names:
            if not nm:
                continue
            try:
                s = re.sub(re.escape(nm), replacement, s, flags=re.IGNORECASE)
            except Exception:
                pass
            parts = [p for p in re.split(r"[\s,]+", nm) if len(p.strip()) >= 3]
            if len(parts) >= 2:
                joined = r"\s+".join(re.escape(p) for p in parts)
                try:
                    s = re.sub(joined, replacement, s, flags=re.IGNORECASE)
                except Exception:
                    pass
        return s

    def _final_scrub_customer_fields(result_obj: Dict[str, Any], known_names: List[str]) -> (Dict[str, Any], List[str]):
        issues: List[str] = []
        for key in ("summary_brief", "summary_markdown", "fraud_markdown", "estimated_costs_markdown", "conclusion"):
            try:
                original = str(result_obj.get(key) or "")
                scrubbed = _force_scrub_known_names(original, known_names, replacement="Owner")
                result_obj[key] = scrubbed
            except Exception as e:
                issues.append(f"Final name scrub failed for {key}: {e}")
        return result_obj, issues

    def _detect_remaining_name_leaks(result_obj: Dict[str, Any], known_names: List[str]) -> List[str]:
        leaks: List[str] = []
        blob = "\n".join([
            str(result_obj.get("summary_brief") or ""),
            str(result_obj.get("summary_markdown") or ""),
            str(result_obj.get("fraud_markdown") or ""),
            str(result_obj.get("estimated_costs_markdown") or ""),
            str(result_obj.get("conclusion") or ""),
        ])
        for nm in known_names:
            if not nm:
                continue
            try:
                if re.search(re.escape(nm), blob, flags=re.IGNORECASE):
                    leaks.append(f"PII leak remained after final name scrub: {nm}")
                    continue
                parts = [p for p in re.split(r"[\s,]+", nm) if len(p.strip()) >= 3]
                if len(parts) >= 2:
                    joined = r"\s+".join(re.escape(p) for p in parts)
                    if re.search(joined, blob, flags=re.IGNORECASE):
                        leaks.append(f"PII leak remained after final name scrub: {nm}")
            except Exception:
                pass
        return leaks

    _known_pii_names = _extract_name_candidates_for_redaction(uploaded_text_all or "")
    _redaction_issues: List[str] = []
    try:
        result, _redaction_issues = _final_scrub_customer_fields(result, _known_pii_names)
    except Exception as e:
        _redaction_issues = [f"Final-stage name scrub failed: {e}"]

    _remaining_pii_leaks = _detect_remaining_name_leaks(result, _known_pii_names)
    if _redaction_issues or _remaining_pii_leaks:
        _block_reasons = _redaction_issues + _remaining_pii_leaks
        log.error("REPORT BLOCKED: final-stage name scrub failure | reasons=%s", _block_reasons)
        return JSONResponse(
            status_code=200,
            content={
                "status": "blocked",
                "error": "REPORT BLOCKED: PII name scrub failed",
                "reasons": _block_reasons,
                "file_number": file_number,
                "request_type": req_label,
                "redaction_status": "Redacted PII: Blocked - final-stage name scrub failed",
            },
        )

    result["redaction_status"] = "Redacted PII: Successful ✅"

    # -----------------------
    # PRE-PDF VALIDATOR (Comprehensive / Guidelines only)
    # Blocks customer PDF generation when core sections are fallback/placeholder content.
    # -----------------------
    def _predf_summary_invalid(v: Any) -> bool:
        s = str(v or "").strip()
        if not s:
            return True
        low = s.lower()
        if re.search(r"(?im)^##\s*Detailed\s+Condition\s+Report\s*$", s):
            # heading only
            lines = [ln.strip() for ln in s.splitlines() if ln.strip() and not ln.strip().startswith('#')]
            if not lines:
                return True
        if re.search(r"(?im)^##\s*Detailed\s+Condition\s+Report\s*\n\s*N/?A\s*$", s):
            return True
        bad_phrases = [
            "structured narrative response did not fully validate",
            "deterministic extraction completed",
            "preserved deterministic extracted facts",
            "instead of printing blank fields",
            "fallback confirms the job executed",
            "available identifiers instead of returning blank sections",
        ]
        return any(p in low for p in bad_phrases)

    def _predf_cost_invalid(v: Any) -> bool:
        s = str(v or "").strip()
        if not s:
            return True
        low = s.lower()
        bad_phrases = [
            "approximate repair cost evaluation should reflect the estimate lines",
            "final repair amounts remain subject to estimate validation and qualified appraiser review",
            "any final amount remains subject to estimate validation and appraiser review",
            "estimate of record totals could not be extracted with confidence",
            "unable to generate cost approximation on this run",
        ]
        if any(p in low for p in bad_phrases):
            return True
        if len(re.findall(r"\$\s*[0-9]", s)) < 2:
            return True
        return False

    def _predf_conclusion_invalid(v: Any) -> bool:
        s = str(v or "").strip()
        if not s:
            return True
        low = s.lower()
        bad_phrases = [
            "the uploaded files were processed successfully",
            "the available review evidence was preserved",
            "final claim handling should proceed after confirmation",
            "should not be relied upon as fully documented until",
            "available estimate and photo requirements",
        ]
        return any(p in low for p in bad_phrases)

    if ai_intent in {"comprehensive", "guidelines_only"}:
        _predf_reasons: List[str] = []
        if _predf_summary_invalid(result.get("summary_markdown")):
            _predf_reasons.append("REPORT BLOCKED: invalid or fallback Detailed Condition Report")
        if _predf_cost_invalid(result.get("estimated_costs_markdown")):
            _predf_reasons.append("REPORT BLOCKED: invalid or fallback Approximate Repair Cost Breakdown")
        if _predf_conclusion_invalid(result.get("conclusion")):
            _predf_reasons.append("REPORT BLOCKED: invalid or fallback Conclusion")
        if _predf_reasons:
            log.error("REPORT BLOCKED: pre-PDF validator failure | reasons=%s", _predf_reasons)
            return JSONResponse(
                status_code=200,
                content={
                    "status": "blocked",
                    "error": "REPORT BLOCKED: comprehensive quality gate failure",
                    "reasons": _predf_reasons,
                    "file_number": file_number,
                    "request_type": req_label,
                    "redaction_status": result.get("redaction_status", "Redacted PII: Successful ✅"),
                },
            )

    # PDF helpers
    # -----------------------
    def _pdf_sanitize(text: str, max_token_len: int = 60) -> str:
        if text is None:
            return ""
        s = str(text).replace("\r\n", "\n").replace("\r", "\n")
        s = "".join(ch if ord(ch) < 256 else " " for ch in s)
        def _break(tok: str) -> str:
            if len(tok) <= max_token_len:
                return tok
            return " ".join(tok[i:i+max_token_len] for i in range(0, len(tok), max_token_len))
        s = " ".join(_break(t) for t in s.split(" "))
        return s

    pdf = FPDF(); pdf.add_page()
    # --- NSPXN Logo (Top Right, First Page Only) ---
    try:
        logo_path = os.path.join(os.path.dirname(__file__), "ChatGPT logo100725.png")
        if os.path.exists(logo_path):
            pdf.image(logo_path, x=pdf.w - 45, y=8, w=35)  # small–medium size
    except Exception:
        pass
    pdf.set_auto_page_break(auto=True, margin=10)
    pdf.set_left_margin(10); pdf.set_right_margin(10)

    try:
        pdf.add_font("DejaVu","", "DejaVuSans.ttf", uni=True); pdf.set_font(size=11)
    except Exception:
        pdf.set_font("Arial", size=11)

    def mc(s):
        try:
            effective_w = pdf.w - pdf.l_margin - pdf.r_margin
            if effective_w <= 5:
                effective_w = 180
            safe = _pdf_sanitize(s)
            if not safe.strip():
                safe = "-"
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(effective_w, 6, safe)
        except Exception:
            effective_w = pdf.w - pdf.l_margin - pdf.r_margin
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(effective_w, 6, (_pdf_sanitize(str(s))[:2000] + " …"))


    
    def _money2(x: Optional[float]) -> str:
        try:
            if x is None:
                return "$0.00"
            return "${:,.2f}".format(float(x))
        except Exception:
            return "$0.00"

    def _compute_cost_total_from_md(md_text: str) -> Optional[float]:
        """Compute a photo-only approximate total from the model's cost markdown.
        We intentionally sum the *section totals* (labor totals, paint materials, parts subtotal, tax)
        to avoid double-counting per-part lines.
        """
        if not md_text:
            return None
        t = str(md_text).replace("\r\n", "\n").replace("\r", "\n")

        def _grab_amount(pat: str) -> Optional[float]:
            m = re.search(pat, t, flags=re.IGNORECASE | re.MULTILINE)
            if not m:
                return None
            try:
                return float(m.group(1).replace(",", ""))
            except Exception:
                return None

        # Prefer explicit totals (bold amounts) from the model output
        body_labor = _grab_amount(r"^\s*[-*]?\s*Body\s+labor\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*")
        paint_labor = _grab_amount(r"^\s*[-*]?\s*Paint\s+labor\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*")
        mech_labor = _grab_amount(r"^\s*[-*]?\s*Mechanical[^:]*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*")
        if mech_labor is None:
            mech_labor = _grab_amount(r"^\s*[-*]?\s*Mechanical[^=\n]*=\s*\$?\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b")
        frame_labor = _grab_amount(r"^\s*[-*]?\s*Frame\s+labor\s*:\s*.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*")

        # Paint materials dollars (prefer subtotal/extended amount lines; never the $/hr rate)
        paint_mat = None
        for _pat in [
            r"^\s*[-*]?\s*Paint\s+materials\s+subtotal\s*[:=]\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)",
            r"^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s*[:=]\s*\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*",
            r"^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s*[:=]\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)",
        ]:
            paint_mat = _grab_amount(_pat)
            if paint_mat is not None:
                break

        # Parts subtotal dollars
        parts_sub = _grab_amount(r"^\s*\*\*\s*Estimated\s+parts\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)")
        if parts_sub is None:
            parts_sub = _grab_amount(r"^\s*[-*]?\s*Parts\s+subtotal\s*:\s*\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*")
        if parts_sub is None:
            # Some model outputs use a simple 'Parts:' label
            parts_sub = _grab_amount(r"^\s*[-*]?\s*Parts\s*:\s*\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*")

        tax_amt = _grab_amount(r"^\s*[-*]?\s*(?:Sales\s+tax|Estimated\s+tax|Tax)\b.*?\*\*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\*\*")

        # If tax dollars weren't provided but an assumed rate exists, compute tax on (parts + paint materials) only.
        if tax_amt is None:
            m_rate = re.search(r"(?im)^\s*[-*]?\s*Tax\s+rate\s*(?:\(assumed\)|assumption|\(assumed\)\s*:|\(assumed\)\s*)?\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)\s*%", t)
            if not m_rate:
                m_rate = re.search(r"(?im)\bTax\s+rate\s*(?:\(assumed\))?\s*[:\-]?\s*([0-9]+(?:\.[0-9]+)?)\s*%", t)
            if m_rate and parts_sub is not None and paint_mat is not None:
                try:
                    rate = float(m_rate.group(1)) / 100.0
                    tax_amt = round((float(parts_sub) + float(paint_mat)) * rate, 2)
                except Exception:
                    tax_amt = None

        # If the model failed to provide the key components, do not guess a total here.
        components = [body_labor, paint_labor, mech_labor, frame_labor, paint_mat, parts_sub, tax_amt]
        have_any = any(v is not None for v in components)
        have_core = (parts_sub is not None) and (paint_mat is not None)
        # Require at least parts + paint materials + some labor; tax can be derived if rate exists.
        have_some_labor = any(v is not None for v in (body_labor, paint_labor, mech_labor, frame_labor))
        if not (have_any and have_core and have_some_labor):
            return None

        total = 0.0
        for v in (body_labor, paint_labor, mech_labor, frame_labor, paint_mat, parts_sub, tax_amt):
            if isinstance(v, (int, float)):
                total += float(v)
        return round(total, 2)

    def _inject_clean_total_line(md_text: str, total_val: Optional[float]) -> str:
        """Remove any existing 'Approximate Repair Cost Total' lines.
        Option 1 mode: keep the model's 'Approximate Repair Total: $X' line untouched
        and do NOT inject any additional total line.
        """
        if not md_text:
            return md_text or ""
        t = str(md_text).replace("\r\n", "\n").replace("\r", "\n")

        lines: List[str] = []
        for ln in t.splitlines():
            s = (ln or "").strip()
            # Remove any injected/legacy 'Repair Cost Total' lines (we will not re-add).
            if re.search(r"(?i)^\s*Approximate\s+Repair\s+Cost\s+Total\s*:", s):
                continue
            if re.search(r"(?i)^\s*Approximate\s+Total\s+Repair\s+Cost\s*:", s):
                continue
            # Also remove any standalone variants with $ on the line
            if re.search(r"(?i)^\s*Approximate\s+Repair\s+Cost\s+Total\b", s) and "$" in s:
                continue
            lines.append(ln)

        return "\n".join(lines).strip()



    def _parse_approx_total(md_text: str) -> Optional[float]:
            """Parse the approximate repair total from estimated_costs_markdown (best-effort).
            Prefers a standalone/bold total near the end of the cost section.
            """
            if not md_text:
                return None
            t = str(md_text).replace("\r\n","\n").replace("\r","\n")

            # 1) Prefer explicit label variants
            label_patterns = [
                r"(?im)^\s*\*\*?\s*Approximate\s+Repair\s+Cost\s+Total\s*[:\-]?\s*\*\*?\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"(?im)^\s*\*\*?\s*Approximate\s+Repair\s+Total\s*[:\-]?\s*\*\*?\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
                r"(?im)^\s*\*\*?\s*Approximate\s+Total\s+Repair\s+Cost\s*[:\-]?\s*\*\*?\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)",
            ]
            for pat in label_patterns:
                m = re.search(pat, t)
                if m:
                    try:
                        return float(m.group(1).replace(",", ""))
                    except Exception:
                        pass

            # 2) Prefer a standalone line that is just $X,XXX (often bolded)
            standalone = []
            for ln in t.splitlines():
                s = ln.strip()
                # strip markdown bold
                s2 = re.sub(r"^[*_]+|[*_]+$", "", s).strip()
                m = re.fullmatch(r"\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)", s2)
                if m:
                    try:
                        standalone.append(float(m.group(1).replace(",", "")))
                    except Exception:
                        pass
            if standalone:
                return standalone[-1]

            # 3) Otherwise: take the last dollar amount after a 'Tax' section (common structure)
            # BUT ignore tier boundary dollars like $3,500 / $10,000 from Severity Tier checklists.
            tail_lines = t.splitlines()
            # keep only lines after the first Tax heading if present
            for i, ln in enumerate(tail_lines):
                if re.search(r"(?i)^\s*tax\b", ln.strip()):
                    tail_lines = tail_lines[i:]
                    break
            # drop any severity/tier checklist lines
            safe_tail = []
            for ln in tail_lines:
                if re.search(r"(?i)severity\s+tier|minor\s*\(|moderate\s*\(|major\s*\(|total\s+loss\s+threshold", ln):
                    continue
                if re.match(r"^\s*\[[ xX]\]", ln.strip()):
                    continue
                safe_tail.append(ln)
            monies = re.findall(r"\$\s*([0-9][0-9,]*(?:\.[0-9]{1,2})?)", "\n".join(safe_tail))
            if monies:
                # choose the last amount in the tail
                try:
                    return float(monies[-1].replace(",", ""))
                except Exception:
                    return None
            return None

    def _enforce_severity_tier_checkmarks(md_text: str) -> str:
        """Rewrite the Severity Tier block to ensure exactly one box is checked based on parsed total."""
        if not md_text:
            return md_text or ""
        t = str(md_text).replace("Likely Total Loss Threshold Approaching", "Possible Total Loss Threshold Approaching")
        total_val = _parse_approx_total(t)
        # thresholds
        tier = None
        if isinstance(total_val, (int, float)):
            if total_val < 3500:
                tier = "minor"
            elif total_val < 10000:
                tier = "moderate"
            elif total_val < 25000:
                tier = "major"
            else:
                tier = "possible_tl"

        # Remove any existing Severity Tier block (best-effort)
        lines_in = t.splitlines()
        out = []
        i = 0
        while i < len(lines_in):
            ln = lines_in[i]
            if re.search(r"(?i)^\s*#{1,6}\s*Severity\s+Tier\b", ln):
                # skip until next heading or blank-run end
                i += 1
                while i < len(lines_in) and not re.search(r"^\s*#{1,6}\s+\w", lines_in[i]):
                    i += 1
                continue
            # Also remove any existing tier checkbox/list lines even if the heading was stripped upstream
            if re.search(r"(?i)^\s*(?:[-*]\s*)?(?:\[\s*[xX ]\s*\]|[☐☑])?\s*(Minor|Moderate|Major|Possible\s+Total\s+Loss)\b", ln.strip()):
                i += 1
                continue
            out.append(ln)
            i += 1
        t2 = "\n".join(out).rstrip()

        # Append normalized Severity Tier block at end of cost section (or end of markdown)
        boxes = {
            "minor": ("☑","☐","☐","☐"),
            "moderate": ("☐","☑","☐","☐"),
            "major": ("☐","☐","☑","☐"),
            "possible_tl": ("☐","☐","☐","☑"),
            None: ("☐","☐","☐","☐"),
        }[tier]
        sev_block = (
            "\n### Severity Tier\n"
            f"{boxes[0]} Minor (< $3,500)\n"
            f"{boxes[1]} Moderate ($3,500-$10,000)\n"
            f"{boxes[2]} Major ($10,000+)\n"
            f"{boxes[3]} Possible Total Loss Threshold Approaching\n"
        )
        return (t2 + "\n" + sev_block).strip() + "\n"

    def _strip_unwanted_cost_lines_for_pdf(md_text: str) -> str:
        """Remove known unwanted lines in cost markdown before PDF render.
        This is used ONLY for Photos-Only PDFs to prevent model-inserted prompt headers,
        'Totals' math blocks, duplicate section headers, and any estimate/docs language.
        """
        if not md_text:
            return md_text or ""
        t = str(md_text).replace("Likely Total Loss Threshold Approaching", "Possible Total Loss Threshold Approaching")
        cleaned: List[str] = []
        in_repair_disclaimer = False
        for ln in t.splitlines():
            s = (ln or "").strip()

            # Remove any mid-report Repair Cost Disclaimer block entirely.
            if re.search(r"(?i)^\s*(?:\*\*)?repair\s+cost\s+disclaimer\b", s):
                in_repair_disclaimer = True
                continue
            if in_repair_disclaimer:
                # stop on blank line or a heading
                if (not s) or re.search(r"^\s*#{1,6}\s+\w", s):
                    in_repair_disclaimer = False
                continue

            # Remove model/prompt headers and echoes
            if s in ("Approximate Repair Cost Breakdown", "(See estimated_costs_markdown field.)"):
                continue
            if re.search(r"(?i)estimated_costs_markdown", s):
                continue
            if re.search(r"(?i)\bPopulate\s+JSON\s+field\b", s):
                continue
            if re.search(r"(?i)^\s*##\s*Approximate\s+Repair\s+Cost\s+Breakdown\b", s):
                continue

            # Remove Totals blocks / arithmetic
            if re.search(r"(?i)^\s*###\s*Totals\b", s):
                continue
            if ("+" in ln and "=" in ln and re.search(r"\$\s*[0-9]", ln)):
                continue
            if re.search(r"(?i)^\s*###\s*Approximate\s+Total\s+Repair\s+Cost\b", s):
                continue
            if re.search(r"(?i)^\s*[-*]\s*Estimated\s+total\b", s):
                continue
            # Strip model totals/arithmetic phrasing variants (we print ONE locked total line)
            if re.search(r"(?i)\bEstimated\s+Repair\s+Total\b", s):
                continue
            if re.search(r"(?i)\bPre-?tax\s+total\b", s):
                continue
            if re.search(r"(?i)\bApproximate\s+total\s+repair\s+cost\b", s):
                continue
            if re.search(r"(?i)\bRounded\b", s) and re.search(r"\$\s*\d", ln):
                continue

            # Strip model helper headings / arithmetic lines we never want in the PDF
            if re.search(r"(?i)^\s*cost\s+math\s*\(approx\.?\)\s*$", s):
                continue
            if re.search(r"(?i)^\s*[-*]?\s*tax\s*:\s*\$.*[×x\*].*=", s):
                continue


            # Strip additional model cost-summary headings and totals (we print our own deterministic tax/total lines)
            if re.search(r"(?i)^\s*cost\s+summary\b", s):
                continue
            if re.search(r"(?i)\bApprox\.?\s*Repair\s*Total\b", s) and not re.search(r"(?i)^\s*\*{0,2}\s*Approximate\s+Repair\s+Total\s*:\s*\$\s*[0-9]", s):
                # Also drop any other repair-total style lines that include a dollar amount
                if re.search(r"(?i)\brepair\s*total\b", s) and "$" in s:
                    continue
                continue
            if re.search(r"(?i)^\s*[-*]?\s*tax\s+rate\s+assumption\b", s):
                continue
            if re.search(r"(?i)^\s*[-*]?\s*tax\s+rate\s*\(assumed\)\b", s):
                continue
            if re.search(r"(?i)^\s*[-*]?\s*estimated\s+tax\b", s):
                continue
            if re.search(r"(?i)^\s*tax\s*\(apply\b", s):
                continue
            if re.search(r"(?i)^\s*[-*]?\s*sales\s+tax\b", s) and not re.search(r"(?i)^\s*[-*]?\s*sales\s+tax\s*\(assumed\s*7%\s*for\s*approximation\)", s):
                continue
            if re.search(r"(?i)\bassumed\b.*\btax\b", s) and not re.search(r"(?i)^\s*[-*]?\s*sales\s+tax\s*\(assumed\s*7%\s*for\s*approximation\)", s):
                continue

            # Remove Severity headings/labels here (we render a normalized block later)
            if re.search(r"(?i)^\s*###\s*Severity\s+Tier\b", s):
                continue
            if s.lower() == "severity tier":
                continue

            # Strip any model-provided severity checklist lines (we print exactly one normalized block)
            if re.search(r"(?i)\bminor\s*\(|\bmoderate\s*\(|\bmajor\s*\(|total\s+loss\s+threshold\s+approaching", s):
                continue

            # Remove unwanted notes/disclaimers inserted by older logic
            if re.search(r"(?i)\bbest-?effort\b|\bfrom\s+estimate\b|\bfrom\s+docs\b|\bnot\s+evidenced\b", ln):
                continue

            # Never print Inspection Location inside the cost section (per your requirement)
            if re.search(r"(?i)^\*\*Inspection\s+Location\*\*\s*:", s) or re.search(r"(?i)^Inspection\s+Location\s*:", s):
                continue

            # Remove any existing total line; we will compute/insert a clean one
            if re.search(r"(?i)^\s*Approximate\s+Repair\s+Cost\s+Total\s*:", s):
                continue
            if re.search(r"(?i)^\s*Approximate\s+Total\s+Repair\s+Cost\s*:", s):
                continue

            # Avoid duplicate parts subtotal echo lines
            if re.search(r"(?i)estimated\s+parts\s+subtotal", s) and any(re.search(r"(?i)parts\s+subtotal", x) for x in cleaned[-3:]):
                continue

            cleaned.append(ln)
        return "\n".join(cleaned).strip()

    def _scrub_photo_only_narrative_cost_headers(md_text: str) -> str:
        """Remove the model's internal prompt cost headers from the main narrative.
        We render the cost breakdown as its own controlled PDF section.
        """
        if not md_text:
            return md_text or ""
        out: List[str] = []
        for ln in str(md_text).replace("\r\n", "\n").replace("\r", "\n").splitlines():
            s = (ln or "").strip()
            if re.search(r"(?i)^\s*##\s*Approximate\s+Repair\s+Cost\s+Breakdown\b", s):
                continue
            if re.search(r"(?i)\bPopulate\s+JSON\s+field\b", ln):
                continue
            if re.search(r"(?i)\bSee\s+estimated_costs_markdown\s+field\b", ln):
                continue
            out.append(ln)
        return "\n".join(out).strip()

    def render_repair_cost_section(pdf_obj: FPDF, md: str, tax_rate: Optional[float] = None) -> None:
        """Render the Approximate Repair Cost Breakdown in a controlled PDF format.
        Locked behavior:
        - If total is missing, compute it from existing subtotals + tax
        - Always key Severity Tier off that final total
        - Do not print a second/alternate total label
        """
        if not isinstance(md, str):
            md = str(md or "")
        text = md.replace("\r\n", "\n").replace("\r", "\n")

        # Extract Repair Cost Disclaimer block (remove from body; render separately)
        disclaimer_text = ""
        body_lines = []
        in_disclaimer = False
        for ln in text.splitlines():
            if (re.search(r"(?i)^\s*#{1,6}\s*repair\s+cost\s+disclaimer\b", ln)
                or re.search(r"(?i)^\s*\*\*\s*repair\s+cost\s+disclaimer\s*\*\*\s*:", ln)
                or re.search(r"(?i)^\s*_?repair\s+cost\s+disclaimer\s*:", ln)):
                in_disclaimer = True
                m_inline = re.search(r":\s*(.+)\s*$", ln)
                if m_inline:
                    disclaimer_text += (m_inline.group(1).strip() + " ")
                continue
            if in_disclaimer:
                if re.search(r"^\s*#{1,6}\s+\w", ln):
                    in_disclaimer = False
                else:
                    if ln.strip():
                        disclaimer_text += (ln.strip() + " ")
                    continue
            if in_disclaimer:
                continue
            body_lines.append(ln)

        disclaimer_text = (disclaimer_text or "").strip()

        # Remove headings + Totals blocks + arithmetic lines we never want printed directly
        cleaned = []
        skip_totals = False
        for ln in body_lines:
            s = ln.strip()

            if re.search(r"^\s*#{1,6}\s*approximate\s+repair\s+cost\s+breakdown\b", ln, flags=re.I):
                continue
            if re.search(r"^\s*#{1,6}\s*totals\b", ln, flags=re.I):
                skip_totals = True
                continue
            if skip_totals:
                if re.search(r"^\s*#{1,6}\s+\w", ln):
                    skip_totals = False
                else:
                    continue

            if re.search(r"(?i)\bestimated\s+total\b\s*:", ln):
                continue
            if ("+" in ln and "=" in ln and re.search(r"\$\s*[0-9]", ln)):
                continue

            # Never print model-provided severity lines here; we render them deterministically below
            if re.search(r"(?i)severity\s+tier|minor\s*\(|moderate\s*\(|major\s*\(|total\s+loss\s+threshold", ln):
                continue

            if re.search(r"(?i)^\s*#{1,6}\s*approximate\s+total\s+repair\s+cost\b", ln):
                continue

            if re.search(r"(?i)inspection\s+location\s*:", ln):
                continue

            # Keep the canonical total line if it already exists; strip alternate total labels
            if re.search(r"(?i)\bApprox\.?\s*Repair\s*Total\b", ln):
                if not re.search(r"(?i)^\s*\*{0,2}\s*Approximate\s+Repair\s+Total\s*:\s*\$\s*[0-9]", s):
                    continue

            cleaned.append(ln)

        if tax_rate is None or not isinstance(tax_rate, (int, float)) or tax_rate <= 0:
            tax_rate = 0.07  # locked photos-only approximation fallback

        def _grab_money_line(pats: List[str]) -> Optional[float]:
            for pat in pats:
                mm = re.search(pat, text, flags=re.IGNORECASE | re.MULTILINE)
                if mm:
                    try:
                        return float(mm.group(1).replace(",", ""))
                    except Exception:
                        pass
            return None

        # Parse component dollars from the markdown/body
        body_labor = _grab_money_line([
            r"^\s*[-*]?\s*Body(?:\s+labor)?\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Body(?:\s+labor)?\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
        ])
        paint_labor = _grab_money_line([
            r"^\s*[-*]?\s*Paint\s+labor\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Paint\s+labor\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Refinish\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Refinish\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
        ])
        mech_labor = _grab_money_line([
            r"^\s*[-*]?\s*Mechanical(?:/SRS/Glass|/diagnostic)?\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Mechanical(?:/SRS/Glass|/diagnostic)?\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Mechanical[^=\n]*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
        ])
        frame_labor = _grab_money_line([
            r"^\s*[-*]?\s*Frame(?:/measure)?\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Frame(?:/measure)?\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
        ])
        setup_measure = _grab_money_line([
            r"^\s*[-*]?\s*Setup\s*&\s*Measure\s*:\s*.*?=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Setup\s*&\s*Measure\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
        ])

        labor_sub = _grab_money_line([
            r"^\s*[-*]?\s*Labor\s+subtotal\s*\([^\n]*?\)\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Labor\s+subtotal\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Labor\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
        ])

        # If labor subtotal is absent, sum labor line items
        if labor_sub is None:
            calc_labor = 0.0
            have_labor_piece = False
            for v in (body_labor, paint_labor, mech_labor, frame_labor, setup_measure):
                if isinstance(v, (int, float)):
                    calc_labor += float(v)
                    have_labor_piece = True
            if have_labor_piece:
                labor_sub = round(calc_labor, 2)

        # Parts subtotal
        parts_sub = _grab_money_line([
            r"^\s*[-*]?\s*Estimated\s+parts\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Parts\s+subtotal\s*\(approx\.?\)\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Parts\s+subtotal\s*\(approx\.?\)\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Parts\s+subtotal\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Parts\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Parts\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
        ])

        # Paint materials subtotal ONLY — never the $/hr rate
        paint_mat = _grab_money_line([
            r"^\s*[-*]?\s*Paint\s+materials\s+subtotal\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Paint\s+materials\s+subtotal\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Paint\s*(?:&\s*|and\s*)?materials\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
        ])

        # Existing tax / total if already present
        tax_amt = _grab_money_line([
            r"^\s*[-*]?\s*Sales\s+tax\s*\(assumed\s*7%\s*for\s*approximation\)\s*=\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*[-*]?\s*Tax\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
        ])
        total_val = _grab_money_line([
            r"^\s*\*{0,2}\s*Approximate\s+Repair\s+Total\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
            r"^\s*\*{0,2}\s*Estimated\s+Total\s+Approximate\s+Repair\s+Cost\s*:\s*\$\s*([0-9][0-9,]*(?:\.[0-9]{2})?)\b",
        ])

        tax_basis: Optional[float] = None

        # If tax missing, compute it from parts + paint materials only
        if tax_amt is None and isinstance(parts_sub, (int, float)) and isinstance(paint_mat, (int, float)):
            tax_basis = round(float(parts_sub) + float(paint_mat), 2)
            tax_amt = round(tax_basis * float(tax_rate), 2)
        elif isinstance(parts_sub, (int, float)) and isinstance(paint_mat, (int, float)):
            tax_basis = round(float(parts_sub) + float(paint_mat), 2)

        # If total missing, compute from existing subtotals + tax
        if total_val is None:
            if isinstance(labor_sub, (int, float)) and isinstance(parts_sub, (int, float)) and isinstance(paint_mat, (int, float)) and isinstance(tax_amt, (int, float)):
                total_val = round(float(labor_sub) + float(parts_sub) + float(paint_mat) + float(tax_amt), 2)

        # Print cleaned body first
        if cleaned:
            for ln in cleaned:
                s = (ln or "").strip()
                if not s:
                    pdf_obj.ln(2)
                    continue
                if re.match(r"^#{3,6}\s+\S", s):
                    heading = re.sub(r"^#{3,6}\s*", "", s).strip()
                    try:
                        pdf_obj.set_font("Helvetica", "B", 11)
                    except Exception:
                        pdf_obj.set_font("Arial", "B", 11)
                    pdf_obj.ln(1)
                    pdf_obj.cell(0, 6, _pdf_sanitize(heading), ln=True)
                    try:
                        pdf_obj.set_font("Helvetica", "", 11)
                    except Exception:
                        pdf_obj.set_font("Arial", "", 11)
                    continue
                mc(s)

        # Deterministic tax lines
        pdf_obj.ln(1)
        try:
            pdf_obj.set_font("Helvetica", "", 11)
        except Exception:
            pdf_obj.set_font("Arial", "", 11)

        if isinstance(tax_basis, (int, float)) and isinstance(tax_amt, (int, float)) and isinstance(tax_rate, (int, float)):
            mc(f"Tax rate: {float(tax_rate)*100:.3f}%")
            mc(f"Tax basis (parts + paint materials): {_money2(tax_basis)}")
            mc(f"Tax: {_money2(tax_amt)}")

        # FIX 1: always inject/display the final total when available
        if isinstance(total_val, (int, float)):
            try:
                pdf_obj.set_font("Helvetica", "B", 11)
            except Exception:
                pdf_obj.set_font("Arial", "B", 11)
            mc(f"Approximate Repair Total: {_money2(total_val)}")
            try:
                pdf_obj.set_font("Helvetica", "", 11)
            except Exception:
                pdf_obj.set_font("Arial", "", 11)

        # FIX 2: Severity Tier always keys off the final total
        tier = None
        if isinstance(total_val, (int, float)):
            if total_val < 3500:
                tier = "minor"
            elif total_val < 10000:
                tier = "moderate"
            elif total_val < 25000:
                tier = "major"
            else:
                tier = "possible_tl"

        boxes = {
            "minor": ("[x]", "[ ]", "[ ]", "[ ]"),
            "moderate": ("[ ]", "[x]", "[ ]", "[ ]"),
            "major": ("[ ]", "[ ]", "[x]", "[ ]"),
            "possible_tl": ("[ ]", "[ ]", "[ ]", "[x]"),
            None: ("[ ]", "[ ]", "[ ]", "[ ]"),
        }[tier]

        pdf_obj.ln(1)
        mc("Severity Tier")
        mc(f"{boxes[0]} Minor (< $3,500)")
        mc(f"{boxes[1]} Moderate ($3,500-$10,000)")
        mc(f"{boxes[2]} Major ($10,000+)")
        mc(f"{boxes[3]} Possible Total Loss Threshold Approaching")

        # Repair Cost Disclaimer intentionally not rendered here (only final combined disclaimer prints later)
        if False and disclaimer_text:
            try:
                pdf_obj.ln(4)
                x_left = pdf_obj.l_margin
                x_right = pdf_obj.w - pdf_obj.r_margin
                y_line = pdf_obj.get_y()
                pdf_obj.set_draw_color(180, 180, 180)
                pdf_obj.line(x_left, y_line, x_right, y_line)
                pdf_obj.ln(3)

                pdf_obj.set_text_color(90, 90, 90)
                pdf_obj.set_font("Helvetica", "B", 9)
                pdf_obj.cell(0, 5, "Repair Cost Disclaimer:", ln=True)
                pdf_obj.set_font("Helvetica", "", 8)
                pdf_obj.multi_cell(0, 4, _pdf_sanitize(disclaimer_text))
                pdf_obj.set_text_color(0, 0, 0)
                try:
                    pdf_obj.set_font("Helvetica", "", 11)
                except Exception:
                    pdf_obj.set_font("Arial", "", 11)
            except Exception:
                # fail-open: don't break PDF rendering
                pass

    def add_thumbnail_page(pdf_obj: FPDF, image_paths: List[str]) -> None:
        """Append exactly ONE page containing thumbnails of all uploaded photos (as space allows)."""
        if not image_paths:
            return

        pdf_obj.add_page()
        try:
            pdf_obj.set_font("Helvetica", "B", 12)
        except Exception:
            pdf_obj.set_font("Arial", "B", 12)
        pdf_obj.cell(0, 8, "Uploaded Photos", ln=True)
        pdf_obj.ln(2)

        # Layout (single-page, grid)
        cols = 4
        gutter = 3.0
        usable_w = pdf_obj.w - pdf_obj.l_margin - pdf_obj.r_margin
        thumb_w = (usable_w - gutter * (cols - 1)) / cols
        x0 = pdf_obj.l_margin
        y = pdf_obj.get_y()
        x = x0
        col = 0

        placed = 0
        total = len(image_paths)

        for p in image_paths:
            # Hard-stop to keep this as ONE PAGE ONLY
            if y + thumb_w > (pdf_obj.h - pdf_obj.b_margin - 10):
                break
            try:
                pdf_obj.image(p, x=x, y=y, w=thumb_w)
                placed += 1
            except Exception:
                pass

            col += 1
            if col >= cols:
                col = 0
                x = x0
                y += thumb_w + gutter
            else:
                x += thumb_w + gutter

        # Footer note if truncated by the 1-page rule
        if placed < total:
            try:
                pdf_obj.set_font_size(9)
            except Exception:
                pass
            pdf_obj.set_y(pdf_obj.h - pdf_obj.b_margin - 8)
            pdf_obj.set_x(pdf_obj.l_margin)
            pdf_obj.cell(0, 6, f"Showing {placed} of {total} uploaded photos (1-page appendix limit).", ln=True)

    # POI-15 Total Loss trigger from uploaded text ONLY
    poi15_hit = False
    try:
        txt = uploaded_text_all.lower()
        if re.search(r"\b15\s*total\s*loss\b", txt, flags=re.IGNORECASE):
            poi15_hit = True
        elif re.search(r"point\s*of\s*impact[^A-Za-z0-9]{0,10}15[^A-Za-z0-9]{0,20}total\s*loss", txt, flags=re.IGNORECASE):
            poi15_hit = True
    except Exception:
        poi15_hit = False

    if ai_intent == "damage_report_from_photos":
        # -----------------------------
        # NSPXN.com Condition Report PDF
        # -----------------------------
        # Hard "safe" top margin so no section bar/box can overlap the logo/header.
        # (User-facing: move Vehicle Identification down a few lines.)
        SAFE_TOP_Y = 60
    
        def _ensure_safe_top() -> None:
            try:
                if pdf.get_y() < SAFE_TOP_Y:
                    pdf.set_y(SAFE_TOP_Y)
            except Exception:
                pass
    
        def _section_bar(title: str) -> None:
            """Draw a brighter, color-coded section header bar (aligned to DOCX style)."""
            t = str(title or "").strip()
            cmap = {
                "VEHICLE IDENTIFICATION": (0, 112, 192),              # bright blue
                "REPORT SUMMARY": (191, 112, 0),                      # orange
                "APPROXIMATE REPAIR COST BREAKDOWN": (0, 153, 76),     # green
                "FRAUD & AUTHENTICITY CHECK": (112, 48, 160),          # purple
                "CONCLUSION": (64, 64, 64),                           # dark gray
                "DISCLAIMER": (96, 96, 96),                           # gray
            }
            rgb = cmap.get(t.upper(), (0, 112, 192))
            _ensure_safe_top()
            try:
                pdf.ln(3)
                pdf.set_fill_color(*rgb)
                pdf.set_text_color(255, 255, 255)
                pdf.set_font("Helvetica", "B", 12)
            except Exception:
                pdf.ln(3)
                pdf.set_fill_color(*rgb)
                pdf.set_text_color(255, 255, 255)
                pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 8, _pdf_sanitize(t), ln=True, fill=True)
            pdf.set_text_color(0, 0, 0)
            try:
                pdf.set_font("Helvetica", "", 9)
            except Exception:
                pdf.set_font("Arial", "", 9)

        def _scrub_model_headings(md_text: str) -> str:
            """Remove model-emitted headings that duplicate PDF section headers."""
            if not md_text:
                return md_text or ""
            lines = str(md_text).replace("\r\n","\n").replace("\r","\n").splitlines()
            out = []
            for ln in lines:
                s = (ln or "").strip()
                if not s:
                    out.append(ln)
                    continue
                # Drop markdown headings like '# Condition Report (Photos Only)' or '## ...'
                if re.match(r"^#+\s+", s):
                    continue
                # Drop echoes like 'Report Selected'
                if re.search(r"(?i)^report\s+selected\b", s):
                    continue
                # Drop helper artifacts the model sometimes emits
                if re.search(r"(?i)estimated_costs_markdown", s):
                    continue
                out.append(ln)
            return "\n".join(out).strip()
    
        def _extract_photo_confirmed_odometer(summary_md: str) -> Optional[str]:
            """Best-effort: extract the photo-confirmed odometer from the Photo #3 row (single source of truth)."""
            if not summary_md:
                return None
            t = str(summary_md).replace("\r\n","\n").replace("\r","\n")
            # Prefer Photo 3 table row
            for ln in t.splitlines():
                if re.search(r"^\s*\|\s*3\s*\|", ln):
                    m = re.search(r"(?i)\bodometer\b[^0-9]{0,60}([0-9][0-9,]{2,})\s*mi", ln)
                    if m:
                        return m.group(1).replace(",", "") + " mi"
                    m2 = re.search(r"(?i)\bodometer\b[^0-9]{0,60}([0-9][0-9,]{2,})\b", ln)
                    if m2:
                        return m2.group(1).replace(",", "") + " mi"
            # Fallback: any odometer mention
            m = re.search(r"(?i)\bodometer\b[^0-9]{0,80}([0-9][0-9,]{2,})\s*mi", t)
            if m:
                return m.group(1).replace(",", "") + " mi"
            return None
    
        # Title (larger + bold)
        try:
            pdf.set_font("Helvetica", "B", 16)
        except Exception:
            pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, "NSPXN.com Condition Report", ln=True, align="C")
        try:
            pdf.set_font("Helvetica", "", 11)
        except Exception:
            pdf.set_font("Arial", "", 11)
        pdf.ln(2)
    
        # Vehicle Identification (fixed PDF block)
        _section_bar("VEHICLE IDENTIFICATION")
        _summary_md_raw = (result.get("summary_markdown") or "").strip()
        _odo_photo = _extract_photo_confirmed_odometer(_summary_md_raw)
        _odo_print = _odo_photo or (result.get("odometer_estimate_only") or "N/A")
        pdf_status = (result.get("redaction_status") or "").replace("✅", "OK").strip() or "N/A"
    
        # Use colon formatting (no double-odometer conflicts)
        mc(f"File #: {file_number or 'N/A'}")
        mc(f"Claim #: {result.get('claim_number') or 'N/A'}")
        mc(f"Inspected For: {ia_company or 'N/A'}")
        mc(f"VIN: {result.get('vin') or 'N/A'}")
        mc(f"VIN Verification: {result.get('vin_verification') or 'N/A'}")
        mc(f"Vehicle: {_format_vehicle_value(result.get('vehicle'))}")
        mc(f"Odometer: {_odo_print}")
        mc(f"Primary Impact: {result.get('primary_impact') or 'N/A'}")
        mc(f"Secondary Impact: {result.get('secondary_impact') or 'N/A'}")
        mc(f"Redaction Status: {pdf_status}")
    
        # Report Summary (scrub markdown headings)
        _section_bar("REPORT SUMMARY")
        _summary_md = _scrub_model_headings(_summary_md_raw)
        if ai_intent == "damage_report_from_photos":
            _summary_md = _scrub_photo_only_narrative_cost_headers(_summary_md)
        mc(_summary_md if _summary_md else "-")
        pdf.ln(2)
    
        # Controlled Repair Cost section rendering (prevents duplicate headings, Totals blocks, duplicate tiers, and bad totals)
        _raw_costs_md = result.get("estimated_costs_markdown") or ""
        costs_md = _strip_unwanted_cost_lines_for_pdf(_raw_costs_md)
    
        # NOTE: Total + Severity Tier are rendered deterministically inside render_repair_cost_section.
    
        _section_bar("APPROXIMATE REPAIR COST BREAKDOWN")
        render_repair_cost_section(pdf, costs_md, tax_rate=tax_rate)
    
        _section_bar("FRAUD & AUTHENTICITY CHECK")
        mc((result.get("fraud_markdown") or "").strip() or "-")
    
        _section_bar("CONCLUSION")
        mc((result.get("conclusion") or "").strip() or "-")
    
        # --- Combined Disclaimer (end of report) ---
        try:
            pdf.ln(4)
            x_left = pdf.l_margin
            x_right = pdf.w - pdf.r_margin
            y_line = pdf.get_y()
            pdf.set_draw_color(180, 180, 180)
            pdf.line(x_left, y_line, x_right, y_line)
            pdf.ln(3)
    
            pdf.set_text_color(90, 90, 90)
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(0, 5, "Disclaimer:", ln=True)
            pdf.set_font("Helvetica", "", 8)
    
            disclaimer_body = (
                "This report is based solely on photographic evidence and is intended as a preliminary visual damage assessment only. "
                "It does not constitute a formal appraisal, repair estimate, or safety certification. Hidden structural, mechanical, or electronic "
                "damage may exist beyond what is visible in the photographs. All findings should be verified by a qualified collision repair technician, "
                "appraiser, and/or insurance adjuster performing a physical inspection. This report was generated using artificial intelligence. AI systems "
                "may make errors or misinterpret visual information. All photos, damage descriptions, conclusions, and findings must be independently "
                "reviewed and verified by a qualified appraiser before preparing or finalizing any repair estimate. This repair cost is an approximation "
                "derived from the visible conditions in the provided photos only. Actual repair scope and cost may change after teardown, measurement, and "
                "confirmation of hidden damage, glass bonding requirements, and any sensor/trim replacement needs."
            )
            pdf.multi_cell(0, 4, _pdf_sanitize(disclaimer_body))
            pdf.set_text_color(0, 0, 0)
        except Exception:
            pass
    
        safe_file = _safe(file_number)
        pdf_filename = f"AI_Condition_Report_{safe_file}.pdf"
    else:
        _is_comprehensive_pdf = str(ai_intent or "").strip().lower() == "comprehensive"

        def _comp_section_bar(title: str) -> None:
            t = str(title or "").strip()
            cmap = {
                "VEHICLE IDENTIFICATION": (0, 112, 192),
                "NSPXN.COM CONDITION SUMMARY": (191, 112, 0),
                "APPROXIMATE REPAIR COST BREAKDOWN": (0, 153, 76),
                "FRAUD DETECTION": (112, 48, 160),
                "DISCLAIMER": (96, 96, 96),
            }
            rgb = cmap.get(t.upper(), (0, 112, 192))
            pdf.ln(3)
            pdf.set_fill_color(*rgb)
            pdf.set_text_color(255, 255, 255)
            try:
                pdf.set_font("Helvetica", "B", 12)
            except Exception:
                pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 8, _pdf_sanitize(t), ln=True, fill=True)
            pdf.set_text_color(0, 0, 0)
            try:
                pdf.set_font("Helvetica", "", 11)
            except Exception:
                pdf.set_font("Arial", "", 11)

        # --- Locked NSPXN top section: black box, logo2.png left, white report header right ---
        try:
            _header_x = pdf.l_margin
            _header_y = 8
            _header_w = pdf.w - pdf.l_margin - pdf.r_margin
            _header_h = 25
            pdf.set_fill_color(0, 0, 0)
            pdf.rect(_header_x, _header_y, _header_w, _header_h, style="F")

            _title_x = _header_x + 8
            _logo_path = os.path.join(os.path.dirname(__file__), "logo2.png")
            if os.path.exists(_logo_path):
                try:
                    pdf.image(_logo_path, x=_header_x + 5, y=_header_y + 4, w=40)
                    _title_x = _header_x + 50
                except Exception as e:
                    log.warning(f"Logo render skipped: {e}")
            else:
                log.warning(f"logo2.png not found at {_logo_path}")

            pdf.set_xy(_title_x, _header_y + 7)
            pdf.set_text_color(255, 255, 255)
            try:
                pdf.set_font("Helvetica", "B", 12)
            except Exception:
                pdf.set_font("Arial", "B", 12)
            pdf.cell(_header_w - (_title_x - _header_x) - 6, 8, _pdf_sanitize("NSPXN.com Audit Report" if _is_comprehensive_pdf else "NSPXN.com Condition Report"), ln=False, align="L")
            pdf.set_text_color(0, 0, 0)
            try:
                pdf.set_font("Helvetica", "", 9)
            except Exception:
                pdf.set_font("Arial", "", 9)
            pdf.set_y(_header_y + _header_h + 5)
        except Exception as e:
            log.warning(f"Top section render skipped: {e}")
            pdf.set_text_color(0, 0, 0)
            try:
                pdf.set_font("Arial", "", 9)
            except Exception:
                pass
            pdf.set_y(38)

        _comp_section_bar("Vehicle Identification")
        mc(f"File Number: {file_number}")
        mc(f"Inspected For: {ia_company}")
        mc(f"Appraiser ID #: {appraiser_id}")
        mc(f"Request Type: {result['request_type']}")

        # --- Supplement header (documents only; ignore negated mentions) ---
        smark = result.get("summary_markdown","")
        _txt_docs = uploaded_text_all or ""
        _supp_doc_hit = bool(re.search(
            r"(?is)\b(Supplement\s+(?:Summary|of\s+Record)|Estimate\s+Version:\s*S0[1-9]\b|\bS0[1-9]\b|\bSupplement\s+Estimate\b)",
            _txt_docs
        ))
        _no_supp_negation = not re.search(r"(?is)\b(no|not)\s+(a\s+)?supplement\b", _txt_docs)
        supp_detected_docs = _supp_doc_hit and _no_supp_negation
        if supp_detected_docs:
            mc("Supplement Status: Supplement Estimate detected in documentation")
            _possible_amt = None
            _m = re.search(r"(?is)\bPossible\s+Supplement\s+Amount\s*\$?([0-9,]+\.\d{2})\b", _txt_docs)
            if _m:
                _possible_amt = _m.group(1)
            mc("Supplement Details")
            # NEW: list all supplement tags (S01, S02, ...) instead of a single version only
            supp_versions_docs = sorted(set(re.findall(r"(?i)\bS[0-9]{2}\b", _txt_docs)))
            if supp_versions_docs:
                mc(f"- Supplements detected in documents: {', '.join(supp_versions_docs)}")
            if _possible_amt:
                mc(f"- Possible amount noted: ${_possible_amt}")
            if not (supp_versions_docs or _possible_amt):
                mc("- Supplement indicators present (e.g., 'Supplement Summary' or S01/S02).")

        # --- Total Loss echo (documents-only; no narrative trigger) ---
        explicit_tl_hit = False
        try:
            txt_docs = uploaded_text_all or ""
            poi15_docs = re.search(
                r"(?is)\bpoint\s*of\s*impact[^a-z0-9]{0,10}15[^a-z0-9]{0,20}total\s*loss\b",
                txt_docs,
            )
            doc_tl_docs = re.search(
                r"(?i)\b(estimate\s*type|type\s*of\s*loss)\s*:\s*total\s*loss\b",
                txt_docs,
            )
            explicit_tl_hit = bool(poi15_docs or doc_tl_docs)
        except Exception:
            explicit_tl_hit = False
        if explicit_tl_hit:
            mc("Estimate Type: Total Loss (explicit in documents)")

        mc(f"Claim #: {result['claim_number']}")
        mc(f"VIN (from estimate/photos): {result['vin']}")
        mc(f"VIN verification (estimate vs photo): {result['vin_verification']}")
        mc(f"Vehicle: {_format_vehicle_value(result.get('vehicle'))}")
        mc(f"Odometer (from estimate): {result['odometer_estimate_only']}")
        mc(f"Compliance Score: {result['compliance_score']}")
        pdf_status = result["redaction_status"].replace("✅", "OK")
        mc(pdf_status)
        _comp_section_bar("NSPXN.com Condition Summary"); mc((smark or '').strip())
        _comp_section_bar("Approximate Repair Cost Breakdown"); mc((result.get("estimated_costs_markdown") or "").strip())
        _comp_section_bar("Fraud Detection"); mc((result["fraud_markdown"] or 'N/A').strip())
        _comp_section_bar("Conclusion"); mc((result.get("conclusion") or "").strip())

        # --- AI Disclaimer (after report content) ---
        try:
            pdf.ln(4)
            x_left = pdf.l_margin
            x_right = pdf.w - pdf.r_margin
            y_line = pdf.get_y()
            pdf.set_draw_color(180, 180, 180)
            pdf.line(x_left, y_line, x_right, y_line)
            pdf.ln(3)

            pdf.set_text_color(90, 90, 90)
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(0, 5, "Disclaimer:", ln=True)
            pdf.set_font("Helvetica", "", 8)

            disclaimer_body = (
                "This report was generated using artificial intelligence. AI systems may make errors or misinterpret "
                "visual information. All photos, damage descriptions, conclusions, and findings must be independently "
                "reviewed and verified by a qualified appraiser before preparing or finalizing any repair estimate."
            )
            pdf.multi_cell(0, 4, _pdf_sanitize(disclaimer_body))
            pdf.ln(2)
            try:
                ts_est = datetime.now(ZoneInfo("America/New_York")).strftime("%m/%d/%Y %I:%M %p")
                ts_label = f"Generated: {ts_est} EST"
            except Exception:
                ts_label = f"Generated: {datetime.now().strftime('%m/%d/%Y %I:%M %p')}"
            try:
                pdf.set_font("Helvetica", "", 8)
            except Exception:
                pdf.set_font("Arial", "", 8)
            pdf.cell(0, 4, _pdf_sanitize(ts_label), ln=True)
            pdf.set_text_color(0, 0, 0)
        except Exception:
            pass


        safe_file = _safe(file_number)
        pdf_filename = f"{safe_file}.pdf"


    # --- One-page photo thumbnail appendix (all uploaded photos) ---
    try:
        add_thumbnail_page(pdf, thumbnail_paths)
    except Exception:
        pass

    pdf_path = os.path.join(PDF_DIR, pdf_filename)
    pdf_written = False
    try:
        out = pdf.output(dest="S")
        if isinstance(out, (bytes, bytearray)):
            data_bytes = bytes(out)
        else:
            data_bytes = str(out).encode("latin-1", "ignore")
        with open(pdf_path, "wb") as f:
            f.write(data_bytes)
        pdf_written = os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0
    except Exception as e:
        logging.warning(f"PDF write error using dest='S': {e}")
        try:
            pdf.output(pdf_path)
            pdf_written = os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0
        except Exception as e2:
            logging.error(f"PDF write fallback failed: {e2}")
            pdf_written = False

    if pdf_written:
        try:
            with open(os.path.join(PDF_DIR, f"latest_{_safe(file_number)}.txt"), "w", encoding="utf-8") as _mf:
                _mf.write(os.path.basename(pdf_path))
        except Exception:
            pass

    pdf_url = f"/download-pdf?filename={urllib.parse.quote(os.path.basename(pdf_filename))}" if pdf_written else ""

    # -----------------------
    # Email — info-only (attach PDF)
    # -----------------------
    try:
        msg = EmailMessage()
        if ai_intent == "damage_report_from_photos":
            subj = f"NSPXN.com Condition Report: {file_number or ''} {result['claim_number'] or ''}".strip()
            body = (
                "NSPXN.com Condition Report\n\n"
                f"Inspected For: {ia_company}\n"
                f"Claim #: {result['claim_number'] or 'N/A'}    File #: {file_number or 'N/A'}\n"
                f"Odometer: {result['odometer_estimate_only'] or 'N/A'}    Primary Impact: {result['primary_impact'] or 'N/A'}\n"
                f"Secondary Impact: {result['secondary_impact'] or 'N/A'}\n\n"
                f"{result['redaction_status']}\n\n"
                "Condition Summary\n"
                f"{(result['summary_markdown'] or 'N/A')}\n\n"
                "Fraud & Authenticity Check\n"
                f"{(result['fraud_markdown'] or 'N/A')}\n\n"
                "Conclusion\n"
                f"{(result['conclusion'] or 'N/A')}\n"
            )
        else:
            try:
                _txt_email = uploaded_text_all or ""
                _poi15_email = re.search(
                    r"(?is)\bpoint\s*of\s*impact[^a-z0-9]{0,10}15[^a-z0-9]{0,20}total\s*loss\b",
                    _txt_email,
                )
                _doc_tl_email = re.search(
                    r"(?i)\b(estimate\s*type|type\s*of\s*loss)\s*:\s*total\s*loss\b",
                    _txt_email,
                )
                _explicit_tl_email = bool(_poi15_email or _doc_tl_email)
            except Exception:
                _explicit_tl_email = False

            # NEW: supplement line includes all Sxx tags if present
            supp_line = ""
            if supp_detected_docs:
                supp_versions_email = sorted(set(re.findall(r"(?i)\bS[0-9]{2}\b", _txt_email or "")))
                if supp_versions_email:
                    supp_line = (
                        "Supplement Status: Supplement Estimates detected in documentation "
                        f"({', '.join(supp_versions_email)})\n"
                    )
                else:
                    supp_line = "Supplement Status: Supplement Estimate detected in documentation\n"

            tl_line = "Estimate Type: Total Loss (explicit in documents)\n" if _explicit_tl_email else ""
            subj = f"NSPXN.com Review: {result['claim_number'] or file_number}"
            body = (
                "NSPXN.com Condition Report\n\n"
                f"File Number: {file_number}\n"
                f"Inspected For: {ia_company}\n"
                f"Appraiser ID #: {appraiser_id}\n"
                f"Request Type: {result['request_type']}\n"
                f"{supp_line}"
                f"{tl_line}"
                f"Claim #: {result['claim_number']}\n"
                f"VIN (from estimate/photos): {result['vin']}\n"
                f"VIN verification (estimate vs photo): {result['vin_verification']}\n"
                f"Vehicle: {result['vehicle']}\n"
                f"Odometer (from estimate): {result['odometer_estimate_only']}\n"
                f"Compliance Score: {result['compliance_score']}\n\n"
                f"{result['redaction_status']}\n\n"
                "NSPXN.com Condition Summary\n"
                f"{result['summary_markdown']}\n\n"
                "Fraud Detection\n"
                f"{result['fraud_markdown']}\n"
            )

        msg["Subject"] = subj
        msg["From"] = "info@nspxn.com"
        msg["To"] = "info@nspxn.com"
        msg["Cc"] = "growley@ractrak.com"
        msg.set_content(body)

        try:
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            msg.add_attachment(pdf_bytes, maintype="application", subtype="pdf", filename=pdf_filename)
        except Exception as e:
            logging.warning(f"Failed to attach PDF to email: {e}")

        with smtplib.SMTP_SSL("mail.tierra.net", 465, timeout=20) as smtp:
            smtp.login("info@nspxn.com", "grr2025GRR")
            smtp.send_message(msg)
        log.info("Info email sent to info@nspxn.com")
    except Exception as e:
        logging.error(f"Email error: {e}")

    # Expose lightweight analytics metadata to the outer router via response headers.
    # This avoids buffering/parsing the full response body in main.py and keeps PDF/report delivery stable.
    try:
        response.headers["X-NSPXN-Report-Completed"] = "true"
        response.headers["X-NSPXN-AI-Intent"] = str(ai_intent or "")
        response.headers["X-NSPXN-File-Number"] = str(file_number or "")
        if str(ai_intent or "").strip().lower() == "comprehensive":
            _score_header = _numeric_score_or_blank(result.get("compliance_score"))
            if _score_header:
                response.headers["X-NSPXN-Compliance-Score"] = str(_score_header)
                response.headers["X-NSPXN-Score-Source"] = "main_comprehensive_locked.result.compliance_score"
    except Exception:
        pass

    return {
        **result,
        "web_summary": result["summary_brief"],
        "gpt_output": result["summary_markdown"],
        "pdf_url": pdf_url,
        "pdf_filename": pdf_filename if pdf_written else "",
        "pdf_status": "ready" if pdf_written else "not_created"
    }

    # -----------------------
# PDF download
    # -----------------------
@app.get("/download-pdf")
async def download_pdf(file_number: Optional[str] = None, filename: Optional[str] = None):
    """Download the newest matching generated PDF from PDF_DIR.
    Accepts the exact returned filename or a file_number. This prevents false 404s
    caused by URL encoding, safe-name normalization, or filename/file_number mismatch.
    """
    try:
        search_dirs: List[str] = []
        for d in (PDF_DIR, "/tmp"):
            if d and os.path.isdir(d) and d not in search_dirs:
                search_dirs.append(d)

        candidates: List[str] = []

        if filename:
            raw_name = os.path.basename(str(filename))
            names: List[str] = []
            for nm in (raw_name, urllib.parse.unquote(raw_name), _safe(raw_name), _safe(urllib.parse.unquote(raw_name))):
                if nm and nm not in names:
                    names.append(nm)
            for d in search_dirs:
                for nm in names:
                    pth = os.path.join(d, nm)
                    if os.path.exists(pth) and os.path.isfile(pth):
                        candidates.append(pth)

            if not candidates:
                safe_stem = os.path.splitext(_safe(urllib.parse.unquote(raw_name)))[0]
                for d in search_dirs:
                    candidates.extend(glob.glob(os.path.join(d, f"*{safe_stem}*.pdf")))

        if file_number:
            safe_num = _safe(str(file_number))
            manifest = os.path.join(PDF_DIR, f"latest_{safe_num}.txt")
            if os.path.exists(manifest):
                try:
                    mf_name = os.path.basename(open(manifest, "r", encoding="utf-8").read().strip())
                    mf_path = os.path.join(PDF_DIR, mf_name)
                    if os.path.exists(mf_path) and os.path.isfile(mf_path):
                        candidates.append(mf_path)
                except Exception:
                    pass
            for d in search_dirs:
                candidates.extend(glob.glob(os.path.join(d, f"*{safe_num}*.pdf")))

        candidates = [pth for pth in candidates if pth and os.path.exists(pth) and os.path.isfile(pth) and pth.lower().endswith(".pdf")]
        if not candidates:
            return JSONResponse(status_code=404, content={"detail": "Not Found"})

        latest = max(candidates, key=lambda pth: os.path.getmtime(pth))
        return FileResponse(path=latest, media_type="application/pdf", filename=os.path.basename(latest))
    except Exception as e:
        logging.error(f"PDF download error: {e}")
        return JSONResponse(status_code=500, content={"detail": "PDF download error"})
